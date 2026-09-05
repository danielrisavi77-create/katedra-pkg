#!/usr/bin/env python3
"""Usporedba rada koji je Katedra izradila s verzijom koju je autor uredio.

Uporaba:
  python3 provjeri_povratak.py IZVORNI.docx VRACENI.docx
  python3 provjeri_povratak.py IZVORNI.docx VRACENI.docx --json .katedra/povratak.json

Zašto postoji
-------------
`katedra-lite/references/povratak.md` i `references/stil_autora.md` opisuju cijeli
mod 7 („malo sam izmijenio, usporedi s onim što si mi dao") i imenuju OVU skriptu
kao njegov motor. Skripte nije bilo. Mod je postojao kao proza, pa se usporedba
radila napamet: regresije koje je autor slučajno unio nisu se vraćale, a njegov
glas se nije pamtio.

Nalazi se dijele na tri hrpe, točno kako povratak.md propisuje:

* **vraćamo** — regresija: autor je nehotice pokvario nešto što je bilo ispravno
  (raspon 133–150 → 133-150, hrvatski navodnici → ravni, obrisan citat, nestao
  obvezni dio). Ovo se javlja autoru i vraća.
* **pamtimo** — glas: rečenica je prepisana bez diranja citata i brojki. Autor je
  u pravu, mi smo pisali njemu strano. Ide u `stil_autora.json`.
* **pitamo** — sve ostalo: promijenjena brojka, dodan sadržaj, promijenjen citat.
  Nijedan alat to ne smije presuditi sam.

Izlazni kod: 1 ako ima regresija ili nestalih obveznih dijelova, inače 0.
"""
from __future__ import annotations

import argparse
import difflib
import json
import os
import re
import sys
import zipfile

try:
    import docx  # noqa: F401
    from docx import Document
except ImportError:
    print("❌ nedostaje python-docx (pip install python-docx --break-system-packages)",
          file=sys.stderr)
    sys.exit(2)

W_T = re.compile(r"<w:t\b[^>]*>(.*?)</w:t>", re.S)

CITAT = re.compile(
    r"\[\d{1,3}(?:[\s,–\-]+\d{1,3})*\]"                       # IEEE
    r"|\([A-ZČĆŠŽĐ][\wčćžšđ\-]+[^)]{0,60}?\d{4}[a-z]?[^)]{0,20}\)"  # autor-godina
)
BROJ = re.compile(r"(?<![\w.,])\d+(?:[.,]\d+)?(?![\w])")

# (opis, uzorak u IZVORNOM, uzorak u VRAĆENOM) — regresija je kad je prvi nestao,
# a drugi se pojavio na istom mjestu.
REGRESIJE = [
    ("raspon: en-crtica zamijenjena spojnicom", r"\d–\d", r"\d-\d"),
    ("hrvatski navodnici zamijenjeni ravnima", r"„", r"\""),
    ("znak množenja × zamijenjen slovom x", r"\d\s*×\s*\d", r"\d\s*[xX]\s*\d"),
    ("nedjeljivi razmak izgubljen", r"\d ", r"\d "),
    ("decimalni zarez zamijenjen točkom", r"\d,\d", r"\d\.\d"),
]


def _tekst(put: str) -> tuple[list[str], str]:
    d = Document(put)
    odlomci = [p.text.strip() for p in d.paragraphs if p.text.strip()]
    for t in d.tables:
        for red in t.rows:
            for c in red.cells:
                if c.text.strip():
                    odlomci.append(c.text.strip())
    fus = ""
    try:
        with zipfile.ZipFile(put) as z:
            for dio in ("word/footnotes.xml", "word/endnotes.xml"):
                try:
                    fus += "".join(W_T.findall(z.read(dio).decode("utf-8")))
                except KeyError:
                    pass
    except (OSError, zipfile.BadZipFile):
        pass
    return odlomci, fus


def _naslovi(put: str) -> list[str]:
    d = Document(put)
    return [p.text.strip() for p in d.paragraphs
            if p.style.name.lower().startswith("heading") and p.text.strip()]


def _kljucni_dijelovi(naslovi: list[str]) -> set[str]:
    trazeni = ("sažetak", "sazetak", "summary", "uvod", "zaključak", "zakljucak",
               "literatura", "izjava", "sadržaj", "sadrzaj", "metodolog")
    return {n for n in naslovi if any(t in n.lower() for t in trazeni)}


def usporedi(izvorni: str, vraceni: str) -> dict:
    a_odl, a_fus = _tekst(izvorni)
    b_odl, b_fus = _tekst(vraceni)
    a_tekst, b_tekst = "\n".join(a_odl) + a_fus, "\n".join(b_odl) + b_fus

    nalaz = {"vracamo": [], "pamtimo": [], "pitamo": [],
             "nestali_dijelovi": [], "statistika": {}}

    # 1) nestali obvezni dijelovi
    a_kljucni = _kljucni_dijelovi(_naslovi(izvorni))
    b_kljucni_l = {n.lower() for n in _kljucni_dijelovi(_naslovi(vraceni))}
    for n in sorted(a_kljucni):
        if n.lower() not in b_kljucni_l:
            nalaz["nestali_dijelovi"].append(n)

    # 2) izgubljeni citati i brojke (na razini cijelog dokumenta)
    a_cit, b_cit = set(CITAT.findall(a_tekst)), set(CITAT.findall(b_tekst))
    for c in sorted(a_cit - b_cit):
        nalaz["vracamo"].append({"vrsta": "izgubljen_citat", "sto": c,
                                 "zasto": "citat postoji u izvorniku, u vraćenoj verziji ga nema"})
    a_br, b_br = set(BROJ.findall(a_tekst)), set(BROJ.findall(b_tekst))
    for b in sorted(a_br - b_br):
        nalaz["pitamo"].append({"vrsta": "nestala_brojka", "sto": b,
                                "zasto": "brojka iz izvornika ne postoji u vraćenoj verziji"})

    # 3) usporedba odlomak po odlomak
    matcher = difflib.SequenceMatcher(None, a_odl, b_odl, autojunk=False)
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            continue
        stari = a_odl[i1:i2]
        novi = b_odl[j1:j2]
        if tag == "delete":
            for s in stari:
                nalaz["pitamo"].append({"vrsta": "obrisan_odlomak", "sto": s[:200],
                                        "zasto": "odlomak iz izvornika je uklonjen"})
            continue
        if tag == "insert":
            for s in novi:
                nalaz["pitamo"].append({"vrsta": "dodan_odlomak", "sto": s[:200],
                                        "zasto": "novi tekst nije prošao evidence gate"})
            continue
        # replace: uparuj po redu i klasificiraj
        for s, n in zip(stari, novi):
            nalaz_odlomka = _klasificiraj(s, n)
            nalaz[nalaz_odlomka[0]].append(nalaz_odlomka[1])
        for s in stari[len(novi):]:
            nalaz["pitamo"].append({"vrsta": "obrisan_odlomak", "sto": s[:200],
                                    "zasto": "odlomak iz izvornika je uklonjen"})
        for n in novi[len(stari):]:
            nalaz["pitamo"].append({"vrsta": "dodan_odlomak", "sto": n[:200],
                                    "zasto": "novi tekst nije prošao evidence gate"})

    nalaz["statistika"] = {
        "odlomaka_izvornik": len(a_odl), "odlomaka_vraceno": len(b_odl),
        "citata_izvornik": len(a_cit), "citata_vraceno": len(b_cit),
        "vracamo": len(nalaz["vracamo"]), "pamtimo": len(nalaz["pamtimo"]),
        "pitamo": len(nalaz["pitamo"]),
    }
    return nalaz


def _klasificiraj(stari: str, novi: str) -> tuple[str, dict]:
    """Regresija, glas ili pitanje — po tome ŠTO se promijenilo, ne koliko."""
    pokvareno = [opis for opis, u_starom, u_novom in REGRESIJE
                 if re.search(u_starom, stari) and re.search(u_novom, novi)
                 and not re.search(u_starom, novi)]
    if pokvareno:
        return "vracamo", {"vrsta": "regresija", "opis": "; ".join(pokvareno),
                           "prije": stari[:180], "poslije": novi[:180],
                           "zasto": "izvornik je bio ispravan; izmjena ga je pokvarila"}

    if set(CITAT.findall(stari)) != set(CITAT.findall(novi)):
        return "pitamo", {"vrsta": "promijenjen_citat",
                          "prije": stari[:180], "poslije": novi[:180],
                          "zasto": "citati se razlikuju; nijedan alat to ne smije presuditi sam"}

    if set(BROJ.findall(stari)) != set(BROJ.findall(novi)):
        return "pitamo", {"vrsta": "promijenjena_brojka",
                          "prije": stari[:180], "poslije": novi[:180],
                          "zasto": "brojka je izvedena iz modela; izmjena u tekstu ju razilazi"}

    return "pamtimo", {"vrsta": "glas_autora",
                       "prije": stari[:180], "poslije": novi[:180],
                       "zasto": "citati i brojke netaknuti — ovo je autorov stil, ne greška"}


def ispisi(n: dict, izvorni: str, vraceni: str) -> None:
    print("=" * 64)
    print("POVRATAK IZ WORDA")
    print(f"  izvornik: {izvorni}")
    print(f"  vraćeno:  {vraceni}")
    print("=" * 64)
    s = n["statistika"]
    print(f"odlomaka {s['odlomaka_izvornik']} → {s['odlomaka_vraceno']} | "
          f"citata {s['citata_izvornik']} → {s['citata_vraceno']}")

    if n["nestali_dijelovi"]:
        print(f"\n⛔ NESTALI OBVEZNI DIJELOVI ({len(n['nestali_dijelovi'])}):")
        for d in n["nestali_dijelovi"]:
            print(f"   • {d}")
        print("   Rad kojemu fali dio pada formalno, prije nego ga itko pročita.")

    for hrpa, znak, naslov, uputa in [
        ("vracamo", "❌", "VRAĆAMO (regresije)",
         "Autor je nehotice pokvario ono što je bilo ispravno. Javi mu i vrati."),
        ("pitamo", "❓", "PITAMO",
         "Odluka je autorova. Ne mijenjaj bez odgovora."),
        ("pamtimo", "✎", "PAMTIMO (glas autora)",
         "Autor je u pravu: pisali smo njemu strano. Ide u stil_autora.json."),
    ]:
        stavke = n[hrpa]
        print(f"\n{znak} {naslov} ({len(stavke)})")
        if not stavke:
            print("   nema")
            continue
        for x in stavke[:12]:
            opis = x.get("opis") or x.get("vrsta")
            print(f"   • {opis}")
            if "prije" in x:
                print(f"       prije:   {x['prije']}")
                print(f"       poslije: {x['poslije']}")
            elif "sto" in x:
                print(f"       {x['sto']}")
        if len(stavke) > 12:
            print(f"   … još {len(stavke) - 12}")
        print(f"   → {uputa}")

    print("\nPopravci idu u RUKOPIS, dokument se gradi iznova. "
          "Vraćeni .docx se nikad ne krpa (željezno pravilo 1).")


def main(argv=None) -> int:
    ap = argparse.ArgumentParser(description="Usporedi izvorni i autorom uređeni .docx.")
    ap.add_argument("izvorni")
    ap.add_argument("vraceni")
    ap.add_argument("--pdf", help="PDF vraćene verzije — za provjeru brojeva stranica")
    ap.add_argument("--json", dest="json_out", metavar="PUT")
    a = ap.parse_args(argv)

    for put in (a.izvorni, a.vraceni):
        if not os.path.isfile(put):
            print(f"❌ nema datoteke: {put}", file=sys.stderr)
            return 2

    n = usporedi(a.izvorni, a.vraceni)
    ispisi(n, a.izvorni, a.vraceni)

    if a.pdf:
        print(f"\nℹ️  brojevi stranica: pokreni provjeri_reference.py {a.pdf}")

    if a.json_out:
        os.makedirs(os.path.dirname(os.path.abspath(a.json_out)), exist_ok=True)
        with open(a.json_out, "w", encoding="utf-8") as fh:
            json.dump({"izvorni": a.izvorni, "vraceni": a.vraceni, **n},
                      fh, ensure_ascii=False, indent=2)
        print(f"\n✔ JSON: {a.json_out}")

    return 1 if (n["vracamo"] or n["nestali_dijelovi"]) else 0


if __name__ == "__main__":
    sys.exit(main())
