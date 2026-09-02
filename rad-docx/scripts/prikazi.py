#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Prikaz kao nedjeljiv blok: natpis + tablica/slika + redak „Izvor:" na istoj stranici.

Radi nad GOTOVIM .docx-om, jer pandoc ne zna ni za keepNext ni za cantSplit.

  · natpis („Tablica 3. …", „Grafikon 1. …")  → keepNext + keepLines
  · svaki red tablice                          → cantSplit
  · svaka ćelija tablice                       → keepNext   (drži tablicu uz izvor)
  · odlomak sa slikom                          → keepNext + keepLines + prored auto
  · redak „Izvor:"                             → keepLines
  · prijelom stranice pred blokom iz --prelomi

Usput izvlači `blokovi.json` (ključ + natpis) koji treba izmjeri.py.

    python3 prikazi.py rad.docx [--prelomi prelomi.json] [--blokovi-out blokovi.json]
                                [--u-mjestu | --izlaz rad2.docx]
"""

import argparse
import json
import re
import shutil
import sys

try:
    import docx
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("nedostaje python-docx")

NATPIS = re.compile(r"^(Tablica|Grafikon|Slika)\s+(\d+)\s*[.:]")
IZVOR = re.compile(r"^\s*Izvor\s*:", re.I)


def kljuc_natpisa(m):
    return f"{m.group(1).lower()}{m.group(2)}"


def pPr(p):
    el = p._p.get_or_add_pPr()
    return el


def postavi(p, *imena):
    """Dodaj prazne pPr zastavice (w:keepNext, w:keepLines) ako ih nema."""
    el = pPr(p)
    for ime in imena:
        tag = qn(f"w:{ime}")
        if el.find(tag) is None:
            e = OxmlElement(f"w:{ime}")
            # zastavice idu na početak pPr — shema propisuje redoslijed
            el.insert(0, e)


def prored_auto(p):
    """Bez lineRule='auto' čitač smije obrezati inline sliku na visinu retka."""
    el = pPr(p)
    sp = el.find(qn("w:spacing"))
    if sp is None:
        sp = OxmlElement("w:spacing")
        el.append(sp)
    sp.set(qn("w:line"), "240")
    sp.set(qn("w:lineRule"), "auto")


def ima_sliku(p):
    return bool(p._p.findall(".//" + qn("a:blip"))) or \
           bool(p._p.findall(".//" + qn("w:drawing")))


def prijelom_prije(p):
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    novi = OxmlElement("w:p")
    novi.append(r)
    p._p.addprevious(novi)


def tijelo(d):
    """Odlomci i tablice u REDOSLIJEDU KAKO STOJE U DOKUMENTU.

    `d.paragraphs` i `d.tables` su dva odvojena popisa, pa se iz njih ne vidi što
    slijedi za čim. Struktura bloka (natpis → prikaz → izvor) bez toga se ne može
    provjeriti."""
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    for dijete in d.element.body.iterchildren():
        if dijete.tag == qn("w:p"):
            yield Paragraph(dijete, d)
        elif dijete.tag == qn("w:tbl"):
            yield Table(dijete, d)


def je_prikaz(el):
    from docx.table import Table
    if isinstance(el, Table):
        return True
    return ima_sliku(el)


def obradi(put, prelomi, blokovi_out):
    d = docx.Document(put)
    elementi = list(tijelo(d))
    blokovi = []
    n_natpis = n_izvor = n_slika = n_prelom = 0
    lazni = []

    for i, el in enumerate(elementi):
        from docx.table import Table
        if isinstance(el, Table):
            continue
        tekst = el.text.strip()
        m = NATPIS.match(tekst)
        if m:
            # Natpis je natpis SAMO ako iza njega stoji prikaz. Bez toga se
            # stavke iz „POPIS TABLICA I GRAFIKONA" broje kao natpisi — to je
            # kvar koji je test uhvatio na stvarnom radu.
            sljedeci = elementi[i + 1] if i + 1 < len(elementi) else None
            if sljedeci is None or not je_prikaz(sljedeci):
                lazni.append(tekst[:40])
                continue
            postavi(el, "keepNext", "keepLines")
            k = kljuc_natpisa(m)
            blokovi.append({"kljuc": k, "natpis": tekst})
            n_natpis += 1
            if k in prelomi:
                prijelom_prije(el)
                n_prelom += 1
            continue
        if IZVOR.match(tekst):
            postavi(el, "keepLines")
            n_izvor += 1
            continue
        if ima_sliku(el):
            postavi(el, "keepNext", "keepLines")
            prored_auto(el)
            n_slika += 1

    for t in d.tables:
        for row in t.rows:
            trPr = row._tr.get_or_add_trPr()
            if trPr.find(qn("w:cantSplit")) is None:
                trPr.insert(0, OxmlElement("w:cantSplit"))
            for cell in row.cells:
                for p in cell.paragraphs:
                    # keepNext u ćelijama zadnjeg reda je ono što drži
                    # tablicu uz redak „Izvor:"
                    postavi(p, "keepNext")

    d.save(put)
    if blokovi_out:
        json.dump(blokovi, open(blokovi_out, "w", encoding="utf-8"),
                  ensure_ascii=False, indent=1)
    return {"natpisa": n_natpis, "izvora": n_izvor, "slika": n_slika,
            "tablica": len(d.tables), "prijeloma": n_prelom, "blokovi": blokovi,
            "lazni": lazni}


def main():
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("docx")
    ap.add_argument("--prelomi", help="JSON popis ključeva blokova pred kojima ide prijelom")
    ap.add_argument("--blokovi-out", default="blokovi.json")
    ap.add_argument("--izlaz", help="zapiši u novu datoteku umjesto u mjestu")
    a = ap.parse_args()

    prelomi = set()
    if a.prelomi:
        try:
            prelomi = set(json.load(open(a.prelomi, encoding="utf-8")))
        except FileNotFoundError:
            pass

    put = a.docx
    if a.izlaz:
        shutil.copy(a.docx, a.izlaz)
        put = a.izlaz

    r = obradi(put, prelomi, a.blokovi_out)
    print(f"natpisa: {r['natpisa']} · izvora: {r['izvora']} · slika: {r['slika']}"
          f" · tablica: {r['tablica']} · umetnutih prijeloma: {r['prijeloma']}")
    if r["lazni"]:
        print(f"   (preskočeno {len(r['lazni'])} redaka koji izgledaju kao natpis a "
              f"nemaju prikaz iza sebe — popis prikaza, spomen u tekstu)")
    if r["natpisa"] != r["izvora"]:
        print(f"⚠️  natpisa {r['natpisa']}, a redaka „Izvor:" f"” {r['izvora']} — "
              "svaki prikaz treba izvor ispod sebe")
    print(f"→ {a.blokovi_out}")


if __name__ == "__main__":
    main()
