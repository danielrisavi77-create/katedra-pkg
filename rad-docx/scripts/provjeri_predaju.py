#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Završni gate pred predajom. Izlazni kod 1 = rad se NE predaje.

Provjerava, po redu težine:

  A. brojke iz modela stoje u radu, i nijedna zastarjela nije ostala
  B. sve komponente koje zadatak izrijekom traži postoje
  C. formalna pravila iz profila fakulteta (A4, margine, font, prored)
  D. Wordova polja: TOC, updateFields, zabilješke i REF uravnoteženi
  E. numeracija: prednji dio bez podnožja, tijelo počinje od zadane stranice
  F. prikazi: natpis iznad, izvor ispod, ništa se ne lomi preko stranica
  G. slike umetnute u zadanoj širini
  H. nema ostataka [PROVJERI STR.] i [TREBA IZVOR]

Sve što profil ne propisuje preskače se i navodi kao ograničenje, ne kao greška.

    python3 provjeri_predaju.py rad.docx [--profil p.json] [--model model.json]
        [--model-prije prethodni.json] [--zadatak zadatak.json]
        [--prelomi prelomi.json] [--grafikoni grafikoni.json]
"""

import argparse
import json
import os
import re
import sys
import zipfile

try:
    import docx
except ImportError:
    sys.exit("nedostaje python-docx")

PLACEHOLDERI = ("[PROVJERI STR.]", "[TREBA IZVOR]", "{{model.")


def hr(x, dec=2):
    return f"{float(x):.{dec}f}".replace(".", ",")


def ravno(d, prefiks=""):
    """Ugniježđeni JSON → {putanja: broj} za sve brojeve."""
    out = {}
    if isinstance(d, dict):
        for k, v in d.items():
            out.update(ravno(v, f"{prefiks}/{k}" if prefiks else str(k)))
    elif isinstance(d, list):
        for i, v in enumerate(d):
            out.update(ravno(v, f"{prefiks}[{i}]"))
    elif isinstance(d, (int, float)) and not isinstance(d, bool):
        out[prefiks] = float(d)
    return out


def decimala(x):
    """Koliko decimala model prikazuje za tu vrijednost."""
    s = repr(round(x, 6)).rstrip("0")
    if "." not in s:
        return 0
    return min(3, len(s.split(".")[1]))


def zapisi(x):
    """Svi vjerojatni zapisi brojke u tekstu.

    Model nosi 0,03, a tekst tipično piše 0,030 — i obrnuto. Zato se provjeravaju
    prirodni zapis te dvo- i trodecimalni."""
    return {hr(x, decimala(x)), hr(x, 2), hr(x, 3)}


def zapis_tocan(x):
    """Najodređeniji zapis: onoliko decimala koliko ih model doista nosi, ali
    najmanje dvije. Samo poklapanje NA NJEMU je dokaz zastarjelosti."""
    return hr(x, max(2, decimala(x)))


def stoji_u(tekst, zapis):
    """Brojka stoji u tekstu samo ako iza nje ne slijedi još jedna cifra.

    Bez toga „2,2" pogađa „2,25" i crna lista lažno prijavljuje zastarjelost."""
    return re.search(r"(?<![\d,.])" + re.escape(zapis) + r"(?![\d])", tekst) is not None


class Provjera:
    def __init__(self):
        self.greske, self.upozorenja, self.ogranicenja = [], [], []
        # Strukturirani nalaz o zadatku. Ispis je za čovjeka; ovo je za rubrika.py,
        # koja je do sada morala pretpostavljati da je provjera prošla jer nalaz
        # nigdje nije postojao u obliku koji se dade pročitati.
        self.zadatak = None

    def g(self, x): self.greske.append(x)
    def u(self, x): self.upozorenja.append(x)
    def o(self, x): self.ogranicenja.append(x)


def tekst_dokumenta(d):
    dijelovi = [p.text for p in d.paragraphs]
    for t in d.tables:
        for r in t.rows:
            for c in r.cells:
                dijelovi.append(c.text)
    return "\n".join(dijelovi)


# ── A. brojke ────────────────────────────────────────────────────────────────
def provjeri_model(P, sve, model, model_prije):
    if not model:
        P.o("nema model.json — brojke nisu provjerene protiv izvora istine")
        return
    vrijednosti = ravno(model)
    nema = []
    for putanja, v in vrijednosti.items():
        if abs(v) < 0.0005 or float(v).is_integer() and abs(v) > 9999:
            continue                      # nule i velike cijele brojke su previše česte
        if not any(stoji_u(sve, z) for z in zapisi(v)):
            nema.append(f"{putanja} = {hr(v, decimala(v))}")
    if nema:
        P.u(f"brojki iz modela koje ne nalazim u radu: {len(nema)}"
            + (f" (npr. {nema[0]})" if nema else ""))

    if model_prije:
        stare = ravno(model_prije)
        crna = {}
        for putanja, staro in stare.items():
            novo = vrijednosti.get(putanja)
            if novo is None or abs(novo - staro) < 1e-9:
                continue
            aktualni = set()
            for v in vrijednosti.values():
                aktualni |= zapisi(v)
            for s in zapisi(staro) - aktualni:
                crna.setdefault(s, (putanja, novo, s == zapis_tocan(staro)))
        for s, (putanja, novo, tocan) in crna.items():
            if not stoji_u(sve, s):
                continue
            poruka = (f"zastarjela vrijednost „{s}" f"” ostala u tekstu — "
                      f"{putanja} je sada {hr(novo, max(2, decimala(novo)))}")
            if tocan:
                P.g(poruka)
            else:
                # Kraći zapis („2,2" za staru vrijednost 2,20) može biti i nepovezana
                # brojka iz teksta. Prijavljuje se, ali ne blokira predaju.
                P.u(poruka + " — ILI nepovezana brojka istog zapisa; provjeri rukom")
    else:
        P.o("nema prethodnog model.json — zastarjele vrijednosti nisu provjerene")


# ── B. komponente zadatka ────────────────────────────────────────────────────
def provjeri_zadatak(P, sve, zadatak):
    if not zadatak:
        P.o("nema zadatak.json — nije provjereno je li rad odgovorio na uputu predmeta")
        P.zadatak = {"provjeren": False, "razlog": "nema zadatak.json"}
        return
    nalazi = []
    for k in zadatak.get("komponente", []):
        igle = k.get("igle")
        if igle and any(i in sve for i in igle):
            nalazi.append({"naziv": k["naziv"], "status": "u_radu",
                           "dokaz": "igla nađena u tekstu rada"})
            continue
        prov = k.get("provjereno")
        if prov:
            # Zahtjev koji se ne da izraziti niskom u dokumentu (npr. „stranica i kod
            # parafraze") provjerava se alatom, a nalaz stoji uz komponentu. Bez ove
            # grane takav je zahtjev uvijek padao kao „u radu nema", jer se tražio
            # doslovan tekst zahtjeva unutar teksta rada.
            P.o(f"zadatak, provjereno alatom ({prov.get('alat', '?')}): "
                f"{k['naziv']} — {prov.get('nalaz', '')}")
            nalazi.append({"naziv": k["naziv"], "status": "provjereno_alatom",
                           "dokaz": f"{prov.get('alat', '?')}: {prov.get('nalaz', '')}"})
            continue
        if not igle:
            P.o(f"zadatak, nije strojno provjerljivo (nema igala ni nalaza provjere): "
                f"{k['naziv']}")
            nalazi.append({"naziv": k["naziv"], "status": "nije_strojno_provjerljivo",
                           "dokaz": "nema ni igala ni zapisanog nalaza provjere"})
            continue
        P.g(f"zadatak traži, a u radu nema: {k['naziv']}")
        nalazi.append({"naziv": k["naziv"], "status": "nema_u_radu",
                       "dokaz": "nijedna igla nije nađena: " + ", ".join(igle)})
    P.zadatak = {"provjeren": True, "komponente": nalazi}


# ── C. formalna pravila iz profila ───────────────────────────────────────────
POCETAK_TIJELA = re.compile(r"^\s*(1\.?\s+\S|UVOD\b|Uvod\b)", re.I)
KRAJ_TIJELA = re.compile(r"^\s*(?:\d+\.?\s*)?(LITERATURA|Literatura|POPIS LITERATURE|"
                         r"Popis literature|BIBLIOGRAFIJA|Bibliografija|REFERENCES)\s*$")


def _tijelo_odlomci(d):
    """Odlomci TIJELA rada: od Uvoda do popisa literature.

    Ograda je obavezna. Bez nje u „tijelo" ulaze bibliografske jedinice, koje su
    tipično jednostruko proredene i 11 pt, pa mjerenje prijavi da tekst odstupa od
    kućnog stila. Nađeno kao regresija na radu koji je prije toga bio čist:
    prored 1,5 „u samo 69 % odlomaka" i 12 pt „u 78 % runova" — oboje je bila
    literatura, ne tijelo."""
    out, poceo = [], False
    for p in d.paragraphs:
        t = (p.text or "").strip()
        if not t:
            continue
        stil = (p.style.name or "") if p.style is not None else ""
        naslov = stil.startswith("Heading")
        if not poceo:
            if naslov and POCETAK_TIJELA.match(t):
                poceo = True
            continue
        if KRAJ_TIJELA.match(t):
            break
        if len(t) < 60 or naslov or stil.startswith("TOC"):
            continue
        if re.match(r"^(Tablica|Grafikon|Slika|Prikaz)\s+\d", t) or t.lower().startswith("izvor"):
            continue
        out.append(p)
    if not out:                     # dokument bez stilova naslova — mjeri sve
        for p in d.paragraphs:
            t = (p.text or "").strip()
            if len(t) >= 60 and not re.match(r"^(Tablica|Grafikon|Slika|Prikaz)\s+\d", t) \
                    and not t.lower().startswith("izvor"):
                out.append(p)
    return out


def _prevladava(vrijednosti):
    """Najčešća vrijednost i njezin udio; None ako nema uzorka.

    PAŽNJA: `None` se ovdje smije predati samo ako je već razrješen u naslijeđenu
    vrijednost. Odbacivanje `None`-ova IZVRĆE rezultat: u Wordovu dokumentu
    velika većina runova nema zadanu veličinu (nasljeđuje ju iz docDefaults), a
    eksplicitno je zadaje samo nekoliko — pa najčešća „zadana" vrijednost bude
    upravo iznimka. Nađeno na stvarnom diplomskom radu: 32 runa sa 11 pt u
    sažetku prijavljena su kao prevladavajuća veličina, protiv 1473 runa koji
    ispravno nasljeđuju 12 pt."""
    vrijednosti = [v for v in vrijednosti if v is not None]
    if not vrijednosti:
        return None, 0.0
    from collections import Counter
    (v, n), = Counter(vrijednosti).most_common(1)
    return v, n / len(vrijednosti)


def _zadana_velicina(dd):
    m = re.search(r'<w:sz w:val="(\d+)"/>', dd)
    return int(m.group(1)) / 2 if m else None


def _zadano_pismo(dd):
    m = re.search(r'<w:rFonts[^>]*w:ascii="([^"]+)"', dd)
    return m.group(1) if m else None


def _efektivne(odlomci, atribut, zadano):
    """Za svaki run vrati STVARNO primijenjenu vrijednost: run → stil → docDefaults."""
    out = []
    for p in odlomci:
        stil_v = None
        try:
            f = p.style.font
            stil_v = f.size.pt if atribut == "size" and f.size is not None else (
                f.name if atribut == "name" else None)
        except Exception:
            pass
        for r in p.runs:
            if atribut == "size":
                v = r.font.size.pt if r.font.size is not None else None
            else:
                v = r.font.name
            out.append(v if v is not None else (stil_v if stil_v is not None else zadano))
    return out


def _iz_docdefaults(docx_put):
    with zipfile.ZipFile(docx_put) as z:
        styles = z.read("word/styles.xml").decode("utf-8")
    m = re.search(r"<w:docDefaults>.*?</w:docDefaults>", styles, re.S)
    return (m.group(0) if m else ""), styles


def provjeri_format(P, docx_put, d, profil):
    sec = d.sections[0]
    if round(sec.page_width.cm, 1) != 21.0 or round(sec.page_height.cm, 1) != 29.7:
        P.g(f"stranica nije A4 ({sec.page_width.cm:.1f} × {sec.page_height.cm:.1f} cm)")

    fmt = (profil or {}).get("format", {})
    marg = fmt.get("margine_cm")
    if marg:
        for naziv, atr in (("gore", "top_margin"), ("dolje", "bottom_margin"),
                           ("lijevo", "left_margin"), ("desno", "right_margin")):
            trazeno = marg.get(naziv)
            if trazeno is None:
                continue
            stvarno = getattr(sec, atr).cm
            if abs(stvarno - trazeno) > 0.05:
                P.g(f"margina {naziv}: {stvarno:.2f} cm, profil traži {hr(trazeno)} cm")
    else:
        P.o("profil ne propisuje margine — nisu provjerene")

    dd, styles = _iz_docdefaults(docx_put)
    odlomci = _tijelo_odlomci(d)
    if not odlomci:
        P.o("ne nalazim odlomke tijela — oblikovanje teksta nije provjereno")
        return

    # ── prored: mjeri se STVARNO primijenjena vrijednost ──
    # Word 1,5 obično piše na SVAKI odlomak, a docDefaults ostavlja prazan.
    # Provjera samo po docDefaults zato lažno prijavljuje da proreda nema.
    trazeni = fmt.get("prored")
    if trazeni:
        prored, udio = _prevladava([p.paragraph_format.line_spacing for p in odlomci])
        pravila, _ = _prevladava([str(p.paragraph_format.line_spacing_rule)
                                  for p in odlomci])
        u_stilu = ('w:line="%d"' % round(trazeni * 240)) in styles
        if prored is None:
            P.u("prored nije zadan ni na odlomcima ni u stilu — nasljeđuje se iz predloška")
        elif abs(float(prored) - trazeni) > 0.01:
            P.g(f"prored je {hr(float(prored), 2)}, profil traži {hr(trazeni)} "
                f"(u {udio * 100:.0f} % odlomaka tijela)")
        elif udio < 0.9:
            P.u(f"prored {hr(trazeni)} u samo {udio * 100:.0f} % odlomaka tijela — "
                "ostatak odstupa")
        if prored is not None and not u_stilu:
            P.u("prored je zadan po odlomcima, a ne u stilu ni u docDefaults — "
                "svaki novi odlomak koji se dopiše past će na drugu vrijednost")
        # `_prevladava` vraća None kad prored nigdje nije zadan, a `str(None)` je
        # niska „None": istinita i bez AUTO/POINT/MULTIPLE, pa je grana ispod
        # okidala GREŠKU „prored je fiksan (None)" na svakom dokumentu koji
        # prored nasljeđuje iz predloška. Poruka je uz to bila obrnuta od
        # istine (nije zadan ≠ fiksan je), a nalaz je blokirao predaju —
        # o istom dokumentu alat je u istom ispisu javljao i „prored nije
        # zadan" (upozorenje) i „prored je fiksan" (greška).
        if pravila and pravila != "None" \
                and "AUTO" not in pravila.upper() and "POINT" not in pravila.upper() \
                and "MULTIPLE" not in pravila.upper():
            P.g(f"prored je fiksan ({pravila}) — inline slike se obrežu na visinu retka")

    # ── font i veličina: prvo runovi tijela, pa stil, pa docDefaults ──
    fontovi = fmt.get("font") or []
    if fontovi:
        font, udio_f = _prevladava(_efektivne(odlomci, "name", _zadano_pismo(dd)))
        if font is None:
            P.g(f"ne mogu utvrditi pismo teksta; profil traži {', '.join(fontovi)}")
        elif font not in fontovi:
            P.g(f"pismo teksta je {font}, profil traži {', '.join(fontovi)}")
        elif udio_f < 0.9:
            P.u(f"{font} u samo {udio_f * 100:.0f} % runova tijela — ostatak odstupa")

    vel = fmt.get("velicina_pt")
    if vel:
        v, udio_v = _prevladava(_efektivne(odlomci, "size", _zadana_velicina(dd)))
        if v is None:
            P.u(f"veličinu pisma ne nalazim ni na runovima, ni u stilu, ni u "
                f"docDefaults (profil traži {hr(vel, 0)} pt)")
        elif abs(v - vel) > 0.01:
            P.g(f"veličina pisma je {hr(v, 0)} pt, profil traži {hr(vel, 0)} pt "
                f"(u {udio_v * 100:.0f} % runova tijela)")
        elif udio_v < 0.9:
            P.u(f"{hr(vel, 0)} pt u {udio_v * 100:.0f} % runova tijela — "
                "ostatak odstupa (sažetak i sličan prednji dio smiju odstupati)")


# ── D. polja ─────────────────────────────────────────────────────────────────
def provjeri_polja(P, docx_put, dxml):
    with zipfile.ZipFile(docx_put) as z:
        try:
            sett = z.read("word/settings.xml").decode("utf-8")
        except KeyError:
            sett = ""
    if "updateFields" not in sett:
        P.g("settings.xml ne traži osvježavanje polja — sadržaj bi ostao prazan")
    if not re.search(r"instrText[^>]*>\s*TOC|fldSimple[^>]*w:instr=\"[^\"]*TOC", dxml):
        P.g("nema Wordovog polja TOC — sadržaj nije generiran iz Worda")

    b = len(re.findall(r'w:fldCharType="begin"', dxml))
    e = len(re.findall(r'w:fldCharType="end"', dxml))
    if b != e:
        P.g(f"neuravnotežena polja: {b} begin, {e} end — Word bi prikazao sirovu instrukciju")

    zab = set(re.findall(r'w:bookmarkStart[^>]*w:name="((?:tablica|grafikon|slika)\d+)"', dxml))
    ref = set(re.findall(r'REF\s+((?:tablica|grafikon|slika)\d+)', dxml))
    visak = ref - zab
    if visak:
        P.g(f"unakrsna referenca bez zabilješke: {', '.join(sorted(visak))} "
            "— Word bi ispisao „Error! Reference source not found.\"")
    return len(zab), len(ref)


# ── E. numeracija ────────────────────────────────────────────────────────────
def provjeri_numeraciju(P, dxml, profil):
    sekcije = re.findall(r"<w:sectPr.*?</w:sectPr>", dxml, re.S)
    trazeno = ((profil or {}).get("format", {}).get("numeracija") or {})
    pocetak = trazeno.get("tijelo_pocinje_od")
    if len(sekcije) < 2:
        if pocetak:
            P.g("dokument ima jednu sekciju, a profil traži numeraciju od tijela rada")
        else:
            P.o("dokument ima jednu sekciju — numeracija po sekcijama nije provjerena")
        return
    prednji, tijelo = sekcije[0], sekcije[-1]
    if trazeno.get("prednji_dio") == "bez" and "footerReference" in prednji:
        P.g("prednji dio ima podnožje, a profil traži da nije numeriran")
    if pocetak and f'w:pgNumType w:start="{pocetak}"' not in tijelo:
        P.g(f"numeracija tijela ne počinje od {pocetak}")
    if "footerReference" not in tijelo:
        P.g("tijelo rada nema podnožje s brojem stranice")


# ── F. prikazi ───────────────────────────────────────────────────────────────
def provjeri_prikaze(P, d, prelomi_put):
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.oxml.ns import qn
    NATPIS = re.compile(r"^(Tablica|Grafikon|Slika)\s+(\d+)\s*[.:]")

    elementi = []
    for dijete in d.element.body.iterchildren():
        if dijete.tag == qn("w:p"):
            elementi.append(Paragraph(dijete, d))
        elif dijete.tag == qn("w:tbl"):
            elementi.append(Table(dijete, d))

    natpisa = 0
    for i, el in enumerate(elementi):
        if isinstance(el, Table):
            continue
        m = NATPIS.match(el.text.strip())
        if not m:
            continue
        sljedeci = elementi[i + 1] if i + 1 < len(elementi) else None
        je_prikaz = isinstance(sljedeci, Table) or (
            sljedeci is not None and not isinstance(sljedeci, Table)
            and bool(sljedeci._p.findall(".//" + qn("w:drawing"))))
        if not je_prikaz:
            continue
        natpisa += 1
        # izvor mora slijediti unutar dva elementa iza prikaza
        rep = elementi[i + 2:i + 4]
        if not any(not isinstance(x, Table) and x.text.strip().lower().startswith("izvor")
                   for x in rep):
            P.g(f"prikaz bez retka „Izvor:" f"” ispod sebe: {el.text.strip()[:50]}")

    if os.path.exists(prelomi_put or ""):
        try:
            prelomi = json.load(open(prelomi_put, encoding="utf-8"))
        except Exception:
            prelomi = []
        if prelomi:
            kljucevi = [p if isinstance(p, str) else p.get("kljuc") for p in prelomi]
            P.u(f"umetnuti prijelomi pred prikazima: {', '.join(map(str, kljucevi))} "
                "(namjerno — prikaz se inače lomi)")
    else:
        P.o("nema prelomi.json — nije izmjereno lome li se prikazi preko stranica")
    return natpisa


# ── G. slike ─────────────────────────────────────────────────────────────────
def provjeri_slike(P, dxml, grafikoni):
    if not grafikoni:
        P.o("nema grafikoni.json — širina umetnutih slika nije provjerena")
        return
    for f, dim in grafikoni.items():
        px = round(dim["w_cm"] * 37.7952755)      # 96 dpi; 28,35 bi dalo sliku 25 % manju
        if f'cx="{px * 9525}"' not in dxml:       # 1 px = 9525 EMU
            P.u(f"{f}: ne nalazim umetanje u širini {dim['w_cm']} cm")


# ── H. placeholderi ──────────────────────────────────────────────────────────
def provjeri_placeholdere(P, sve):
    for ph in PLACEHOLDERI:
        n = sve.count(ph)
        if n:
            P.g(f"u tekstu je ostalo {n}× „{ph}" f"”")


def main():
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("docx")
    ap.add_argument("--profil")
    ap.add_argument("--model")
    ap.add_argument("--model-prije", help="prethodni model.json — za crnu listu")
    ap.add_argument("--zadatak")
    ap.add_argument("--prelomi", default="prelomi.json")
    ap.add_argument("--grafikoni")
    ap.add_argument("--json", dest="json_out",
                    help="zapiši nalaz kao JSON (npr. .katedra/predaja.json) — "
                         "rubrika.py ga čita kao dokaz za kriterij „zadatak\"")
    a = ap.parse_args()

    ucitaj = lambda p: json.load(open(p, encoding="utf-8")) if p and os.path.exists(p) else None
    profil, model = ucitaj(a.profil), ucitaj(a.model)
    model_prije, zadatak = ucitaj(a.model_prije), ucitaj(a.zadatak)
    grafikoni = ucitaj(a.grafikoni)

    d = docx.Document(a.docx)
    sve = tekst_dokumenta(d)
    with zipfile.ZipFile(a.docx) as z:
        dxml = z.read("word/document.xml").decode("utf-8")

    P = Provjera()
    provjeri_model(P, sve, model, model_prije)
    provjeri_zadatak(P, sve, zadatak)
    provjeri_format(P, a.docx, d, profil)
    n_zab, n_ref = provjeri_polja(P, a.docx, dxml)
    provjeri_numeraciju(P, dxml, profil)
    natpisa = provjeri_prikaze(P, d, a.prelomi)
    provjeri_slike(P, dxml, grafikoni)
    provjeri_placeholdere(P, sve)

    print("=" * 74)
    print("PROVJERA PRED PREDAJU —", os.path.basename(a.docx))
    print("=" * 74)
    print(f"odlomaka: {len(d.paragraphs)} · tablica: {len(d.tables)} · "
          f"slika: {len(d.inline_shapes)} · prikaza s natpisom: {natpisa}")
    print(f"zabilješki: {n_zab} · unakrsnih referenci: {n_ref}")
    print(f"profil: {a.profil or '—'} · model: {a.model or '—'} · zadatak: {a.zadatak or '—'}")
    print()
    if P.greske:
        print(f"❌ GREŠKE ({len(P.greske)}) — rad se ne predaje:")
        for x in P.greske:
            print("   ·", x)
    else:
        print("✅ nijedna greška")
    if P.upozorenja:
        print(f"\n⚠️  UPOZORENJA ({len(P.upozorenja)}):")
        for x in P.upozorenja:
            print("   ·", x)
    if P.ogranicenja:
        print(f"\nℹ️  OGRANIČENJA — nije provjereno ({len(P.ogranicenja)}):")
        for x in P.ogranicenja:
            print("   ·", x)
    if a.json_out:
        with open(a.json_out, "w", encoding="utf-8") as f:
            json.dump({"rad": os.path.basename(a.docx),
                       "prosao": not P.greske,
                       "greske": P.greske,
                       "upozorenja": P.upozorenja,
                       "ogranicenja": P.ogranicenja,
                       "zadatak": P.zadatak},
                      f, ensure_ascii=False, indent=2)
        print(f"nalaz zapisan: {a.json_out}")

    print()
    sys.exit(1 if P.greske else 0)


if __name__ == "__main__":
    main()
