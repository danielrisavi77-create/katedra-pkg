#!/usr/bin/env python3
"""
check_rules.py — usklađenost STVARNOG .docx-a s pravilima fakulteta.

Ne provjerava kako bi rad trebao izgledati po opisu, nego što doista piše u
datoteci: theme i docDefaults, stil Normal, formatiranje pojedinih runova,
margine sekcija, prijelomi pred poglavljima, natpisi prikaza i oblik citata.

Font se u Wordu razrješava u četiri sloja (theme → docDefaults → stil → run) i
svaki od njih može reći nešto drugo. Zato se ispisuju sva četiri: rad koji u
stilu ima Times New Roman, a u polovici runova Calibri, izgleda uredno dok ga
mentor ne otvori na drugom računalu.

    python3 <KATEDRA_SKILL>/scripts/check_rules.py ./rad.docx --fakultet efzg --tip zavrsni
    python3 <KATEDRA_SKILL>/scripts/check_rules.py ./rad.docx --profil <KATEDRA_SKILL>/references/fakulteti/efzg.json --tip zavrsni
    python3 <KATEDRA_SKILL>/scripts/check_rules.py ./rad.docx --fakultet fpzg --tip diplomski --json ./.katedra/pravila.json

Izlazni kodovi:
  0  sve u redu
  1  ima nalaza
  2  greška (nema profila, nema datoteke, datoteka nije .docx)
"""
import argparse
import json
import os
import re
import sys
import zipfile
from collections import Counter
from xml.etree import ElementTree as ET

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import hr_text as H
import citation_dialects as C
from profile_rules import ProfileRuleError, resolve_profile, resolve_work_type

HERE = os.path.dirname(os.path.abspath(__file__))
FAKULTETI = os.path.join(HERE, "..", "references", "fakulteti")

TOLERANCIJA_CM = 0.05
# Procjena stranica iz broja riječi: A4, TNR 12, prored 1,5, margine ~2,5 cm.
# Raspon je namjerno širok — bez rendera se točan broj ne može znati.
RIJECI_PO_STRANICI = (270, 330)

NS_A = "http://schemas.openxmlformats.org/drawingml/2006/main"

OK, UPOZ, LOSE = "✅", "⚠️", "❌"


# ------------------------------------------------------------------- profil

def nadi_profil(fakultet, putanja, tip=None):
    """Vrati (profil, putanja) ili izađi s kodom 2.

    Izravni --profil ostaje backward-compatible flat profile. --fakultet ide
    kroz B07 generated registry + compositional resolver, pa alias može nositi
    i programme context (npr. RFIR → efzg + programme=rfir).
    """
    if putanja:
        if not os.path.isfile(putanja):
            greska(f"profil ne postoji: {putanja}")
        return ucitaj_json(putanja), putanja
    if not fakultet:
        greska("navedi --fakultet <slug/alias> ili --profil <put/do/profila.json>")
    try:
        resolved = resolve_profile(
            fakultet,
            faculty_dir=FAKULTETI,
            work_type=tip,
        )
    except ProfileRuleError as exc:
        greska(str(exc))
    slug = resolved.context["faculty"]
    # v1.9 (nalaz 4): provenance po pravilu ostaje dostupan Izvjestaju da ❌
    # zadrže samo pravila koja stvarno stoje u službenim uputama (type=explicit).
    global _PROVENANCE
    _PROVENANCE = dict(getattr(resolved, "provenance", None) or {})
    return resolved.profile, os.path.join(FAKULTETI, f"{slug}.json")


# v1.9 (nalaz 4): sidecar provenance za trenutačni profil — puni ga nadi_profil
# (resolver) ili, za --profil, datoteka `<ime>.provenance.json` uz profil
# (tako .katedra/resolved_profile.json nosi resolved_profile.provenance.json).
_PROVENANCE = {}


def _ucitaj_provenance_sidecar(put_profila):
    if not put_profila or not str(put_profila).endswith(".json"):
        return {}
    sidecar = str(put_profila)[:-5] + ".provenance.json"
    if not os.path.isfile(sidecar):
        return {}
    try:
        with open(sidecar, encoding="utf-8") as f:
            data = json.load(f)
        return data if isinstance(data, dict) else {}
    except (OSError, json.JSONDecodeError):
        return {}


def _normaliziraj_provenance(prov):
    """Svedi oba oblika provenance zapisa na {pointer: meta}.

    Registry resolver daje ravni rječnik {pointer: meta}; profil po
    `_schema.json` (i samostojni put resolvera, npr. hks-fzs) daje
    {"default": meta, "rules": {pointer: meta}} — vrijednosti izvan `rules`
    su nizovi/metapodaci, ne pravila. Nedostajući `type` u pravilu nasljeđuje
    se iz `default`.
    """
    if not isinstance(prov, dict):
        return {}
    if isinstance(prov.get("rules"), dict):
        default = prov.get("default") if isinstance(prov.get("default"), dict) else {}
        out = {}
        for pointer, meta in prov["rules"].items():
            if not isinstance(meta, dict):
                continue
            out[pointer] = {**default, **meta}
        return out
    return {k: v for k, v in prov.items() if isinstance(v, dict)}


def ucitaj_json(put):
    try:
        with open(put, encoding="utf-8") as f:
            return json.load(f)
    except (OSError, json.JSONDecodeError) as e:
        greska(f"{put}: {e}")


def greska(poruka):
    print(f"❌ {poruka}", file=sys.stderr)
    sys.exit(2)


def norm(s):
    """Bez dijakritika, mala slova, bez numeracije i interpunkcije."""
    s = H.bez_dijakritika(str(s or "")).lower()
    s = re.sub(r"^\s*(?:[ivxlc]+|\d+)(?:\.\d+)*\.?\s+", " ", s)
    s = re.sub(r"[^a-z0-9]+", " ", s)
    return re.sub(r"\s+", " ", s).strip()


# Naslovi koji pripadaju prednjem/završnom aparatu rada, a nisu sadržajna
# poglavlja. Klasifikacija se koristi usko: za broj poglavlja i za
# prepoznavanje popisa prikaza. Ne mijenja redoslijed/obavezne dijelove.
# v1.1-fix (Q1a/Q1b): doslovni popisi nizova zamijenjeni su normaliziranim
# predikatom — hrvatski naslov stoji u padežu i ima desetke varijanti, pa je
# svaki „allowlist" ponovno otvarao iste lažne nalaze.
_PRIKAZ_RIJEC = r"(?:tablic\w*|slik\w*|grafikon\w*|graf\w*|ilustracij\w*|prikaz\w*|shem\w*)"
POPIS_PRIKAZA_RE = re.compile(
    rf"^(?:popis|kazalo|pregled)\s+{_PRIKAZ_RIJEC}(?:\s+(?:i|te)?\s*{_PRIKAZ_RIJEC})*$")
# Jednorječne (ili ustaljene) cjeline završnog aparata — traži se CIJELI naslov,
# da „Izvori financiranja poduzeća" ostane sadržajno poglavlje.
BACK_MATTER_RE = re.compile(
    r"^(?:literatur\w*|bibliograf\w*|referenc\w*|izvori|"
    r"zivotopis\w*|curriculum vitae|prilo(?:g|zi)|dodat(?:ak|ci)|"
    r"privit(?:ak|ci)|sazet\w*|summary|abstract)$")
# v1.1-fix (Q1 dorada): „počinje na popis/kazalo" je bilo preširoko — sadržajno
# poglavlje „2. POPIS I ANALIZA MJERA" proglašavalo se završnim aparatom, pa je
# broj poglavlja padao na 1, a opseg se rezao na tom poglavlju. Završni aparat je
# KRATAK, SAMOSTALAN naslov popisa stvari („POPIS LITERATURE", „KAZALO POJMOVA"),
# pa se traži CIJELI naslov, a ne samo njegov početak.
_POPISNA_STVAR = (r"(?:literatur\w*|izvor\w*|bibliograf\w*|referenc\w*|"
                  r"tablic\w*|slik\w*|grafikon\w*|graf\w*|prikaz\w*|shem\w*|"
                  r"ilustracij\w*|kratic\w*|pojm\w*|simbol\w*|oznak\w*|"
                  r"prilo(?:g\w*|z\w*)|dodat\w*|privit\w*)")
_POPISNI_PRIDJEV = r"(?:koriste\w*|upotrijebljen\w*|uporabljen\w*)"
_POPISNA_STAVKA = rf"(?:{_POPISNI_PRIDJEV}\s+)?{_POPISNA_STVAR}"
POPIS_STVARI_RE = re.compile(
    rf"^(?:popis|kazalo)(?:\s+{_POPISNA_STAVKA}"
    rf"(?:\s+(?:i|te)?\s*{_POPISNA_STAVKA})*)?$")



# --------------------------------------------- strukturno: numeriran naslov
# Treći pokušaj klasifikacije naslova, i prvi koji ne ovisi o rječniku.
#
# Povijest je poučna: doslovni allowlist (preúzak — «SAŽETAK», «PRILOZI» ispali
# iz aparata), pa prefiks „popis|kazalo" (preširok — sadržajno poglavlje
# „2. POPIS I ANALIZA MJERA" postalo aparat), pa zatvoreni skup imenica (opet
# preúzak — „POPIS SKRAĆENICA", „KAZALO IMENA", „POPIS SLIKA, TABLICA I
# GRAFIKONA"). Svaki krug je pomicao istu granicu i svaki je promašio s druge
# strane, jer je rječnik hrvatskih naslova praktički neomeđen.
#
# Struktura je pouzdanija od rječnika: u hrvatskom radu SADRŽAJNA poglavlja su
# NUMERIRANA („1. UVOD" … „7. ZAKLJUČAK"), a prednji i završni aparat nije
# („SAŽETAK", „POPIS LITERATURE", „ŽIVOTOPIS"). Provjereno na stvarnom EFZG
# završnom radu: svih 7 sadržajnih poglavlja numerirano, svih 6 dijelova
# aparata nenumerirano, uključujući „POPIS SLIKA, TABLICA I GRAFIKONA" koji bi
# svaki zatvoreni skup imenica opet promašio.
#
# Rječnički predikat OSTAJE kao pričuva za rad koji uopće ne numerira poglavlja.
BROJ_NASLOVA_RE = re.compile(r"^\s*\d+(?:\.\d+)*\.?\s+\S")


def je_numeriran_naslov(tekst):
    return bool(BROJ_NASLOVA_RE.match(str(tekst or "")))


def rad_numerira_poglavlja(naslovi1, prag=2):
    """Numerira li ovaj rad svoja poglavlja? Prag 2 da jedan „1. UVOD" ne odluči."""
    return sum(1 for t in naslovi1 if je_numeriran_naslov(t)) >= prag


def je_back_matter_naslov(tekst):
    """Pripada li naslov prednjem/završnom aparatu (nije sadržajno poglavlje)."""
    n = norm(tekst)
    if not n:
        return False
    if POPIS_STVARI_RE.match(n):
        return True
    return bool(BACK_MATTER_RE.match(n))


def je_popis_prikaza_naslov(tekst):
    return bool(POPIS_PRIKAZA_RE.match(norm(tekst)))


def je_uvodni_naslov(tekst):
    """„UVOD", „1. Uvod", „I. UVOD", „UVOD U PROBLEMATIKU", „UVODNE NAPOMENE"."""
    n = norm(tekst)
    return n == "uvod" or n.startswith("uvod ") or n.startswith("uvodn")


def _je_toc_redak(tekst):
    """Redak sadržaja: točkasti vodič ili tabulator pa broj stranice."""
    t = str(tekst or "")
    return bool(re.search(r"\.{3,}", t) or re.search(r"\t[\s.]*\d+\s*$", t))


def _je_toc_stil(stil):
    return bool(re.match(r"(?i)^(?:toc|sadrzaj|sadržaj|contents)", (stil or "").strip()))


def _je_toc_stil_unosa(stil):
    """Stil Wordova UNOSA u sadržaj („TOC 1"…„TOC 9", „Sadržaj 2").

    Namjerno UŽE od `_je_toc_stil`: stil „TOC Heading" nosi naslov „Sadržaj",
    a taj naslov JEST naslov rada i mora ostati kandidat za obavezni dio
    „sadržaj" — točno tako je oblikovan stvarni EFZG završni rad. Razlikuje ih
    broj razine iza imena stila; to je Wordova konvencija, ne popis riječi.
    """
    return bool(re.match(r"(?i)^(?:toc|sadrzaj|sadržaj|contents)\s*\d",
                         (stil or "").strip()))


def _je_popisni_naslov(tekst):
    """SADRŽAJ / KAZALO / POPIS nečega — iza njih slijedi popis, ne tijelo rada."""
    n = norm(tekst)
    return bool(re.match(r"^(?:sadrzaj|kazalo|contents|popis)\b", n))


def je_pocetak_glavnog_teksta(tekst):
    """Praktični marker početka tijela — isti kriterij koji koristi hr_text._iz_docx.

    v1.1-fix (zajednička specifikacija): redak sadržaja („1. UVOD.........3")
    NIJE početak tijela, a rimski broj („I. UVOD") i prošireni naslov
    („UVOD U PROBLEMATIKU", „UVODNE NAPOMENE") jesu.
    """
    t = str(tekst or "")
    if not t.strip() or _je_toc_redak(t):
        return False
    if je_uvodni_naslov(t):
        return True
    # arapski ili rimski broj poglavlja pa riječ; broj je najviše dvoznamenkast
    # da godina na početku rečenice („1998. godine…") ne prođe kao naslov
    return bool(re.match(r"^\s*(?:\d{1,2}(?:\.\d{1,2})*\.?|[IVXLC]{1,4}\.)\s+\S", t))


def _outline_razina(p):
    """Razina iz w:outlineLvl (1..9) ili 0 ako odlomak nije u strukturi naslova."""
    try:
        ppr = p._p.find(qn("w:pPr"))
    except AttributeError:
        return 0
    if ppr is None:
        return 0
    el = ppr.find(qn("w:outlineLvl"))
    if el is None:
        return 0
    try:
        razina = int(el.get(qn("w:val")))
    except (TypeError, ValueError):
        return 0
    return razina + 1 if 0 <= razina <= 8 else 0


def indeks_pocetka_tijela(redoslijed):
    """Indeks bloka na kojem počinje glavni tekst, ili None ako se ne prepozna.

    Prednost ima odlomak u stilu Heading (ili s postavljenim outlineLvl).
    Nijedan kandidat unutar sadržaja/popisa se ne prihvaća.
    """
    u_popisu = False
    pricuva = None
    for i, (vrsta, blok) in enumerate(redoslijed):
        if vrsta != "p":
            continue
        t = tekst_bloka(blok).strip()
        if not t:
            continue
        r = razina_naslova(blok) or _outline_razina(blok)
        if r:
            u_popisu = _je_popisni_naslov(t)
        if u_popisu or _je_toc_stil(stil_imena(blok)) or _je_toc_redak(t):
            continue
        if not je_pocetak_glavnog_teksta(t):
            continue
        if r:
            return i
        if pricuva is None:
            pricuva = i
    return pricuva


def tekst_glavnog_dijela(redoslijed):
    """(tekstovi glavnog teksta, opis opsega) — od početka tijela do popisa izvora.

    v1.1-fix (Q2): broj riječi se mjeri na tijelu rada. Naslovnica, sažetak,
    sadržaj i popis izvora nisu tekst koji fakultet broji u opseg.
    """
    svi = [tekst_bloka(b).strip() for v, b in redoslijed
           if v == "p" and tekst_bloka(b).strip()]
    poc = indeks_pocetka_tijela(redoslijed)
    if poc is None:
        return svi, ("brojan je CIJELI dokument — početak tijela („UVOD\" / „1. …\") "
                     "nije prepoznat, pa su uključeni i naslovnica, sadržaj i popis izvora")
    tijelo, prvi, zadnji = [], None, None
    for i, (vrsta, blok) in enumerate(redoslijed):
        if vrsta != "p" or i < poc:
            continue
        t = tekst_bloka(blok).strip()
        if not t:
            continue
        if H.NASLOV_LIT.match(t) or je_back_matter_naslov(t):
            zadnji = t
            break
        if prvi is None:
            prvi = t
        tijelo.append(t)
    opis = ("brojan je samo glavni tekst: od „%s\" do %s; naslovnica, izjava, "
            "sažetak, sadržaj i završni aparat nisu uključeni"
            % ((prvi or "—")[:40],
               ("„%s\"" % zadnji[:40]) if zadnji else "kraja dokumenta"))
    return tijelo, opis


# --------------------------------------------------------------- čitanje docx

def qn(tag):
    from docx.oxml.ns import qn as _qn
    return _qn(tag)


# v1.1-fix (Q19): strop za raspakiranu veličinu dijela .docx-a. `z.read()` je
# raspakiravao koliko god je zapis najavio, pa je .docx s napuhanim theme1.xml
# (nekoliko KB u arhivi, gigabajti raspakirano) rušio provjeru na memoriji
# umjesto da uredno prijavi nalaz. Theme je popis fontova i realno je reda
# veličine desetak kB; 8 MB je već tri reda veličine iznad svega viđenog.
#
# v1.1-fix (Q19, drugi krug): prvi popravak je (a) vjerovao NAJAVLJENOJ veličini
# iz zaglavlja zipa — polje koje pošiljatelj sam upisuje, pa ne štiti ni od čega
# — i (b) stajao samo nad `word/theme/*`, dok se `word/document.xml` parsirao
# bez ikakvog stropa, i to PRIJE nego se do theme uopće dođe (`Document(rad)` u
# main()). Mjereno: arhiva od 54 kB koja najavljuje 12 MB document.xml popela je
# RSS na 303 MB. Strop je zato prebačen na STVARNO PROČITANE bajtove
# (`hr_text.procitaj_dio_zipa`), a glavni dijelovi se provjeravaju prije
# otvaranja dokumenta (`hr_text.prevelik_dio_docxa`).
MAX_DIO_BAJTOVA = H.MAX_XML_BAJTOVA


def _procitaj_dio(z, ime, strop=MAX_DIO_BAJTOVA):
    """Sadržaj dijela arhive, ili None ako STVARNI sadržaj prelazi strop."""
    return H.procitaj_dio_zipa(z, ime, strop)


def tema_fontovi(put):
    """{'major': ..., 'minor': ...} iz word/theme/theme1.xml."""
    out = {}
    try:
        with zipfile.ZipFile(put) as z:
            imena = [n for n in z.namelist() if n.startswith("word/theme/")]
            if not imena:
                return out
            sirovo = _procitaj_dio(z, sorted(imena)[0])
            if sirovo is None:
                return out
            korijen = ET.fromstring(sirovo)
    except (zipfile.BadZipFile, ET.ParseError, KeyError, OSError):
        return out
    for oznaka, kljuc in (("majorFont", "major"), ("minorFont", "minor")):
        el = korijen.find(f".//{{{NS_A}}}{oznaka}/{{{NS_A}}}latin")
        if el is not None and el.get("typeface"):
            out[kljuc] = el.get("typeface")
    return out


def font_iz_rpr(rpr, tema):
    """(ime, velicina_pt) iz jednog w:rPr, uz razrješavanje theme-referenci."""
    if rpr is None:
        return None, None
    ime = None
    rf = rpr.find(qn("w:rFonts"))
    if rf is not None:
        # v1.1-fix (Q19): w:hAnsi je slot koji u mnogim Wordovim dokumentima
        # doista nosi hrvatske dijakritičke znakove; sam w:ascii nije dovoljan.
        ime = rf.get(qn("w:ascii")) or rf.get(qn("w:hAnsi"))
        if not ime:
            t = rf.get(qn("w:asciiTheme")) or rf.get(qn("w:hAnsiTheme")) or ""
            if t:
                ime = tema.get("major" if t.startswith("major") else "minor")
    vel = None
    sz = rpr.find(qn("w:sz"))
    if sz is not None and sz.get(qn("w:val")):
        try:
            vel = float(sz.get(qn("w:val"))) / 2.0
        except ValueError:
            vel = None
    return ime, vel


def _elementi_tijela(el):
    """Djeca tijela dokumenta, uz prolaz kroz `w:sdt` omotače.

    v1.1-fix: Wordov automatski sadržaj (Insert → Table of Contents) i svako
    polje uneseno kao content control stoje unutar `w:sdt`, pa ih obilazak koji
    gleda samo izravnu djecu tijela ne vidi. Posljedica nije bila kozmetička:
    na stvarnom EFZG završnom radu s automatskim sadržajem naslov „Sadržaj"
    nije postojao ni za jednu provjeru, pa je obavezni dio „sadržaj" prijavljen
    kao „nema naslova za: sadržaj" (❌) na radu koji sadržaj uredno ima.
    `w:sdt` je omotač, ne sadržaj — zato se raspakirava, a ne broji.
    """
    sdt, sdt_content = qn("w:sdt"), qn("w:sdtContent")
    for dijete in el.iterchildren():
        if dijete.tag == sdt:
            sadrzaj = dijete.find(sdt_content)
            if sadrzaj is not None:
                yield from _elementi_tijela(sadrzaj)
        else:
            yield dijete


def blokovi(doc):
    """Odlomci i tablice u redoslijedu u kojem stoje u tijelu dokumenta."""
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    for dijete in _elementi_tijela(doc.element.body):
        if dijete.tag == qn("w:p"):
            yield "p", Paragraph(dijete, doc)
        elif dijete.tag == qn("w:tbl"):
            yield "tbl", Table(dijete, doc)


def stil_imena(p):
    try:
        return (p.style.name or "").strip()
    except Exception:
        return ""


def razina_naslova(p):
    """1 za Heading 1 / Naslov 1, 2.. za niže, 0 ako nije naslov."""
    m = re.match(r"(?i)^(?:heading|naslov)\s*(\d)$", stil_imena(p))
    if m:
        return int(m.group(1))
    return 0


def tekst_bloka(blok):
    """Vidljivi tekst odlomka, uključujući INLINE `w:sdt` (v. hr_text).

    `Paragraph.text` gleda samo izravnu djecu `w:r`, pa citat ubačen preko
    References → Insert Citation (Word ga sprema kao inline `w:sdt`) tiho
    nestane iz teksta koji provjere vide.
    """
    return H.tekst_odlomka(blok)


# ------------------------------------------- popis sadržaja nije izvor naslova
# v1.1-fix (drugi krug, PREKOREKCIJA): raspakiravanje `w:sdt` učinilo je Wordov
# automatski sadržaj vidljivim — i time su njegovi UNOSI („3. ZAKLJUČAK",
# „LITERATURA") postali kandidati za naslove. Posljedica je bila obrnuta od
# namjeravane: rad kojemu zaključak i literatura NE POSTOJE prolazio je kao
# „4 nađeno, 0 fali ✅" jer ih sadržaj nabraja. Prije popravka isti rad je
# ispravno dobivao ❌.
#
# Prijašnji filtar `_je_toc_redak` hvatao je samo redak s tabulatorom i brojem
# stranice, a Word ima kvačicu „Show page numbers" (i opciju hiperveza umjesto
# brojeva) — takav je sadržaj posve legitiman. Zato se popis sadržaja prepoznaje
# po STRUKTURI, u tri neovisna oblika, a ne po izgledu retka.
#
# Naslov samog sadržaja („SADRŽAJ") se NE odbacuje — on je jedini pravi naslov
# u tom bloku i profil ga traži kao obavezni dio.
def _je_sdt_sadrzaja(sdt):
    """Je li ovaj `w:sdt` blok popis sadržaja (a ne bilo koji content control)."""
    # 1. Wordov automatski sadržaj: docPartGallery = „Table of Contents".
    for el in sdt.iter(qn("w:docPartGallery")):
        if "table of contents" in (el.get(qn("w:val")) or "").lower():
            return True
    # 2. Polje TOC ili PAGEREF unutar bloka.
    for el in sdt.iter(qn("w:instrText")):
        s = " ".join((el.text or "").upper().split())
        if s.startswith("TOC") or " TOC " in f" {s} " or "PAGEREF" in s:
            return True
    prvi = None
    for p in sdt.iter(qn("w:p")):
        t = H.tekst_odlomka(p).strip()
        # 3. Blok sadrži unos u „toc N" stilu ili redak s vodičem/brojem stranice.
        if _je_toc_stil_unosa(_stil_iz_xml(p)) or _je_toc_redak(t):
            return True
        if t and prvi is None:
            prvi = t
    # 4. Blok koji POČINJE naslovom „SADRŽAJ" / „POPIS …" jest taj popis.
    return bool(prvi and _je_popisni_naslov(prvi))


def _stil_iz_xml(p_el):
    """Ime stila iz `w:pStyle` (radi na golom XML elementu, bez python-docx)."""
    ppr = p_el.find(qn("w:pPr"))
    if ppr is None:
        return ""
    st = ppr.find(qn("w:pStyle"))
    return (st.get(qn("w:val")) or "") if st is not None else ""


def _u_sdt_sadrzaja(p):
    """Leži li odlomak unutar `w:sdt` bloka koji je popis sadržaja."""
    el = getattr(p, "_p", p)
    sdt_tag = qn("w:sdt")
    roditelj = el.getparent()
    while roditelj is not None:
        if roditelj.tag == sdt_tag and _je_sdt_sadrzaja(roditelj):
            return True
        roditelj = roditelj.getparent()
    return False


def je_unos_sadrzaja(p, tekst):
    """Je li odlomak UNOS u popisu sadržaja (dakle ne i naslov u radu).

    PREKOREKCIJA (3. krug recenzije): prva verzija je `_je_toc_redak` primjenjivala
    bezuvjetno, a taj uzorak hvata svaki tekst s tri uzastopne točke ili s
    tabulatorom pa brojem. Odlomak sa STILOM naslova time je mogao ispasti iz
    popisa naslova — dakle rad koji poglavlje uredno ima izgubio bi ga iz svake
    provjere. Stil naslova je jači dokaz od izgleda retka: pravi naslov IZVAN
    bloka sadržaja nikad nije unos u sadržaju, bez obzira kako izgleda.
    """
    t = str(tekst or "").strip()
    if not t:
        return False
    u_sadrzaju = _u_sdt_sadrzaja(p)
    pravi_naslov = bool(razina_naslova(p) or _outline_razina(p))
    if pravi_naslov and not u_sadrzaju:
        return False
    if _je_toc_redak(t):
        return True
    if _je_toc_stil_unosa(stil_imena(p)):
        return True
    if not u_sadrzaju:
        return False
    # Unutar bloka sadržaja: naslov samog popisa i pravi Heading ostaju naslovi,
    # sve ostalo je nabrajanje tuđih naslova.
    return not (_je_popisni_naslov(t) or pravi_naslov)


def _prijelom_u_runovima(p, samo_ispred_teksta=False):
    """Ima li odlomak w:br type=page (po potrebi samo prije prvog runa s tekstom)."""
    for r in p._p.iter():
        if r.tag == qn("w:br") and r.get(qn("w:type")) == "page":
            return True
        if samo_ispred_teksta and r.tag == qn("w:t") and (r.text or "").strip():
            return False
    return False


def _prijelom_iza_teksta(p):
    """Prijelom stranice na KRAJU odlomka (Ctrl+Enter iza zadnjeg znaka)."""
    zadnji_tekst = -1
    prijelomi = []
    for i, el in enumerate(p._p.iter()):
        if el.tag == qn("w:t") and (el.text or "").strip():
            zadnji_tekst = i
        elif el.tag == qn("w:br") and el.get(qn("w:type")) == "page":
            prijelomi.append(i)
    return any(i > zadnji_tekst for i in prijelomi)


def _prekid_sekcije(p):
    """Prekid sekcije tipa page/oddPage zapisan u pPr ovog odlomka."""
    ppr = p._p.find(qn("w:pPr"))
    if ppr is None:
        return False
    sekt = ppr.find(qn("w:sectPr"))
    if sekt is None:
        return False
    tip = sekt.find(qn("w:type"))
    if tip is None:
        return True          # bez w:type Word podrazumijeva nextPage
    return tip.get(qn("w:val")) in ("nextPage", "page", "oddPage", "evenPage")


def _stil_ima_prijelom(p, dubina=6):
    """Nasljeđuje li odlomak `pageBreakBefore` iz svog stila (ili osnovnog stila).

    Prati `w:basedOn` lanac jer se u praksi pravilo često postavi na izvedeni
    stil, a nasljeđuje ga više naslova.
    """
    try:
        stil = p.style
    except Exception:
        return False
    while stil is not None and dubina > 0:
        try:
            if stil.paragraph_format.page_break_before:
                return True
        except Exception:
            pass
        try:
            stil = stil.base_style
        except Exception:
            return False
        dubina -= 1
    return False


def ima_prijelom_prije(p, prethodni=None):
    """Je li pred ovim odlomkom stvarni prijelom stranice.

    v1.1-fix (D8): Wordov Ctrl+Enter NIKAD ne završi u samom naslovu — prijelom
    ostane u prethodnom odlomku ili u praznom odlomku-razmaknici ispred naslova.
    Zato se gleda unatrag: pageBreakBefore na naslovu, prijelom u naslovu prije
    prvog runa s tekstom, prijelom na kraju prethodnog odlomka, prazna
    razmaknica čiji je jedini sadržaj prijelom, ili prekid sekcije.
    """
    try:
        if p.paragraph_format.page_break_before:
            return True
    except Exception:
        pass
    # v1.1-fix (dopuna D8): `pageBreakBefore` postavljen na STILU (npr. na
    # „Heading 1") naslijeđen je na svaki naslov i Word ga poštuje — to je i
    # preporučeni način da se pravilo primijeni na cijeli rad odjednom, umjesto
    # ručno po odlomku. Provjera je gledala samo razinu odlomka, pa je dokument
    # koji pravilo primjenjuje ISPRAVNO, preko stila, prijavljivala kao
    # „0/N poglavlja ima prijelom". Uhvaćeno kad je generator (build_docx.py)
    # proizveo dokument po ovom istom profilu i pao na vlastitoj provjeri.
    if _stil_ima_prijelom(p):
        return True
    if _prijelom_u_runovima(p, samo_ispred_teksta=True):
        return True
    for vrsta, blok in reversed(list(prethodni or [])):
        if vrsta != "p":
            return False                      # tablica ili drugi blok — stani
        if _prekid_sekcije(blok):
            return True
        tekst = tekst_bloka(blok).strip()
        if not tekst:
            if _prijelom_u_runovima(blok):
                return True
            continue                          # prazna razmaknica — gledaj dalje
        return _prijelom_iza_teksta(blok)
    return False


# ------------------------------------------------------------------- nalazi

# Strojni ugovor izvještaja (audit: „machine contract"). `stanje` ostaje emoji
# radi ljudske tablice; `severity` je enum koji potrošač smije parsirati.
SEVERITY = {OK: "u_skladu", UPOZ: "za_provjeru", LOSE: "krsenje"}


def _rule_id_iz_naziva(pravilo):
    """Rezervni identifikator kad pozivatelj nije zadao `rule_id`.

    Ne bi se smio koristiti — svi pozivi u ovom modulu zadaju `rule_id`
    eksplicitno; ovo postoji da vanjski/naknadni poziv ne ostane bez ključa.
    """
    osnova = re.sub(r"\s*\(.*?\)\s*", "", str(pravilo or "")).strip().lower()
    osnova = H.bez_dijakritika(osnova)
    return "nepoznato." + (re.sub(r"[^a-z0-9]+", "_", osnova).strip("_") or "pravilo")


# ------------------------------------------------ Q15: pravilo bez vrijednosti
# v1.1-fix (Q15a, drugi krug): dogovor o „pravilo postoji, ali mu je vrijednost
# neupotrebljiva (null ili izvan enuma)" već postoji u repozitoriju —
# profile_rules.UNKNOWN_RULES_KEY i unknown_rule_pointers(), a
# references/fakulteti/_resolved_schema.json izrijekom traži da ga potrošač
# prikaže kao UPOZORENJE, ne kao „profil ne propisuje". check_rules je takav
# potrošač i taj je ključ dosad ignorirao: profil kojemu je netko eksplicitno
# poništio `format.prored` dobivao je zeleno „profil ne propisuje ✅", što je
# suprotno od istine — pravilo postoji, samo ga ne razumijemo.
#
# Nijedan isporučeni profil (efzg.json, fpzg.json) nema null, pa ovo ne može
# proizvesti nalaz na legitimnom ulazu.
def _pointer_pogadja(pointer, rule_id):
    """Odgovara li JSON Pointer („/format/prored") ovom rule_id-u.

    Podudaranje je po SUFIKSU segmenata jer se dva imenovanja povijesno
    razlikuju u korijenu („/struktura/prikazi/natpis" ↔ „prikazi.natpis"), i po
    PREFIKSU jer poništen roditelj („/format/odlomak") pokriva svako pravilo
    ispod sebe („format.odlomak.min_recenica").
    """
    tockasti = pointer.strip("/").replace("/", ".")
    if not tockasti or not rule_id:
        return False
    return (tockasti == rule_id
            or tockasti.endswith("." + rule_id)
            or rule_id.startswith(tockasti + "."))


def nepoznata_pravila(profil):
    """Pointeri pravila koja profil postavlja, a Katedra ih ne zna provjeriti."""
    from profile_rules import UNKNOWN_RULES_KEY, unknown_rule_pointers
    zapisano = profil.get(UNKNOWN_RULES_KEY)
    if isinstance(zapisano, list):
        return list(zapisano)          # resolver ga je već izračunao
    try:
        return unknown_rule_pointers(profil)   # ravan --profil, bez resolvera
    except Exception:
        return []


class Izvjestaj:
    def __init__(self, profil, put_profila):
        self.redci = []
        self.profil = profil
        self.put_profila = put_profila
        self.izvor = (profil.get("izvor") or {}).get("dokument") \
            or f"profil {os.path.basename(put_profila)} (bez navedenog izvora)"
        self.za_potvrdu = profil.get("status") == "nepotvrdeno"
        # Fakultet koji nije prošao faculty_scale_gate: pravila su IZVEDENA, ne pročitana
        # iz službenih strojno dostupnih uputa. Nalazi se i dalje računaju i ispisuju —
        # jer je alternativa ručna provjera, što željezno pravilo 8 ne dopušta — ali ne
        # blokiraju predaju, jer bi to bilo blokiranje na temelju pretpostavke.
        self.advisory = (profil.get("admisija") == "nije-admitiran"
                         or profil.get("nalazi") == "advisory")
        self.nepoznata = nepoznata_pravila(profil)
        # v1.9 (nalaz 4): provenance po pravilu — iz resolvera ili iz sidecara
        # uz profil. Bez njega na nepotvrđenom profilu nijedno pravilo nije
        # „potvrđeno" pa sve daje ⚠️ (pravilo 18 paketa).
        self.provenance = _normaliziraj_provenance(
            dict(_PROVENANCE) or _ucitaj_provenance_sidecar(put_profila))

    def pravilo_potvrdeno(self, rule_id):
        """Stoji li pravilo stvarno u službenim uputama (provenance type=explicit).

        Na potvrđenom profilu vrijedi sve; na nepotvrđenom samo ono što
        provenance izrijekom potvrđuje. Pointer i rule_id se uspoređuju kao
        segmenti (jedan je uređeni podniz drugoga), jer se korijeni povijesno
        razlikuju („/struktura/prikazi/natpis" ↔ „prikazi.natpis") i jer opseg
        nosi tip rada u pointeru („/struktura/opseg/seminarski/rijeci").
        """
        if not self.za_potvrdu:
            return True

        def podniz(kraci, dulji):
            it = iter(dulji)
            return all(any(s == d for d in it) for s in kraci)

        rid = [s for s in str(rule_id or "").split(".") if s]
        # točan/podnizni pogodak, pa najbliži roditelj („format.odlomak" za
        # „format.odlomak.min_recenica" — provenance je tamo po retcima)
        for pointer, meta in self.provenance.items():
            if (meta or {}).get("type") != "explicit":
                continue
            seg = [s for s in pointer.strip("/").split("/") if s]
            if not seg or not rid:
                continue
            if podniz(seg, rid) or podniz(rid, seg):
                return True
            if len(rid) > 2 and (seg == rid[:-1] or seg[-len(rid) + 1:] == rid[:-1]):
                return True
        return False

    def dodaj(self, pravilo, trazeno, nadjeno, stanje, detalji=None,
              rule_id=None, lokacije=None):
        """Zapiši jedan redak izvještaja.

        `rule_id` je STABILAN strojni identifikator pravila i namjerno je oblika
        putanje u profil (`format.prored`, `prikazi.izvor_ispod`) — isto što
        provenance sloj već koristi kao JSON Pointer. Do audita jedini ugovor
        prema agentu bila je hrvatska prikazna niska (`pravilo`), s vrstom rada
        interpoliranom u nju („broj poglavlja (seminarski)") i emojijem kao
        statusom. Posljedica nije bila kozmetička: test koji redak traži po
        prozi tiho prestane išta provjeravati čim se poruka preformulira, i
        upravo su tako regresije u ovom modulu prolazile zelene.

        `severity` je enum izveden iz `stanje` da potrošač ne mora raditi
        hrvatski NLP nad vlastitim izlazom. Prikazna polja ostaju nepromijenjena.
        """
        rid = rule_id or _rule_id_iz_naziva(pravilo)
        detalji = list(detalji or [])
        pogodeni = [p for p in self.nepoznata if _pointer_pogadja(p, rid)]
        if pogodeni:
            # Q15a: pravilo POSTOJI u profilu, ali mu vrijednost nije upotrebljiva.
            # Ne smije izgledati ni kao „profil ne propisuje" ni kao uredan nalaz.
            stanje = UPOZ if stanje == OK else stanje
            trazeno = "profil postavlja, vrijednost neupotrebljiva"
            detalji.append(
                "profil postavlja " + ", ".join(pogodeni)
                + " na null ili na vrijednost izvan dopuštenog skupa, pa se to "
                  "pravilo NIJE provjerilo (v1.1-fix Q15a)")
            detalji.append("Što napraviti: ispravi tu vrijednost u profilu "
                           "fakulteta ili je ukloni ako fakultet pravilo ne "
                           "propisuje — prazno i „ne propisuje\" nije isto.")
        if stanje == LOSE and not self.pravilo_potvrdeno(rid):
            # v1.9 (nalaz 4), pravilo 18: crveno samo za pravila koja stvarno
            # stoje u službenim uputama; nepotvrđeno pravilo daje ⚠️, ne ❌.
            stanje = UPOZ
            detalji.append("pravilo nije potvrđeno u službenim uputama (status "
                           "profila: nepotvrdeno) — ⚠️ umjesto ❌, ne blokira")
        self.redci.append({
            "rule_id": rid,
            "pravilo": pravilo,
            "trazeno": "" if trazeno is None else str(trazeno),
            "nadjeno": "" if nadjeno is None else str(nadjeno),
            "stanje": stanje,
            "severity": SEVERITY.get(stanje, "nepoznato"),
            "detalji": detalji,
            "lokacije": list(lokacije or []),
            "izvor": self.izvor,
            "za_potvrdu": self.za_potvrdu,
        })

    def zakljuci_nepoznata(self):
        """Q15a: pravila koja nijedna provjera nije ni dotaknula.

        Provjera koja se pri praznoj vrijednosti tiho vrati (npr.
        `provjeri_odlomke` bez `min_recenica`) inače ne ostavlja nijedan redak,
        pa poništeno pravilo nestane iz izvještaja bez traga.
        """
        for pointer in self.nepoznata:
            if any(_pointer_pogadja(pointer, r["rule_id"]) for r in self.redci):
                continue
            self.dodaj(pointer.strip("/").replace("/", "."), None,
                       "nije provjereno", UPOZ,
                       rule_id=pointer.strip("/").replace("/", "."))

    @property
    def nalazi(self):
        return [r for r in self.redci if r["stanje"] != OK]


# ------------------------------------------------------------------ provjere

def font_stila(stil, tema, dubina=0):
    """(ime, velicina) iz stila, uz nasljeđivanje kroz w:basedOn."""
    if stil is None or dubina > 6:
        return None, None
    try:
        el = stil.element
    except AttributeError:
        return None, None
    ime, vel = font_iz_rpr(el.find(qn("w:rPr")), tema)
    if ime is None or vel is None:
        try:
            roditelj = stil.base_style
        except AttributeError:
            roditelj = None
        r_ime, r_vel = font_stila(roditelj, tema, dubina + 1)
        ime = ime or r_ime
        vel = vel if vel is not None else r_vel
    return ime, vel


def provjeri_font(iz, doc, put, profil):
    """Font i veličina kroz sva četiri sloja + što je stvarno na tekstu.

    Sloj niže u lancu (theme, docDefaults) ne mora biti kršenje ako ga viši
    nadjačava, ali jest rizik: tablice i okviri često ostanu na defaultu.
    Zato se ❌ dodjeljuje samo EFEKTIVNOM fontu, a razlika među slojevima je ⚠️.
    """
    fmt = profil.get("format") or {}
    trazeni = fmt.get("font") or []
    tr_vel = fmt.get("velicina_pt")
    tema = tema_fontovi(put)

    slojevi, slojevi_vel = {}, {}
    if tema:
        slojevi["theme"] = tema.get("minor") or tema.get("major")

    dd_ime = dd_vel = None
    dd = doc.styles.element.find(qn("w:docDefaults"))
    if dd is not None:
        rpr_def = dd.find(qn("w:rPrDefault"))
        dd_ime, dd_vel = font_iz_rpr(
            rpr_def.find(qn("w:rPr")) if rpr_def is not None else None, tema)
        slojevi["docDefaults"] = dd_ime
        slojevi_vel["docDefaults"] = dd_vel

    n_ime = n_vel = None
    try:
        n_ime, n_vel = font_stila(doc.styles["Normal"], tema)
    except (KeyError, AttributeError):
        pass
    slojevi["stil Normal"] = n_ime
    slojevi_vel["stil Normal"] = n_vel

    pricuva_ime = n_ime or dd_ime or slojevi.get("theme")
    pricuva_vel = n_vel if n_vel is not None else dd_vel

    # Razriješi runove GLAVNE PROZE, ne naslovnicu, naslove, retke "Izvor:"
    # ni završne popise. Ako dokument nema prepoznatljiv marker tijela, zadrži
    # stari fallback i mjeri sve nenatpisne odlomke da ne izgubimo detekciju.
    redoslijed = list(blokovi(doc))
    poc_tijela = indeks_pocetka_tijela(redoslijed)
    u_tijelu = poc_tijela is None
    ef_font, ef_vel, run_font, run_vel = Counter(), Counter(), Counter(), Counter()
    for i, (vrsta, blok) in enumerate(redoslijed):
        if vrsta != "p":
            continue
        tekst = tekst_bloka(blok).strip()
        if poc_tijela is not None and i == poc_tijela:
            u_tijelu = True
        if u_tijelu and (H.NASLOV_LIT.match(tekst) or je_back_matter_naslov(tekst)):
            break
        if not u_tijelu or razina_naslova(blok) or H.IZVOR.match(tekst):
            continue
        s_ime, s_vel = font_stila(getattr(blok, "style", None), tema)
        for r in blok.runs:
            n = len((r.text or "").strip())
            if not n:
                continue
            ime, vel = font_iz_rpr(r._r.find(qn("w:rPr")), tema)
            if ime:
                run_font[ime] += n
            if vel:
                run_vel[vel] += n
            ef_font[ime or s_ime or pricuva_ime or "—"] += n
            ef_vel[vel if vel is not None else
                   (s_vel if s_vel is not None else pricuva_vel) or 0] += n

    efektivni = [k for k, _ in ef_font.most_common() if k and k != "—"] \
        or [x for x in [pricuva_ime] if x]
    opis = " · ".join(f"{k}: {v or '—'}" for k, v in slojevi.items())
    if run_font:
        opis += " · runovi: " + ", ".join(f"{k} ({v} zn.)"
                                          for k, v in run_font.most_common(4))
    else:
        opis += " · runovi: bez vlastitog fonta (nasljeđuju stil)"
    # theme se ocjenjuje samo posredno: ako ga itko doista koristi (rFonts
    # asciiTheme), razriješen je već u docDefaults ili u stilu.
    razliciti_slojevi = {v for v in (dd_ime, n_ime) if v} | set(run_font)

    if not trazeni:
        iz.dodaj("font", "profil ne propisuje", ", ".join(efektivni) or "—",
                 UPOZ if len(efektivni) > 1 else OK, [opis],
                 rule_id="format.font")
    else:
        dopusteni = {norm(f) for f in trazeni}
        krivi = [v for v in efektivni if norm(v) not in dopusteni]
        neusklađeni = [v for v in sorted(razliciti_slojevi, key=str.lower)
                       if norm(v) not in dopusteni and v not in krivi]
        det = [opis]
        if krivi:
            stanje = LOSE
            det.append("izvan propisa u tekstu: " + ", ".join(krivi))
        elif neusklađeni:
            stanje = UPOZ
            det.append("niži sloj izvan propisa: " + ", ".join(neusklađeni)
                       + " — stil Normal to nadjačava za prozu, ali tablice, okviri i "
                         "novi odlomci mogu pasti na taj font")
        elif len(efektivni) > 1:
            stanje = UPOZ
            det.append("svi su dopušteni, ali dokument miješa više fontova — ujednači")
        else:
            stanje = OK
        iz.dodaj("font", " ili ".join(trazeni), ", ".join(efektivni) or "—", stanje, det,
                 rule_id="format.font")

    # ---- veličina
    ef_velicine = sorted({v for v in ef_vel if v}) or \
        [v for v in [pricuva_vel] if v]
    opis_v = [f"{k}: {v:g} pt" for k, v in slojevi_vel.items() if v]
    if run_vel:
        opis_v.append("runovi: " + ", ".join(f"{k:g} pt ({v} zn.)"
                                             for k, v in run_vel.most_common(4)))
    nadjeno_v = ", ".join(f"{v:g} pt" for v in ef_velicine) or "—"
    if tr_vel is None:
        iz.dodaj("veličina fonta", "profil ne propisuje", nadjeno_v,
                 UPOZ if len(ef_velicine) > 1 else OK, opis_v,
                 rule_id="format.velicina_pt")
        return
    if not ef_velicine:
        iz.dodaj("veličina fonta", f"{tr_vel:g} pt", "nije zapisana nigdje", UPOZ,
                 ["dokument se oslanja na Wordov default — postavi je izrijekom"],
                 rule_id="format.velicina_pt")
        return
    naslov_pt = fmt.get("naslov_poglavlja_pt")
    krive = [v for v in ef_velicine if abs(v - float(tr_vel)) > 0.01]
    sloj_krivi = [v for v in sorted({x for x in slojevi_vel.values() if x})
                  if abs(v - float(tr_vel)) > 0.01 and v not in krive]
    det = list(opis_v)
    if krive:
        stanje = LOSE
        det.append("izvan propisa u glavnoj prozi: "
                   + ", ".join(f"{v:g} pt" for v in krive))
        if naslov_pt:
            det.append(f"(naslov poglavlja smije biti {naslov_pt:g} pt — naslovi nisu "
                       "uključeni u ovo mjerenje)")
    else:
        # Niži Word sloj nije stvarna veličina proze ako ga svi izmjereni runovi/
        # stilovi tijela nadjačavaju. Zadrži ga u detaljima radi dijagnostike, ali
        # ne prijavljuj false positive samo zbog docDefaults/theme vrijednosti.
        stanje = OK
        if sloj_krivi:
            det.append("niži sloj izvan propisa (informativno, nadjačan u glavnoj prozi): "
                       + ", ".join(f"{v:g} pt" for v in sloj_krivi))
    iz.dodaj("veličina fonta", f"{tr_vel:g} pt", nadjeno_v, stanje, det,
                 rule_id="format.velicina_pt")


def provjeri_prored_i_poravnanje(iz, doc, profil):
    fmt = profil.get("format") or {}
    try:
        pf = doc.styles["Normal"].paragraph_format
    except (KeyError, AttributeError):
        pf = None

    tr_prored = fmt.get("prored")
    nadjen = None
    if pf is not None:
        ls = pf.line_spacing
        if ls is not None:
            nadjen = float(ls) if isinstance(ls, float) else round(ls.pt / 12.0, 2)
    if tr_prored is None:
        iz.dodaj("prored (stil Normal)", "profil ne propisuje",
                 nadjen if nadjen else "nije zadan", OK,
                 rule_id="format.prored")
    elif nadjen is None:
        iz.dodaj("prored (stil Normal)", tr_prored, "nije zadan u stilu Normal", UPOZ,
                 ["prored je vjerojatno postavljen ručno po odlomcima — "
                  "postavi ga u stilu da vrijedi za cijeli rad"],
                 rule_id="format.prored")
    else:
        stanje = OK if abs(nadjen - float(tr_prored)) < 0.01 else LOSE
        iz.dodaj("prored (stil Normal)", tr_prored, f"{nadjen:g}", stanje,
                 rule_id="format.prored")

    tr_por = fmt.get("poravnanje")
    imena = {0: "lijevo", 1: "sredina", 2: "desno", 3: "obostrano"}
    nad = None
    if pf is not None and pf.alignment is not None:
        nad = imena.get(int(pf.alignment), str(pf.alignment))
    if tr_por is None:
        iz.dodaj("poravnanje (stil Normal)", "profil ne propisuje", nad or "nije zadano", OK,
                 rule_id="format.poravnanje")
    elif nad is None:
        iz.dodaj("poravnanje (stil Normal)", tr_por, "nije zadano u stilu Normal", UPOZ,
                 ["Word tada poravnava lijevo; postavi obostrano u stilu Normal"
                  if tr_por == "obostrano" else "postavi izrijekom u stilu Normal"],
                 rule_id="format.poravnanje")
    else:
        iz.dodaj("poravnanje (stil Normal)", tr_por, nad,
                 OK if nad == tr_por else LOSE,
                 rule_id="format.poravnanje")


def provjeri_margine(iz, doc, profil):
    trazene = ((profil.get("format") or {}).get("margine_cm")) or {}
    # Q15 (3. krug): poništena POJEDINA margina (`format.margine_cm.gore: null`)
    # dolazila je do `float(None)` i rušila alat neuhvaćenom iznimkom PRIJE nego
    # što bi sloj nepoznatih pravila stigao ispisati svoje ⚠️. Ostalih 37
    # pointera je to podnosilo; ova četiri i `struktura.opseg.<tip>` nisu.
    # Vrijednost bez broja se ovdje samo preskače — redak ⚠️ „profil postavlja,
    # vrijednost neupotrebljiva" emitira `Izvjestaj.dodaj`/`zakljuci_nepoznata`.
    trazene = {k: v for k, v in trazene.items() if isinstance(v, (int, float))}
    if not trazene:
        return
    polja = [("gore", "top_margin"), ("dolje", "bottom_margin"),
             ("lijevo", "left_margin"), ("desno", "right_margin")]
    for i, sek in enumerate(doc.sections, 1):
        odstupanja, opis = [], []
        for naziv, atr in polja:
            if naziv not in trazene:
                continue
            vrijednost = getattr(sek, atr)
            stvarno = None if vrijednost is None else round(vrijednost.cm, 2)
            opis.append(f"{naziv} {stvarno if stvarno is not None else '—'}")
            if stvarno is None or abs(stvarno - float(trazene[naziv])) > TOLERANCIJA_CM:
                odstupanja.append(f"{naziv}: traženo {trazene[naziv]:g} cm, "
                                  f"nađeno {stvarno if stvarno is not None else '—'} cm")
        oznaka = "margine" if len(doc.sections) == 1 else f"margine (sekcija {i})"
        iz.dodaj(oznaka,
                 ", ".join(f"{k} {trazene[k]:g}" for k, _ in polja if k in trazene) + " cm",
                 ", ".join(opis) + " cm",
                 LOSE if odstupanja else OK,
                 odstupanja + ([f"tolerancija {TOLERANCIJA_CM:g} cm"] if odstupanja else []),
                 rule_id="format.margine_cm")


#: Format papira: naziv → (širina cm, visina cm). Zadano je A4 jer na njemu
#: počiva i procjena broja stranica niže u ovom modulu.
FORMATI_PAPIRA = {"a4": (21.0, 29.7), "letter": (21.59, 27.94), "a5": (14.8, 21.0)}


def provjeri_format_stranice(iz, doc, profil):
    """Format papira, jer se opseg procjenjuje uz pretpostavku A4 (kvar 44).

    Margine su se mjerile, papir nije. Rad na US Letteru prolazio je sve provjere
    jer su margine bile točne, a razlika je 1,76 cm visine — dovoljno da se
    prijelom, pa i broj stranica, promijene.
    """
    trazeni = str(((profil.get("format") or {}).get("stranica")) or "a4").lower()
    ocekivano = FORMATI_PAPIRA.get(trazeni)
    if not ocekivano:
        return
    sirina_t, visina_t = ocekivano
    for i, sek in enumerate(doc.sections, 1):
        w, h = sek.page_width, sek.page_height
        stvarno_w = None if w is None else round(w.cm, 2)
        stvarno_h = None if h is None else round(h.cm, 2)
        opis = (f"{stvarno_w} × {stvarno_h} cm" if stvarno_w and stvarno_h else "—")
        krivo = (stvarno_w is None or stvarno_h is None
                 or abs(stvarno_w - sirina_t) > TOLERANCIJA_CM
                 or abs(stvarno_h - visina_t) > TOLERANCIJA_CM)
        detalji = []
        if krivo:
            nadjen = next((n.upper() for n, (a, b) in FORMATI_PAPIRA.items()
                           if stvarno_w and stvarno_h
                           and abs(stvarno_w - a) <= TOLERANCIJA_CM
                           and abs(stvarno_h - b) <= TOLERANCIJA_CM), None)
            detalji.append(f"traženo {trazeni.upper()} ({sirina_t:g} × {visina_t:g} cm), "
                           f"nađeno {opis}" + (f" — to je {nadjen}" if nadjen else ""))
            detalji.append("procjena broja stranica u ovom alatu pretpostavlja A4")
        oznaka = ("format stranice" if len(doc.sections) == 1
                  else f"format stranice (sekcija {i})")
        iz.dodaj(oznaka, f"{trazeni.upper()} {sirina_t:g} × {visina_t:g} cm", opis,
                 LOSE if krivo else OK, detalji, rule_id="format.stranica")


def _razina_opsega(opseg, tip, profil):
    """Koliko je težak nalaz o opsegu za ovaj tip rada.

    Raspon koji ne stoji u službenim uputama ne smije davati ❌. Na FPZG-u je
    raspon seminarskog rada izrijekom naslijeđena procjena („opseg ovisi o
    silabusu pojedinoga kolegija"), a esej kao vrsta rada ondje uopće ne postoji —
    pa je svaki esej završavao crvenim kršenjem pravila kojega nema. Crveno je
    rezervirano za ono što se može pokazati prstom u dokumentu fakulteta.
    """
    razine = {"obavijest": None, "za_potvrdu": UPOZ, "krsenje": LOSE}
    zadano = (opseg.get("_razina_nalaza") or {}).get(tip)
    if zadano in razine:
        return razine[zadano]
    return UPOZ if profil.get("status") == "nepotvrdeno" else LOSE


def provjeri_opseg(iz, profil, tip, broj_rijeci, opis_opsega=None,
                   zadano_korisnikom=None):
    """`zadano_korisnikom` je (min, max) koji je korisnik izrijekom zadao.

    Ako postoji, on je mjerilo: student piše po uputi kolegija, a ne po
    ovojnici koju je profil naslijedio iz ranijeg skilla.
    """
    opseg = ((profil.get("struktura") or {}).get("opseg")) or {}
    if zadano_korisnikom:
        mn, mx = zadano_korisnikom
        iz.dodaj("broj riječi (zadao korisnik)", f"{mn}–{mx}", broj_rijeci,
                 OK if mn <= broj_rijeci <= mx else LOSE,
                 ["Raspon je zadao korisnik i ima prednost pred profilom fakulteta."]
                 + ([opis_opsega] if opis_opsega else []),
                 rule_id="struktura.opseg.rijeci")
        return
    if not opseg:
        return
    # Q15 (3. krug): `struktura.opseg.<tip>: null` je dolazio do `pravilo.get(...)`
    # i rušio alat neuhvaćenom iznimkom. Poništen blok se ovdje ponaša kao da ga
    # nema; da ga profil postavlja, kaže sloj nepoznatih pravila svojim ⚠️.
    if not isinstance(opseg.get(tip), dict):
        iz.dodaj(f"opseg ({tip})", ", ".join(sorted(opseg)) or "—",
                 f"profil nema pravila za tip „{tip}\"", UPOZ,
                 ["provjeri --tip; opseg i broj poglavlja nisu provjereni"],
                 rule_id="struktura.opseg")
        return
    pravilo = opseg[tip]

    # v1.1-fix (Q2): opseg mjerenja se ispisuje izrijekom — broj bez opsega
    # nije provjerljiv.
    opseg_det = [opis_opsega] if opis_opsega else []

    razina = _razina_opsega(opseg, tip, profil)
    if pravilo.get("rijeci") and razina is None:
        opseg_det = list(opseg_det) + [
            pravilo.get("napomena")
            or "Opseg za ovaj tip rada zadaje silabus kolegija; profil ga ne mjeri."]
        iz.dodaj(f"broj riječi ({tip})", "zadaje silabus", broj_rijeci, OK, opseg_det,
                 rule_id="struktura.opseg.rijeci")
    elif pravilo.get("rijeci"):
        mn, mx = pravilo["rijeci"]
        u_rasponu = mn <= broj_rijeci <= mx
        iz.dodaj(f"broj riječi ({tip})", f"{mn}–{mx}", broj_rijeci,
                 OK if u_rasponu else razina, list(opseg_det),
                 rule_id="struktura.opseg.rijeci")
    elif pravilo.get("napomena"):
        iz.dodaj(f"broj riječi ({tip})", "zadaje silabus", broj_rijeci, OK,
                 list(opseg_det) + [pravilo["napomena"]],
                 rule_id="struktura.opseg.rijeci")

    if pravilo.get("stranice"):
        mn, mx = pravilo["stranice"]
        p_min = broj_rijeci / RIJECI_PO_STRANICI[1]
        p_max = broj_rijeci / RIJECI_PO_STRANICI[0]
        preklapa = p_max >= mn and p_min <= mx
        iz.dodaj(f"broj stranica ({tip})", f"{mn}–{mx}",
                 f"PROCJENA {p_min:.0f}–{p_max:.0f}",
                 OK if preklapa else UPOZ,
                 ["PROCJENA, ne mjerenje: .docx ne sadrži broj stranica — izveden je iz "
                  f"{broj_rijeci} riječi uz {RIJECI_PO_STRANICI[0]}–{RIJECI_PO_STRANICI[1]} "
                  "riječi po stranici (A4, prored 1,5).",
                  "Točan broj daje samo render (Word ili LibreOffice → PDF)."]
                 + opseg_det,
                 rule_id="struktura.opseg.stranice")


def sadrzajna_poglavlja(naslovi1):
    """Sadržajna poglavlja = Heading 1 OD uvoda DO prvog naslova završnog aparata.

    v1.1-fix (Q1a): oduzimanje „allowlist" naslova je istom radu istodobno
    tražilo SAŽETAK i ŽIVOTOPIS (obavezni dijelovi) i kažnjavalo ih (broj
    poglavlja). Raspon rješava oboje bez popisa doslovnih nizova.
    """
    poc = next((i for i, t in enumerate(naslovi1) if je_uvodni_naslov(t)), None)

    # Strukturni put (v. BROJ_NASLOVA_RE gore): kad rad numerira poglavlja,
    # sadržajni dio je Uvod plus niz numeriranih naslova iza njega, a prvi
    # nenumerirani naslov zatvara tijelo. Ne ovisi o tome kako se aparat zove.
    if rad_numerira_poglavlja(naslovi1):
        if poc is None:
            return [t for t in naslovi1 if je_numeriran_naslov(t)]
        izlaz = [naslovi1[poc]]
        for naslov in naslovi1[poc + 1:]:
            if not je_numeriran_naslov(naslov):
                break
            izlaz.append(naslov)
        return izlaz

    # Pričuva: rad bez numeracije — ostaje rječnički predikat.
    if poc is None:
        return [t for t in naslovi1 if not je_back_matter_naslov(t)]
    kraj = next((i for i in range(poc + 1, len(naslovi1))
                 if je_back_matter_naslov(naslovi1[i])), len(naslovi1))
    return naslovi1[poc:kraj]


def provjeri_poglavlja(iz, profil, tip, naslovi1):
    opseg = (((profil.get("struktura") or {}).get("opseg")) or {}).get(tip) or {}
    if not opseg.get("poglavlja"):
        return
    mn, mx = opseg["poglavlja"]
    n = len(naslovi1)
    iz.dodaj(f"broj poglavlja ({tip})", f"{mn}–{mx}", n,
             OK if mn <= n <= mx else LOSE,
             ["Heading 1: " + " · ".join(t[:40] for t in naslovi1)] if naslovi1
             else ["nema nijednog odlomka u stilu Heading 1 — Katedra ne vidi strukturu"],
                 rule_id="struktura.opseg.poglavlja")


def provjeri_prijelome(iz, profil, h1):
    fmt = profil.get("format") or {}
    if "prijelom_pred_poglavljem" not in fmt:
        return
    trazi = bool(fmt["prijelom_pred_poglavljem"])
    if not h1:
        iz.dodaj("prijelom pred poglavljem",
                 "da" if trazi else "ne", "nema Heading 1", UPOZ,
                 rule_id="format.prijelom_pred_poglavljem")
        return
    prvi = h1[0][0]
    krivi = []
    for redni, (idx, tekst, prijelom) in enumerate(h1):
        if idx == prvi and trazi:
            continue          # prvo poglavlje ne treba prijelom ispred sebe
        if prijelom != trazi:
            krivi.append(tekst[:48])
    iz.dodaj("prijelom pred poglavljem", "da" if trazi else "ne",
             f"{sum(1 for _, _, b in h1 if b)}/{len(h1)} poglavlja ima prijelom",
             OK if not krivi else LOSE,
             ([("bez prijeloma: " if trazi else "s prijelomom (profil ga ne traži): ")
               + " · ".join(krivi)] if krivi else []),
                 rule_id="format.prijelom_pred_poglavljem")


SINONIMI = {
    "popis izvora": ["popis izvora", "literatura", "bibliografija", "izvori"],
    "popis literature": ["popis literature", "literatura", "popis izvora", "bibliografija"],
    "sadrzaj": ["sadrzaj", "kazalo"],
    "sazetak i kljucne rijeci hrvatski": ["sazetak", "kljucne rijeci"],
    "summary and keywords engleski": ["summary", "abstract", "keywords"],
    "izjava o akademskoj cestitosti": ["akademskoj cestitosti", "izjava o izvornosti",
                                       "izjava o autorstvu"],
    "zivotopis": ["zivotopis", "curriculum vitae"],
    "popis slika i grafikona": ["popis slika", "popis grafikona", "popis ilustracija"],
    "teorijski okvir": ["teorijski okvir", "teorijsk"],
    "metodologija": ["metodolog", "metoda istrazivanja"],
    "analiza": ["analiz", "empirij", "studija slucaja", "rezultat"],
}
# v1.1-fix (Q1c): korijeni sinonima. Hrvatski naslov stoji u padežu
# („POPIS LITERATURE", „POPISU IZVORA"), pa nominativna usporedba pada baš na
# najčešćem obliku naslova popisa izvora.
SINONIMI_KORIJENI = {
    "popis izvora": ["literatur", "izvor", "bibliograf", "referenc"],
    "popis literature": ["literatur", "izvor", "bibliograf", "referenc"],
    "sadrzaj": ["sadrzaj", "kazalo"],
    "sazetak i kljucne rijeci hrvatski": ["sazet"],
    "sazetak": ["sazet"],
    "zakljucak": ["zakljuc"],
    "zivotopis": ["zivotopis", "curriculum vitae"],
}
# Dijelovi koji se iz teksta ne mogu pouzdano prepoznati (nemaju vlastiti naslov).
NEPROVJERLJIVO = ("naslovnica",)


def _simetricni_sinonimi(tablica):
    """Sinonimi vrijede u OBA smjera.

    Tablica je pisana kao „ključ profila → drugi oblici", pa je vrijedila samo u
    jednom smjeru: profil koji dio zove „izjava o autorstvu" nije nalazio rad
    naslovljen „Izjava o akademskoj čestitosti", iako je upravo taj par ovdje
    zapisan. FPZG profil koristi baš te nazive i zato je uredno složen rad
    prijavljivao pet nepostojećih nedostataka (nalaz iz zakrpe-katedra.md).
    Sinonimija je simetrična relacija; tablica je sada samo njezin zapis.
    """
    prosireno = {}
    for kljuc, oblici in tablica.items():
        skupina = {kljuc, *oblici}
        for clan in skupina:
            prosireno.setdefault(clan, set()).update(skupina)
    return {k: sorted(v) for k, v in prosireno.items()}


SINONIMI = _simetricni_sinonimi(SINONIMI)

# Dio čije ime nosi oznaku neobaveznosti („(ako postoje)", „(po potrebi)") nije
# kršenje kad ga nema. Provjera je tu oznaku odbacivala zajedno sa zagradom i
# tražila dio kao obavezan, pa je profil govorio jedno, a alat drugo.
_NEOBAVEZNO_RE = re.compile(r"\(\s*(?:ako\b|po\s+potrebi|opcion|nije\s+obvez)[^)]*\)",
                            re.IGNORECASE)


def je_neobavezan_dio(naziv):
    return bool(_NEOBAVEZNO_RE.search(str(naziv or "")))


def _korijen_pogadja(korijen, kandidat):
    """Riječ kandidata počinje korijenom, uz najviše 3 znaka padežnog nastavka."""
    if " " in korijen:
        return korijen in kandidat
    for rijec in kandidat.split():
        if rijec.startswith(korijen) and len(rijec) - len(korijen) <= 3:
            return True
    return False


# v1.1-fix (Q1c dorada): pogodak za obavezni dio koji je POPIS (popis izvora,
# popis tablica…). Prije je bilo dovoljno da se pojam ili korijen pojavi bilo
# gdje u naslovu, pa je sadržajno poglavlje „2. IZVORI FINANCIRANJA PODUZEĆA"
# zadovoljavalo „popis izvora" i nedostatak popisa literature ostajao je skriven.
_POPISNI_UVOD_RE = re.compile(r"^(?:popis|kazalo|spisak)\b\s*")
_POPISNI_PRIDJEV_RE = re.compile(rf"^{_POPISNI_PRIDJEV}$")


def _popisne_stavke(kandidat):
    """Stavke popisnog naslova („popis korištenih izvora" → ['izvora']).

    None ako naslov uopće nije popisni (ne počinje s „popis"/„kazalo"/„spisak").
    """
    if not _POPISNI_UVOD_RE.match(kandidat):
        return None
    ostatak = _POPISNI_UVOD_RE.sub("", kandidat)
    return [r for r in ostatak.split()
            if r not in ("i", "te") and not _POPISNI_PRIDJEV_RE.match(r)]


def _pojam_pogadja(oblici, korijeni, rijec):
    """Imenuje li pojedina riječ naslova traženi dio (oblik ili korijen)."""
    return any(rijec == o or o.endswith(" " + rijec) for o in oblici) \
        or any(_korijen_pogadja(kor, rijec) for kor in korijeni)


def _popisni_naslov_pogadja(oblici, korijeni, kandidat):
    """Zadovoljava li naslov obavezni POPIS — traži se cijeli naslov, ne dio."""
    if any(kandidat == o for o in oblici):
        return True
    stavke = _popisne_stavke(kandidat)
    if stavke is not None:
        # „POPIS TABLICA I SLIKA" — dovoljna je jedna stavka koja odgovara
        return any(_pojam_pogadja(oblici, korijeni, s) for s in stavke)
    # nije popisni naslov: mora ga CIJELOG imenovati pojam („IZVORI", „LITERATURA")
    rijeci = [r for r in kandidat.split() if r not in ("i", "te")]
    return bool(rijeci) and all(_pojam_pogadja(oblici, korijeni, r) for r in rijeci)


def _prozni_pogodak(oblici, korijeni, prozni):
    """Cjeloviti pogodak USIDREN na početku proznog retka (ne bilo gdje u retku)."""
    for k in prozni:
        if any(k == o or k.startswith(o + " ") for o in oblici):
            return True
        prva = k.split()[0] if k.split() else ""
        for kor in korijeni:
            if " " in kor:
                if k.startswith(kor):
                    return True
            elif prva.startswith(kor) and len(prva) - len(kor) <= 3:
                return True
    return False


def provjeri_obavezne(iz, profil, naslovi, kratki, h1_tekstovi, proza=None):
    """v1.1-fix (D5): obavezni dio se traži u NASLOVIMA (stil Heading ili
    strukturno naslovu sličan redak). Obična rečenica („Zaključak ove analize
    je jasan.") više ne prolazi kao poglavlje — najviše postane „možda"."""
    trazeni = ((profil.get("struktura") or {}).get("obavezni_dijelovi")) or []
    if not trazeni:
        return
    kandidati = [norm(t) for t in naslovi] + [norm(t) for t in kratki]
    kandidati = [k for k in kandidati if k]
    prozni = [k for k in (norm(t) for t in (proza or [])) if k]

    nadjeni, falni, mozda, rucno, neobavezni = [], [], [], [], []
    for dio in trazeni:
        n = norm(dio)
        if any(n.startswith(x) for x in NEPROVJERLJIVO):
            rucno.append(dio)
            continue
        # „tijelo teksta" je isti dio kao „razrada" — nema vlastiti naslov nego
        # se prepoznaje po poglavljima između uvoda i zaključka. FPZG profil ga
        # zove tako, pa je bez ovoga tražen kao doslovan naslov kojeg nijedan rad
        # nema.
        if n.startswith("razrada") or n.startswith("tijelo"):
            # razrada = ima li poglavlja između uvoda i zaključka
            sredina = [t for t in h1_tekstovi
                       if "uvod" not in norm(t) and "zakljuc" not in norm(t)
                       and not H.NASLOV_LIT.match(t)]
            (nadjeni if sredina else falni).append(dio)
            continue
        oblici = SINONIMI.get(n, []) + [n, re.sub(r"\s*\(.*", "", dio).strip().lower()]
        korijeni = list(SINONIMI_KORIJENI.get(n, []))
        # Ime dijela često nabraja stavke jedne te iste sekcije („sažetak i ključne
        # riječi", „popis slika i grafikona"). Naslov u radu nosi samo jednu od njih
        # („SAŽETAK"), pa je usporedba cijelog imena padala na urednom radu. Svaki
        # član nabrajanja je zato vlastiti oblik, sa svojim korijenima.
        for dio_imena in re.split(r"\s+(?:i|te|and)\s+", re.sub(r"\s*\(.*", "", n)):
            dio_imena = dio_imena.strip()
            if len(dio_imena) > 2:
                oblici.append(dio_imena)
                korijeni.extend(SINONIMI_KORIJENI.get(dio_imena, []))
        oblici = [norm(o) for o in oblici if norm(o)]
        # v1.1-fix (Q1c dorada): za dio koji je POPIS traži se cijeli naslov;
        # za sadržajne dijelove („uvod", „zaključak") ostaje šira usporedba jer
        # „ZAKLJUČNA RAZMATRANJA" jest zaključak.
        if _POPISNI_UVOD_RE.match(n):
            pogodak = any(_popisni_naslov_pogadja(oblici, korijeni, k)
                          for k in kandidati)
        else:
            pogodak = any(any(o in k for k in kandidati) for o in oblici) \
                or any(any(_korijen_pogadja(kor, k) for k in kandidati)
                       for kor in korijeni)
        if pogodak:
            nadjeni.append(dio)
        elif je_neobavezan_dio(dio):
            neobavezni.append(dio)
        elif _prozni_pogodak(oblici, korijeni, prozni):
            mozda.append(dio)
        else:
            falni.append(dio)

    detalji = []
    if neobavezni:
        detalji.append("nije nađeno, ali profil to i ne traži obavezno: "
                       + "; ".join(neobavezni))
    if falni:
        detalji.append("nema naslova za: " + "; ".join(falni))
    if mozda:
        detalji.append("nema vlastiti naslov, pojam se pojavljuje samo u prozi "
                       "(možda, provjeri ručno): " + "; ".join(mozda)
                       + ". Što napraviti: dodaj poglavlje s naslovom u stilu "
                         "Heading 1 — obična rečenica nije dio strukture rada.")
    if rucno:
        detalji.append("ne može se utvrditi iz teksta (nema vlastiti naslov), "
                       "provjeri ručno: " + "; ".join(rucno))
    detalji.append("usporedba je neosjetljiva na velika slova i dijakritiku; "
                   "traže se naslovi (stil Heading ili kratak samostalan redak "
                   "bez završne točke, podebljan/centriran/verzalom)")
    iz.dodaj("obavezni dijelovi", f"{len(trazeni)}",
             f"{len(nadjeni)} nađeno, {len(falni)} fali"
             + (f", {len(mozda)} možda" if mozda else "")
             + (f", {len(neobavezni)} neobavezno" if neobavezni else "")
             + (f", {len(rucno)} za ručnu provjeru" if rucno else ""),
             LOSE if (falni or mozda) else (UPOZ if rucno else OK), detalji,
                 rule_id="struktura.obavezni_dijelovi")


def je_naslovu_slican(p, tekst):
    """Odlomak koji nije u stilu Heading, ali se strukturno ponaša kao naslov:
    kratak, samostalan, bez završne interpunkcije, podebljan/centriran/verzalom."""
    t = (tekst or "").strip()
    if not t or len(t) >= 80:
        return False
    if t[-1] in ".!?,;:":
        return False
    if _je_toc_redak(t):
        return False          # redak sadržaja nije naslov poglavlja
    if _outline_razina(p):
        return True
    try:
        centriran = p.paragraph_format.alignment is not None \
            and int(p.paragraph_format.alignment) == 1
    except (TypeError, ValueError, AttributeError):
        centriran = False
    verzal = t == t.upper() and any(c.isalpha() for c in t)
    runovi = [r for r in p.runs if (r.text or "").strip()]
    podebljan = bool(runovi) and all(r.bold for r in runovi)
    return bool(centriran or verzal or podebljan)


def provjeri_odlomke(iz, profil, odlomci, napomena=None):
    min_rec = ((profil.get("format") or {}).get("odlomak") or {}).get("min_recenica")
    if not min_rec:
        return
    if not odlomci:
        iz.dodaj("rečenica po odlomku", f"najmanje {min_rec}",
                 "nije pronađen prozni tekst", UPOZ,
                 ["Dokument nema odlomaka koje Katedra prepoznaje kao prozu.",
                  "Provjeri koriste li se stilovi Heading za naslove i ima li rad tijelo."],
                 rule_id="format.odlomak.min_recenica")
        return
    # Redak koji završava dvotočkom NAJAVLJUJE ono što slijedi (popis, citat,
    # tablicu) i po naravi je jedna rečenica. Mjeriti ga kao odlomak znači
    # kažnjavati uobičajen akademski oblik — vidjelo se čim je Katedra počela
    # sastavljati rad iz rukopisa, na svakom popisu.
    mjereni = [o for o in odlomci if not o.rstrip().endswith(":")]
    najave = len(odlomci) - len(mjereni)
    prekratki = []
    for i, o in enumerate(mjereni, 1):
        n = len(H.recenice(o))
        if n < min_rec:
            prekratki.append((i, n, o))
    odlomci = mjereni
    udio = (len(prekratki) / len(odlomci) * 100) if odlomci else 0
    detalji = [f"{i}. odlomak ({n} reč.): {o[:70]}…" for i, n, o in prekratki[:6]]
    if len(prekratki) > 6:
        detalji.append(f"… i još {len(prekratki) - 6}")
    if najave:
        detalji.append(f"{najave} najavna retka (završavaju dvotočkom) nisu mjerena "
                       "— oni uvode popis, citat ili prikaz")
    if napomena:
        detalji.append(napomena)
    iz.dodaj("rečenica po odlomku", f"najmanje {min_rec}",
             f"{len(prekratki)} od {len(odlomci)} ispod minimuma ({udio:.0f} %)",
             OK if not prekratki else (UPOZ if udio < 10 else LOSE), detalji,
                 rule_id="format.odlomak.min_recenica")


# v1.1-fix (D7): puni broj natpisa (i „2.1"), razdjelnik smije biti točka,
# dvotočje ili crtica, a „Shema" je sad i u VRSTE_PRIKAZA (prije je bila mrtva
# grana regexa).
VRSTE_PRIKAZA = ("Tablica", "Slika", "Grafikon", "Shema")
# v1.1-fix (drugi krug, PREKOREKCIJA): usporedba je bila osjetljiva na velika
# slova, pa verzalni natpis („TABLICA 1. Kretanje noćenja") nije bio natpis.
# Sam po sebi to je bio propušten nalaz, ali u sprezi s eskalacijom „dokument
# ima w:tbl a nijedan natpis → ❌" davao je LAŽNI ❌ posve usklađenom radu:
# verzalni natpisi su uobičajena kućna norma i nisu pogreška. Veličina slova je
# stvar oblikovanja, a ne strukture, pa se usporedba radi bez nje; kanonski
# oblik vrste vraća se natrag da KORIJEN_PRIKAZA i poruke ostanu iste.
NATPIS_RE = re.compile(
    r"(?i)^\s*(tablica|slika|grafikon|shema)\s*(\d+(?:\.\d+)*)\s*[.:–—-]\s*(.*)$")
KANON_VRSTE = {v.lower(): v for v in VRSTE_PRIKAZA}
# Korijeni za poziv u tekstu; „slika" u lokativu glasi „slici" (k → c).
KORIJEN_PRIKAZA = {"Tablica": r"tablic", "Slika": r"sli(?:k\w*|ci)",
                   "Grafikon": r"grafikon", "Shema": r"shem"}


def natpis(t):
    """(vrsta, broj) ako je odlomak NATPIS prikaza, inače None.

    „Tablica 1. Kretanje noćenja" je natpis; „Tablica 1. pokazuje da…" je
    rečenica koja se na tablicu poziva. Razlikuje ih prvo slovo iza broja.

    Vrsta prikaza prepoznaje se bez obzira na veličinu slova („TABLICA 1. …"),
    a vraća se u kanonskom obliku („Tablica").
    """
    m = NATPIS_RE.match(t or "")
    if not m:
        return None
    vrsta = KANON_VRSTE.get(m.group(1).lower())
    if not vrsta:
        return None
    ostatak = m.group(3).strip()
    if ostatak and not (ostatak[0].isupper() or ostatak[0].isdigit()
                        or ostatak[0] in "„\"'"):
        return None                      # poziv u tekstu, ne natpis
    return vrsta, m.group(2)


def provjeri_prikaze(iz, profil, doc):
    pravila = ((profil.get("struktura") or {}).get("prikazi")) or {}
    if not pravila:
        return
    redoslijed = list(blokovi(doc))
    natpisi = []      # (index_u_redoslijedu, vrsta, broj, tekst)
    spomeni = []      # tekst u kojem se traži poziv na prikaz (bez natpisa i izvora)
    u_popisu_prikaza = False
    for i, (vrsta, blok) in enumerate(redoslijed):
        if vrsta != "p":
            continue
        t = tekst_bloka(blok).strip()
        if not t:
            continue
        if je_popis_prikaza_naslov(t):
            u_popisu_prikaza = True
            continue
        # v1.1-fix (D7 dorada): zasun se otvara na SVAKOM naslovu, ne samo na
        # stilu Heading. Rad s ručno oblikovanim naslovima (podebljano/centrirano)
        # inače nikad ne izlazi iz „POPISA TABLICA" pa mu se progutaju svi natpisi
        # i dobije lažno kršenje „nijedan prikaz nema natpis".
        # v1.1-fix (D7b, treći krug): dva otpuštanja, oba STRUKTURNA.
        #
        # Drugi krug je otpuštao samo na stilu Heading. To je popravilo predloške
        # koji unose u popisu grupiraju pod podnaslovima „TABLICE" / „SLIKE"
        # (podebljan verzalni redak nije naslov poglavlja), ali je vratilo izvorni
        # kvar radovima pisanima BEZ stilova: takav rad iz „POPISA TABLICA" nikad
        # ne izađe, progutaju mu se svi natpisi i dobije lažno „nijedan prikaz nema
        # natpis". Leksički popis riječi („TABLICE", „SLIKE", …) bio bi četvrti krug
        # istog kvara, pa se otpušta po OBLIKU odlomka:
        #
        #   1. stvarni naslov (Heading stil / outline razina) — popis je gotov;
        #   2. obična proza — redak koji nije ni natpis, ni naslovu sličan, ni
        #      redak sadržaja. Popis prikaza sadrži samo unose i eventualne
        #      podnaslove; prva prozna rečenica znači da smo u tijelu rada.
        #
        # Podnaslovi unutar popisa i dalje ne otpuštaju zasun jer nisu proza.
        prozni_redak = (not natpis(t)
                        and not je_naslovu_slican(blok, t)
                        and not _je_toc_redak(t))
        if (not natpis(t) and razina_naslova(blok)) or prozni_redak:
            u_popisu_prikaza = False
        if u_popisu_prikaza:
            continue
        m = natpis(t)
        if m:
            natpisi.append((i, m[0], m[1], t))
        elif not H.IZVOR.match(t) and not razina_naslova(blok):
            spomeni.append(t)
    tijelo_tekst = " ".join(spomeni)

    if not natpisi:
        # v1.1-fix (D7): ako dokument ima stvarne tablice ili slike, a nijedan
        # natpis, to nije „nema se što provjeriti" nego kršenje.
        ima_tablica = any(vrsta == "tbl" for vrsta, _ in redoslijed)
        ima_sliku = any(True for _ in doc.element.body.iter(qn("w:drawing"))) \
            or any(True for _ in doc.element.body.iter(qn("w:pict")))
        if ima_tablica or ima_sliku:
            sto = " i ".join([x for x in ["tablice" if ima_tablica else "",
                                          "slike" if ima_sliku else ""] if x])
            iz.dodaj("prikazi", "natpis " + str(pravila.get("natpis", "—")),
                     f"dokument sadrži {sto}, ali nijedan prikaz nema natpis", LOSE,
                     ["natpisi nisu u obliku „Tablica 1. Naziv\" pa se ni izvor ni "
                      "poziv u tekstu ne mogu provjeriti",
                      "Što napraviti: iznad svakog prikaza dodaj natpis "
                      "„Tablica 1. Naziv\" (ili „Slika 1. Naziv\") i ispod njega "
                      "redak „Izvor: …\"."],
                 rule_id="prikazi.natpis")
        else:
            iz.dodaj("prikazi", "natpis " + str(pravila.get("natpis", "—")),
                     "nema nijedne tablice, slike ni grafikona s natpisom", UPOZ,
                     ["ako rad ima prikaze, natpisi nisu u obliku „Tablica 1. Naziv\" "
                      "pa se ne mogu provjeriti"],
                 rule_id="prikazi.natpis")
        return

    if pravila.get("mora_biti_spomenut_u_tekstu"):
        nespomenuti = []
        for _, vrsta, broj, tekst in natpisi:
            korijen = KORIJEN_PRIKAZA[vrsta]
            # puni identifikator + negativni pogled unaprijed: „2.1" ne smije
            # pogoditi „2.13", a „2" ne smije pogoditi „2.1"
            uzorak = re.compile(
                rf"(?i)\b{korijen}\w*\s*(?:br\.?\s*)?{re.escape(broj)}(?!\d)(?!\.\d)")
            spomen = [r for r in uzorak.finditer(tijelo_tekst)]
            if not spomen:
                nespomenuti.append(f"{vrsta} {broj}.")
        iz.dodaj("prikaz spomenut u tekstu", "svaki prikaz",
                 f"{len(natpisi) - len(nespomenuti)}/{len(natpisi)} spomenuto",
                 OK if not nespomenuti else LOSE,
                 ([f"bez spomena u tekstu: {', '.join(nespomenuti)}"] if nespomenuti else [])
                 + ["traži se spominjanje izvan samog natpisa, u bilo kojem padežu "
                    "(„u Tablici 1.\", „prema Grafikonu 2.\")"],
                 rule_id="prikazi.mora_biti_spomenut_u_tekstu")

    if pravila.get("izvor_ispod"):
        bez_izvora = []
        for i, vrsta, broj, _ in natpisi:
            nasao = False
            for _, blok in redoslijed[i + 1:i + 5]:
                if isinstance(blok, str):
                    continue
                tekst = getattr(blok, "text", None)
                if tekst is None:      # tablica
                    # preko hr_text-a, da se vidi i redak „Izvor:" koji je u
                    # ćeliji omotan u content control (`w:sdt`)
                    celije = " ".join(H.celije_tablica([blok]))
                    if H.IZVOR.search(celije or ""):
                        nasao = True
                        break
                    continue
                if H.IZVOR.match((tekst or "").strip()):
                    nasao = True
                    break
            if not nasao:
                bez_izvora.append(f"{vrsta} {broj}.")
        iz.dodaj("„Izvor:\" ispod prikaza", "svaki prikaz",
                 f"{len(natpisi) - len(bez_izvora)}/{len(natpisi)} ima izvor",
                 OK if not bez_izvora else LOSE,
                 ([f"bez retka „Izvor:\": {', '.join(bez_izvora)}"] if bez_izvora else [])
                 + ["traži se redak koji počinje s „Izvor:\" najviše 4 bloka ispod natpisa"],
                 rule_id="prikazi.izvor_ispod")


# B10: oblik citata koristi centralni citation dialect parser.
def provjeri_citate(iz, profil, tekst):
    cit = profil.get("citiranje") or {}
    stil = cit.get("stil") or "autor-godina"
    try:
        dialect = C.resolve_dialect(stil)
    except C.CitationDialectError as e:
        # v1.1-fix (Q15): nepoznat citatni stil je prije značio TIHI izostanak
        # retka o citatima — izvještaj je izgledao kao da citiranje nije ni
        # trebalo provjeriti. Nepoznato pravilo mora ostati vidljivo, inače
        # pogrešan (ili ručno prepravljen) profil daje uredan izvještaj.
        iz.dodaj("točka iza godine u citatu", str(stil),
                 "citatni stil nije prepoznat — citati nisu provjereni", UPOZ,
                 [str(e),
                  "Što napraviti: ispravi „citiranje.stil\" u profilu na jedan od "
                  "dopuštenih stilova pa ponovi provjeru."],
                 rule_id="citiranje.stil")
        return
    # Točka iza godine pripada samo author-year obitelji. IEEE i legal-footnote
    # imaju vlastitu sintaksu i ovdje ne smiju dobiti lažni author-year nalaz.
    if dialect != "author-year" or "tocka_iza_godine" not in cit:
        return
    trazi_tocku = bool(cit["tocka_iza_godine"])
    refs = C.parse_citations(tekst, stil)
    s_tockom = [r.raw[:60] for r in refs if r.year_has_dot is True]
    bez_tocke = [r.raw[:60] for r in refs if r.year_has_dot is False]

    ukupno = len(s_tockom) + len(bez_tocke)
    krivi = bez_tocke if trazi_tocku else s_tockom
    primjer = cit.get("u_tekstu") or ("(Prezime, 2011.)" if trazi_tocku else "(Prezime, 2011)")
    if not ukupno:
        iz.dodaj("točka iza godine u citatu", primjer, "nema prepoznatih author-year citata u tijelu",
                 UPOZ, [f"deklarirani stil je „{stil}”; provjeri postoje li citati u tijelu rada"],
                 rule_id="citiranje.tocka_iza_godine")
        return
    detalji = []
    if krivi:
        det = Counter(krivi)
        detalji.append(("bez točke: " if trazi_tocku else "s točkom: ")
                       + " · ".join(f"{k}{' ×' + str(v) if v > 1 else ''}"
                                    for k, v in det.most_common(5)))
        detalji.append("profil traži " + ('„2011.” (s točkom)' if trazi_tocku
                                          else '„2011” (bez točke)'))
    iz.dodaj("točka iza godine u citatu", primjer,
             f"{ukupno - len(krivi)}/{ukupno} citata u traženom obliku",
             OK if not krivi else LOSE, detalji,
                 rule_id="citiranje.tocka_iza_godine")


# --------------------------------------------------------------------- ispis

def sirina(s):
    return len(str(s))


def skrati(s, n):
    s = str(s)
    return s if sirina(s) <= n else s[:n - 1] + "…"


def ispis(iz, rad, tip):
    p = iz.profil
    print("=" * 78)
    print(f"PRAVILA FAKULTETA — {p.get('naziv', p.get('slug', '?'))}")
    print("=" * 78)
    print(f"rad: {rad}")
    print(f"tip rada: {tip} · profil: {os.path.relpath(iz.put_profila)}")
    print(f"izvor pravila: {iz.izvor}")
    if iz.advisory:
        print()
        print("ℹ️  PROFIL NIJE ADMITIRAN (faculty_scale_gate) — NALAZI SU SAVJETODAVNI.")
        print("    Pravila su izvedena, ne pročitana iz službenih strojno dostupnih uputa,")
        print("    pa nalazi NE blokiraju predaju. Za blokirajuće ponašanje: --strogo.")
        print("    Svaki nalaz prije predaje provjeri u službenim uputama fakulteta.")
    if iz.za_potvrdu:
        print()
        print("⚠️  STATUS PROFILA: NEPOTVRĐENO.")
        print("    Pravila nisu očitana iz službenih uputa nego prenesena ili izvedena.")
        print("    SVAKI nalaz ispod nosi oznaku „za potvrdu\" i prije predaje se mora")
        print("    provjeriti u službenim uputama fakulteta.")
    print()

    st = [30, 20, 24]
    print(f"{'pravilo':<{st[0]}} {'traženo':<{st[1]}} {'nađeno':<{st[2]}} stanje")
    print("-" * 78)
    for r in iz.redci:
        print(f"{skrati(r['pravilo'], st[0]):<{st[0]}} "
              f"{skrati(r['trazeno'], st[1]):<{st[1]}} "
              f"{skrati(r['nadjeno'], st[2]):<{st[2]}} {r['stanje']}")
    print("-" * 78)

    nalazi = iz.nalazi
    if nalazi:
        print("\nNALAZI\n")
        for r in nalazi:
            oznaka = " [za potvrdu]" if r["za_potvrdu"] else ""
            print(f"{r['stanje']} {r['pravilo']}{oznaka}")
            print(f"    traženo: {r['trazeno']}   ·   nađeno: {r['nadjeno']}")
            for d in r["detalji"]:
                print(f"    → {d}")
            print(f"    izvor pravila: {r['izvor']}")
            print()

    lose = sum(1 for r in iz.redci if r["stanje"] == LOSE)
    upoz = sum(1 for r in iz.redci if r["stanje"] == UPOZ)
    dobro = sum(1 for r in iz.redci if r["stanje"] == OK)
    print("SAŽETAK")
    print(f"  {OK} u skladu: {dobro}   {UPOZ} za provjeru: {upoz}   {LOSE} kršenja: {lose}")
    if iz.za_potvrdu:
        print("  Profil je NEPOTVRĐEN — sve gore vrijedi „za potvrdu\".")
    if not nalazi:
        print("  Nema nalaza: dokument odgovara profilu u svemu što se iz .docx-a može očitati.")
    elif not lose:
        print("  Nema kršenja (❌). Izlazni kod je 0 — stavke za provjeru (⚠️) "
              "traže oko, ne nužno ispravak, i ne blokiraju (v1.9, nalaz 4).")
    if any(r["pravilo"].startswith("broj stranica") for r in iz.redci):
        print("  Broj stranica nije mjeren nego procijenjen — .docx ga ne sadrži.")
    print()


# ---------------------------------------------------------------------- main

def main():
    ap = argparse.ArgumentParser(
        description="Usklađenost .docx-a s profilom fakulteta.")
    ap.add_argument("rad", help="putanja do .docx-a")
    ap.add_argument("--fakultet", help="slug ili alias iz references/fakulteti/")
    ap.add_argument("--profil", help="izravna putanja do <slug>.json")
    ap.add_argument("--tip", default=None,
                    help="seminarski | zavrsni | diplomski (default: jedini iz profila)")
    ap.add_argument("--strogo", action="store_true",
                    help="nalazi blokiraju i kad profil nije admitiran "
                         "(vraća izlazni kod 1)")
    ap.add_argument("--json", dest="json_out", metavar="PUT",
                    help="zapiši nalaze i za user_profile.py / .katedra/")
    a = ap.parse_args()

    if not os.path.isfile(a.rad):
        greska(f"rad ne postoji: {a.rad}")
    if not a.rad.lower().endswith(".docx"):
        greska("očekuje se .docx (za .md/.txt nema formatiranja za provjeru)")

    # Q19 (drugi krug): strop se mora provjeriti PRIJE `Document(a.rad)` —
    # parsiranje je upravo ono što troši memoriju, pa strop koji se primjenjuje
    # tek u provjeri fonta ne štiti ništa.
    prevelik = H.prevelik_dio_docxa(a.rad)
    if prevelik:
        greska(f"{os.path.basename(a.rad)}: {prevelik[0]} raspakiran prelazi "
               f"{H.MAX_XML_BAJTOVA // (1024 * 1024)} MB — to nije rad nego "
               f"oštećen ili namjerno napuhan zip.\n"
               f"   Što napraviti: uzmi datoteku koju si stvarno predao/la, ili "
               f"je otvori u Wordu i spremi ponovno kao .docx.")

    profil, put_profila = nadi_profil(a.fakultet, a.profil, a.tip)

    try:
        from docx import Document
    except ImportError:
        greska("treba python-docx:  pip install python-docx --break-system-packages")
    try:
        doc = Document(a.rad)
    except Exception as e:
        greska(f"{a.rad} se ne može otvoriti kao .docx: {e}")

    try:
        tip = resolve_work_type(a.tip, profil)
    except ProfileRuleError as e:
        greska(str(e))

    # struktura dokumenta
    redoslijed = list(blokovi(doc))
    h1, naslovi, kratki, proza = [], [], [], []
    for i, (vrsta, blok) in enumerate(redoslijed):
        if vrsta != "p":
            continue
        t = tekst_bloka(blok).strip()
        if not t:
            continue
        if je_unos_sadrzaja(blok, t):
            # Popis sadržaja NABRAJA naslove — on ih ne stvara. Bez ovoga rad
            # bez zaključka i literature prolazi kao uredan samo zato što ih
            # njegov (zastarjeli ili automatski) sadržaj spominje.
            continue
        r = razina_naslova(blok)
        if r:
            naslovi.append(t)
        if r == 1:
            h1.append((i, t, ima_prijelom_prije(blok, redoslijed[:i])))
        elif not r and len(t) < 80:
            # v1.1-fix (D5): kratak redak je kandidat za naslov samo ako se
            # strukturno ponaša kao naslov; ostalo je proza.
            (kratki if je_naslovu_slican(blok, t) else proza).append(t)

    # Geometrija proze: natuknice i blok-citati NISU odlomci (v. hr_text).
    odlomci, _ = H.ucitaj(a.rad, bez_nabrajanja=True)
    napomena_tijela = None
    if not odlomci:
        # hr_text tijelo počinje na „UVOD" / „1. …"; bez toga uzmi cijeli dokument
        odlomci, _ = H.ucitaj(a.rad, samo_tijelo=False, bez_nabrajanja=True)
        if odlomci:
            napomena_tijela = ("tijelo rada nije prepoznato po naslovu „UVOD\" / „1. …\" "
                               "pa su mjereni svi odlomci, uključujući predtekst")
    tijelo = " ".join(odlomci)
    tijelo_tekstovi, opis_opsega = tekst_glavnog_dijela(redoslijed)
    broj_rijeci = len(H.rijeci(" ".join(tijelo_tekstovi)))

    iz = Izvjestaj(profil, put_profila)
    provjeri_font(iz, doc, a.rad, profil)
    provjeri_prored_i_poravnanje(iz, doc, profil)
    provjeri_margine(iz, doc, profil)
    provjeri_format_stranice(iz, doc, profil)
    provjeri_opseg(iz, profil, tip, broj_rijeci, opis_opsega)
    provjeri_poglavlja(iz, profil, tip, sadrzajna_poglavlja([t for _, t, _ in h1]))
    provjeri_prijelome(iz, profil, h1)
    provjeri_obavezne(iz, profil, naslovi, kratki, [t for _, t, _ in h1], proza)
    provjeri_odlomke(iz, profil, odlomci, napomena_tijela)
    provjeri_prikaze(iz, profil, doc)
    provjeri_citate(iz, profil, tijelo)
    iz.zakljuci_nepoznata()

    ispis(iz, os.path.basename(a.rad), tip)

    if a.json_out:
        os.makedirs(os.path.dirname(os.path.abspath(a.json_out)) or ".", exist_ok=True)
        with open(a.json_out, "w", encoding="utf-8") as f:
            json.dump({
                "alat": "check_rules",
                "rad": os.path.abspath(a.rad),
                "fakultet": profil.get("slug"),
                "tip": tip,
                "status_profila": profil.get("status"),
                "izvor_pravila": iz.izvor,
                "rijeci": broj_rijeci,
                "redci": iz.redci,
                "broj_nalaza": len(iz.nalazi),
                "admisija": profil.get("admisija") or "admitiran",
                "nalazi_su_advisory": bool(iz.advisory and not a.strogo),
                "broj_krsenja": sum(1 for r in iz.redci if r["stanje"] == LOSE),
            }, f, ensure_ascii=False, indent=1)
        print(f"[nalazi → {a.json_out}]")

    # v1.9 (nalaz 4): izlazni kod 1 samo zbog ❌ (potvrđeno kršenje); ⚠️ „za
    # potvrdu" / nepotvrđen profil ne blokira (pravilo 18 paketa).
    if not any(r["stanje"] == LOSE for r in iz.redci):
        return 0
    if iz.advisory and not a.strogo:
        print()
        print(f"ℹ️  nalaza: {len(iz.nalazi)} — SAVJETODAVNI, ne blokiraju predaju "
              "(profil nije admitiran).")
        print("   Izlazni kod je 0 namjerno. Za blokirajuće ponašanje dodaj --strogo.")
        return 0
    return 1


if __name__ == "__main__":
    sys.exit(main())
