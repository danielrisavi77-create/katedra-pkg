#!/usr/bin/env python3
"""Pravila HKS-FZS uputa za diplomski koja profil paketa ne može nositi — provjera nad .docx-om.

Sve što stoji ovdje ima lokator u „Posebne upute za pisanje diplomskog rada FZS HKS" (veljača 2026.).
Izlaz je savjetodavan: profil je `nepotvrdeno` (pročitan sažimačem), pa nijedan nalaz nije ❌ nego
⚠️ [za potvrdu] dok se PDF uputa ne pročita izravno (pravilo 18).

Od v1.9 (nalaz 9) opseg po dijelovima, obaveznost, redoslijed dijelova i podsekcija te razmak
odlomaka nosi PROFIL (`struktura.opseg.diplomski.dijelovi`, `format.odlomak.razmak_pt` u
references/fakulteti/hks-fzs.json) i provjerava ih opća skripta `scripts/provjeri_dijelove.py`.
OVDJE ostaje samo ono što shema ne nosi: font/veličina u ćelijama tablica, „Tablica N." s točkom u
tekstu, Vancouver interpunkcija (citat prije zareza, „i sur." nakon 6 autora), N u prvoj rečenici
Rezultata sažetka, broj ključnih riječi (MeSH), redni brojevi u Zaključku, upućivanje na tablice u
Raspravi, životopis u natuknicama, PAGE polje u podnožju. Provjere opsega/dijelova su zadržane radi
usporedbe (iste zaključke na rad.docx daje provjeri_dijelove.py), ali izvor istine je profil.

Provjere (lokator):
  opseg      Uvod ≥ 3000 riječi i ≤ 1/3 teksta (2.1, 3.10); Rasprava ≥ 1000 (2.1); Sažetak ≤ 1800 znakova (3.6)
  sažetak    podnaslovi Uvod/Cilj/Metode/Rezultati/Zaključak; prva rečenica Rezultata nosi N; ≥3 ključne riječi (3.6)
  summary    Title + Background/Aim/Methods/Results/Conclusion (3.7)
  dijelovi   TDK, Basic Documentation Card, Sažetak, Summary, Sadržaj prisutni (3.4–3.8)
  metode     redoslijed podsekcija: Etika → Ustroj → Mjesto i vrijeme → Jedinica analize → Postupci → Statistički (3.13)
  rasprava   naslovi „Prednosti i nedostatci istraživanja" i „Smjernice za buduća istraživanja" (3.15); bez upućivanja na tablice
  zaključak  rezultati označeni rednim brojevima (3.16)
  tablice    natpis „Tablica N." iznad; u tekstu „(Tablica N)"/„u Tablici N" BEZ točke (5.5); sadržaj tablica 11 pt
  citati     broj u zagradi PRIJE zareza/točke, ne poslije (5.7); ≤6 autora pa „i sur." u popisu (5.7)
  format     Normal: Book Antiqua 11, prored 1,5, razmak iza odlomka 6 pt (2.3); prijelom pred Heading 1 (2.7)
  numeracija brojevi stranica u podnožju, tijelo počinje od 1 na Uvodu (2.6)
  životopis  natuknice, ne rečenice (3.19)

Uporaba:  python3 provjeri_hks_fzs.py ./rad.docx [--json .katedra/hks_fzs.json]
"""
from __future__ import annotations

import argparse
import json
import re
import sys
import zipfile

from docx import Document

RE_HEAD1 = re.compile(r"^\s*(\d{1,2})\.\s+(.+)$")
RE_TAB_REF_TOCKA = re.compile(r"\b(?:u\s+)?Tablic[ai]\s+\d+\.(?=\s+[a-zčćžšđ]|\))")  # točka pa malo slovo/zagrada = nije kraj rečenice
RE_TAB_REF = re.compile(r"\bTablic[ai]\s+\d+")
RE_CIT_POSLIJE = re.compile(r"[.,;]\s?\(\d{1,3}(?:\s*[,–-]\s*\d{1,3})*\)")   # „.(12)" ili „, (3)"
RE_CIT = re.compile(r"\(\d{1,3}(?:\s*[,–-]\s*\d{1,3})*\)")


def _n(s):
    return s.strip().lower()


def blokovi(doc):
    """Odlomci tijela + oznaka poglavlja (Heading 1 broj) i podpoglavlja (Heading 2/3 tekst)."""
    out, h1, h2 = [], None, None
    for p in doc.paragraphs:
        st = p.style.name.lower()
        if st.startswith("heading 1"):
            m = RE_HEAD1.match(p.text)
            h1 = (int(m.group(1)), m.group(2).strip()) if m else (None, p.text.strip())
            h2 = None
        elif st.startswith("heading 2") or st.startswith("heading 3"):
            h2 = p.text.strip()
        out.append((h1, h2, st, p))
    return out


def rijeci(txt):
    return len(re.findall(r"\S+", txt))


def analiza(put):
    doc = Document(put)
    B = blokovi(doc)
    n = []  # nalazi: (razina, sekcija, poruka)
    def nalaz(sekcija, poruka, ok=False):
        n.append({"sekcija": sekcija, "ok": ok, "poruka": poruka})

    # ---- opseg po dijelovima
    def tekst_pogl(broj):
        return "\n".join(p.text for h1, h2, st, p in B if h1 and h1[0] == broj and not st.startswith("heading"))
    naslovi = {h1[0]: h1[1] for h1, _, _, _ in B if h1 and h1[0]}
    def broj_po_nazivu(kljuc):
        for k, v in naslovi.items():
            if kljuc in _n(v):
                return k
        return None
    b_uvod, b_rasp, b_zakl, b_lit = broj_po_nazivu("uvod"), broj_po_nazivu("rasprav"), broj_po_nazivu("zaključ"), broj_po_nazivu("literatur")
    uvod_r = rijeci(tekst_pogl(b_uvod)) if b_uvod else 0
    rasp_r = rijeci(tekst_pogl(b_rasp)) if b_rasp else 0
    tijelo_r = sum(rijeci(tekst_pogl(k)) for k in naslovi if b_lit is None or k < b_lit)
    nalaz("opseg", f"Uvod {uvod_r} riječi (traženo ≥ 3000; odj. 2.1)", ok=uvod_r >= 3000)
    if tijelo_r:
        udio = uvod_r / tijelo_r
        nalaz("opseg", f"Uvod je {udio:.0%} tijela (traženo ≤ 33 %; odj. 3.10)", ok=udio <= 1 / 3)
    nalaz("opseg", f"Rasprava {rasp_r} riječi (traženo ≥ 1000; odj. 2.1)", ok=rasp_r >= 1000)

    # ---- prednji dio: TDK, BDC, sažetak, summary, sadržaj (ne-heading odlomci prije 1. UVOD)
    prednji = [p.text.strip() for h1, _, _, p in B if h1 is None and p.text.strip()]
    pred_txt = "\n".join(prednji)
    for kljuc, naziv in (("temeljna dokumentacijska kartica", "TDK"), ("basic documentation card", "Basic Documentation Card"),
                         ("sažetak", "Sažetak"), ("summary", "Summary"), ("sadržaj", "Sadržaj")):
        nalaz("dijelovi", f"{naziv}: {'nađen' if kljuc in pred_txt.lower() else 'NIJE nađen u prednjem dijelu'} (odj. 3.4–3.8)",
              ok=kljuc in pred_txt.lower())

    # sažetak: blok između retka „SAŽETAK" i „SUMMARY"
    def izmedju(a, b):
        idx = [i for i, t in enumerate(prednji) if _n(t) == a]
        if not idx:
            return []
        i0 = idx[0] + 1
        j = next((i for i in range(i0, len(prednji)) if _n(prednji[i]) == b), len(prednji))
        return prednji[i0:j]
    saz = izmedju("sažetak", "summary")
    if saz:
        saz_txt = "\n".join(saz)
        znak = len(re.sub(r"\s+", " ", saz_txt))
        nalaz("sažetak", f"Sažetak {znak} znakova (traženo ≤ 1800; odj. 3.6)", ok=znak <= 1800)
        for pod in ("uvod", "cilj", "metode", "rezultati", "zaključak"):
            ima = any(_n(t).startswith(pod) for t in saz)
            nalaz("sažetak", f"podnaslov „{pod.capitalize()}\": {'da' if ima else 'NE'} (odj. 3.6)", ok=ima)
        kr = next((t for t in saz if _n(t).startswith("ključne riječi")), None)
        if kr:
            broj = len([x for x in re.split(r"[;,]", kr.split(":", 1)[-1]) if x.strip()])
            nalaz("sažetak", f"ključne riječi: {broj} (traženo ≥ 3, iz MeSH-a — MeSH provjeriti ručno; odj. 3.6)", ok=broj >= 3)
        else:
            nalaz("sažetak", "ključne riječi: redak „Ključne riječi:\" nije nađen (odj. 3.6)")
        rez_i = next((i for i, t in enumerate(saz) if _n(t).startswith("rezultati")), None)
        if rez_i is not None:
            blok = saz[rez_i].split(":", 1)[1].strip() if ":" in saz[rez_i] else (saz[rez_i + 1] if rez_i + 1 < len(saz) else "")
            prva = re.split(r"(?<=\.)\s", blok)[0]
            ima_n = bool(re.search(r"\d", prva))
            nalaz("sažetak", "prva rečenica Rezultata nosi N: " + ("da" if ima_n else "NE") + " — „" + prva[:90] + "…\" (odj. 3.6)", ok=ima_n)
    else:
        nalaz("sažetak", "blok SAŽETAK … SUMMARY nije nađen")
    summ = izmedju("summary", "sadržaj")
    if summ:
        for pod in ("title", "background", "aim", "methods", "results", "conclusion"):
            ima = any(_n(t).startswith(pod) for t in summ)
            nalaz("summary", f"„{pod.capitalize()}\": {'da' if ima else 'NE'} (odj. 3.7)", ok=ima)

    # ---- metode: redoslijed podsekcija
    b_met = broj_po_nazivu("metod")
    if b_met:
        pods = []
        for h1, h2, st, p in B:
            if h1 and h1[0] == b_met and st.startswith("heading 2"):
                pods.append(p.text.strip())
        prop = ["eti", "ustroj", "mjesto", "jedinica", "postup", "statisti"]
        poz = []
        for k in prop:
            i = next((i for i, t in enumerate(pods) if k in _n(t)), None)
            poz.append(i)
        redoslijed_ok = all(x is not None for x in poz) and poz == sorted(poz)
        cisti = [re.sub(r"^\d+(\.\d+)*\.?\s*", "", t) for t in pods]
        nalaz("metode", "redoslijed podsekcija: " + " → ".join(cisti)
              + "  |  propisano: Etika → Ustroj → Mjesto i vrijeme → Jedinica analize → Postupci → Statistički (odj. 3.13)", ok=redoslijed_ok)

    # ---- rasprava
    if b_rasp:
        pods = [p.text.strip() for h1, h2, st, p in B if h1 and h1[0] == b_rasp and st.startswith("heading 2")]
        ima_pn = any("prednosti" in _n(t) for t in pods)
        ima_sm = any("smjernice" in _n(t) for t in pods)
        nalaz("rasprava", f"odjeljak „Prednosti i nedostatci istraživanja\": {'da' if ima_pn else 'NE'}"
              + ("  ⚠ pravopis u Uputama je „nedostatci\"" if any("nedostaci" in _n(t) for t in pods) else ""), ok=ima_pn)
        nalaz("rasprava", f"odjeljak „Smjernice za buduća istraživanja\": {'da' if ima_sm else 'NE'} (odj. 3.15)", ok=ima_sm)
        rt = tekst_pogl(b_rasp)
        upu = len(RE_TAB_REF.findall(rt))
        nalaz("rasprava", f"upućivanje na tablice u Raspravi: {upu}× (Upute: ne upućuje se na tablice; odj. 3.15)", ok=upu == 0)
        prvi = next((p.text for h1, h2, st, p in B if h1 and h1[0] == b_rasp and not st.startswith("heading") and p.text.strip()), "")
        nalaz("rasprava", f"prvi odlomak sažima rezultate? provjeri ručno — „{prvi[:110]}…\" (odj. 3.15)", ok=True)

    # ---- zaključak: redni brojevi
    if b_zakl:
        zt = [p.text.strip() for h1, h2, st, p in B if h1 and h1[0] == b_zakl and not st.startswith("heading") and p.text.strip()]
        num = sum(1 for t in zt if re.match(r"^(\d+[.)]|[a-z][.)]|[IVX]+\.)\s", t))
        nalaz("zaključak", f"odlomci označeni rednim brojevima: {num}/{len(zt)} (Upute: rezultate označiti rednim brojevima; odj. 3.16)", ok=num >= 1)

    # ---- tablice
    svi = [p.text for _, _, st, p in B if not st.startswith("heading")]
    txt_all = "\n".join(t for t in svi if not re.match(r"^\s*Tablica\s+\d+\.\s", t))  # bez natpisa
    s_tockom = RE_TAB_REF_TOCKA.findall(txt_all)
    nalaz("tablice", f"upućivanje „Tablica N.\" s točkom u tekstu: {len(s_tockom)}× (Upute: „(Tablica 1)\" bez točke; odj. 5.5) npr. {s_tockom[:4]}", ok=len(s_tockom) == 0)
    natpisi = [t for t in svi if re.match(r"^\s*Tablica\s+\d+\.\s", t)]
    nalaz("tablice", f"natpisi „Tablica N.\" (s točkom): {len(natpisi)} (odj. 5.5)", ok=len(natpisi) > 0)
    # veličina fonta u tablicama
    vel = {}
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                for p in c.paragraphs:
                    for r in p.runs:
                        if r.text.strip():
                            v = r.font.size.pt if r.font.size else (p.style.font.size.pt if p.style.font.size else None)
                            vel[v] = vel.get(v, 0) + 1
    nalaz("tablice", f"veličina fonta u ćelijama tablica: {dict(sorted(vel.items(), key=lambda x: -x[1]))} (Upute: 11 kao u tekstu; odj. 5.5)",
          ok=all(k in (11, None) for k in vel))

    # ---- citati: prije interpunkcije; „i sur." nakon 6 autora
    poslije = RE_CIT_POSLIJE.findall(txt_all)
    nalaz("citati", f"citat POSLIJE zareza/točke („.(12)\"): {len(poslije)}× (Upute: prije; odj. 5.7) npr. {poslije[:4]}", ok=len(poslije) == 0)
    if b_lit:
        lit = [p.text.strip() for h1, h2, st, p in B if h1 and h1[0] == b_lit and re.match(r"^\d{1,3}\.\s", p.text)]
        losi = []
        for t in lit:
            autori = t.split(".", 1)[1].split(".")[0] if "." in t else ""
            # autori do prve točke iza inicijala: „Prezime AB, Prezime C, … i sur"
            aut = re.split(r",\s*", autori)
            if len(aut) > 7 or (len(aut) == 7 and "i sur" not in autori):
                losi.append(t[:60])
        nalaz("citati", f"popis: jedinica s više od 6 autora bez „i sur.\": {len(losi)} od {len(lit)} (odj. 5.7)" + (f" npr. {losi[:2]}" if losi else ""), ok=not losi)

    # ---- format: Normal stil, razmak iza odlomka, prijelomi
    normal = doc.styles["Normal"]
    f = normal.font
    pf = normal.paragraph_format
    nalaz("format", f"stil Normal: font {f.name}, {f.size.pt if f.size else '?'} pt, prored {pf.line_spacing}, razmak iza {pf.space_after.pt if pf.space_after else 0} pt (Upute: Book Antiqua 11, 1,5, 6 pt; odj. 2.3)",
          ok=(f.name == "Book Antiqua" and f.size and f.size.pt == 11 and pf.line_spacing in (1.5,) and pf.space_after and pf.space_after.pt == 6))
    paras = [p for _, _, _, p in B]
    h1s = [i for i, (_, _, st, _) in enumerate(B) if st.startswith("heading 1")]
    bez_prijeloma = []
    for i in h1s:
        p = paras[i]
        pb = p.paragraph_format.page_break_before
        xml = p._p.xml
        # prijelom kao <w:br w:type="page"/> u prethodnom (nepraznom ili praznom) odlomku, ili sectPr
        prev_break = i > 0 and ('w:type="page"' in paras[i - 1]._p.xml or "<w:sectPr" in paras[i - 1]._p.xml)
        if not pb and not prev_break and "pageBreakBefore" not in xml:
            bez_prijeloma.append(p.text.strip()[:40])
    nalaz("format", f"Heading 1 bez prijeloma stranice ispred: {len(bez_prijeloma)}/{len(h1s)} (Upute: svako poglavlje na novoj stranici; odj. 2.7) {bez_prijeloma[:5]}",
          ok=not bez_prijeloma)

    # ---- numeracija stranica: sectPr pgNumType / footer
    z = zipfile.ZipFile(put)
    docxml = z.read("word/document.xml").decode("utf8", "ignore")
    starts = re.findall(r'<w:pgNumType[^>]*w:start="(\d+)"', docxml)
    footers = [nm for nm in z.namelist() if nm.startswith("word/footer")]
    ima_page = any(b"PAGE" in z.read(nm) for nm in footers)
    nalaz("numeracija", f"sekcija: {docxml.count('<w:sectPr')}, pgNumType start: {starts or 'nema'}, PAGE polje u podnožju: {'da' if ima_page else 'NE'} (Upute: donji desni ugao, Uvod = 1; odj. 2.6) — poravnanje desno provjeriti u Wordu",
          ok=ima_page and "1" in starts)

    # ---- životopis
    b_ziv = broj_po_nazivu("životopis")
    if b_ziv:
        zt = [p.text.strip() for h1, h2, st, p in B if h1 and h1[0] == b_ziv and not st.startswith("heading") and p.text.strip()]
        duge = [t for t in zt if len(t.split()) > 25]
        nalaz("životopis", f"odlomaka: {len(zt)}, dužih od 25 riječi (nalik rečenicama, ne natuknicama): {len(duge)} (odj. 3.19)", ok=not duge)

    return {"datoteka": put, "profil": "hks-fzs (nepotvrdeno, sažimač)", "nalazi": n,
            "brojke": {"uvod_rijeci": uvod_r, "rasprava_rijeci": rasp_r, "tijelo_rijeci": tijelo_r}}


def ispisi(r):
    print(f"HKS-FZS UPUTE — {r['datoteka']}   [profil {r['profil']}]")
    print("=" * 64)
    sek = None
    for x in r["nalazi"]:
        if x["sekcija"] != sek:
            sek = x["sekcija"]; print(f"\n{sek.upper()}")
        print(("  ✅ " if x["ok"] else "  ⚠️ ") + x["poruka"])
    ok = sum(1 for x in r["nalazi"] if x["ok"]); print(f"\n{ok}/{len(r['nalazi'])} u skladu · ostalo [za potvrdu] — profil nije potvrđen iz pdftotext-a (pravilo 18: ⚠️, ne ❌).")


def main(argv=None):
    ap = argparse.ArgumentParser(description="HKS-FZS pravila koja profil paketa ne nosi.")
    ap.add_argument("rad"); ap.add_argument("--json", dest="kao_json")
    a = ap.parse_args(argv)
    r = analiza(a.rad); ispisi(r)
    if a.kao_json:
        json.dump(r, open(a.kao_json, "w", encoding="utf-8"), ensure_ascii=False, indent=1)
    return 0


if __name__ == "__main__":
    sys.exit(main())
