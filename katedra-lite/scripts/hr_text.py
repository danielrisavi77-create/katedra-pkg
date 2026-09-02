#!/usr/bin/env python3
"""
Zajednički sloj za hrvatski akademski tekst.

Učitavanje (.docx / .md / .txt), segmentacija na odlomke i rečenice koja
poštuje hrvatske redne brojeve i kratice, te hrvatska abecedna kolacija.

Ovo je jedina točka koja zna razlikovati kraj rečenice od točke iza rednog
broja — bez toga svako mjerenje ritma i duljine odlomka daje krive brojke.
"""
import argparse
import re
import sys
import unicodedata
import zipfile

# --------------------------------------------------------------- učitavanje

STRUKTURNI = ("#", "*", "---", "|", ">")

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_W_PREFIKS = "{" + W_NS + "}"

# ------------------------------------------------- strop na PROČITANIM bajtovima
# v1.1-fix (Q19, drugi krug): prvi popravak je gledao `ZipInfo.file_size`, tj.
# veličinu koju arhiva SAMA O SEBI TVRDI u zaglavlju. To polje upisuje onaj tko
# pravi datoteku, pa ono ne štiti ni od čega: dovoljno je najaviti 1 kB i
# poslati gigabajt. Uz to je strop stajao samo nad `word/theme/*` (nekoliko kB
# fontova), dok se `word/document.xml` — jedini dio koji uistinu može biti
# ogroman — parsirao bez ikakvog ograničenja i PRIJE nego se do stropa dođe.
#
# Zato se ovdje broje STVARNO pročitani bajtovi i čitanje se prekida čim prijeđu
# strop; najavljena veličina služi još samo kao jeftin predfiltar. Memorija je
# time omeđena stropom, a ne maštom pošiljatelja.
#
# 8 MB je ~20× više od `word/document.xml` stvarnog EFZG završnog rada
# (399 954 B), pa nijedan pravi rukopis ne može pasti na tom pragu. lxml stablo
# naraste ~25× u odnosu na XML, pa je najgori dopušteni slučaj reda 200 MB RSS —
# neugodno, ali preživljivo; bez stropa je isti ulaz gigabajtni (mjereno:
# arhiva od 54 kB koja najavljuje 12 MB dala je 303 MB RSS, a ~450 kB koja
# najavljuje 100 MB dala bi ~2,5 GB).
MAX_XML_BAJTOVA = 8 * 1024 * 1024
# Zbroj svih XML dijelova: sto dijelova po 7 MB prolazi kroz strop po dijelu, a
# python-docx ih pri `Document()` parsira sve.
MAX_UKUPNO_XML_BAJTOVA = 64 * 1024 * 1024
# PREKOREKCIJA (3. krug recenzije): prva verzija je štitila zatvoren popis od
# četiri dijela, pa je bomba u `word/numbering.xml` prolazila kroz zaštitu i
# reproducirala točno isti kvar. Popis dijelova nije obrana — python-docx pri
# `Document()` parsira SVE XML dijelove paketa, pa i zaštita mora biti nad svima.
# Ovo ime ostaje kao zadana vrijednost za pozivatelje koji ga izričito traže.
GLAVNI_DIJELOVI_DOCXA = ("word/document.xml", "word/styles.xml",
                         "word/footnotes.xml", "word/endnotes.xml")


def xml_dijelovi_docxa(z):
    """Svi dijelovi arhive koje `Document()` parsira (XML i .rels)."""
    return [ime for ime in z.namelist()
            if ime.lower().endswith((".xml", ".rels"))]


def procitaj_dio_zipa(z, ime, strop):
    """Sadržaj dijela arhive ili None ako STVARNI sadržaj prijeđe `strop`.

    Zajednički pomoćnik: koriste ga `check_rules` (theme, document.xml) i
    `evidence`/`diff` sloj. Ne vjeruje zaglavlju zipa — dekompresira u komadima
    i staje čim zbroj prijeđe strop, pa potrošnja memorije nikad ne prelazi
    strop bez obzira na to što arhiva o sebi tvrdi.
    """
    try:
        if z.getinfo(ime).file_size > strop:
            return None          # jeftin predfiltar; nije jamstvo
    except KeyError:
        return None
    komadi, ukupno = [], 0
    try:
        with z.open(ime) as f:
            while True:
                komad = f.read(256 * 1024)
                if not komad:
                    break
                ukupno += len(komad)
                if ukupno > strop:
                    return None
                komadi.append(komad)
    except (zipfile.BadZipFile, OSError, EOFError):
        return None
    return b"".join(komadi)


def prevelik_dio_docxa(put, strop=MAX_XML_BAJTOVA, dijelovi=None,
                       ukupni_strop=MAX_UKUPNO_XML_BAJTOVA):
    """(ime_dijela, pročitano_bajtova) prvog dijela iznad stropa, ili None.

    Poziva se PRIJE `Document(put)`, jer poslije je kasno — parsiranje je
    upravo ono što troši memoriju. `dijelovi=None` znači SVI XML dijelovi
    paketa: zatvoren popis nije obrana jer bomba onda samo preseli u dio koji
    nije na popisu (`word/numbering.xml`), a `Document()` parsira sve.
    """
    try:
        with zipfile.ZipFile(put) as z:
            imena = list(dijelovi) if dijelovi is not None else xml_dijelovi_docxa(z)
            postojeci = set(z.namelist())
            svi = 0
            for ime in imena:
                if ime not in postojeci:
                    continue
                ukupno = 0
                try:
                    with z.open(ime) as f:
                        while True:
                            komad = f.read(256 * 1024)
                            if not komad:
                                break
                            ukupno += len(komad)
                            svi += len(komad)
                            if ukupno > strop:
                                return ime, ukupno
                            if svi > ukupni_strop:
                                return ime, svi
                except (zipfile.BadZipFile, OSError, EOFError):
                    # Oštećen zapis (npr. kriva CRC suma) NIJE „prevelik" —
                    # inače bi neispravna arhiva dobila krivu poruku umjesto
                    # uredne „ne može se otvoriti kao .docx". Do granice ionako
                    # nismo došli, jer se ona provjerava u svakom koraku.
                    return None
    except (zipfile.BadZipFile, OSError):
        return None
    return None


def ucitaj(putanja, samo_tijelo=True, ukljuci_tablice=False, bez_nabrajanja=False):
    """Vrati (odlomci, markeri). Markeri su naslovi i natpisi prikaza.

    `ukljuci_tablice` određuje ulaze li tekstovi tabličnih ćelija u `odlomci`.
    Default je False jer većina potrošača mjeri GEOMETRIJU PROZE — broj rečenica
    po odlomku (`check_rules.provjeri_odlomke`, `check_paragraphs`), ritam
    (`check_ai_style`) — a ćelija poput „Godina" ili „41 %" nije odlomak i
    obarala bi ta mjerenja lažnim ❌.

    `bez_nabrajanja` izbacuje stavke popisa i blok-citate iz `odlomci`. Isti
    razlog: natuknica je jedna rečenica, pa pravilo „najmanje dvije rečenice po
    odlomku" inače proglašava kršenjem svaki rad koji ima popis. Potrošači koji
    mjere sadržaj ih moraju zadržati — natuknica JEST tekst rada.

    Potrošači koji provjeravaju OČUVANJE SADRŽAJA moraju ih dobiti i zato
    prosljeđuju True: `verify_rewrite` (obrisan citat ili preokrenuta brojka u
    tablici inače prolazi kao „sadržaj očuvan" — nalaz D2 audita) i
    `originality_check` (doslovno prepisan tekst u tablici je jednako prepisan).
    """
    p = str(putanja)
    if p.endswith(".docx"):
        return _iz_docx(p, samo_tijelo, ukljuci_tablice, bez_nabrajanja)
    tekst = open(p, encoding="utf8").read()
    return _iz_markdowna(tekst)


def _iz_markdowna(tekst):
    odlomci, markeri = [], []
    for red in tekst.split("\n"):
        t = red.strip()
        if not t:
            continue
        if t.startswith(STRUKTURNI):
            markeri.append(t)
        else:
            odlomci.append(t)
    return odlomci, markeri


# v1.9 (nalaz 6/7): uz osnovne oblike i „Popis citirane literature", „Citirana
# literatura", „Korištena literatura", „Literatura i izvori", „Popis referenci",
# „Bibliografija", „Reference(s)". Bez toga je check_argument HKS-FZS popis od
# 75 stavki brojao kao poglavlje tijela s 2115 riječi i „1 citatom", a
# verify_sources javljao da popisa nema. provjeri_literaturu.py nosi svoje
# lokalno proširenje istog smisla (NASLOV_LIT_PROSIREN) i ostaje netaknut.
NASLOV_LIT = re.compile(
    r"(?i)^\s*(?:\d+\.?\s*)?("
    r"(?:POPIS\s+)?(?:CITIRAN[AE]\s+|KORI[SŠ]TEN[AE]\s+)?"
    r"(?:LITERATURA|LITERATURE)(?:\s+I\s+IZVOR[AI])?"
    r"|POPIS\s+REFERENC[AEI]|REFERENCES?|REFERENCIJE|BIBLIOGRAFIJA"
    r"|POPIS\s+IZVORA|IZVORI(?:\s+I\s+LITERATURA)?"
    r")\s*$"
)
# v1.1-fix (drugi krug): i ovdje bez osjetljivosti na velika slova — verzalni
# natpis („TABLICA 1. Kretanje noćenja") inače nije marker nego prolazi kao
# proza, pa se broji u geometriju odlomaka i obara mjerenja. Isti propis kao
# check_rules.NATPIS_RE.
# Razdjelnik je i točka i dvotočka: FPZG Upute na jednom mjestu daju „Tablica 1:
# Naziv", pa je s uzorkom koji traži samo točku obranjen FPZG rad prijavljivao da
# „nema nijedan prikaz" (nalaz 4 iz fpzg-diplomski/references/zakrpe-katedra.md).
# check_rules.NATPIS_RE je razdjelnike već proširio; ovdje su zaostali, pa su dvije
# datoteke istog paketa različito čitale isti natpis.
NATPIS = re.compile(r"(?i)^\s*(Tablica|Grafikon|Slika|Shema|Prilog)\s+\d+(?:\.\d+)*\s*[.:]")
IZVOR = re.compile(r"^\s*(Izvor|Napomena)\s*:")


# ---------------------------------------------- početak glavnog teksta (tijela)
# v1.0.1-fix (D2 / body-start): isti propis vrijedi i u check_rules.py.
# Prije je redak sadržaja „1. UVOD.........3" otvarao tijelo usred SADRŽAJA
# (pa je cijeli sadržaj ulazio kao lažni odlomci), a „I. UVOD" i
# „UVOD U PROBLEMATIKU" se uopće nisu prepoznavali.
TOC_STIL = re.compile(r"(?i)^(?:toc|sadr[žz]aj)")
TOC_REDAK = re.compile(r"\.{3,}|\t\s*\d+\s*$")
SADRZAJ_NASLOV = re.compile(r"(?i)^\s*(?:sadr[žz]aj|kazalo|popis\s+\S+.*)$")
POGLAVLJE_BROJ = re.compile(r"(?i)^\s*(?:\d+(?:\.\d+)*\.?|[ivx]+\.)\s+\S")
UVOD_NASLOV = re.compile(r"(?i)^\s*(?:\d+(?:\.\d+)*\.?\s*)?uvod[^.!?]*$")


def je_redak_sadrzaja(tekst, stil=""):
    """Redak popisa sadržaja: TOC stil, dot-leader ili tabulator + broj stranice."""
    return bool(TOC_STIL.match((stil or "").strip())
                or TOC_REDAK.search(str(tekst or "")))


def je_pocetak_tijela(tekst, stil="", naslov=False, u_sadrzaju=False):
    """Otvara li ovaj odlomak glavni tekst rada.

    `naslov` je True za odlomke u Heading stilu ili s postavljenom outline
    razinom — samo se za njih prihvaćaju rimski brojevi i prošireni naslovi
    uvoda; obični odlomci zadržavaju uži nasljedni uzorak.
    """
    t = str(tekst or "").strip()
    if not t or u_sadrzaju or je_redak_sadrzaja(t, stil):
        return False
    if naslov:
        return bool(UVOD_NASLOV.match(t) or POGLAVLJE_BROJ.match(t))
    return bool(re.match(r"^\s*1\.?\s+\S", t) or UVOD_NASLOV.match(t))


def _stil_i_razina(p):
    """(stil, je_naslov, je_h1) za python-docx odlomak."""
    try:
        stil = p.style.name or ""
    except Exception:
        stil = ""
    je_naslov = stil.startswith("Heading")
    je_h1 = bool(re.match(r"^Heading\s*1\b", stil))
    lvl = None
    try:
        pPr = p._p.pPr
        if pPr is not None and pPr.outlineLvl is not None:
            lvl = int(pPr.outlineLvl.val)
    except Exception:
        lvl = None
    if lvl is not None:
        je_naslov = True
        if lvl == 0:
            je_h1 = True
    return stil, je_naslov, je_h1


# Elementi koji u odlomku nose vidljivi znak, i njihov tekstualni prikaz.
_ZNAKOVNI_ELEMENTI = {"tab": "\t", "ptab": "\t", "cr": "\n",
                      "noBreakHyphen": "-", "softHyphen": ""}
# Podstabla koja se NE čitaju: svojstva odlomka/runa (u `w:pPr/w:tabs` stoje
# definicije tabulatorskih pozicija — one nisu tabulatori u tekstu!), kod polja
# i tekst obrisan praćenjem promjena.
_NETEKSTUALNI = {"pPr", "rPr", "sdtPr", "sdtEndPr", "tblPr", "tblGrid",
                 "instrText", "delInstrText", "delText", "del", "fldChar"}


def tekst_odlomka(p):
    """Cijeli vidljivi tekst odlomka, uključujući INLINE `w:sdt` omotače.

    v1.1-fix (drugi krug): raspakiravanje `w:sdt` bilo je izvedeno samo na
    BLOK razini (dijete `w:body`). Word, međutim, sprema citat ubačen preko
    References → Insert Citation, plain-text/rich-text content control i
    Document Property polje kao INLINE `w:sdt` UNUTAR odlomka. `Paragraph.runs`
    vraća samo izravnu djecu `w:r`, pa je `p.text` takav citat tiho ispuštao:
    odlomak „Autori tvrde da je tako (Čavlek, 1998., str. 81)…" dolazio je do
    provjere citata bez citata, i rad s urednim citiranjem dobivao je
    „nema prepoznatih author-year citata u tijelu ⚠️".

    Tabulatori i prijelomi zadržavaju se doslovno jer o njima ovisi
    prepoznavanje retka sadržaja („1. UVOD\\t3").
    """
    el = getattr(p, "_p", None)
    if el is None:
        el = getattr(p, "_tc", None) or getattr(p, "_element", p)
    dijelovi = []
    _skupi_tekst(el, dijelovi)
    return "".join(dijelovi)


def _skupi_tekst(el, out):
    for x in el:
        tag = x.tag
        if not isinstance(tag, str) or not tag.startswith(_W_PREFIKS):
            continue                      # crteži, matematika, tuđi namespace
        ime = tag[len(_W_PREFIKS):]
        if ime in _NETEKSTUALNI:
            continue
        if ime == "t":
            out.append(x.text or "")
        elif ime == "br":
            # Prijelom STRANICE ili stupca nije znak u tekstu (isto tako i
            # python-docx); samo prijelom retka daje novi red.
            vrsta = x.get(_W_PREFIKS + "type")
            out.append("\n" if vrsta in (None, "textWrapping") else "")
        elif ime in _ZNAKOVNI_ELEMENTI:
            out.append(_ZNAKOVNI_ELEMENTI[ime])
        else:
            _skupi_tekst(x, out)


def odlomci_celije(c):
    """Odlomci ćelije, uz raspakiravanje `w:sdt` omotača unutar ćelije.

    `Cell.paragraphs` gleda samo izravnu djecu `w:tc`, pa je odlomak omotan u
    content control (npr. polje predloška u tablici) bio nevidljiv i ćelija je
    izgledala prazna.
    """
    from docx.text.paragraph import Paragraph
    for dijete in _djeca_tijela(c._tc):
        if dijete.tag.endswith("}p"):
            yield Paragraph(dijete, c)


def celije_tablica(tablice):
    """Tekst ćelija svih tablica, uključujući ugniježđene.

    v1.0.1-fix (D2): tablice su prije bile nevidljive — obrisan citat ili
    prevrnuta brojka u ćeliji prolazili su kao „sadržaj očuvan". Ćelija je
    uvijek tijelo rada, nikad marker.
    """
    out = []
    for t in tablice:
        for red in t.rows:
            for c in red.cells:
                for p in odlomci_celije(c):
                    tx = tekst_odlomka(p).strip()
                    if tx:
                        out.append(tx)
                out.extend(celije_tablica(c.tables))
    return out


def _blokovi_u_redoslijedu(d):
    """Odlomci i tablice u STVARNOM redoslijedu dokumenta.

    `Document.paragraphs` i `Document.tables` su dvije odvojene liste bez
    informacije o međusobnom položaju. Zbog toga je prva verzija D2 popravka
    dodavala sve ćelije tablica na kraj, bez obzira na to jesu li prije ili
    poslije početka tijela — pa je tablica s naslovnice („Mentor:", „prof. dr.
    sc. …") curila u tijelo rada. Ovdje se prolazi kroz XML tijela dokumenta pa
    je položaj tablice poznat.
    """
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    for dijete in _djeca_tijela(d.element.body):
        if dijete.tag.endswith("}p"):
            yield "p", Paragraph(dijete, d)
        elif dijete.tag.endswith("}tbl"):
            yield "tbl", Table(dijete, d)


def _djeca_tijela(el):
    """Djeca tijela dokumenta, uz raspakiravanje `w:sdt` omotača.

    v1.1-fix: Wordov automatski sadržaj i svako polje uneseno kao content
    control (`w:sdt`) omataju svoje odlomke, pa ih obilazak izravne djece
    tijela uopće ne vidi — a s njima ni odlomke rada koji su u takvom omotaču.
    Isti propis vrijedi u check_rules._elementi_tijela; ta dva obilaska moraju
    vidjeti isti dokument, inače se mjerenja dvaju alata razilaze.
    """
    for dijete in el.iterchildren():
        if dijete.tag.endswith("}sdt"):
            sadrzaj = next((x for x in dijete.iterchildren()
                            if x.tag.endswith("}sdtContent")), None)
            if sadrzaj is not None:
                yield from _djeca_tijela(sadrzaj)
        else:
            yield dijete


def je_nabrajanje_ili_citat(p, stil=""):
    """Stavka popisa ili blok-citat — NIJE prozni odlomak.

    Otkriveno kad je Katedra počela SASTAVLJATI rad iz markdownskog rukopisa: prva
    natuknica u popisu je jedna rečenica, pa je pravilo „najmanje dvije rečenice po
    odlomku" proglašavalo kršenjem svaki rad koji ima popis ili citat — a to je
    većina radova. Signal je strukturan (Wordov `w:numPr`, odnosno stil „List …" /
    „Quote"), ne duljina teksta.
    """
    from docx.oxml.ns import qn
    ime = (stil or "").strip().lower()
    if ime.startswith(("list", "popis", "quote", "citat", "intense quote")):
        return True
    try:
        ppr = p._p.find(qn("w:pPr"))
    except AttributeError:
        return False
    return ppr is not None and ppr.find(qn("w:numPr")) is not None


def _prodji_docx(d, samo_tijelo, prisili_tijelo=False, ukljuci_tablice=False,
                 bez_nabrajanja=False):
    """Jedan prolaz kroz dokument. `prisili_tijelo` znači: tijelo je već otvoreno."""
    odlomci, markeri = [], []
    u_literaturi = False
    u_sadrzaju = False
    poceo = prisili_tijelo or not samo_tijelo
    for vrsta, blok in _blokovi_u_redoslijedu(d):
        if vrsta == "tbl":
            # Ćelije ulaze u tijelo samo ako ih potrošač traži I ako je tijelo
            # već otvoreno — inače je to tablica naslovnice ili popisa.
            if ukljuci_tablice and poceo and not (samo_tijelo and u_literaturi):
                try:
                    odlomci.extend(celije_tablica([blok]))
                except Exception:
                    pass
            continue
        p = blok
        t = tekst_odlomka(p).strip()
        if not t:
            continue
        stil, je_naslov, je_h1 = _stil_i_razina(p)
        if not poceo:
            # v1.0.1-fix: popis sadržaja zatvara BILO KOJI naslov, ne samo
            # Heading 1. Prva verzija je otvarala „u_sadrzaju" na temelju samog
            # teksta („SADRŽAJ", „POPIS TABLICA"), a zatvarala ga isključivo na
            # Heading 1 — pa na dokumentu bez ijednog Word Heading stila (vrlo
            # čest studentski .docx) ili s naslovima u Heading 2 zasun se nikad
            # nije otpuštao i tijelo se nije otvaralo. Rezultat je bio NULA
            # odlomaka, na što svaka daljnja provjera javlja „uredno".
            if je_naslov or je_h1:
                u_sadrzaju = False
            if SADRZAJ_NASLOV.match(t):
                u_sadrzaju = True
            elif not je_naslov and len(t) >= 80 and t.endswith((".", "!", "?")):
                # Dugačka proza s završnom interpunkcijom nije redak sadržaja;
                # ako smo do nje došli, popis sadržaja je sigurno gotov.
                u_sadrzaju = False
            if je_pocetak_tijela(t, stil, je_naslov, u_sadrzaju):
                poceo = True
            markeri.append(t)
            continue
        if NASLOV_LIT.match(t):
            u_literaturi = True
            markeri.append(t)
            continue
        if samo_tijelo and u_literaturi:
            continue
        if stil.startswith("Heading") or NATPIS.match(t) or IZVOR.match(t):
            markeri.append(t)
            continue
        if bez_nabrajanja and je_nabrajanje_ili_citat(p, stil):
            markeri.append(t)
            continue
        # kratki blokovi bez završne interpunkcije nisu proza (ćelije, natpisi)
        if len(t) < 40 and not t.endswith((".", "!", "?")):
            markeri.append(t)
            continue
        odlomci.append(t)
    return odlomci, markeri


def _iz_docx(putanja, samo_tijelo, ukljuci_tablice=False, bez_nabrajanja=False):
    try:
        from docx import Document
    except ImportError:
        sys.exit("Treba python-docx:  pip install python-docx --break-system-packages")
    prevelik = prevelik_dio_docxa(putanja)
    if prevelik:
        sys.exit(f"❌ {putanja}: {prevelik[0]} raspakiran prelazi "
                 f"{MAX_XML_BAJTOVA // (1024 * 1024)} MB — to nije rad nego oštećen "
                 f"ili namjerno napuhan zip.\n"
                 f"   Što napraviti: uzmi datoteku koju si stvarno predao/la, ili je "
                 f"otvori u Wordu i spremi ponovno kao .docx.")
    d = Document(putanja)
    odlomci, markeri = _prodji_docx(d, samo_tijelo, ukljuci_tablice=ukljuci_tablice,
                                    bez_nabrajanja=bez_nabrajanja)

    # SIGURNOSNA MREŽA: `samo_tijelo` nikad ne smije tiho vratiti prazno tijelo
    # za dokument koji očito ima prozu. Prazan rezultat znači da svaka daljnja
    # provjera (check_rules, check_ai_style, check_argument, originality_check,
    # verify_rewrite) radi nad ničim i javlja „uredno" — točno onaj obrazac
    # „zeleno svjetlo na crveno stanje" zbog kojeg ovaj popravak postoji. Ako
    # heuristika početka tijela ne nađe ništa, pada se na staro ponašanje
    # (cijeli dokument kao tijelo), isto kao dokumentirani legacy fallback u
    # check_rules.provjeri_font.
    if samo_tijelo and not odlomci:
        rezervni, rezervni_markeri = _prodji_docx(
            d, samo_tijelo, prisili_tijelo=True, ukljuci_tablice=ukljuci_tablice)
        if rezervni:
            return rezervni, rezervni_markeri or markeri
    return odlomci, markeri


# ------------------------------------------------------------ segmentacija

KRATICE = (r"str|izd|sur|dr|mr|prof|doc|br|tj|npr|god|mln|mlrd|tis|cca|"
           r"odn|usp|op|cit|itd|sl|tzv|čl|st|sv|god|Nar|nov|"
           # v1.0.1-fix (Q3a): akademske titule i uredničke kratice
           r"sc|ur|prev|mag|univ|spec|akad|vol")

# v1.0.1-fix (Q3a): riječi iza kojih redni broj NIJE kraj rečenice
# („prema članku 5. Zakona…"). Rješava se LIJEVIM kontekstom, jer se
# „2019. Bila je…" i „članku 5. Zakona…" razlikuju samo po onome ŠTO JE PRIJE
# točke, a ne po znaku iza nje.
REDNI_GOVERNORI = (r"čl|član(?:ak|ka|ku|kom)?|stavak|stavk(?:a|u|om)|"
                   r"točk(?:a|e|i|u)|tablic(?:a|e|i|u)|slik(?:a|e|i|u)|"
                   r"poglavlj(?:e|a|u|em)|br|str")

_ZASTITA = ""


# Nastavak koji smije slijediti „članku 5." a da to NIJE nova rečenica: naziv
# propisa u genitivu. Whitelist je namjerno uzak — ako se propis ne prepozna,
# pada se na staro (pre-popravak) ponašanje, tj. rečenica se razdvoji. To je
# svjestan izbor smjera pogreške: prerano razdvajanje samo napuhuje broj
# rečenica i time maskira prekršaj, dok pogrešno SPAJANJE stvara lažni ❌ na
# ispravnom odlomku. Nezavisna revizija je pokazala da je prva verzija ovog
# pravila (bezuvjetna zaštita točke iza „governor + broj") spajala sasvim
# obične rečenice — „Vidi tablicu 3. Rezultati su jasni." — i time obarala
# `provjeri_odlomke` na usklađenom radu.
_PROPIS_NASTAVAK = (r"Zakon|Ustav|Pravilnik|Uredb|Direktiv|Odluk|Statut|Ugovor|"
                    r"Konvencij|Sporazum|Poslovnik|Naredb|Kodeks|Protokol|"
                    r"Smjernic|Uput|Strategij|Program|Plan")


def _zastiti(t):
    """Zamijeni točke koje NISU kraj rečenice privremenim znakom."""
    # Točka iza „članak 5." štiti se SAMO ako slijedi naziv propisa (veliko
    # slovo). Nastavak malim slovom već pokriva općenito pravilo za znamenku
    # niže, pa ovdje nije potreban.
    t = re.sub(r"(?i)\b(" + REDNI_GOVERNORI + r")(\.?\s+\d+)\."
               r"(?=\s+(?:" + _PROPIS_NASTAVAK + r"))",
               r"\1\2" + _ZASTITA, t)
    # v1.1-fix: kratica se prepoznavala samo malim slovom, pa je ona na POČETKU
    # rečenice, kad iza nje slijedi velika riječ („Usp. Čavlek…", „Ur. Marić…",
    # „Prof. Ivić tvrdi…"), ostajala nezaštićena i lomila se u zasebnu
    # „rečenicu" („Usp."). Time je broj rečenica u odlomku bio napuhan, a
    # napuhan broj skriva prekratak odlomak od `provjeri_odlomke`. Kratica je
    # ista riječ bez obzira na položaj u rečenici, pa je usporedba neosjetljiva
    # na veličinu slova.
    t = re.sub(r"(?i)\b(" + KRATICE + r")\.", r"\1" + _ZASTITA, t)
    # redni broj / godina: točka iza znamenke, a iza nje NIJE velika slova+razmak
    # koji započinju novu rečenicu -> ostavlja se npr. "2020. godine"
    t = re.sub(r"(\d)\.(?=\s*[a-zčćžšđ(])", r"\1" + _ZASTITA, t)
    t = re.sub(r"(\d)\.(?=\s*\d)", r"\1" + _ZASTITA, t)
    t = re.sub(r"(\d)\.(?=\s*[–—-])", r"\1" + _ZASTITA, t)
    # inicijali: "N. Čavlek"
    t = re.sub(r"\b([A-ZČĆŽŠĐ])\.(?=\s+[A-ZČĆŽŠĐ])", r"\1" + _ZASTITA, t)
    return t


# v1.0.1-fix (Q3a): rečenica smije završiti zatvarajućim navodnikom ili
# zagradom („…važno.” Sljedeća…) i smije započeti otvarajućim navodnikom ili
# zagradom. Bez toga se odlomak s citatom u navodnicima broji kao jedna
# rečenica manje, pa provjeri_odlomke prijavljuje lažni pad.
_ZATVARANJE = r"”“\"»«›’'\)\]"
_OTVARANJE = r"A-ZČĆŽŠĐ„“\"»(\["
_GRANICA = (r"(?:(?<=[.!?])|(?<=[.!?][" + _ZATVARANJE + r"]))"
            r"\s+(?=[" + _OTVARANJE + r"])")


def recenice(tekst):
    """Podijeli na rečenice, uz hrvatske redne brojeve i kratice."""
    t = _zastiti(tekst)
    dijelovi = re.split(_GRANICA, t)
    return [d.replace(_ZASTITA, ".").strip() for d in dijelovi if d.strip()]


def rijeci(tekst):
    return [w for w in re.split(r"\s+", tekst.strip()) if w]


# --------------------------------------------------- hrvatska abeceda (kolacija)

_HR = {
    "a": "a", "b": "b", "c": "c1", "č": "c2", "ć": "c3", "d": "d1",
    "dž": "d2", "đ": "d3", "e": "e", "f": "f", "g": "g", "h": "h",
    "i": "i", "j": "j", "k": "k", "l": "l1", "lj": "l2", "m": "m",
    "n": "n1", "nj": "n2",
    "o": "o", "p": "p", "q": "q", "r": "r", "s": "s1", "š": "s2",
    "t": "t", "u": "u", "v": "v", "w": "w", "x": "x", "y": "y",
    "z": "z1", "ž": "z2",
    # ð (U+00F0) je lookalike hrvatskog đ iz CP1250 konverzije — sortira se na
    # isto mjesto, inače isto prezime završi na dva mjesta u popisu literature.
    "ð": "d3",
}
# v1.0.1-fix (Q3b): hrvatski digrafi sortiraju se kao JEDNO slovo
# (L < LJ < M, N < NJ < O, D < DŽ < Đ). Ograničenje: kad l/n i j pripadaju
# različitim slogovima (npr. „injekcija", „konjugacija", „nadživjeti"),
# ovo ih svejedno čita kao digraf — to je uobičajena konvencija u popisima
# literature i ne postoji način da se razlikuje bez rječnika.
_DIGRAFI = ("dž", "lj", "nj")
# strani dijakritici koji se u hrvatskom popisu tretiraju kao osnovno slovo
_STRANI = {"ö": "o", "ó": "o", "ô": "o", "ø": "o", "ü": "u", "ú": "u",
           "ä": "a", "á": "a", "à": "a", "å": "a", "é": "e", "è": "e",
           "ê": "e", "ë": "e", "í": "i", "î": "i", "ı": "i", "ğ": "g",
           "ń": "n", "ñ": "n", "ł": "l", "ç": "c", "ß": "ss", "ý": "y"}


def _osnovno(baza):
    """Preslikaj osnovna latinična slova u hrvatske kolacijske kodove."""
    return "".join(_HR.get(c, c) for c in baza)


def hr_kljuc(s):
    """Ključ za sortiranje po hrvatskoj abecedi: C < Č < Ć, S < Š, Z < Ž, D < Đ,
    uz digrafe DŽ, LJ i NJ kao zasebna slova."""
    s = s.lower()
    out, i = [], 0
    while i < len(s):
        if s[i:i + 2] in _DIGRAFI:
            out.append(_HR[s[i:i + 2]]); i += 2; continue
        z = s[i]
        if z in _HR:
            out.append(_HR[z])
        elif z in _STRANI:
            out.append(_osnovno(_STRANI[z]))
        elif z.isalnum():
            # v1.0.1-fix (Q3b): nepoznati strani dijakritik (ę, ř, ă, ș, ő…)
            # inače završi IZA slova ž; svodi se na osnovno slovo.
            baza = bez_dijakritika(z).lower()
            out.append(_osnovno(baza) if baza and baza != z else z)
        else:
            out.append(" ")
        i += 1
    return "".join(out)


# v1.0.1-fix (Q3c): znakovi bez NFD dekompozicije (đ, ø, ł, ß, æ, œ) preživjeli
# su uklanjanje combining marksa, pa ih je nizvodni [^a-z0-9]+ pretvarao u
# razmak („među" -> „me u"). Napomena: ß -> ss MIJENJA duljinu niza, pa za taj
# jedini slučaj ne vrijedi 1:1 duljinska očuvanost na koju se poziva
# originality_check._prozor_oko_shinglea; svi ostali unosi su 1:1.
_PREDFOLD = {
    "đ": "d", "Đ": "D", "ø": "o", "Ø": "O", "ł": "l", "Ł": "L",
    "ß": "ss", "æ": "ae", "Æ": "Ae", "œ": "oe", "Œ": "Oe",
    # v1.1-fix (drugi krug): Ð/ð (U+00D0/U+00F0, islandski eth) NIJE hrvatsko
    # Đ/đ (U+0110/U+0111), ali izgleda identično i redovito nastaje pri
    # konverziji iz CP1250 / starih .doc-ova i pri kopiranju iz PDF-a. Bez ovog
    # unosa „Ðurđica" i „Đurđica" daju različit ključ, pa se isto prezime u
    # popisu literature sortira na dva mjesta, a `originality_check` ga vidi kao
    # dva različita niza. NFD ga ne razlaže (nema dekompoziciju), pa bi ga
    # nizvodni [^a-z0-9]+ pretvorio u razmak — isti kvar kao Q3c.
    "ð": "d", "Ð": "D",
}


def bez_dijakritika(s):
    s = "".join(_PREDFOLD.get(c, c) for c in s)
    return "".join(c for c in unicodedata.normalize("NFD", s)
                   if unicodedata.category(c) != "Mn")


# ---------------------------------------------------------------- citiranje
# B10: parser je centraliziran u citation_dialects.py. Ova imena ostaju kao
# compatibility API za stare alate i testove.
try:
    from . import citation_dialects as C
except ImportError:  # izravno pokretanje scripts/hr_text.py
    import citation_dialects as C

CITAT_ZAGRADNI = C.CITAT_ZAGRADNI
CITAT_NARATIVNI = C.CITAT_NARATIVNI
LOKATOR = C.LOKATOR


def bez_lokatora(tekst):
    return C.bez_lokatora(tekst)


def godine_u_citatima(tekst):
    """Multiskup godina author-year citata (legacy API)."""
    from collections import Counter
    return Counter(r.year[:4] for r in C.parse_citations(tekst or "", "autor-godina") if r.year)


def kljucevi_citata(tekst):
    """Multiskup (prvo_prezime, godina) author-year citata (legacy API)."""
    from collections import Counter
    return Counter((r.author, r.year) for r in C.parse_citations(tekst or "", "autor-godina") if r.author and r.year)


def main(argv=None):
    ap = argparse.ArgumentParser(
        description="Sažmi strukturu i osnovne tekstualne metrike hrvatskog akademskog rada."
    )
    ap.add_argument("rad", help="ulazni .docx, .md ili .txt dokument")
    args = ap.parse_args(argv)

    odl, mark = ucitaj(args.rad)
    rec = [r for o in odl for r in recenice(o)]
    print(f"odlomaka: {len(odl)} | markera: {len(mark)} | rečenica: {len(rec)}")
    print(f"citiranih godina: {sum(godine_u_citatima(' '.join(odl)).values())}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
