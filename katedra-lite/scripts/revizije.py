#!/usr/bin/env python3
"""Tri operacije nad `.docx` kroz životni ciklus jedne revizije: prihvati, redline, toc.

Zajedno u jednoj datoteci namjerno (paketni proračun broja datoteka je zategnut,
v. `tests/unit/test_kontekstni_proracun.py`) — sve tri dijele isti kontekst: rad
koji dolazi s neprihvaćenim izmjenama, pa se prepisuje, pa treba pokazati što se
promijenilo, pa treba TOC koji nije očito zastario prije nego korisnik stigne do
Worda i ručno pritisne Update Field.

## prihvati — Word Track Changes → čist tekst, na razini XML-a

`python-docx`-ov `Paragraph.text` čita samo runove koji su IZRAVNA djeca `<w:p>`.
Umetnuti tekst živi unutar `<w:ins>`, obrisani unutar `<w:del>` — oboje jedan
stupanj dublje, pa ih `.text` PRESKAČE. Rad s i najmanjom neprihvaćenom izmjenom
kroz `python-docx` izgleda krnj: nedostaju riječi, rečenice se raspadaju usred
misli. Svaka dijagnoza koja krene od takvog čitanja polazi od krivog polazišta.
Radi ono što Word radi s „Review → Accept All Changes", programatski: `<w:ins>`
se raspakira (tekst ostaje), `<w:del>` se briše sa sadržajem. Wordovi komentari
(`<w:comment...>`) NISU praćene izmjene i ovim se ne dirju.

    python3 <KATEDRA_SKILL>/scripts/revizije.py prihvati ulaz.docx izlaz.docx [--json]

## redline — obojena usporedba dvije verzije, za čovjeka

`diff_versions.py` uspoređuje verzije radi INTERNOG praćenja (sha256, izgubljeni
citati) i ne proizvodi dokument za vizualno prelistavanje. Ovo radi to: diff na
razini odlomaka pa riječi, treći `.docx` gdje je obrisan tekst crven i precrtan,
dodan tekst crven bez precrtavanja. Boja je izravno na fontu, NE Wordov
`<w:ins>`/`<w:del>` mehanizam (ondje boja ovisi o postavkama recenzenta u Wordu
pa „crveno" nije zajamčeno). Premješteni odlomci pojavljuju se dva puta (brisanje
na starom mjestu, dodavanje na novom) — to je točan opis diffa na razini teksta,
ne bug.

    python3 <KATEDRA_SKILL>/scripts/revizije.py redline prije.docx poslije.docx izlaz.docx

## toc — procjena stranica keširanog TOC polja

Sadržaj je Wordovo TOC polje (`TOC \\o "1-2" \\h \\z \\u`), ali PAMTI POSLJEDNJI
IZRAČUNATI REZULTAT kao obične odlomke — dok korisnik u Wordu ne pritisne
„Update Field" (F9), ti odlomci ostaju stari čak i kad se naslovi promijene ili
poglavlje naraste. Ovo NE zamjenjuje Update Field (koji jedini poznaje stvarne
fontove i prijelom korisnikova računala) — PROCJENJUJE brojeve stranica preko
LibreOffice headless renderiranja + pretrage teksta po stranicama, pa keširane
retke prepiše na tu procjenu. Razlika prema Wordu tipično ±1 stranica.

    python3 <KATEDRA_SKILL>/scripts/revizije.py toc ulaz.docx izlaz.docx [--preskoci-stranice N]

Izlazni kodovi (svaki podnaredba): 0 gotovo · 1 ulaz nečitljiv/ovisnost nedostaje
· 2 (samo `toc`) Sadržaj nije prepoznat kao TOC polje.
"""
from __future__ import annotations

import argparse
import difflib
import json
import os
import shutil
import subprocess
import sys
import tempfile
import zipfile

try:
    import docx
    from docx.shared import RGBColor
except ImportError:  # pragma: no cover
    print("❌ nedostaje python-docx (pip install python-docx --break-system-packages)", file=sys.stderr)
    sys.exit(1)

try:
    from lxml import etree
except ImportError:  # pragma: no cover
    print("❌ nedostaje lxml (pip install lxml --break-system-packages)", file=sys.stderr)
    sys.exit(1)


# ───────────────────────── prihvati ─────────────────────────

REMOVE_ENTIRELY = {
    "del", "moveFrom", "moveFromRangeStart", "moveFromRangeEnd",
    "rPrChange", "pPrChange", "tblPrChange", "trPrChange", "tcPrChange",
    "sectPrChange", "numberingChange",
}
UNWRAP = {"ins", "moveTo", "moveToRangeStart", "moveToRangeEnd"}
PARTS_TO_PROCESS = ("document.xml", "footer", "header", "footnotes.xml", "endnotes.xml")


def _process_element(root) -> None:
    changed = True
    while changed:
        changed = False
        for el in list(root.iter()):
            if not isinstance(el.tag, str):
                continue
            local = etree.QName(el).localname
            parent = el.getparent()
            if parent is None:
                continue
            if local in REMOVE_ENTIRELY:
                parent.remove(el)
                changed = True
            elif local in UNWRAP:
                idx = list(parent).index(el)
                for i, child in enumerate(list(el)):
                    parent.insert(idx + i, child)
                parent.remove(el)
                changed = True


def _process_xml_bytes(data: bytes) -> bytes:
    tree = etree.fromstring(data)
    _process_element(tree)
    return etree.tostring(tree, xml_declaration=True, encoding="UTF-8", standalone=True)


def accept_revisions(src: str, dst: str) -> dict:
    """Vrati sažetak {part: {ins: N, del: N}} za svaki obrađeni dio paketa."""
    try:
        zin = zipfile.ZipFile(src, "r")
    except (OSError, zipfile.BadZipFile) as exc:
        raise SystemExit(f"❌ ne mogu pročitati {src}: {exc}")
    summary = {}
    try:
        zout = zipfile.ZipFile(dst, "w", zipfile.ZIP_DEFLATED)
    except OSError as exc:
        raise SystemExit(f"❌ ne mogu pisati {dst}: {exc}")
    for item in zin.infolist():
        data = zin.read(item.filename)
        is_target = item.filename.startswith("word/") and item.filename.endswith(".xml") \
            and any(key in item.filename for key in PARTS_TO_PROCESS)
        if is_target:
            n_ins = data.count(b"<w:ins ") + data.count(b"<w:ins>")
            n_del = data.count(b"<w:del ") + data.count(b"<w:del>")
            if n_ins or n_del:
                data = _process_xml_bytes(data)
                summary[item.filename] = {"ins": n_ins, "del": n_del}
        zout.writestr(item, data)
    zin.close()
    zout.close()
    return summary


def cmd_prihvati(a) -> int:
    summary = accept_revisions(a.ulaz, a.izlaz)
    if a.json:
        print(json.dumps({"ulaz": a.ulaz, "izlaz": a.izlaz, "obradeno": summary}, ensure_ascii=False, indent=2))
        return 0
    if not summary:
        print(f"ℹ️  {a.ulaz} nije imao praćenih izmjena — {a.izlaz} je identičan sadržajno.")
    else:
        total_ins = sum(v["ins"] for v in summary.values())
        total_del = sum(v["del"] for v in summary.values())
        print(f"✅ prihvaćeno: {total_ins} umetanja, {total_del} brisanja, u {len(summary)} dijelova paketa")
        for part, counts in summary.items():
            print(f"   {part}: ins={counts['ins']} del={counts['del']}")
    print(f"   napisano: {a.izlaz}")
    print("   napomena: komentari (Wordovi review-comments) nisu dirani — extract_comments.py radi normalno na izlazu.")
    return 0


def count_revisions(src: str) -> dict:
    """Prebroji praćene izmjene bez ijednog zapisa. Vraća {part: {ins, del}}.

    Zašto postoji: pravilo 30 traži prihvat praćenih izmjena PRIJE prve
    ekstrakcije, a jedini alat koji ih je vidio ujedno je i pisao novu datoteku.
    Provjera koja mijenja ulaz ne može biti prvi korak gatea.
    """
    try:
        zin = zipfile.ZipFile(src, "r")
    except (OSError, zipfile.BadZipFile) as exc:
        raise SystemExit(f"❌ ne mogu pročitati {src}: {exc}")
    summary = {}
    with zin:
        for item in zin.infolist():
            if not (item.filename.startswith("word/") and item.filename.endswith(".xml")
                    and any(key in item.filename for key in PARTS_TO_PROCESS)):
                continue
            data = zin.read(item.filename)
            n_ins = data.count(b"<w:ins ") + data.count(b"<w:ins>")
            n_del = data.count(b"<w:del ") + data.count(b"<w:del>")
            n_mv = data.count(b"<w:moveFrom ") + data.count(b"<w:moveTo ")
            if n_ins or n_del or n_mv:
                summary[item.filename] = {"ins": n_ins, "del": n_del, "premjesteno": n_mv}
    return summary


def cmd_provjeri(a) -> int:
    summary = count_revisions(a.ulaz)
    uk_ins = sum(v["ins"] for v in summary.values())
    uk_del = sum(v["del"] for v in summary.values())
    uk_mv = sum(v["premjesteno"] for v in summary.values())
    nalaz = bool(summary)
    if a.json_out:
        with open(a.json_out, "w", encoding="utf-8") as fh:
            json.dump({"rad": a.ulaz, "ima_pracene_izmjene": nalaz,
                       "ins": uk_ins, "del": uk_del, "premjesteno": uk_mv,
                       "dijelovi": summary}, fh, ensure_ascii=False, indent=2)
    if not nalaz:
        print(f"✅ {a.ulaz} nema praćenih izmjena — ekstrakcija čita ono što je u dokumentu.")
        return 0
    print(f"❌ {a.ulaz} ima praćene izmjene: {uk_ins} umetanja, {uk_del} brisanja, "
          f"{uk_mv} premještanja, u {len(summary)} dijelova paketa.")
    for part, c in summary.items():
        print(f"   {part}: ins={c['ins']} del={c['del']} mv={c['premjesteno']}")
    print("   Svaka ekstrakcija nad ovim dokumentom čita KRIVI tekst: obrisani odlomci")
    print("   se broje, umetnuti ne. Opseg, sažetak i citati bit će pogrešni.")
    print(f"   Prvo: python3 revizije.py prihvati {a.ulaz} rad-prihvaceno.docx")
    return 1


# ───────────────────────── redline ─────────────────────────

RED = RGBColor(0xC0, 0x00, 0x00)


def _load_docx(path):
    try:
        return docx.Document(path)
    except Exception as exc:  # noqa: BLE001
        raise SystemExit(f"❌ ne mogu otvoriti {path} kao .docx: {exc}")


def _word_diff_paragraph(out_doc, style, old_text, new_text):
    p = out_doc.add_paragraph(style=style)
    old_words = old_text.split(" ")
    new_words = new_text.split(" ")
    wsm = difflib.SequenceMatcher(None, old_words, new_words, autojunk=False)
    for tag, i1, i2, j1, j2 in wsm.get_opcodes():
        if tag == "equal":
            chunk = " ".join(old_words[i1:i2])
            if chunk:
                p.add_run(chunk + " ")
        elif tag == "delete":
            chunk = " ".join(old_words[i1:i2])
            if chunk:
                r = p.add_run(chunk + " ")
                r.font.color.rgb = RED
                r.font.strike = True
        elif tag == "insert":
            chunk = " ".join(new_words[j1:j2])
            if chunk:
                r = p.add_run(chunk + " ")
                r.font.color.rgb = RED
        elif tag == "replace":
            old_chunk = " ".join(old_words[i1:i2])
            new_chunk = " ".join(new_words[j1:j2])
            if old_chunk:
                r = p.add_run(old_chunk + " ")
                r.font.color.rgb = RED
                r.font.strike = True
            if new_chunk:
                r = p.add_run(new_chunk + " ")
                r.font.color.rgb = RED
    return p


def _marked_paragraph(out_doc, style, text, deleted):
    p = out_doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        r.font.color.rgb = RED
        if deleted:
            r.font.strike = True
    return p


def build_redline(baseline_path: str, final_path: str, out_path: str) -> int:
    base_doc = _load_docx(baseline_path)
    final_doc = _load_docx(final_path)

    base_paras = [(p.style.name, p.text) for p in base_doc.paragraphs]
    final_paras = [(p.style.name, p.text) for p in final_doc.paragraphs]
    base_texts = [t for _, t in base_paras]
    final_texts = [t for _, t in final_paras]

    sm = difflib.SequenceMatcher(None, base_texts, final_texts, autojunk=False)
    opcodes = sm.get_opcodes()

    out_doc = final_doc  # naslijedi stilove/sekcije/fontove iz finalne verzije
    for p in list(out_doc.paragraphs):
        p._element.getparent().remove(p._element)

    n_changed_paragraphs = 0
    for tag, i1, i2, j1, j2 in opcodes:
        if tag == "equal":
            for k in range(i2 - i1):
                style, text = final_paras[j1 + k]
                p = out_doc.add_paragraph(style=style)
                if text:
                    p.add_run(text)
        elif tag == "delete":
            for k in range(i1, i2):
                style, text = base_paras[k]
                _marked_paragraph(out_doc, style, text, deleted=True)
                n_changed_paragraphs += 1
        elif tag == "insert":
            for k in range(j1, j2):
                style, text = final_paras[k]
                _marked_paragraph(out_doc, style, text, deleted=False)
                n_changed_paragraphs += 1
        elif tag == "replace":
            n_old, n_new = i2 - i1, j2 - j1
            if n_old == n_new:
                for off in range(n_old):
                    b_style, b_text = base_paras[i1 + off]
                    f_style, f_text = final_paras[j1 + off]
                    if not b_text and not f_text:
                        out_doc.add_paragraph(style=f_style)
                    else:
                        _word_diff_paragraph(out_doc, f_style, b_text, f_text)
                    n_changed_paragraphs += 1
            else:
                for k in range(i1, i2):
                    style, text = base_paras[k]
                    _marked_paragraph(out_doc, style, text, deleted=True)
                for k in range(j1, j2):
                    style, text = final_paras[k]
                    _marked_paragraph(out_doc, style, text, deleted=False)
                n_changed_paragraphs += n_old + n_new

    out_doc.save(out_path)
    return n_changed_paragraphs


def cmd_redline(a) -> int:
    n = build_redline(a.prije, a.poslije, a.izlaz)
    print(f"✅ redline napisan: {a.izlaz}")
    print(f"   izmijenjenih/premještenih odlomaka: {n}")
    print("   napomena: premješteni odlomci prikazuju se DVA puta (brisanje na starom mjestu, "
          "dodavanje na novom) — to je očekivano ponašanje diffa na razini teksta, ne bug.")
    return 0


# ───────────────────────── toc ─────────────────────────

def _render_pdf(docx_path: str, out_dir: str) -> str:
    exe = shutil.which("soffice") or shutil.which("libreoffice")
    if not exe:
        print("❌ LibreOffice (soffice) nije dostupan u ovom okruženju.", file=sys.stderr)
        sys.exit(1)
    subprocess.run(
        [exe, "--headless", "--convert-to", "pdf", "--outdir", out_dir, docx_path],
        check=True, capture_output=True, timeout=120,
    )
    base = os.path.splitext(os.path.basename(docx_path))[0]
    pdf_path = os.path.join(out_dir, base + ".pdf")
    if not os.path.isfile(pdf_path):
        print("❌ LibreOffice nije proizveo PDF (provjeri je li datoteka ispravan .docx).", file=sys.stderr)
        sys.exit(1)
    return pdf_path


def _find_toc_paragraphs(doc):
    """Vrati indekse odlomaka oblika 'naslov<TAB>broj' — keširani TOC redci."""
    out = []
    for i, p in enumerate(doc.paragraphs):
        t = p.text
        if "\t" in t:
            head, _, tail = t.rpartition("\t")
            if head.strip() and tail.strip().isdigit():
                out.append((i, head.strip()))
    return out


def estimate_toc(docx_path: str, out_path: str, skip_pages: int = 4) -> int:
    try:
        from pypdf import PdfReader
    except ImportError:  # pragma: no cover
        print("❌ nedostaje pypdf (pip install pypdf --break-system-packages)", file=sys.stderr)
        sys.exit(1)

    doc = docx.Document(docx_path)
    toc_entries = _find_toc_paragraphs(doc)
    if not toc_entries:
        print("❌ Nema prepoznatljivih TOC redaka (\"naslov<TAB>broj\"). "
              "Je li Sadržaj uopće Wordovo TOC polje?", file=sys.stderr)
        return 2

    with tempfile.TemporaryDirectory() as tmp:
        pdf_path = _render_pdf(docx_path, tmp)
        reader = PdfReader(pdf_path)
        found = {}
        for page_idx, page in enumerate(reader.pages):
            if page_idx < skip_pages:
                continue  # preskoči naslovnicu/Sadržaj samog sebe (lažni pogoci)
            text = page.extract_text() or ""
            # Traži cijelu stranicu, ne samo početak — naslov se često nalazi
            # na sredini stranice (prethodni odlomak se nastavlja s prošle
            # stranice pa tek onda počinje novo poglavlje).
            for _, heading in toc_entries:
                if heading in found:
                    continue
                needle = heading[: min(30, len(heading))]
                if needle in text:
                    found[heading] = page_idx + 1

        missing = [h for _, h in toc_entries if h not in found]
        if missing:
            print(f"⚠️  {len(missing)} naslova nije pronađeno u renderiranom PDF-u "
                  f"(provjeri odgovaraju li naslovi TOC redaka stvarnim naslovima u tekstu):")
            for h in missing[:10]:
                print(f"     - {h}")

        first_heading_page = found.get(toc_entries[0][1])
        offset = (first_heading_page - 1) if first_heading_page else 0

        for idx, heading in toc_entries:
            page = found.get(heading)
            if page is None:
                continue
            printed = max(1, page - offset)
            p = doc.paragraphs[idx]
            template = next((r for r in p.runs if r.text), p.runs[-1] if p.runs else None)
            new_text = f"{heading}\t{printed}"
            if template is None:
                p.add_run(new_text)
            else:
                bold, italic, underline = template.bold, template.italic, template.underline
                size, font_name = template.font.size, template.font.name
                for r in list(p.runs):
                    r.text = ""
                template.text = new_text
                template.bold, template.italic, template.underline = bold, italic, underline
                if size:
                    template.font.size = size
                if font_name:
                    template.font.name = font_name

    doc.save(out_path)
    return 0


def cmd_toc(a) -> int:
    rc = estimate_toc(a.ulaz, a.izlaz, skip_pages=a.preskoci_stranice)
    if rc == 0:
        print(f"✅ TOC redci procijenjeni i upisani: {a.izlaz}")
        print("   OBAVEZNO: ovo je procjena (LibreOffice paginacija), ne stvarni Wordov izračun.")
        print("   Prije predaje: otvori u Wordu → klik desnom na Sadržaj → Update Field.")
    return rc


# ───────────────────────── CLI ─────────────────────────

def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    sub = ap.add_subparsers(dest="cmd", required=True)

    p_prihvati = sub.add_parser("prihvati", help="prihvati sve praćene izmjene (Track Changes), na razini XML-a")
    p_prihvati.add_argument("ulaz")
    p_prihvati.add_argument("izlaz")
    p_prihvati.add_argument("--json", action="store_true", help="strojno čitljiv sažetak na stdout")

    p_provjeri = sub.add_parser("provjeri", help="ima li dokument praćenih izmjena (ne mijenja ga)")
    p_provjeri.add_argument("ulaz")
    p_provjeri.add_argument("--json", dest="json_out", metavar="PUT",
                            help="zapiši nalaz u JSON")

    p_redline = sub.add_parser("redline", help="obojena usporedba dvije verzije .docx-a, za čovjeka")
    p_redline.add_argument("prije", help="bazna (starija) verzija .docx-a")
    p_redline.add_argument("poslije", help="nova (finalna) verzija .docx-a")
    p_redline.add_argument("izlaz", help="izlazna redline .docx datoteka")

    p_toc = sub.add_parser("toc", help="procijeni i upiši stranice keširanog TOC polja")
    p_toc.add_argument("ulaz")
    p_toc.add_argument("izlaz")
    p_toc.add_argument("--preskoci-stranice", type=int, default=4,
                        help="broj početnih PDF stranica koje se ignoriraju pri pretrazi, zadano 4")

    a = ap.parse_args()
    if a.cmd == "prihvati":
        return cmd_prihvati(a)
    if a.cmd == "provjeri":
        return cmd_provjeri(a)
    if a.cmd == "redline":
        return cmd_redline(a)
    if a.cmd == "toc":
        return cmd_toc(a)
    return 2


if __name__ == "__main__":
    sys.exit(main())
