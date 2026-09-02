#!/usr/bin/env python3
"""Vancouver (n) citati protiv numeriranog popisa literature — što nijedan dijalekt u paketu ne vidi.

Zašto postoji
-------------
`citation_dialects.py` poznaje autor-godina, IEEE `[n]` i legal-footnote. Zdravstveni fakulteti
(HKS-FZS, MEF, ZVU, Stomatološki…) gotovo isključivo traže Vancouver: `(n)`, `(n, m)`, `(n–m)`,
`i sur.` u popisu, numeracija **po redoslijedu prvog pojavljivanja**. Na takvom radu faza B
rad-audita prijavi „unknown, 0 citata" i jedan lažni kritični nalaz — a 147 stvarnih citata
ostane neprovjereno.

Što provjerava
--------------
1. **Siročad** — broj u popisu koji se nigdje ne citira.
2. **Citat bez reference** — broj u tekstu veći od duljine popisa.
3. **Redoslijed prvog pojavljivanja** — Vancouver traži da se brojevi pojave uzlazno; prvi skok
   (npr. 1, 2, 5, 3) znači da je popis prenumeriran ili je citat naknadno umetnut.
4. **Rasponi** — `(5-7)` sa spojnicom umjesto en-crtice; `(5, 6, 7)` koji bi trebao biti `(5–7)`.
5. **Popis** — prazni brojevi, dupli brojevi, brojevi izvan niza 1..N.

Ne provjerava sadržaj reference (to je `verify_sources.py`) ni „i sur." pravilo (to je kućni stil).

Uporaba
-------
  python3 provjeri_vancouver.py ./rad.docx [--json .katedra/vancouver.json] [--od-naslova "POPIS CITIRANE LITERATURE"]
Izlazni kod 1 ako ima siročadi ili citata bez reference; 0 inače (redoslijed i rasponi su savjet).
"""
from __future__ import annotations

import argparse
import json
import re
import sys

from docx import Document

RE_CIT = re.compile(r"\((\d{1,3}(?:\s*[,;]\s*\d{1,3}|\s*[–-]\s*\d{1,3})*)\)")
RE_POPIS_STAVKA = re.compile(r"^\s*(\d{1,3})\.\s+\S")
NASLOVI_POPISA = ("POPIS CITIRANE LITERATURE", "POPIS LITERATURE", "LITERATURA", "REFERENCE", "REFERENCES")


def odlomci(put):
    d = Document(put)
    out = []
    for p in d.paragraphs:
        out.append(p.text)
    for t in d.tables:
        for row in t.rows:
            for c in row.cells:
                out.append(c.text)
    return out, [p.text for p in d.paragraphs]


def razdvoji(paras, od_naslova):
    """(tijelo, popis) — popis počinje na naslovu popisa, završava na sljedećem Heading-1 nalik retku."""
    start = None
    for i, t in enumerate(paras):
        s = t.strip().upper()
        if od_naslova and s.endswith(od_naslova.upper()):
            start = i
            break
        if not od_naslova and any(s.endswith(n) for n in NASLOVI_POPISA) and len(s) < 60:
            start = i
            break
    if start is None:
        return paras, []
    kraj = len(paras)
    for j in range(start + 1, len(paras)):
        s = paras[j].strip()
        if re.match(r"^\d{1,2}\.\s+[A-ZČĆŽŠĐ ]{4,}$", s):  # sljedeći numerirani Heading 1 (npr. „9. PRILOZI")
            kraj = j
            break
    return paras[:start] + paras[kraj:], paras[start + 1:kraj]


def brojevi_iz_citata(grupa):
    nums = []
    for dio in re.split(r"\s*[,;]\s*", grupa):
        m = re.match(r"(\d+)\s*[–-]\s*(\d+)$", dio)
        if m:
            a, b = int(m.group(1)), int(m.group(2))
            nums.extend(range(a, b + 1) if a <= b else [a, b])
        elif dio.strip().isdigit():
            nums.append(int(dio))
    return nums


def analiza(put, od_naslova=None):
    svi, samo_paras = odlomci(put)
    tijelo, popis = razdvoji(samo_paras, od_naslova)
    # citati i iz tablica/ćelija (svi) ali bez popisa
    popis_set = set(popis)
    tekst_za_citate = [t for t in svi if t not in popis_set]

    redoslijed, citati_sirovo, raspon_spojnica, nabrajanje_umjesto_raspona, bez_razmaka = [], [], [], [], []
    for t in tekst_za_citate:
        for m in RE_CIT.finditer(t):
            g = m.group(1)
            # tablične ćelije „n (%)": „158 (77,8)" — decimala bez razmaka iza zareza, ispred zagrade brojka
            # decimala („77,8", „50,0") vs. citat bez razmaka („67,68"): decimala ima jednu znamenku iza
            # zareza ILI joj u istoj ćeliji prethodi brojka („158 (77,8)"); ostalo je citat, ali se bilježi
            # kao stilski nalaz jer Vancouver traži razmak iza zareza.
            if re.fullmatch(r"\d{1,3},\d", g) or (re.fullmatch(r"\d{1,3},\d{2}", g) and re.search(r"\d\s*$", t[:m.start()])):
                continue
            if re.search(r"\d,\d", g):
                bez_razmaka.append(g)
            citati_sirovo.append(g)
            nums = brojevi_iz_citata(g)
            redoslijed.extend(nums)
            if re.search(r"\d\s*-\s*\d", g):
                raspon_spojnica.append(g)
            ns = [int(x) for x in re.findall(r"\d+", g)]
            if len(ns) >= 3 and "-" not in g and "–" not in g and ns == list(range(ns[0], ns[0] + len(ns))):
                nabrajanje_umjesto_raspona.append(g)

    citirani = set(redoslijed)
    stavke = {}
    dupli, prazni = [], []
    for t in popis:
        m = RE_POPIS_STAVKA.match(t)
        if m:
            n = int(m.group(1))
            if n in stavke:
                dupli.append(n)
            stavke[n] = t.strip()
    N = max(stavke) if stavke else 0
    prazni = [n for n in range(1, N + 1) if n not in stavke]

    sirocad = sorted(n for n in stavke if n not in citirani)
    bez_reference = sorted(n for n in citirani if n not in stavke)

    # redoslijed prvog pojavljivanja
    prvi_put, videno = [], set()
    for n in redoslijed:
        if n not in videno:
            videno.add(n)
            prvi_put.append(n)
    skokovi = []
    ocekivan = 1
    for n in prvi_put:
        if n > ocekivan:
            skokovi.append((n, ocekivan))
        ocekivan = max(ocekivan, n + 1) if n >= ocekivan else ocekivan
    return {
        "datoteka": put,
        "popis_stavki": len(stavke), "N_max": N, "popis_prazni_brojevi": prazni, "popis_dupli": sorted(set(dupli)),
        "citata_u_tekstu": len(citati_sirovo), "razlicitih_citiranih": len(citirani),
        "sirocad": sirocad, "citat_bez_reference": bez_reference,
        "prvo_pojavljivanje": prvi_put[:200],
        "skokovi_redoslijeda": skokovi[:20],
        "raspon_sa_spojnicom": raspon_spojnica, "nabrajanje_umjesto_raspona": nabrajanje_umjesto_raspona,
        "citat_bez_razmaka": bez_razmaka,
    }


def ispisi(r):
    print(f"VANCOUVER — {r['datoteka']}")
    print("=" * 56)
    print(f"popis: {r['popis_stavki']} stavki (1..{r['N_max']})"
          + (f"  ❌ prazni brojevi: {r['popis_prazni_brojevi']}" if r['popis_prazni_brojevi'] else "")
          + (f"  ❌ dupli: {r['popis_dupli']}" if r['popis_dupli'] else ""))
    print(f"citata u tekstu: {r['citata_u_tekstu']} · različitih brojeva: {r['razlicitih_citiranih']}")
    print(("✅" if not r['sirocad'] else "❌") + f" siročad (u popisu, necitirano): {r['sirocad'] or 'nema'}")
    print(("✅" if not r['citat_bez_reference'] else "❌") + f" citat bez reference: {r['citat_bez_reference'] or 'nema'}")
    if r['skokovi_redoslijeda']:
        s = r['skokovi_redoslijeda'][0]
        print(f"⚠️ redoslijed prvog pojavljivanja nije uzlazan: prvi skok na {s[0]} (očekivano {s[1]}); ukupno skokova {len(r['skokovi_redoslijeda'])}")
        print(f"   prvih 30 brojeva po pojavljivanju: {r['prvo_pojavljivanje'][:30]}")
    else:
        print("✅ redoslijed prvog pojavljivanja uzlazan")
    if r['raspon_sa_spojnicom']:
        print(f"⚠️ raspon sa spojnicom umjesto en-crtice: {len(r['raspon_sa_spojnicom'])}× npr. {r['raspon_sa_spojnicom'][:5]}")
    if r.get('citat_bez_razmaka'):
        print("⚠️ citat bez razmaka iza zareza (Vancouver: „(67, 68)\"): " + str(len(r['citat_bez_razmaka'])) + "× npr. " + str(r['citat_bez_razmaka'][:6]))
    if r['nabrajanje_umjesto_raspona']:
        print(f"⚠️ tri i više uzastopnih brojeva nabrojani umjesto raspona: {r['nabrajanje_umjesto_raspona'][:5]}")
    print("\nOvo provjerava brojeve, ne sadržaj izvora (za to verify_sources.py). Redoslijed i rasponi su savjet.")


def main(argv=None):
    ap = argparse.ArgumentParser(description="Vancouver (n) citati protiv numeriranog popisa.")
    ap.add_argument("rad")
    ap.add_argument("--od-naslova", help="točan naslov popisa literature (zadano: automatski)")
    ap.add_argument("--json", dest="kao_json")
    a = ap.parse_args(argv)
    r = analiza(a.rad, a.od_naslova)
    ispisi(r)
    if a.kao_json:
        with open(a.kao_json, "w", encoding="utf-8") as f:
            json.dump(r, f, ensure_ascii=False, indent=1)
    return 1 if (r["sirocad"] or r["citat_bez_reference"]) else 0


if __name__ == "__main__":
    sys.exit(main())
