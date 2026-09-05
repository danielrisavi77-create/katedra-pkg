#!/usr/bin/env python3
"""Obranjeni rad kao mjerilo oblika — izmjeri ga, ne procjenjuj.

Zašto postoji
-------------
Željezno pravilo 17 kaže da je **uzorak jači od profila**: rad koji je mentor dao
kao mjerilo mjeri se i upisuje u profil kao `primjerak`, pa gate javlja „odstupa
od primjerka X”, a ne „krši pravilo” za pravilo kojega u službenim Uputama nema.

Pravilo je postojalo, a **nabava uzorka nije**: ovisila je o tome ima li student
slučajno rad koji mu je netko dao. A obranjeni radovi hrvatskih ustanova javno
stoje u repozitorijima (Dabar i repozitoriji fakulteta). Student skine tri rada
sa svojeg odsjeka, ovaj ih alat izmjeri, i pravilo 17 prestaje biti prigodno.

Što ovo NIJE
------------
**Ne skida ništa s interneta.** Rad skida student, alat mjeri datoteku. Ne
ocjenjuje je li uzorak dobar — obranjeni rad nije norma nego **opservacija**, i
tako se i zapisuje.

Granica prema `check_rules.py`
------------------------------
`check_rules` provjerava dokument **protiv profila**. Ovdje se mjeri dokument
**bez profila**, da bi se dobila opservacija koja u profil tek ulazi. Mjerni
primitivi (tema fontova, razina naslova, normalizacija) uzimaju se iz
`check_rules`, ne prepisuju — pravila ostaju ondje gdje jesu.
"""
from __future__ import annotations

import argparse
import json
import os
import re
import sys
from datetime import date

SKRIPTE = os.path.dirname(os.path.abspath(__file__))
if SKRIPTE not in sys.path:
    sys.path.insert(0, SKRIPTE)

import context  # noqa: E402
import hr_text as H  # noqa: E402
from check_rules import (  # noqa: E402
    font_stila,
    norm,
    razina_naslova,
    stil_imena,
    tema_fontovi,
)

EMU_CM = 360000
STANJE = "primjerci.json"

OK, UPOZ, RAZLIKA = "✅", "⚠️", "≠"


def _cm(v):
    return round(v / EMU_CM, 2) if v else None


def _mod(vrijednosti):
    """Najčešća vrijednost — uzorak, ne prosjek: prosjek fonta nema značenje."""
    vrijednosti = [v for v in vrijednosti if v is not None]
    if not vrijednosti:
        return None
    return max(set(vrijednosti), key=vrijednosti.count)


# Stilovi koji NISU tijelo, iako nisu ni naslovi: natpis prikaza, unosi sadržaja
# i popisa prikaza. Znaju biti dulji od 80 znakova i nose vlastitu veličinu pisma
# (natpis 11 pt, unos sadržaja 11 pt), pa upadnu u mjerenje tijela. Na uzorku s
# ocjenom 5 tijelo je imalo 59 odlomaka BEZ izričite veličine (nasljeđuje 12 pt iz
# docDefaults), a jedinih pet odlomaka s izričitom veličinom bili su natpisi od
# 11 pt — mod je zato ispao 11 pt, upisao se u profil kao primjerak i po pravilu
# 17 nadjačao profil, pa je check_rules poslije blokirao rad zbog „11 pt".
_NIJE_TIJELO = ("caption", "table of figures", "toc", "sadržaj", "sadrzaj",
                "header", "footer", "footnote", "endnote", "bibliography",
                "natpis", "opis slike")


def _zadana_pt(d):
    """Veličina pisma iz docDefaults (w:sz je u polovicama točke)."""
    try:
        import re as _re
        xml = d.styles.element.xml
        m = _re.search(r"<w:docDefaults>.*?<w:sz w:val=\"(\d+)\"", xml, _re.S)
        return int(m.group(1)) / 2 if m else None
    except Exception:  # noqa: BLE001
        return None


def _nije_tijelo(stil):
    s = (stil or "").strip().lower()
    return any(s.startswith(x) or s == x for x in _NIJE_TIJELO)


def izmjeri(put):
    """Opservacije s jednog obranjenog rada. Ništa se ne tumači kao pravilo."""
    import docx
    d = docx.Document(put)
    tema = tema_fontovi(put)
    m = {}

    s = d.sections[0]
    m["margine_cm"] = {"gore": _cm(s.top_margin), "dolje": _cm(s.bottom_margin),
                       "lijevo": _cm(s.left_margin), "desno": _cm(s.right_margin)}

    tijelo_font, tijelo_pt, proredi, poravnanja = [], [], [], []
    n1_pt, n1_bold, n2_pt, n2_bold = [], [], [], []
    naslovi1, naslovi2 = [], []

    for p in d.paragraphs:
        t = (p.text or "").strip()
        if not t:
            continue
        razina = razina_naslova(p)
        stil = stil_imena(p)
        pt = None
        for r in p.runs:
            if r.font.size:
                pt = r.font.size.pt
                break
        if pt is None:
            try:
                pt = d.styles[stil].font.size.pt if d.styles[stil].font.size else None
            except Exception:  # noqa: BLE001
                pt = None
        bold = any(r.bold for r in p.runs) or None

        if razina == 1:
            naslovi1.append(t)
            n1_pt.append(pt)
            n1_bold.append(bool(bold))
        elif razina == 2:
            naslovi2.append(t)
            n2_pt.append(pt)
            n2_bold.append(bool(bold))
        elif not razina and len(t) > 80 and not _nije_tijelo(stil):
            # None znači „nasljeđuje iz docDefaults", a ne „nema veličinu":
            # bez ove zamjene odlomci tijela ne ulaze u mod i mjeri se natpis.
            tijelo_pt.append(pt if pt is not None else _zadana_pt(d))
            try:
                # font_stila vraća (ime, veličina) — za primjerak treba samo ime
                f = font_stila(d.styles[stil], tema)
                tijelo_font.append(f[0] if isinstance(f, (tuple, list)) else f)
            except Exception:  # noqa: BLE001
                pass
            pf = p.paragraph_format
            if pf.line_spacing:
                proredi.append(round(float(pf.line_spacing), 2))
            if pf.alignment is not None:
                poravnanja.append(str(pf.alignment).split()[0].lower())

    m["font"] = _mod(tijelo_font)
    m["velicina_pt"] = _mod(tijelo_pt)
    m["prored"] = _mod(proredi)
    m["poravnanje"] = _mod(poravnanja)
    m["naslov_1_pt"] = _mod(n1_pt)
    m["naslov_2_pt"] = _mod(n2_pt)
    m["naslov_1_bold"] = _mod(n1_bold)
    m["naslov_2_bold"] = _mod(n2_bold)
    m["broj_poglavlja_1"] = len(naslovi1)
    m["numerirani_pododjeljci"] = ("X.Y" if any(re.match(r"^\d+\.\d+", x)
                                                for x in naslovi2) else None)
    m["redoslijed_naslova"] = [x[:60] for x in naslovi1][:20]

    # citat u tekstu — uzorak, ne pravilo
    proza = " ".join((p.text or "") for p in d.paragraphs
                     if not razina_naslova(p))
    citati = re.findall(r"\([A-ZČĆŽŠĐ][^()]{2,60}?\d{4}\.?[^()]{0,20}\)", proza)
    m["citat_primjeri"] = citati[:5]
    m["citat_broj"] = len(citati)
    if citati:
        m["citat_dvotocka_pred_stranicom"] = bool(
            re.search(r"\d{4}\s*:\s*\d", " ".join(citati)))
        m["citat_tocka_iza_godine"] = bool(
            re.search(r"\b\d{4}\.\s*[,)]", " ".join(citati)))

    # popis literature — oblik jedinice i uvlaka
    jedinice, u_popisu = [], False
    for p in d.paragraphs:
        t = (p.text or "").strip()
        if not t:
            continue
        if H.NASLOV_LIT.match(t):
            u_popisu = True
            continue
        if u_popisu and razina_naslova(p) == 1:
            break
        if u_popisu and len(t) > 15:
            jedinice.append(p)
    if jedinice:
        m["popis_primjeri"] = [(p.text or "").strip()[:120] for p in jedinice[:3]]
        m["popis_jedinica"] = len(jedinice)
        uvucenih = sum(
            1 for p in jedinice
            if (p.paragraph_format.first_line_indent
                and p.paragraph_format.first_line_indent < 0)
            or (p.paragraph_format.left_indent and p.paragraph_format.left_indent > 0))
        m["uvlaka_u_popisu"] = uvucenih > len(jedinice) / 2
        m["zavrsna_tocka_u_popisu"] = sum(
            1 for p in jedinice if (p.text or "").strip().endswith(".")
        ) > len(jedinice) / 2

    m["rijeci"] = sum(len((p.text or "").split()) for p in d.paragraphs)
    return m


# --------------------------------------------------------------------------- #
# usporedba s profilom — „odstupa od primjerka”, nikad „krši pravilo”
# --------------------------------------------------------------------------- #

PAROVI = [
    ("font", ("format", "font"), "font tijela"),
    ("velicina_pt", ("format", "velicina_pt"), "veličina pisma"),
    ("prored", ("format", "prored"), "prored"),
    ("uvlaka_u_popisu", ("citiranje", "uvlaka_u_popisu"), "uvlaka u popisu"),
]


def _iz_profila(profil, put):
    cvor = profil
    for k in put:
        if not isinstance(cvor, dict) or k not in cvor:
            return None
        cvor = cvor[k]
    return cvor


def usporedi(mjereno, profil):
    """Razlike između izmjerenog uzorka i onoga što profil deklarira.

    Nijedna razlika NIJE kršenje. Obranjeni rad je opservacija; službene Upute su
    norma. Kad se razilaze, obje ostaju zapisane (pravilo 17).
    """
    out = []
    for kljuc, put, naziv in PAROVI:
        u = mjereno.get(kljuc)
        p = _iz_profila(profil, put)
        if u is None or p is None:
            continue
        # Profil dopušteno navodi VIŠE vrijednosti („font: [Times New Roman,
        # Calibri]”). Uzorak koji pogađa bilo koju od njih nije odstupanje.
        if isinstance(p, (list, tuple)):
            podudara = any(str(u).lower() == str(x).lower() for x in p)
        elif isinstance(p, str):
            podudara = (str(u).lower() in p.lower() or p.lower() in str(u).lower())
        else:
            podudara = (u == p)
        if not podudara:
            out.append({"svojstvo": naziv, "uzorak": u, "profil": p})
    m = mjereno.get("margine_cm") or {}
    pm = _iz_profila(profil, ("format", "margine_cm")) or {}
    for strana in ("gore", "dolje", "lijevo", "desno"):
        u, p = m.get(strana), pm.get(strana)
        if u is not None and p is not None and abs(u - p) > 0.06:
            out.append({"svojstvo": f"margina {strana}", "uzorak": u, "profil": p})
    return out


def _ispisi_mjereno(m):
    print("IZMJERENO NA UZORKU")
    print("=" * 20)
    mg = m.get("margine_cm") or {}
    print(f"  margine cm   gore {mg.get('gore')} · dolje {mg.get('dolje')} · "
          f"lijevo {mg.get('lijevo')} · desno {mg.get('desno')}")
    print(f"  tijelo       {m.get('font')} {m.get('velicina_pt')} pt · "
          f"prored {m.get('prored')} · poravnanje {m.get('poravnanje')}")
    print(f"  naslovi      H1 {m.get('naslov_1_pt')} pt"
          f"{' bold' if m.get('naslov_1_bold') else ''} · "
          f"H2 {m.get('naslov_2_pt')} pt"
          f"{' bold' if m.get('naslov_2_bold') else ''} · "
          f"pododjeljci {m.get('numerirani_pododjeljci') or '—'}")
    print(f"  opseg        {m.get('rijeci')} riječi · "
          f"{m.get('broj_poglavlja_1')} naslova prve razine")
    if m.get("citat_primjeri"):
        print(f"  citat        {m['citat_broj']} u tekstu; primjeri: "
              f"{' · '.join(m['citat_primjeri'][:3])}")
    if m.get("popis_primjeri"):
        print(f"  popis        {m.get('popis_jedinica')} jedinica · "
              f"uvlaka {m.get('uvlaka_u_popisu')} · "
              f"završna točka {m.get('zavrsna_tocka_u_popisu')}")
        for x in m["popis_primjeri"]:
            print(f"               {x}")
    if m.get("redoslijed_naslova"):
        print("  redoslijed   " + " → ".join(m["redoslijed_naslova"][:8]))


def cmd_izmjeri(args):
    m = izmjeri(args.rad)
    _ispisi_mjereno(m)
    if args.profil:
        profil = json.load(open(args.profil, encoding="utf-8"))
        raz = usporedi(m, profil)
        print("\nRAZLIKE PREMA PROFILU")
        print("=" * 21)
        if not raz:
            print(f"  {OK} uzorak se slaže s profilom u svemu što se dade usporediti")
        for r in raz:
            print(f"  {RAZLIKA} {r['svojstvo']}: uzorak {r['uzorak']} · "
                  f"profil {r['profil']}")
        if raz:
            print("\n  Razlika NIJE kršenje. Obranjeni rad je opservacija, službene")
            print("  Upute su norma. Kad se razilaze, obje ostaju zapisane, a gate")
            print("  javlja „odstupa od primjerka X” (željezno pravilo 17).")
    if args.kao_json:
        with open(args.kao_json, "w", encoding="utf-8") as f:
            json.dump(m, f, ensure_ascii=False, indent=1)
    return 0


def cmd_upisi(args):
    m = izmjeri(args.rad)
    zapis = {
        "id": args.id or (os.path.splitext(os.path.basename(args.rad))[0]
                          .lower().replace(" ", "-")),
        "vrsta": args.vrsta,
        "izvor": args.izvor,
        "zabiljezeno": date.today().isoformat(),
        "mjereno": m,
    }
    put = context.resolve_state_file(STANJE, kat=args.kat,
                                     project_root=args.project_root)
    postojeci = []
    if os.path.exists(put):
        with open(put, encoding="utf-8") as f:
            postojeci = (json.load(f) or {}).get("primjerci") or []
    postojeci = [x for x in postojeci if x.get("id") != zapis["id"]]
    postojeci.append(zapis)
    context.atomic_write_json(put, {"schema_version": 1, "primjerci": postojeci})
    print(f"✅ {put} — {len(postojeci)} primjeraka")
    print("\nZapis za profil fakulteta (`/primjerci`) — dodaje ga MAINTAINER u")
    print("references/fakulteti/<slug>.json; runtime ne mijenja instalirani skill:")
    print(json.dumps(zapis, ensure_ascii=False, indent=1)[:1400])
    return 0


def cmd_popis(args):
    put = context.resolve_state_file(STANJE, kat=args.kat,
                                     project_root=args.project_root)
    if not os.path.exists(put):
        print("➖ nema izmjerenih primjeraka. Skini 2–3 obranjena rada sa svojeg")
        print("   odsjeka (Dabar ili repozitorij fakulteta) i pokreni:")
        print("   primjerci.py upisi rad.docx --vrsta diplomski --izvor \"…\"")
        return 0
    with open(put, encoding="utf-8") as f:
        pr = (json.load(f) or {}).get("primjerci") or []
    print(f"PRIMJERCI ({len(pr)})")
    print("=" * 14)
    for x in pr:
        m = x.get("mjereno") or {}
        print(f"  {x['id']} · {x.get('vrsta')} · {x.get('zabiljezeno')}")
        print(f"     {x.get('izvor')}")
        print(f"     {m.get('font')} {m.get('velicina_pt')} pt · prored "
              f"{m.get('prored')} · {m.get('rijeci')} riječi")
    if len(pr) == 1:
        print("\n⚠️  jedan primjerak nije uzorak. Dva rada s istog odsjeka koja se")
        print("    slažu jači su dokaz od jednoga; kad se razilaze, to je samo po")
        print("    sebi nalaz — znači da kućni stil ondje nije ustaljen.")
    return 0


def main(argv=None) -> int:
    ap = argparse.ArgumentParser(
        description="Obranjeni rad kao mjerilo oblika (željezno pravilo 17).")
    ap.add_argument("--kat")
    ap.add_argument("--project-root", dest="project_root")
    sub = ap.add_subparsers(dest="naredba", required=True)

    i = sub.add_parser("izmjeri", help="izmjeri rad i usporedi s profilom")
    i.add_argument("rad", help=".docx obranjenog rada")
    i.add_argument("--profil", help="resolved_profile.json za usporedbu")
    i.add_argument("--json", dest="kao_json")
    i.set_defaults(f=cmd_izmjeri)

    u = sub.add_parser("upisi", help="izmjeri i upiši u .katedra/primjerci.json")
    u.add_argument("rad")
    u.add_argument("--vrsta", required=True,
                   choices=("seminarski", "zavrsni", "diplomski"))
    u.add_argument("--izvor", required=True,
                   help="odakle je rad (repozitorij, mentor, odsjek + godina)")
    u.add_argument("--id")
    u.set_defaults(f=cmd_upisi)

    p = sub.add_parser("popis", help="izmjereni primjerci u projektu")
    p.set_defaults(f=cmd_popis)

    args = ap.parse_args(argv)
    if getattr(args, "rad", None) and not os.path.exists(args.rad):
        print(f"❌ nema datoteke: {args.rad}", file=sys.stderr)
        return 2
    try:
        return args.f(args)
    except Exception as e:  # noqa: BLE001
        print(f"❌ mjerenje nije uspjelo: {type(e).__name__}: {e}", file=sys.stderr)
        return 2


if __name__ == "__main__":
    sys.exit(main())
