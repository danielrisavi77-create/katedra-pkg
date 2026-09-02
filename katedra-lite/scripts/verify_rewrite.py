#!/usr/bin/env python3
"""
verify_rewrite.py — dokazuje da prepisivanje nije izgubilo sadržaj.

Ovo je ono što uopće čini sigurnim delegiranje prepisivanja subagentu. Bez
njega se pravilo „sve verificiraj neovisno" ne može provesti — nema čime.

Uspoređuje MULTISKUPOVE (ne skupove — dvostruko navođenje iste brojke mora
ostati dvostruko):

    brojke      svaki brojčani token
    citati      dialect-aware reference fingerprint (author-year / IEEE / legal-footnote)
    markeri     naslovi i natpisi prikaza, doslovno i istim redoslijedom
    rečenice    nijedna ne smije nestati (izmijenjene se prijavljuju posebno)
    odlomci     broj se smije smanjiti samo ako je to bila svrha zahvata

NAČIN ZAHVATA određuje što je blokirajuće — mora odgovarati koraku pipelinea:

    --zahvat stil         korak 1-2: preoblikovanje dopušteno, riječi se smiju
                          mijenjati; brojke, citati i markeri moraju ostati
    --zahvat lomljenje    korak 3: rečenice se smiju dijeliti, ali skup RIJEČI
                          mora ostati (osim veznika koji postaje prilog)
    --zahvat geometrija   korak 4: rečenice moraju ostati DOSLOVNE, samo se
                          odlomci spajaju

    python3 <KATEDRA_SKILL>/scripts/verify_rewrite.py prije.md poslije.md --zahvat stil --profil .katedra/resolved_profile.json \
      --evidence-gate --claims .katedra/claims.jsonl --evidence .katedra/evidence.jsonl
    python3 <KATEDRA_SKILL>/scripts/verify_rewrite.py rad_prije.docx rad_poslije.docx --zahvat geometrija \
      --require-snapshot --evidence-gate --project-root .

Izlazni kod 1 = blokirajući nalaz, prepisano NE primjenjuj.
"""
import argparse
import difflib
import hashlib
import json
import math
import os
import re
import sys
import unicodedata
import zipfile
from pathlib import Path
from collections import Counter

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import hr_text as H
import citation_dialects as C
from context import resolve_state_dir
from evidence_gate import evaluate_files



def _sha256(path):
    h = hashlib.sha256()
    with open(path, "rb") as fh:
        for chunk in iter(lambda: fh.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def snapshot_status(before_paths, kat):
    versions = Path(kat) / "verzije.json"
    if not versions.is_file():
        return False, f"snapshot metadata ne postoji: {versions}"
    try:
        payload = json.loads(versions.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        return False, f"snapshot metadata nije čitljiva: {versions}: {exc}"
    snapshots = payload.get("snapshoti")
    if not isinstance(snapshots, list):
        return False, f"snapshot metadata nema listu `snapshoti`: {versions}"
    known = {}
    for x in snapshots:
        if isinstance(x, dict):
            known.setdefault(str(x.get("sha256") or ""), []).append(x)
    missing = []
    for path in before_paths:
        try:
            digest = _sha256(path)
        except OSError as exc:
            return False, f"ne mogu izračunati snapshot hash za {path}: {exc}"
        if digest not in known:
            missing.append(str(path))
            continue
        # v1.1-fix D15: hash u verzije.json nije dokaz da rollback postoji —
        # snimljena datoteka mora biti na disku i sadržajno ista (isto što --popis
        # već provjerava). Bez toga gate potvrđuje snapshot kojeg nema.
        ok, poruka = _snapshot_datoteka_ok(known[digest], kat, digest)
        if not ok:
            return False, poruka
    if missing:
        return False, ("nema snapshot hasha za pre-rewrite datoteku/e: " + ", ".join(missing)
                       + "\n                 Što napraviti: napravi snapshot prije prepisivanja "
                         "(diff_versions.py --snapshot ./rad.docx --biljeska \"prije zahvata\").")
    return True, f"snapshot potvrđen za {len(before_paths)} pre-rewrite datoteku/e"


def _snapshot_datoteka_ok(zapisi, kat, digest):
    """v1.1-fix D15: postoji li snimljena datoteka i je li joj sadržaj netaknut."""
    problemi = []
    for zapis in zapisi:
        rel = str(zapis.get("datoteka") or "")
        vid = str(zapis.get("id") or "?")
        if not rel:
            problemi.append(f"snapshot {vid} nema zapisanu datoteku u verzije.json")
            continue
        put = os.path.join(kat, rel)
        if not os.path.isfile(put):
            problemi.append(f"snapshot {vid} je zapisan, ali datoteka nedostaje: {put}")
            continue
        try:
            if _sha256(put) != digest:
                problemi.append(f"snapshot {vid} ne odgovara zapisanom sha256: {put}")
                continue
        except OSError as exc:
            problemi.append(f"snapshot {vid} se ne može pročitati: {put}: {exc}")
            continue
        return True, ""
    return False, (problemi[0] if problemi else "snapshot datoteka nije pronađena") + \
        "\n                 Što napraviti: ponovi snapshot prije prepisivanja " \
        "(diff_versions.py --snapshot ./rad.docx --biljeska \"prije zahvata\") ili vrati " \
        "mapu .katedra/verzije/ iz sigurnosne kopije."


def run_evidence_gate(claims, evidence, sources=None):
    try:
        payload = evaluate_files(claims, evidence, sources_path=sources, policy="strict")
    except (OSError, ValueError) as exc:
        return False, f"evidence gate structural error: {exc}"
    blocked = payload.get("summary", {}).get("blocked", 0)
    if not payload.get("passed"):
        return False, f"evidence gate blokira {blocked} claim(s)"
    return True, f"evidence gate PASS: {payload.get('summary', {}).get('claims', 0)} claim(s)"

def brojke(t):
    """Svi brojčani tokeni, bez pratećih . i , (redni broj vs. decimala)."""
    return Counter(m.group(0).rstrip(".,") for m in re.finditer(r"\d[\d.,]*", t))


def prezimena(t):
    return Counter(k[0] for k in H.kljucevi_citata(t).elements())


def norm_rec(r):
    return re.sub(r"\s+", " ", r).strip()


def _nfc(s):
    """v1.1-fix Q19: NFC vs. NFD je vizualno isti tekst — usporedba se radi nad NFC."""
    return unicodedata.normalize("NFC", s or "")


def _nfc_counter(c):
    out = Counter()
    for k, v in c.items():
        out[_nfc(k)] += v
    return out


def _docx_dijelovi(put):
    """v1.1-fix D2: dijelovi .docx-a koje hr_text._iz_docx ne vraća.

    Vraća (odlomci, ćelije tablica, tekstovi fusnota). Čim čitač u hr_text.py počne
    vraćati tablice, `_dodatne_celije` ih prestaje dodavati (dedup), pa se ovo može
    ukloniti.
    """
    if not str(put).lower().endswith(".docx"):
        return [], [], []
    try:
        from docx import Document
        d = Document(put)
    except Exception:
        return [], [], []
    odlomci = [_nfc((p.text or "").strip()) for p in d.paragraphs]
    odlomci = [t for t in odlomci if t]
    celije = []
    try:
        for t in d.tables:
            for red in t.rows:
                for c in red.cells:
                    tekst = _nfc((c.text or "").strip())
                    if tekst:
                        celije.append(tekst)
    except Exception:
        pass
    try:
        fusnote = [_nfc(x) for x in C.extract_docx_footnotes(put).values() if x]
    except Exception:
        fusnote = []
    return odlomci, celije, fusnote


def _dodatne_celije(celije, odlomci, markeri):
    """Ćelije koje učitani odlomci/markeri već ne sadrže (izbjegava dvostruko brojanje)."""
    poznato = Counter(norm_rec(x) for x in list(odlomci) + list(markeri))
    dodatak = []
    for c in celije:
        k = norm_rec(c)
        if poznato.get(k):
            poznato[k] -= 1
            continue
        dodatak.append(c)
    return dodatak


def _mediji(put):
    """v1.1-fix D2: multiskup ugrađenih prikaza (slike, grafikoni, objekti) po sadržaju.

    Obrisan grafikon ostavlja natpis i „Izvor:" bez prikaza; tekstualna usporedba
    to ne vidi, pa se dijelovi broje i hashiraju.
    """
    if not str(put).lower().endswith(".docx"):
        return Counter()
    c = Counter()
    try:
        with zipfile.ZipFile(put) as z:
            for ime in z.namelist():
                if ime.startswith(("word/media/", "word/charts/", "word/embeddings/")):
                    c[hashlib.sha256(z.read(ime)).hexdigest()[:16]] += 1
    except (OSError, KeyError, zipfile.BadZipFile):
        return Counter()
    return c


def _tokeni(s):
    return set(re.findall(r"\w{4,}", s.lower()))


def _indeks_rijeci(kandidati):
    """v1.1-fix D3: obrnuti indeks dužih riječi — bez njega usporedba svake nestale
    rečenice sa svakom preostalom traje minutama na radu od 60 stranica."""
    indeks, duljine = {}, []
    for j, s in enumerate(kandidati):
        tokeni = _tokeni(s)
        duljine.append(len(tokeni) or 1)
        for w in tokeni:
            indeks.setdefault(w, []).append(j)
    return indeks, duljine


def _najbliza(s, kandidati, indeks, koliko=20):
    """Najbliža preostala rečenica (difflib, cutoff 0.75) uz jeftin predizbor.

    Predizbor rangira po zajedničkim riječima (rjeđa riječ nosi više) i tek najboljih
    `koliko` šalje u difflib. Promašiti može samo rečenicu koja ne dijeli nijednu dužu
    riječ; takva se prijavljuje kao IZGUBLJENA — smjer pogreške je na stranu opreza.
    """
    postings, duljine = indeks
    tokeni = _tokeni(s)
    if not tokeni:
        suzeni = kandidati
    else:
        bodovi = Counter()
        for t in tokeni:
            popis = postings.get(t, ())
            tezina = 1.0 / math.log(2 + len(popis))
            for j in popis:
                bodovi[j] += tezina
        najbolji = sorted(bodovi.items(),
                          key=lambda kv: -kv[1] / (len(tokeni) + duljine[kv[0]]))[:koliko]
        suzeni = [kandidati[j] for j, _ in najbolji]
    return difflib.get_close_matches(s, suzeni, n=1, cutoff=0.75) if suzeni else []


def _citati(put, stil, odlomci, celije, fusnote):
    """Citatni fingerprint nad NFC tekstom: odlomci + ćelije tablica + fusnote.

    v1.1-fix D2/Q19: parser se ne dira (B10), ali mu se predaje NFC-normaliziran tekst
    i dijelovi koje čitač preskače — pod autor-godina/IEEE fusnote nitko nije čitao,
    a citat iz tablice ili fusnote inače nestane bez traga.
    """
    p = Path(put)
    if p.suffix.lower() != ".docx":
        try:
            tekst = _nfc(p.read_text(encoding="utf-8"))
        except OSError:
            return Counter()
        return _nfc_counter(C.citation_fingerprint_text(tekst, stil))
    if C.resolve_dialect(stil) != "legal-footnote":
        tekst = _nfc(" ".join(odlomci + celije + fusnote))
        return _nfc_counter(C.citation_fingerprint_text(tekst, stil))
    # legal-footnote: fusnote i tipizaciju izvora radi citation_dialects nad datotekom
    c = C.citation_fingerprint_file(p, stil)
    if celije:
        c = c + C.citation_fingerprint_text(_nfc(" ".join(celije)), stil)
    return _nfc_counter(c)


def usporedi(prije_p, poslije_p, zahvat='geometrija', citatni_stil='autor-godina'):
    o_a, m_a = H.ucitaj(prije_p, ukljuci_tablice=True)
    o_b, m_b = H.ucitaj(poslije_p, ukljuci_tablice=True)
    o_a, m_a = [_nfc(x) for x in o_a], [_nfc(x) for x in m_a]
    o_b, m_b = [_nfc(x) for x in o_b], [_nfc(x) for x in m_b]
    # v1.1-fix D2: sadržaj tablica ulazi u usporedbu (brojke, citati, rečenice)
    par_a, cel_a, fus_a = _docx_dijelovi(prije_p)
    par_b, cel_b, fus_b = _docx_dijelovi(poslije_p)
    dod_a = _dodatne_celije(cel_a, o_a, m_a)
    dod_b = _dodatne_celije(cel_b, o_b, m_b)
    tijelo_a, tijelo_b = o_a + dod_a + fus_a, o_b + dod_b + fus_b
    t_a, t_b = " ".join(tijelo_a), " ".join(tijelo_b)
    nalazi = []

    # --- markeri: doslovno i istim redoslijedom
    if m_a != m_b:
        raz = [(x, y) for x, y in zip(m_a, m_b) if x != y]
        nalazi.append(("x", f"markeri (naslovi/natpisi) se razlikuju: "
                            f"{len(m_a)} → {len(m_b)}, {len(raz)} izmijenjenih",
                       [f"- {x[:80]}\n      + {y[:80]}" for x, y in raz[:5]]))
    else:
        nalazi.append(("ok", f"markeri: {len(m_a)} doslovno identičnih", []))

    # --- brojke
    ba, bb = brojke(t_a), brojke(t_b)
    izg, dod = ba - bb, bb - ba
    if izg or dod:
        det = ([f"izgubljeno: {dict(list(izg.items())[:12])}"] if izg else []) + \
              ([f"dodano: {dict(list(dod.items())[:12])}"] if dod else [])
        nalazi.append(("x", f"brojke odstupaju ({sum(ba.values())} → {sum(bb.values())})", det))
    else:
        nalazi.append(("ok", f"brojke: {sum(ba.values())} tokena, multiskup identičan", []))

    # --- citati (B10: dialect-aware fingerprint)
    ca = _citati(prije_p, citatni_stil, par_a, cel_a, fus_a)
    cb = _citati(poslije_p, citatni_stil, par_b, cel_b, fus_b)
    izg_c, dod_c = ca - cb, cb - ca
    if izg_c or dod_c:
        det = ([f"izgubljeno: {dict(izg_c)}"] if izg_c else []) + \
              ([f"dodano: {dict(dod_c)}"] if dod_c else [])
        nalazi.append(("x", f"citati odstupaju ({sum(ca.values())} → {sum(cb.values())})", det))
    else:
        nalazi.append(("ok", f"citati: {sum(ca.values())} referenci, identično [{citatni_stil}]", []))

    # --- ugrađeni prikazi (v1.1-fix D2: obrisan grafikon ostavlja natpis bez slike)
    ma, mb = _mediji(prije_p), _mediji(poslije_p)
    if ma or mb:
        izg_m, dod_m = ma - mb, mb - ma
        if izg_m:
            nalazi.append(("x", f"ugrađeni prikazi nedostaju "
                                f"({sum(ma.values())} → {sum(mb.values())}), "
                                f"{sum(izg_m.values())} izgubljenih",
                           ["Natpis i „Izvor:” ostaju bez prikaza.",
                            "Što napraviti: vrati prikaz iz snapshota "
                            "(diff_versions.py --popis pa --vrati)."]))
        elif dod_m:
            nalazi.append(("!", f"ugrađenih prikaza ima više nego prije "
                                f"({sum(ma.values())} → {sum(mb.values())})", []))
        else:
            nalazi.append(("ok", f"ugrađeni prikazi: {sum(ma.values())}, "
                                 f"sadržajno identični", []))

    # --- riječi (blokirajuće samo pri lomljenju rečenica)
    if zahvat == "lomljenje":
        wa = Counter(w.lower().strip(".,;:()„”") for w in H.rijeci(t_a))
        wb = Counter(w.lower().strip(".,;:()„”") for w in H.rijeci(t_b))
        izg_w = {k: v for k, v in (wa - wb).items() if len(k) > 3}
        if izg_w:
            nalazi.append(("x", f"izgubljene riječi pri lomljenju: {len(izg_w)}",
                           [", ".join(list(izg_w)[:15])]))
        else:
            nalazi.append(("ok", f"riječi: {sum(wa.values())} → {sum(wb.values())}, "
                                 f"ništa sadržajno nije izgubljeno", []))

    # --- rečenice
    ra = Counter(norm_rec(r) for o in tijelo_a for r in H.recenice(o))
    rb = Counter(norm_rec(r) for o in tijelo_b for r in H.recenice(o))
    nestale = ra - rb
    if nestale:
        # v1.1-fix D3: odluka se donosi nad CIJELIM multiskupom, ne nad prvih pet
        # ispisanih primjera — inače izgubljena šesta rečenica nikad ne blokira
        kandidati = list(rb)
        indeks = _indeks_rijeci(kandidati)
        izmijenjene, izgubljene = [], []
        for s, n in nestale.items():
            blizu = _najbliza(s, kandidati, indeks)
            (izmijenjene if blizu else izgubljene).extend(
                [(s, blizu[0]) if blizu else s] * n)
        det = [f"IZGUBLJENA: {s[:110]}" for s in izgubljene[:5]]
        if len(izgubljene) > 5:
            det.append(f"… još {len(izgubljene) - 5} izgubljenih (ispis je skraćen)")
        det += [f"IZMIJENJENA:\n      - {s[:95]}\n      + {b[:95]}"
                for s, b in izmijenjene[:5]]
        if zahvat == "geometrija":
            # blokira samo stvarni gubitak; jako preoblikovana rečenica (bliska
            # podudarnost) ostaje upozorenje da „stil" ne bi lažno padao
            razina = "x" if izgubljene else "!"
        else:
            razina = "info"      # preoblikovanje je svrha zahvata
        nalazi.append((razina, f"{sum(nestale.values())} rečenica nije doslovno "
                               f"preneseno ({sum(ra.values())} → {sum(rb.values())}): "
                               f"{len(izmijenjene)} izmijenjenih, "
                               f"{len(izgubljene)} izgubljenih",
                       det if razina != "info" else det[:2]))
    else:
        nalazi.append(("ok", f"rečenice: svih {sum(ra.values())} doslovno preneseno", []))

    # --- odlomci
    if len(o_a) != len(o_b):
        spajanje_ok = zahvat == "geometrija" and len(o_b) < len(o_a)
        razina = "ok" if spajanje_ok else "x"
        poruka = f"odlomaka {len(o_a)} → {len(o_b)}"
        poruka += " (spajanje je svrha zahvata)" if spajanje_ok else " — broj se promijenio"
        nalazi.append((razina, poruka, []))
    else:
        nalazi.append(("ok", f"odlomaka: {len(o_a)}, nepromijenjeno", []))

    # --- higijena
    ravni = t_b.count('"')
    if ravni:
        nalazi.append(("!", f"ravni navodnici u novoj verziji: {ravni}",
                       ["Hrvatski su „ (U+201E) i ” (U+201D)."]))
    return nalazi


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("datoteke", nargs="+", help="parovi: prije poslije [prije poslije ...]")
    ap.add_argument("--zahvat", choices=["stil", "lomljenje", "geometrija"],
                    default="geometrija",
                    help="korak pipelinea — određuje što je blokirajuće (v. docstring)")
    ap.add_argument("--citatni-stil", choices=C.CITATION_STYLES,
                    help="citatni stil; nadjačava profil.citiranje.stil")
    ap.add_argument("--profil", help="resolved profile JSON s citiranje.stil")
    ap.add_argument("--evidence-gate", action="store_true",
                    help="strict B13 claim/evidence gate prije prihvaćanja rewritea")
    ap.add_argument("--claims", help="claim ledger JSONL; default <state>/.katedra/claims.jsonl")
    ap.add_argument("--evidence", help="evidence ledger JSONL; default <state>/.katedra/evidence.jsonl")
    ap.add_argument("--sources", help="optional verify_sources JSON; default <state>/.katedra/izvori.json ako postoji")
    ap.add_argument("--require-snapshot", action="store_true",
                    help="blokiraj rewrite ako pre-rewrite hash nije u verzije.json")
    ap.add_argument("--kat", help="eksplicitna .katedra mapa")
    ap.add_argument("--project-root", help="korijen projekta za .katedra state")
    args = ap.parse_args()
    if len(args.datoteke) % 2:
        sys.exit("Datoteke se navode u parovima: prije poslije")

    try:
        citatni_stil, _, _ = C.resolve_style(args.citatni_stil, args.profil)
    except C.CitationDialectError as exc:
        sys.exit(f"❌ {exc}")

    kat = resolve_state_dir(args.kat, args.project_root)
    before_paths = [args.datoteke[i] for i in range(0, len(args.datoteke), 2)]

    if args.require_snapshot:
        ok, message = snapshot_status(before_paths, kat)
        print(f"SNAPSHOT GATE: {'✓' if ok else '✗'} {message}")
        if not ok:
            print("REZULTAT: ✗ snapshot precondition nije zadovoljen — NE primjenjuj prepisano")
            return 1

    if args.evidence_gate:
        claims = args.claims or os.path.join(kat, "claims.jsonl")
        evidence = args.evidence or os.path.join(kat, "evidence.jsonl")
        sources = args.sources
        if not sources:
            candidate = os.path.join(kat, "izvori.json")
            sources = candidate if os.path.isfile(candidate) else None
        ok, message = run_evidence_gate(claims, evidence, sources)
        print(f"EVIDENCE GATE: {'✓' if ok else '✗'} {message}")
        if not ok:
            print("REZULTAT: ✗ evidence precondition nije zadovoljen — NE primjenjuj prepisano")
            return 1

    blokira = 0
    gubitak = 0
    for i in range(0, len(args.datoteke), 2):
        a, b = args.datoteke[i], args.datoteke[i + 1]
        print("=" * 72)
        print(f"{os.path.basename(a)}  →  {os.path.basename(b)}   [zahvat: {args.zahvat}]")
        print("=" * 72)
        for razina, poruka, det in usporedi(a, b, args.zahvat, citatni_stil):
            znak = {"ok": "✓", "!": "⚠", "x": "✗", "info": "·"}[razina]
            print(f"  {znak} {poruka}")
            for d in det:
                print(f"      {d}")
                if d.startswith("IZGUBLJENA"):
                    gubitak += 1
            if razina == "x":
                blokira += 1
        print()

    print("=" * 72)
    # v1.1-fix: verdikt se ne smije u istom dahu pozvati na „sadržaj očuvan" i
    # prijaviti izgubljenu rečenicu — brojač blokada je odvojen od tvrdnje o sadržaju
    if blokira:
        print(f"REZULTAT: ✗ {blokira} blokirajućih nalaza — NE primjenjuj prepisano")
    elif gubitak:
        print("REZULTAT: ⚠ nema blokirajućih nalaza, ali rečenice su prijavljene kao "
              "IZGUBLJENE — provjeri ih prije primjene")
    else:
        print("REZULTAT: ✓ nema blokirajućih nalaza — prepisano je sigurno primijeniti")
    return 1 if blokira else 0


if __name__ == "__main__":
    sys.exit(main())
