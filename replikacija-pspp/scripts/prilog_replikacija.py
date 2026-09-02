# -*- coding: utf-8 -*-
"""Prilozi o replikaciji u predajni dokument, po kućnom stilu.

Prilog 2 je tablica „ovako piše u radu, ovako ispisuje PSPP, poklapa se”.
Prilog 3 su snimke prozora programa. Oboje se umeće ispred zadanog sidra, a
oblikovanje se preslikava iz samog dokumenta pa se kućni stil ne mora opisivati
dvaput.

    python3 prilog_replikacija.py --conf replikacija.json [--dokument rad.docx]

Ništa se u dokumentu ne mijenja osim onoga što se dodaje. Prije spremanja se
provjerava da nijedan zahvat ne pregazi sidro fusnote.
"""
import argparse
import csv
import json
import os
import re
import shutil
import sys

import docx
from PIL import Image, ImageChops, ImageOps
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ZADANO = {
    "sidro": "Sažetak",
    "naslov_prilog2": "Prilog 2. Replikacija izračuna u programu PSPP",
    "naslov_prilog3": "Prilog 3. Ispis analiza iz programa PSPP",
    "natpis_tablice": "Tablica P2.1. Usporedba vrijednosti iz rada i iz programa PSPP",
    "izvor_tablice": "Izvor: izradio autor.",
    "izvor_slike": "Izvor: snimka prozora programa PSPP.",
    "zaglavlje": ["Oznaka", "Poglavlje", "Statistika", "U radu",
                  "Naredba u PSPP-u", "Iz PSPP-a", "Slaže se"],
    "sirine_cm": [1.7, 1.8, 4.2, 1.7, 3.1, 1.9, 1.5],
    "sirina_teksta_cm": 15.92,
    "dpi_min": 200,
    "pt_natpis": 11,
    "pt_izvor": 10,
    "pt_tablica": 10,
    "rub_slike": [169, 179, 191],
}


# ── slike ───────────────────────────────────────────────────────────────────
def obrezi(sl):
    """Ukloni jednoličan rub. Referenca je boja gornjeg lijevog piksela."""
    sl = sl.convert("RGB")
    razlika = ImageChops.difference(sl, Image.new("RGB", sl.size, sl.getpixel((0, 0))))
    okv = razlika.getbbox()
    if okv is None:
        return sl
    l, g, d, dn = okv
    return sl.crop((max(0, l - 2), max(0, g - 2),
                    min(sl.width, d + 2), min(sl.height, dn + 2)))


def uokviri(sl, rub):
    """Bijeli razmak pa tanak sivi rub; u tisku slika inače pliva po stranici."""
    return ImageOps.expand(ImageOps.expand(sl, border=6, fill=(255, 255, 255)),
                           border=1, fill=tuple(rub))


# ── oblikovanje ─────────────────────────────────────────────────────────────
def kao_tijelo(p, uzor):
    p.style = uzor.style
    pf, up = p.paragraph_format, uzor.paragraph_format
    for a in ("alignment", "line_spacing", "space_after", "space_before",
              "first_line_indent", "left_indent"):
        setattr(pf, a, getattr(up, a))


def bez_uvlake(p, prije=0, poslije=2, prored=1.0):
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(0)
    pf.space_before = Pt(prije)
    pf.space_after = Pt(poslije)
    pf.line_spacing = prored
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT


def nosi_fusnotu(p):
    return any(r._element.findall(qn("w:footnoteReference")) for r in p.runs)


# ── Prilog 2 ────────────────────────────────────────────────────────────────
def tablica_usporedbe(d, sidro, redci, N, uzor_tijelo):
    zag = N["zaglavlje"]
    sirine = [Cm(x) for x in N["sirine_cm"]]
    tbl = d.add_table(rows=1 + len(redci), cols=len(zag))
    try:
        tbl.style = d.tables[0].style
    except Exception:
        pass

    polja = ("oznaka", "gdje_u_radu", "statistika", "vrijednost_u_radu",
             "pspp_naredba", "vrijednost_iz_pspp", "poklapa_se")
    for j, h in enumerate(zag):
        tbl.rows[0].cells[j].text = h
    for i, red in enumerate(redci, 1):
        for j, k in enumerate(polja):
            tbl.rows[i].cells[j].text = str(red.get(k, "") or "")

    # Bez zadanih širina Word razvuče stupac s nazivom naredbe i svaki redak
    # rastegne na dva reda.
    tbl.autofit = False
    lay = OxmlElement("w:tblLayout")
    lay.set(qn("w:type"), "fixed")
    tbl._tbl.tblPr.append(lay)
    grid = tbl._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for gc, s in zip(grid.findall(qn("w:gridCol")), sirine):
            gc.set(qn("w:w"), str(int(s.twips)))

    for ri, row in enumerate(tbl.rows):
        for c, s in zip(row.cells, sirine):
            c.width = s
        trPr = row._tr.get_or_add_trPr()
        cs = OxmlElement("w:cantSplit")
        cs.set(qn("w:val"), "true")
        trPr.append(cs)
        if ri == 0:                       # zaglavlje se ponavlja na svakoj stranici
            th = OxmlElement("w:tblHeader")
            th.set(qn("w:val"), "true")
            trPr.append(th)
        for c in row.cells:
            for pp in c.paragraphs:
                bez_uvlake(pp, 2, 2)
                for r in pp.runs:
                    r.font.size = Pt(N["pt_tablica"])
                    r.font.bold = (ri == 0)
    sidro._p.addprevious(tbl._tbl)


def prilog2(d, sidro, N, uzor_tijelo, uzor_naslov, usporedba):
    with open(usporedba, encoding="utf-8") as f:
        redci = list(csv.DictReader(f))
    slaze = sum(1 for r in redci if r["poklapa_se"] == "da")

    nas = sidro.insert_paragraph_before(N["naslov_prilog2"])
    nas.style = uzor_naslov.style
    nas.paragraph_format.page_break_before = True

    for t in N.get("uvod_prilog2", []):
        kao_tijelo(sidro.insert_paragraph_before(
            t.replace("{ukupno}", str(len(redci))).replace("{slaze}", str(slaze))),
            uzor_tijelo)

    natpis = sidro.insert_paragraph_before(N["natpis_tablice"])
    kao_tijelo(natpis, uzor_tijelo)
    bez_uvlake(natpis, 10, 4)
    natpis.paragraph_format.keep_with_next = True
    for r in natpis.runs:
        r.font.size = Pt(N["pt_natpis"])
        r.font.bold = True

    tablica_usporedbe(d, sidro, redci, N, uzor_tijelo)

    izv = sidro.insert_paragraph_before(N["izvor_tablice"])
    kao_tijelo(izv, uzor_tijelo)
    bez_uvlake(izv, 2, 12)
    for r in izv.runs:
        r.font.size = Pt(N["pt_izvor"])
        r.font.italic = False             # izvor ide običnim slogom

    if N.get("zakljucak_prilog2"):
        kao_tijelo(sidro.insert_paragraph_before(
            N["zakljucak_prilog2"].replace("{ukupno}", str(len(redci)))
            .replace("{slaze}", str(slaze))), uzor_tijelo)
    return len(redci), slaze


# ── Prilog 3 ────────────────────────────────────────────────────────────────
def prilog3(d, sidro, N, uzor_tijelo, uzor_naslov, mapa, natpisi, radna):
    snimke = sorted(f for f in os.listdir(mapa) if f.lower().endswith((".png", ".jpg")))
    if not snimke:
        return 0, ["nema nijedne snimke"]

    obradene = os.path.join(radna, "_slike")
    os.makedirs(obradene, exist_ok=True)

    nas = sidro.insert_paragraph_before(N["naslov_prilog3"])
    nas.style = uzor_naslov.style
    nas.paragraph_format.page_break_before = True
    if N.get("uvod_prilog3"):
        kao_tijelo(sidro.insert_paragraph_before(N["uvod_prilog3"]), uzor_tijelo)

    sirina_teksta = Cm(N["sirina_teksta_cm"])
    upozorenja, vidjeni = [], set()
    for i, ime in enumerate(snimke, 1):
        kljuc = (re.match(r"(\d+)", ime) or re.match(r"()", ime)).group(1)
        opis = natpisi.get(kljuc) or re.sub(r"^\d+[_\-]?", "", os.path.splitext(ime)[0]).replace("_", " ")
        if kljuc in vidjeni:
            opis += " (nastavak)"          # ista analiza na više stranica ispisa
        vidjeni.add(kljuc)

        put = os.path.join(obradene, f"{i:02d}.png")
        uokviri(obrezi(Image.open(os.path.join(mapa, ime))), N["rub_slike"]).save(put)
        px = Image.open(put).width
        sirina = sirina_teksta
        dpi = px / (sirina.cm / 2.54)
        if dpi < N["dpi_min"]:
            # radije uža slika nego mutna
            sirina = Cm(px / N["dpi_min"] * 2.54)
            upozorenja.append(f"{ime}: {dpi:.0f} dpi na punoj širini, umetnuto uže "
                              f"({sirina.cm:.1f} cm)")

        p_sl = sidro.insert_paragraph_before()
        bez_uvlake(p_sl, 12, 4)
        p_sl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sl.paragraph_format.keep_with_next = True     # natpis ne smije otpasti
        p_sl.add_run().add_picture(put, width=sirina)

        nat = sidro.insert_paragraph_before(f"Slika P3.{i}. {opis}")
        kao_tijelo(nat, uzor_tijelo)
        bez_uvlake(nat, 0, 2)
        nat.paragraph_format.keep_with_next = True
        for r in nat.runs:
            r.font.size = Pt(N["pt_natpis"])
            r.font.bold = True

        izv = sidro.insert_paragraph_before(N["izvor_slike"])
        kao_tijelo(izv, uzor_tijelo)
        bez_uvlake(izv, 0, 10)
        for r in izv.runs:
            r.font.size = Pt(N["pt_izvor"])
            r.font.italic = False
    return len(snimke), upozorenja


# ════════════════════════════════════════════════════════════════════════════
def main():
    p = argparse.ArgumentParser()
    p.add_argument("--conf", default="replikacija.json")
    p.add_argument("--dokument")
    args = p.parse_args()

    k = json.load(open(args.conf, encoding="utf-8"))
    N = {**ZADANO, **k.get("prilog", {})}
    radna = k.get("izlaz", "replikacija")
    dokument = args.dokument or N.get("dokument")
    if not dokument or not os.path.exists(dokument):
        sys.exit("nije zadan postojeći .docx (--dokument ili prilog.dokument)")

    izlazni = re.sub(r"\.docx$", "", dokument) + "_prilozi.docx"
    shutil.copy(dokument, izlazni)
    d = docx.Document(izlazni)
    P = d.paragraphs

    sidro = next((x for x in P if x.text.strip() == N["sidro"]), None)
    if sidro is None:
        sys.exit(f"nije nađen odlomak „{N['sidro']}”; prilozi idu neposredno prije njega")
    if nosi_fusnotu(sidro):
        sys.exit("sidro nosi fusnotu; odaberi drugo")
    uzor_tijelo = next(x for x in P if len(x.text) > 200)
    uzor_naslov = next((x for x in P if x.text.strip().startswith("Prilog")), uzor_tijelo)

    ukupno, slaze = prilog2(d, sidro, N, uzor_tijelo, uzor_naslov,
                            os.path.join(radna, "usporedba.csv"))
    natpisi = {(re.match(r"(\d+)", a["ime"]) or re.match(r"()", a["ime"])).group(1):
               a.get("natpis", a["ime"]) for a in k["analize"]}
    broj_slika, upoz = prilog3(d, sidro, N, uzor_tijelo, uzor_naslov,
                               os.path.join(radna, "snimke"), natpisi, radna)
    d.save(izlazni)

    print(f"✅ {izlazni}")
    print(f"   Prilog 2: {slaze}/{ukupno} vrijednosti se poklapa")
    print(f"   Prilog 3: {broj_slika} snimki")
    for u in upoz:
        print("   ⚠ ", u)
    print("   napomena: prikazi u prilozima ne ulaze u popis ilustracija")


if __name__ == "__main__":
    main()
