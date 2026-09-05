#!/usr/bin/env python3
"""Generator fixturea `bolesni_rad.docx` — rad s po jednim primjerkom svake
klase pogreške koju lanac tvrdi da hvata.

Zašto postoji
-------------
Do v1.9.4 je test_all.py mjerio isključivo da alat NE prijavljuje lažne nalaze
nad zdravim tekstom. Nijedan test nije tvrdio da rad S POGREŠKOM mora pasti,
pa je „87/87 prošlo" govorilo o odsutnosti šuma, ne o prisutnosti hvatanja.
Provjera koja ne može pasti nije provjera; test koji ne može pasti nije test.

Svaka posijana pogreška ima svoj ključ u SIJANO. Test `test_bolesni.py` traži
da audit padne (izlazni kod ≠ 0) i da svaki ključ bude prisutan u nalazima.
Kad neka klasa nije pokrivena nijednim alatom, to se OVDJE vidi kao pad, a ne
tek na tuđem radu tri mjeseca kasnije.
"""
import os
import sys

from docx import Document

SIJANO = {
    "citat_bez_reference": "Kovačević (2019) je citiran, a nije u popisu literature.",
    "sirotan": "Jedinica u popisu koju tekst nikad ne citira.",
    "uzorak_nesuglasan": "147 ispitanika u Metodologiji, 152 ispitanika u Rezultatima.",
    "zbroj_kategorija": "Kategorije se zbrajaju u cjelinu bez uputnice po kategoriji.",
    "duga_crtica": "U+2014 u umetnutom položaju.",
    "postotak_nedosljedan": "45% i 62 % u istom radu.",
    "trotocka": "Tri točke umjesto znaka U+2026.",
    "ravni_navodnici": 'Ravni " umjesto hrvatskih.',
    "placeholder": "[TREBA IZVOR] ostao u tijelu teksta.",
    "x_kao_mnozenje": "80 x 80 mm umjesto 80 × 80 mm.",
}


def sagradi(put: str) -> str:
    d = Document()

    d.add_heading("1. UVOD", level=1)
    # naslov iznad narativnog citata: kvar 68 je od ovoga radio ključ ('uvod','2019')
    d.add_paragraph(
        "Prema Kovačević (2019) obvezno glasanje podiže izlaznost. "
        "Isti nalaz potvrđuje i Marković (2021), dok Markov (2021) nudi drukčije "
        "tumačenje — razlika koju rad ne razrješava."
    )

    d.add_heading("2. METODOLOGIJA", level=1)
    d.add_paragraph(
        "Istraživanje je obuhvatilo 147 ispitanika iz tri županije. "
        "Anketa je provedena na uzorku koji pokriva 45% populacije, "
        "a terenski dio je trajao od ožujka do lipnja... "
        "Mjerni instrument bio je formata 80 x 80 mm."
    )

    d.add_heading("3. REZULTATI", level=1)
    d.add_paragraph(
        "U analizu je ušlo 152 ispitanika. Od ukupno 152 odgovora, "
        "63 su bila potvrdna, 47 niječna, 28 suzdržana i 14 nevažećih. "
        "Udio potvrdnih iznosi 62 % ukupnog uzorka."
    )
    d.add_paragraph(
        'Autor navodi da je "razlika statistički značajna" [TREBA IZVOR].'
    )

    d.add_heading("4. ZAKLJUČAK", level=1)
    d.add_paragraph(
        "Rezultati potvrđuju polaznu tezu (Marković, 2021)."
    )

    d.add_heading("POPIS LITERATURE", level=1)
    d.add_paragraph("Marković, I. (2021). Izlaznost i obveza. Zagreb: Naklada A.")
    d.add_paragraph("Markov, P. (2021). Druga strana obveze. Split: Naklada B.")
    d.add_paragraph("Horvat, S. (2018). Nikad citirana jedinica. Rijeka: Naklada C.")

    d.save(put)
    return put


def sagradi_zdravi(put: str) -> str:
    """Negativna kontrola: isti oblik rada, bez ijedne posijane pogreške.

    Provjera koja se pali i na zdravom tekstu nije provjera nego šum, pa test
    mora dokazati OBA smjera: da bolesni rad pada i da zdravi prolazi.
    """
    d = Document()
    d.add_heading("1. UVOD", level=1)
    d.add_paragraph(
        "Prema Marković (2021) obvezno glasanje podiže izlaznost. "
        "Isti nalaz potvrđuje i Horvat (2018), koji analizira dulje razdoblje."
    )
    d.add_heading("2. METODOLOGIJA", level=1)
    d.add_paragraph(
        "Istraživanje je obuhvatilo 147 ispitanika iz tri županije. "
        "Anketa je provedena na uzorku koji pokriva 45 % populacije, "
        "a terenski dio je trajao od ožujka do lipnja. "
        "Mjerni instrument bio je formata 80 × 80 mm."
    )
    d.add_heading("3. REZULTATI", level=1)
    d.add_paragraph(
        "U analizu je ušlo 147 ispitanika, što je cijeli planirani uzorak. "
        "Udio potvrdnih odgovora iznosi 62 % (Marković, 2021)."
    )
    d.add_heading("4. ZAKLJUČAK", level=1)
    d.add_paragraph("Rezultati potvrđuju polaznu tezu (Horvat, 2018).")
    d.add_heading("POPIS LITERATURE", level=1)
    d.add_paragraph("Marković, I. (2021). Izlaznost i obveza. Zagreb: Naklada A.")
    d.add_paragraph("Horvat, S. (2018). Duga razdoblja. Rijeka: Naklada C.")
    d.save(put)
    return put


if __name__ == "__main__":
    izlaz = sys.argv[1] if len(sys.argv) > 1 else "bolesni_rad.docx"
    os.makedirs(os.path.dirname(os.path.abspath(izlaz)), exist_ok=True)
    print(sagradi(izlaz))
