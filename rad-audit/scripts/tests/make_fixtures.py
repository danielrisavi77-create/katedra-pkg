#!/usr/bin/env python3
"""Generira minimalne sintetičke .docx fixture za test_all.py.

Uporaba:  python3 make_fixtures.py <izlazni_folder>

Sve fixture su namjerno male i pokrivaju TOČNO onu logiku koja je najrizičnija
za regresiju (nova/izmijenjena logika u rad-audit skriptama), ne cijeli
mogući ulazni prostor. Ne ovisi ni o čemu osim python-docx + standardnoj lib.
"""
import os
import sys
import zipfile
import shutil
import tempfile
from docx import Document


def _rezip_with_part(src_docx, out_docx, part_path, part_content):
    """Kopiraj src_docx u out_docx, dodavši/zamijenivši jedan XML dio."""
    tmp = tempfile.mkdtemp()
    with zipfile.ZipFile(src_docx) as z:
        z.extractall(tmp)
    full = os.path.join(tmp, *part_path.split("/"))
    os.makedirs(os.path.dirname(full), exist_ok=True)
    with open(full, "w", encoding="utf-8") as f:
        f.write(part_content)
    if os.path.exists(out_docx):
        os.remove(out_docx)
    zf = zipfile.ZipFile(out_docx, "w", zipfile.ZIP_DEFLATED)
    for root, _dirs, files in os.walk(tmp):
        for fn in files:
            fp = os.path.join(root, fn)
            zf.write(fp, os.path.relpath(fp, tmp))
    zf.close()
    shutil.rmtree(tmp, ignore_errors=True)


def _inject_into_body(src_docx, out_docx, xml_snippet):
    tmp = tempfile.mkdtemp()
    with zipfile.ZipFile(src_docx) as z:
        z.extractall(tmp)
    docp = os.path.join(tmp, "word", "document.xml")
    xml = open(docp, encoding="utf-8").read()
    xml = xml.replace("<w:body>", "<w:body>" + xml_snippet, 1)
    open(docp, "w", encoding="utf-8").write(xml)
    if os.path.exists(out_docx):
        os.remove(out_docx)
    zf = zipfile.ZipFile(out_docx, "w", zipfile.ZIP_DEFLATED)
    for root, _dirs, files in os.walk(tmp):
        for fn in files:
            fp = os.path.join(root, fn)
            zf.write(fp, os.path.relpath(fp, tmp))
    zf.close()
    shutil.rmtree(tmp, ignore_errors=True)


def build(out_dir):
    os.makedirs(out_dir, exist_ok=True)
    src = os.path.join(out_dir, "izvori")
    os.makedirs(src, exist_ok=True)

    # 1) navodnici: par, inč-oznaka, nesparen, razdvojen kroz 3 runa
    doc = Document()
    p1 = doc.add_paragraph()
    p1.add_run('On je rekao '); p1.add_run('"pozdrav"'); p1.add_run(' i otišao.')
    doc.add_paragraph('Sidro promjera 12" ugrađeno je na mjestu spoja.')
    p3 = doc.add_paragraph()
    p3.add_run('Ovo je odlomak s nesparenim navodnikom '); p3.add_run('"greška u tekstu.')
    p4 = doc.add_paragraph()
    p4.add_run('Prvi '); p4.add_run('"razdvojen '); p4.add_run('navodnik" u tri runa.')
    doc.save(os.path.join(out_dir, "quotes.docx"))

    # 2) IEEE citiranje s numeriranim naslovom LITERATURE, bez rupa/siročadi
    doc = Document()
    doc.add_paragraph('U radu se koriste izvori [1] i [2] te kasnije [1, 3].')
    doc.add_paragraph('7. LITERATURA')
    doc.add_paragraph('[1] Autor, A. Naslov. Zagreb, 2020.')
    doc.add_paragraph('[2] Autor, B. Naslov. Zagreb, 2019.')
    doc.add_paragraph('[3] Autor, C. Naslov. Zagreb, 2018.')
    doc.save(os.path.join(out_dir, "ieee_numbered_heading.docx"))

    # 2b) Vancouver (N): siroče 5, citat bez reference 6, redoslijed prekršen (3 prije 2),
    #     decimale u tablici „158 (77,8)" NISU citati, „(67,68)" jest (bez razmaka),
    #     „Rec(2003)24" nije autor-godina citat, stavka 4 ima 7 autora bez „i sur."
    doc = Document()
    doc.add_paragraph('Palijativna skrb je pristup (1). Prema Recommendation Rec(2003)24 to vrijedi (3). '
                      'Udio od 12,5 % (2) je nizak, a drugi to potvrđuju (3–4) i (67,68). Nedostaje (6).')
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = '158 (77,8)'
    t.rows[0].cells[1].text = '12 (5,9)'
    doc.add_paragraph('8. POPIS CITIRANE LITERATURE')
    doc.add_paragraph('1. Kellehear A. Compassionate cities. Prog Palliat Care. 2020;28(2):115-9.')
    doc.add_paragraph('2. Connor SR, urednik. Global atlas. 2. izd. London: WHPCA; 2020.')
    doc.add_paragraph('3. Knaul FM, Arreola H, Kwete XJ, Bhadelia A, Rodriguez NM, Vargas V i sur. Evolution. Lancet. 2020;8(1):e1.')
    doc.add_paragraph('4. Aa B, Cc D, Ee F, Gg H, Ii J, Kk L, Mm N. Sedam autora bez i sur. Časopis. 2019;3(4):5-6.')
    doc.add_paragraph('5. Abel J, Kellehear A. Nikad citirano. Ann Palliat Med. 2018;7(2):S3-14.')
    doc.add_paragraph('67. Prvi P. Naslov. Časopis. 2021;1:1.')
    doc.add_paragraph('68. Drugi D. Naslov. Časopis. 2021;1:2.')
    doc.add_paragraph('9. PRILOZI')
    doc.add_paragraph('1. Prilog koji nije referenca.')
    doc.save(os.path.join(out_dir, "vancouver.docx"))

    # 3) autor-godina: siroče (Horvat 2018) + citat bez reference (Kovač 2021)
    doc = Document()
    doc.add_paragraph('Prema istraživanju (Ivić, 2020) utvrđeno je da postoji značajan utjecaj.')
    doc.add_paragraph('Slično zaključuju i drugi autori (Perić i Marić, 2019; Kovač, 2021).')
    doc.add_paragraph('LITERATURA')
    doc.add_paragraph('Ivić, A. (2020). Naslov rada o temi X. Zagreb: Nakladnik.')
    doc.add_paragraph('Perić, B. i Marić, C. (2019). Drugi naslov o temi Y. Split: Izdavač.')
    doc.add_paragraph('Horvat, D. (2018). Treći rad, nikad citiran u tekstu. Rijeka: Neki nakladnik.')
    doc.save(os.path.join(out_dir, "author_year.docx"))

    # 3b) autor-godina bez osobnog autora: mediji/platforme, institucija i
    #      prefiks/sufiks unutar citata. Sve jedinice jesu citirane i nijedna
    #      ne smije postati lažno siroče ili citat bez reference.
    doc = Document()
    doc.add_paragraph(
        'Medijski korpus čine objave (danas.hr, 2025; Index.hr, 2025; '
        'Ministarstvo znanosti i obrazovanja, 2024; UNESCO, 2021). '
        'Za usporedbu vidi (usp. Tonković, Krolo i Marcelić, 2014, za analizu).')
    doc.add_paragraph('LITERATURA')
    doc.add_paragraph('danas.hr (2025). Naslov članka. https://danas.hr/tekst')
    doc.add_paragraph('Index.hr (2025). Drugi članak. https://www.index.hr/tekst')
    doc.add_paragraph('Ministarstvo znanosti i obrazovanja (2024). Strategija obrazovanja.')
    doc.add_paragraph('UNESCO. (2021). Global education report. Paris: UNESCO.')
    doc.add_paragraph(
        'Tonković, Željka, Krolo, Krešimir i Marcelić, Sven (2014). Kulturna potrošnja.')
    doc.save(os.path.join(out_dir, "author_year_institutions.docx"))

    # 4) elektro domena (za auto-detekciju)
    doc = Document()
    doc.add_paragraph('Sustav radi na naponu 400 V i struji 16 A, uz snagu 6 kW.')
    doc.add_paragraph('Razvodni ormar ima zaštitu IP44 i uzemljenje prema propisu.')
    doc.save(os.path.join(out_dir, "elektro.docx"))
    with open(os.path.join(src, "elektro_spec.txt"), "w", encoding="utf-8") as f:
        f.write("Mjerenja pokazuju napon od 400 V i struju 16 A, snaga iznosi 6 kW. IP44 potvrđeno.\n")

    # 5) cross-check lažni pozitivac (substring preko granice rečenice)
    doc = Document()
    doc.add_paragraph('Sidro je dimenzija 40 t nosivosti prema tehničkoj specifikaciji.')
    doc.save(os.path.join(out_dir, "cross_fp.docx"))
    with open(os.path.join(src, "cross_fp_src.txt"), "w", encoding="utf-8") as f:
        f.write("Iznos od 40 tvrtke X je nepovezan podatak koji se ovdje spominje radi testa.\n")

    # 6) verbatim-copy: neoznačeno + označeno + parafrazirano
    doc = Document()
    doc.add_paragraph(
        'Prema dostupnim podacima sustav za automatsko upravljanje temperaturom u '
        'industrijskim postrojenjima znatno smanjuje potrošnju energije tijekom cijele '
        'godine bez dodatnih troškova održavanja.')
    doc.add_paragraph(
        'Kako se navodi: „Prema dostupnim podacima sustav za automatsko upravljanje '
        'temperaturom u industrijskim postrojenjima znatno smanjuje potrošnju energije '
        'tijekom cijele godine bez dodatnih troškova održavanja" [1].')
    doc.add_paragraph(
        'Različita istraživanja pokazuju da pametna regulacija grijanja u tvornicama može '
        'donijeti uštede, no rezultati uvelike ovise o uvjetima svakog pogona.')
    doc.save(os.path.join(out_dir, "overlap.docx"))
    with open(os.path.join(src, "overlap_src.txt"), "w", encoding="utf-8") as f:
        f.write(
            'Prema dostupnim podacima sustav za automatsko upravljanje temperaturom u '
            'industrijskim postrojenjima znatno smanjuje potrošnju energije tijekom cijele '
            'godine bez dodatnih troškova održavanja, što potvrđuju mjerenja na terenu.\n')

    # 7) tracked changes (w:ins) neprihvaćeno
    base = os.path.join(out_dir, "quotes.docx")
    ins = ('<w:p><w:ins w:id="1" w:author="X" w:date="2024-01-01T00:00:00Z">'
           '<w:r><w:t>Umetnuti tekst.</w:t></w:r></w:ins></w:p>')
    _inject_into_body(base, os.path.join(out_dir, "tracked_changes.docx"), ins)

    # 8) footnote s autor-godina citatom (za load_supplementary_text)
    doc = Document()
    doc.add_paragraph('Glavni tekst bez ijednog vidljivog citata u tijelu.')
    doc.save(os.path.join(out_dir, "_footnote_base.docx"))
    fn_xml = ('<?xml version="1.0"?><w:footnotes '
              'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
              '<w:footnote w:id="1"><w:p><w:r><w:t>Prema (Ivić, 2020) mjerenje pokazuje 40 t.'
              '</w:t></w:r></w:p></w:footnote></w:footnotes>')
    _rezip_with_part(os.path.join(out_dir, "_footnote_base.docx"),
                      os.path.join(out_dir, "footnote.docx"),
                      "word/footnotes.xml", fn_xml)

    # 9) IEEE: citat [2] samo u ćeliji tablice + godina [2020] u zagradi
    doc = Document()
    doc.add_paragraph('U tekstu se citira samo [1], a norma je iz [2020]. godine.')
    t = doc.add_table(rows=1, cols=1)
    t.rows[0].cells[0].text = 'Vrijednost prema [2] iznosi 40 t.'
    doc.add_paragraph('LITERATURA')
    doc.add_paragraph('[1] Autor, A. Naslov. Zagreb, 2020.')
    doc.add_paragraph('[2] Autor, B. Naslov. Split, 2019.')
    doc.save(os.path.join(out_dir, "table_cite.docx"))

    # 10) autor-godina: JEDINI citat u fusnoti, popis u body-ju (end-to-end)
    doc = Document()
    doc.add_paragraph('Glavni tekst bez ijednog vidljivog citata u tijelu.')
    doc.add_paragraph('LITERATURA')
    doc.add_paragraph('Ivić, A. (2020). Naslov rada. Zagreb: Nakladnik.')
    doc.save(os.path.join(out_dir, "_fn_ay_base.docx"))
    _rezip_with_part(os.path.join(out_dir, "_fn_ay_base.docx"),
                      os.path.join(out_dir, "fn_ay_cite.docx"),
                      "word/footnotes.xml", fn_xml)

    # 11) run-level w:spacing u rPr (test da --no-indent ne kontaminira rPr)
    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run('Odlomak s razmakom slova u runu i dovoljno teksta za prozu.')
    from docx.oxml.ns import qn
    rpr = r._element.get_or_add_rPr()
    rpr.append(rpr.makeelement(qn('w:spacing'), {qn('w:val'): '20'}))
    doc.add_paragraph('Drugi normalan odlomak proze bez ičega posebnog u sebi.')
    doc.save(os.path.join(out_dir, "rpr_spacing.docx"))

    # 12) engleski “…” par (bez „) + njemački „…“ par
    doc = Document()
    doc.add_paragraph('Abstract: the system uses “smart control” for regulation.')
    doc.add_paragraph('Hrvatski odlomak s njemačkim parom „citat“ koji treba popraviti.')
    doc.save(os.path.join(out_dir, "eng_quotes.docx"))

    # 13) hex literali + prava multiplikacija
    doc = Document()
    doc.add_paragraph('Registar 0x41 i adresa 0xFF00, a dimenzija je 80x80 mm.')
    doc.save(os.path.join(out_dir, "hex.docx"))

    # 14) Normal BEZ firstLine + drugi stil S firstLine (cross-style test)
    doc = Document()
    doc.add_paragraph('Tekst.')
    doc.save(os.path.join(out_dir, "_styles_base.docx"))
    tmp2 = tempfile.mkdtemp()
    with zipfile.ZipFile(os.path.join(out_dir, "_styles_base.docx")) as z:
        z.extractall(tmp2)
    sp = os.path.join(tmp2, "word", "styles.xml")
    s = open(sp, encoding="utf-8").read()
    s = s.replace("</w:styles>",
                   '<w:style w:type="paragraph" w:styleId="Citat9"><w:name w:val="Citat9"/>'
                   '<w:pPr><w:ind w:firstLine="709"/></w:pPr></w:style></w:styles>')
    open(sp, "w", encoding="utf-8").write(s)
    outp = os.path.join(out_dir, "styles_extra.docx")
    if os.path.exists(outp):
        os.remove(outp)
    zf = zipfile.ZipFile(outp, "w", zipfile.ZIP_DEFLATED)
    for root, _dirs, files in os.walk(tmp2):
        for fn in files:
            fp = os.path.join(root, fn)
            zf.write(fp, os.path.relpath(fp, tmp2))
    zf.close()
    shutil.rmtree(tmp2, ignore_errors=True)

    # 15) inline textbox u sredini odlomka, navodnici i PRIJE i POSLIJE njega
    doc = Document()
    doc.add_paragraph('Placeholder')
    doc.save(os.path.join(out_dir, "_txbx_base.docx"))
    txbx_para = (
        '<w:p><w:r><w:t>Prije okvira "prvi citat" ide tekst. </w:t></w:r>'
        '<w:r><w:pict><v:shape xmlns:v="urn:schemas-microsoft-com:vml"><v:textbox><w:txbxContent>'
        '<w:p><w:r><w:t>Tekst u okviru.</w:t></w:r></w:p>'
        '</w:txbxContent></v:textbox></v:shape></w:pict></w:r>'
        '<w:r><w:t>Poslije okvira "drugi citat" nastavlja se.</w:t></w:r></w:p>')
    _inject_into_body(os.path.join(out_dir, "_txbx_base.docx"),
                       os.path.join(out_dir, "txbx.docx"), txbx_para)

    # 16) elektro s V/Hz/%/° i SUKOBOM vrijednosti (230 V vs 400 V uz 'napon')
    doc = Document()
    doc.add_paragraph('Sustav radi na naponu 400 V i frekvenciji 50 Hz, uz struju 16 A.')
    doc.add_paragraph('Kasnije se navodi da napon iznosi 230 V, što rad ne pomiruje.')
    doc.add_paragraph('Udio gubitaka je 45 %, a nagib kabela 10°.')
    doc.save(os.path.join(out_dir, "elektro_konflikt.docx"))

    # 17) istinski mixed stil (IEEE i autor-godina u sličnoj mjeri)
    doc = Document()
    doc.add_paragraph('Prema [1] i [2] te (Ivić, 2020), (Perić, 2019) i (Kovač, 2021) o temi.')
    doc.add_paragraph('LITERATURA')
    doc.add_paragraph('[1] Autor, A. Naslov. Zagreb, 2020.')
    doc.add_paragraph('[2] Autor, B. Naslov. Split, 2019.')
    doc.save(os.path.join(out_dir, "mixed_style.docx"))

    print(f"Fixture spremljene u: {out_dir}")


if __name__ == "__main__":
    build(sys.argv[1] if len(sys.argv) > 1 else "/tmp/rad_audit_fixtures")
