#!/usr/bin/env python3
"""Svaka brojka o VLASTITOM istraživanju mora postojati u Rezultatima.

Uporaba:
  python3 brojke_iz_rasprave.py rad.docx
  python3 brojke_iz_rasprave.py rad.docx --rezultati "5. REZULTATI:6. RASPRAVA" \\
                                          --tvrdnje "6. RASPRAVA:8. POPIS"

Zašto postoji. Rasprava i Zaključak prepričavaju Rezultate, i to je mjesto gdje
brojka najlakše odluta: prepisivanjem, zaokruživanjem ili sjećanjem. Recenzent to
gleda prvo, jer je mehanički provjerljivo. Alat radi upravo tu provjeru.

Kako razlikuje vlastitu brojku od tuđe. Brojka u čijoj REČENICI stoji citat —
"(54)", "(58,62)" — pripisana je literaturi i ne mora biti u Rezultatima. Prozor
od N znakova ovdje ne radi: citat u hrvatskoj akademskoj rečenici stoji na kraju,
često 150+ znakova iza brojke. Sam citatni par "(67,68)" nije brojka nego citat i
preskače se. Brojka bez citata u rečenici tvrdnja je o vlastitom uzorku i MORA se
naći u Rezultatima ili tablicama; ako je nema ondje, a ima je drugdje u radu
(tipično u Metodama), prijavljuje se blaže, kao mjesto koje treba potvrditi.

Uz to provjerava metodološku tvrdnju o zaokruživanju vrijednosti p: ako Statistički
postupci kažu "na dvije decimale", a tablice nose p = 0,002, jedno od toga je
netočno. (Nađeno na stvarnom radu; nijedna druga faza to ne vidi.)
"""
import argparse
import re
import sys

try:
    from common import load_docx_text
except ImportError:
    from docx import Document

    def load_docx_text(path, include_tables=True):
        d = Document(path)
        body = "\n".join(p.text for p in d.paragraphs)
        cells = []
        if include_tables:
            for t in d.tables:
                for r in t.rows:
                    for c in r.cells:
                        cells.append(c.text)
        return body, cells, None

BROJ = re.compile(r"\d{1,3},\d{1,3}")          # 54,1  13,53  3,76 — decimalni zarez
CITAT_BLIZU = re.compile(r"\(\s*\d{1,3}\s*(?:[,;–\-]\s*\d{1,3}\s*)*\)")
RIJECI_DECIMALA = {"jednu": 1, "dvije": 2, "tri": 3, "četiri": 4}


def _odsjecak(t, spec):
    od, _, do = spec.partition(":")
    i = t.rfind(od)
    if i < 0:
        return ""
    j = t.rfind(do)
    return t[i:j] if j > i else t[i:]


def provjeri_p_decimale(body, cells):
    izvor = body + "\n" + "\n".join(cells)
    # Rečenica smije deklarirati VIŠE zaokruženja ("na dvije decimale, na tri kad
    # su manje od 0,01"). Uzimanje prvog broja daje lažni nalaz na ispravnom radu,
    # pa se gleda najveći deklarirani broj decimala u toj rečenici.
    mrec = re.search(r"[^.!?]*zaokružen\w*[^.!?]*[.!?]", body, re.I | re.S)
    if not mrec:
        return None
    tvrdnje = []
    for m in re.finditer(r"\b(\w+)\s+decimal", mrec.group(0), re.I):
        r = m.group(1).lower()
        v = RIJECI_DECIMALA.get(r) or (int(r) if r.isdigit() else None)
        if v:
            tvrdnje.append(v)
    if not tvrdnje:
        return None
    tvrdi = max(tvrdnje)
    stvarno = {len(d) for d in re.findall(r"\bp\s*[=<]\s*0,(\d+)", izvor)}
    prekrsaji = sorted(d for d in stvarno if d > tvrdi)
    return {"tvrdi": tvrdi, "nadjeno": sorted(stvarno), "prekrsaji": prekrsaji}


def main(path, rezultati_spec=None, tvrdnje_spec=None):
    body, cells, _ = load_docx_text(path, include_tables=True)
    rez_spec = rezultati_spec or "5. REZULTATI:6. RASPRAVA"
    tvr_spec = tvrdnje_spec or "6. RASPRAVA:8. POPIS"

    rez = _odsjecak(body, rez_spec) + "\n" + "\n".join(cells)
    tvr = _odsjecak(body, tvr_spec)

    print("=" * 62)
    print("BROJKE IZ RASPRAVE —", path)
    print("=" * 62)
    if not tvr.strip():
        print(f"⚠ odsjecak '{tvr_spec}' nije nađen — zadaj --tvrdnje ručno")
        return 2
    if not rez.strip():
        print(f"⚠ odsjecak '{rez_spec}' nije nađen — zadaj --rezultati ručno")
        return 2

    u_rezultatima = set(BROJ.findall(rez))
    u_radu = set(BROJ.findall(body + "\n" + "\n".join(cells)))
    recenice = [(mm.start(), mm.group(0)) for mm in
                re.finditer(r"[^.!?]*[.!?]", tvr, re.S)]

    def recenica_od(poz):
        pog = ""
        for start, tekst in recenice:
            if start <= poz < start + len(tekst):
                pog = tekst
                break
        return pog

    vlastite, tude = [], 0
    for m in BROJ.finditer(tvr):
        prije = tvr[m.start() - 1:m.start()]
        poslije = tvr[m.end():m.end() + 1]
        if prije == "(" and poslije == ")":          # (67,68) je citat, ne brojka
            tude += 1
            continue
        if CITAT_BLIZU.search(recenica_od(m.start())):
            tude += 1
            continue
        vlastite.append((m.group(0), tvr[max(0, m.start() - 55):m.end() + 12].replace("\n", " ")))

    nedostaju = [(b, k) for b, k in vlastite if b not in u_rezultatima and b not in u_radu]
    drugdje = [(b, k) for b, k in vlastite if b not in u_rezultatima and b in u_radu]
    print(f"Brojki u Rezultatima i tablicama: {len(u_rezultatima)}")
    print(f"Brojki u Raspravi/Zaključku: {len(vlastite) + tude}"
          f"  (pripisano literaturi: {tude}, o vlastitom uzorku: {len(vlastite)})")
    if nedostaju:
        print(f"\n⚠ NEMA POKRIĆA U REZULTATIMA: {len(nedostaju)}")
        for b, k in nedostaju[:20]:
            print(f"   {b:>7}   …{k.strip()}")
        if len(nedostaju) > 20:
            print(f"   … i još {len(nedostaju) - 20}")
        print("\n   Za svaku: ili je u Rezultatima pod drugim zaokruženjem, ili je")
        print("   tvrdnja iz literature bez citata, ili je pogrešna.")
    else:
        print("\n✓ svaka brojka o vlastitom uzorku ima pokriće u Rezultatima")
    if drugdje:
        print(f"\n· nije u Rezultatima, ali jest drugdje u radu (potvrdi): {len(drugdje)}")
        for b, k in drugdje[:8]:
            print(f"   {b:>7}   …{k.strip()}")

    pd = provjeri_p_decimale(body, cells)
    if pd:
        if pd["prekrsaji"]:
            print(f"\n⚠ Statistički postupci tvrde najviše {pd['tvrdi']} decimale, "
                  f"a u radu postoje p-vrijednosti s {pd['prekrsaji']} decimala "
                  f"(nađene duljine: {pd['nadjeno']})")
        else:
            print(f"\n✓ zaokruživanje p ({pd['tvrdi']} decimale) odgovara vrijednostima u radu")

    ok = not nedostaju and not (pd and pd["prekrsaji"])
    print("\nREZULTAT:", "✓ brojke pokrivene" if ok else "⚠ ima nalaza (v. gore)")
    return 0 if ok else 1


if __name__ == "__main__":
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("docx")
    ap.add_argument("--rezultati", default=None, help='"OD:DO"')
    ap.add_argument("--tvrdnje", default=None, help='"OD:DO"')
    a = ap.parse_args()
    sys.exit(main(a.docx, a.rezultati, a.tvrdnje))
