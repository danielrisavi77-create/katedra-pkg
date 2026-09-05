#!/usr/bin/env python3
"""Mjeri koliko se NOVA verzija rada stvarno razlikuje od STARE.

Uporaba:
  python3 parafraza.py stari.docx novi.docx
  python3 parafraza.py stari.docx novi.docx --poglavlja "1. UVOD:2. CILJ" "6. RASPRAVA:7. ZAKLJUČAK"
  python3 parafraza.py stari.docx novi.docx --n 8 --prag 20 --json out.json

Zašto postoji. Kad mentor javi previsoku podudarnost i traži prepisivanje, jedino
pitanje koje vrijedi je "koliko je teksta STVARNO novo". Dojam nije mjera: odlomak
kojemu su zamijenjene tri riječi izgleda prepisano, a n-gramima je 90 % isti.
Alat NE zamjenjuje Turnitin (ne zna za vanjske izvore) — mjeri samo razliku prema
prethodnoj verziji istoga rada, što je ono na što autor može utjecati.

Mjera: udio n-grama nove verzije koji postoje i u staroj (zadano n = 8 riječi).
  ≤ 20 %  prepisano                  (zeleno)
  ≤ 40 %  djelomično                 (žuto — provjeri koje odlomke)
  > 40 %  nije prepisano             (crveno)

Ostatak koji nikad ne pada na nulu: nazivi ustanova, naslovi dokumenata, nazivi
mjernih instrumenata i podnaslovi. Zato se ispisuje i najduži zajednički niz —
ako je on sve sami nazivi, visok postotak nije problem.
"""
import argparse
import difflib
import json
import re
import sys

try:
    from common import load_docx_text
except ImportError:                                     # samostalna uporaba
    from docx import Document

    def load_docx_text(path, include_tables=True):
        d = Document(path)
        body = "\n".join(p.text for p in d.paragraphs)
        cells = []
        if include_tables:
            for t in d.tables:
                for r in t.rows:
                    for c in r.cells:
                        cells.append(c.text)
        return body, cells, None


def _rijeci(t):
    return re.findall(r"\w+", t.lower(), flags=re.UNICODE)


def _ngrami(t, n):
    w = _rijeci(t)
    return set(tuple(w[i:i + n]) for i in range(len(w) - n + 1)), len(w)


def _odsjecak(tekst, od, do):
    i = tekst.rfind(od)
    if i < 0:
        return ""
    j = tekst.rfind(do)
    return tekst[i:j] if j > i else tekst[i:]


def usporedi(stari, novi, n=8):
    gs, _ = _ngrami(stari, n)
    gn, wn = _ngrami(novi, n)
    if not gn:
        return None
    zajednicki = gs & gn
    udio = 100.0 * len(zajednicki) / len(gn)
    ws, wnn = re.findall(r"\S+", stari), re.findall(r"\S+", novi)
    sm = difflib.SequenceMatcher(None, ws, wnn, autojunk=False)
    nizovi = sorted(((m.size, " ".join(wnn[m.b:m.b + m.size]))
                     for m in sm.get_matching_blocks() if m.size >= 10),
                    reverse=True)[:5]
    return {"ngrama_novih": len(gn), "zajednickih": len(zajednicki),
            "udio_posto": round(udio, 1), "rijeci_novih": wn,
            "najduzi_zajednicki": [{"rijeci": s, "tekst": t} for s, t in nizovi]}


def ocjena(udio, prag):
    if udio <= prag:
        return "✓ prepisano"
    if udio <= prag * 2:
        return "⚠ djelomično prepisano"
    return "✗ nije prepisano"


def main(stari_path, novi_path, poglavlja=None, n=8, prag=20.0, json_out=None):
    bs, cs, _ = load_docx_text(stari_path, include_tables=True)
    bn, cn, _ = load_docx_text(novi_path, include_tables=True)

    dijelovi = []
    if poglavlja:
        for spec in poglavlja:
            od, _, do = spec.partition(":")
            dijelovi.append((od, _odsjecak(bs, od, do), _odsjecak(bn, od, do)))
    else:
        dijelovi.append(("CIJELI RAD", bs, bn))

    print("=" * 62)
    print(f"PARAFRAZA — {novi_path}")
    print(f"  naspram   {stari_path}   (n-grami: {n} riječi, prag {prag:.0f} %)")
    print("=" * 62)

    nalazi, lose = {}, 0
    for ime, so, no in dijelovi:
        if not no.strip():
            print(f"\n{ime}: ⚠ odsjecak nije nađen u novoj verziji — provjeri granice")
            continue
        r = usporedi(so, no, n)
        if r is None:
            print(f"\n{ime}: ⚠ premalo teksta za mjeru")
            continue
        oc = ocjena(r["udio_posto"], prag)
        if oc.startswith("✗"):
            lose += 1
        r["ocjena"] = oc
        nalazi[ime] = r
        print(f"\n{ime}")
        print(f"  identično sa starom verzijom: {r['udio_posto']} % "
              f"({r['zajednickih']}/{r['ngrama_novih']} n-grama)   {oc}")
        print(f"  duljina novog teksta: {r['rijeci_novih']} riječi")
        for niz in r["najduzi_zajednicki"][:3]:
            print(f"    ostalo isto ({niz['rijeci']} riječi): {niz['tekst'][:96]}")

    print("\nNAPOMENA: mjeri se razlika prema PRETHODNOJ verziji istog rada, ne prema")
    print("vanjskim izvorima. Nazivi ustanova, naslovi dokumenata i podnaslovi ostaju")
    print("isti po definiciji — provjeri gornje nizove prije nego zaključiš da je loše.")

    if json_out:
        json.dump(nalazi, open(json_out, "w", encoding="utf-8"),
                  ensure_ascii=False, indent=2)
        print(f"\n✔ JSON: {json_out}")

    print("\nREZULTAT:", "✓ sve iznad praga prepisano" if not lose
          else f"⚠ {lose} dio(dijelova) nije dovoljno prepisano")
    return 0 if not lose else 1


if __name__ == "__main__":
    ap = argparse.ArgumentParser(add_help=True, description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("stari")
    ap.add_argument("novi")
    ap.add_argument("--poglavlja", nargs="*", default=None,
                    help='"OD:DO" po dijelu, npr. "1. UVOD:2. CILJ"')
    ap.add_argument("--n", type=int, default=8)
    ap.add_argument("--prag", type=float, default=20.0)
    ap.add_argument("--json", dest="json_out", default=None)
    a = ap.parse_args()
    sys.exit(main(a.stari, a.novi, a.poglavlja, a.n, a.prag, a.json_out))
