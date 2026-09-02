#!/usr/bin/env python3
"""Propagacija popravka: nalazi tvrdnje koje na dva mjesta stoje različito.

Rukopis je izvor istine za prozu, ali natpisi, grafikoni, tablice, sažetak i alt-tekstovi
žive u drugim slojevima. Popravak u jednom sloju ne dira druga tri.

Nastalo iz ciklusa u kojemu se isti obrazac ponovio četiri puta u tri kruga popravaka:
statistika ispravljena u tijelu teksta ostala je kriva u grafikonu, tablici, sažetku i
zaključku; ograde uz dva predmeta dodane u tekst nisu ušle u tablicu ni u sliku.

    python3 propagacija.py --rukopis .katedra/poglavlja
    python3 propagacija.py --rukopis .katedra/poglavlja --docx rad.docx \\
        --generator postprocess.py make_charts.py --json propagacija.json
"""
from __future__ import annotations

import argparse
import json
import pathlib
import re
import sys
from collections import defaultdict

# Brojka s neobveznim postotkom/tisućicom; hvata i 83,4 i 2.500 i 161
RE_BROJ = re.compile(r"(?<![\w.,])(\d{1,3}(?:[. ]\d{3})*(?:,\d+)?)(?![\w])")
# Predmet: kurziv ili velika početna riječ uz „predmet"/„presuda"
RE_PREDMET = re.compile(r"\*([A-ZČĆŠŽĐ][\w\s.]{2,40}?)\*|predmet[uai]?\s+([A-ZČĆŠŽĐ]\w+)")
RE_GODINA = re.compile(r"\b(19\d{2}|20\d{2})\.")

SLOJEVI = ("rukopis", "natpis", "tablica", "izvor", "sazetak", "alt", "generator")


def _rukopis(mapa: pathlib.Path) -> dict[str, list[str]]:
    out: dict[str, list[str]] = defaultdict(list)
    for put in sorted(mapa.glob("*.md")):
        for red in put.read_text(encoding="utf-8").split("\n"):
            g = red.strip()
            if not g:
                continue
            if g.startswith("|"):
                out["tablica"].append(g)
            elif re.match(r"^(Tablica|Grafikon|Slika|Shema|Prilog)\s+\d+[.:]", g):
                out["natpis"].append(g)
            elif re.match(r"^(Izvor|Napomena)\s*:", g, re.I):
                out["izvor"].append(g)
            else:
                out["rukopis"].append(g)
    return out


def _iz_docxa(put: pathlib.Path) -> dict[str, list[str]]:
    out: dict[str, list[str]] = defaultdict(list)
    try:
        from docx import Document
    except ImportError:
        return out
    d = Document(str(put))
    for p in d.paragraphs:
        g = p.text.strip()
        if not g:
            continue
        if p.style.name == "Caption":
            out["natpis"].append(g)
        elif g.startswith("Izvor:") or g.startswith("Napomena:"):
            out["izvor"].append(g)
        else:
            out["sazetak" if len(out["sazetak"]) == 0 and "Esej" in g[:20] else "rukopis"].append(g)
    for tb in d.tables:
        for red in tb.rows:
            out["tablica"].append(" | ".join(c.text.strip() for c in red.cells))
    import zipfile
    try:
        xml = zipfile.ZipFile(put).read("word/document.xml").decode()
        for m in re.finditer(r'descr="([^"]{4,})"', xml):
            out["alt"].append(m.group(1))
    except Exception:
        pass
    return out


def _generator(putovi: list[str]) -> dict[str, list[str]]:
    out: dict[str, list[str]] = defaultdict(list)
    for s in putovi:
        p = pathlib.Path(s)
        if not p.exists():
            continue
        for m in re.finditer(r'"([^"\n]{25,})"|\'([^\'\n]{25,})\'', p.read_text(encoding="utf-8")):
            out["generator"].append(m.group(1) or m.group(2))
    return out


def _vrijednosti(redci: list[str]) -> set[str]:
    v: set[str] = set()
    for r in redci:
        v |= {x for x in RE_BROJ.findall(r)}
    return v


def _konteksti(redci: list[str], vrijednost: str) -> list[str]:
    izl = []
    for r in redci:
        for m in re.finditer(re.escape(vrijednost), r):
            lijevo = r[max(0, m.start() - 40):m.start()]
            desno = r[m.end():m.end() + 25]
            izl.append((lijevo + "«" + vrijednost + "»" + desno).strip())
    return izl


def usporedi(slojevi: dict[str, list[str]], min_duljina: int = 2) -> list[dict]:
    """Vrati brojke koje se pojavljuju u više slojeva, s kontekstima za ručnu prosudbu."""
    po_sloju = {k: _vrijednosti(v) for k, v in slojevi.items() if v}
    sve = set().union(*po_sloju.values()) if po_sloju else set()
    nalazi = []
    for vr in sorted(sve, key=lambda x: (-len(x), x)):
        if len(vr.replace(",", "").replace(".", "")) < min_duljina:
            continue
        gdje = [k for k, s in po_sloju.items() if vr in s]
        if len(gdje) < 2:
            continue
        nalazi.append({
            "vrijednost": vr,
            "slojevi": gdje,
            "konteksti": {k: _konteksti(slojevi[k], vr)[:3] for k in gdje},
        })
    return nalazi


def osamljene(slojevi: dict[str, list[str]]) -> list[dict]:
    """Brojke koje postoje SAMO u prikazu ili samo u generatoru — kandidati za nepropagirano."""
    po_sloju = {k: _vrijednosti(v) for k, v in slojevi.items() if v}
    tijelo = po_sloju.get("rukopis", set())
    izl = []
    for sloj in ("natpis", "tablica", "izvor", "alt", "generator", "sazetak"):
        for vr in sorted(po_sloju.get(sloj, set())):
            if len(vr.replace(",", "").replace(".", "")) < 2:
                continue
            if vr not in tijelo:
                izl.append({"vrijednost": vr, "sloj": sloj,
                            "kontekst": (_konteksti(slojevi[sloj], vr) or [""])[0]})
    return izl


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--rukopis", default=".katedra/poglavlja")
    ap.add_argument("--docx")
    ap.add_argument("--generator", nargs="*", default=[])
    ap.add_argument("--json")
    a = ap.parse_args()

    slojevi: dict[str, list[str]] = defaultdict(list)
    mapa = pathlib.Path(a.rukopis)
    if mapa.exists():
        for k, v in _rukopis(mapa).items():
            slojevi[k] += v
    if a.docx and pathlib.Path(a.docx).exists():
        for k, v in _iz_docxa(pathlib.Path(a.docx)).items():
            slojevi[k] += v
    for k, v in _generator(a.generator).items():
        slojevi[k] += v

    if not slojevi:
        print("nema ulaza — provjeri --rukopis / --docx")
        return 2

    zaj = usporedi(slojevi)
    osam = osamljene(slojevi)

    print("=" * 74)
    print("PROPAGACIJA — vrijednosti kroz slojeve")
    print("=" * 74)
    print("slojevi:", ", ".join(f"{k} ({len(v)})" for k, v in slojevi.items()))
    print()
    print(f"vrijednosti u dva ili više sloja: {len(zaj)}")
    for n in zaj[:25]:
        print(f"  {n['vrijednost']:>10}  →  {', '.join(n['slojevi'])}")
        for sloj, kon in n["konteksti"].items():
            for k in kon[:1]:
                print(f"{'':>14}[{sloj}] {k[:88]}")
    print()
    print(f"⚠ vrijednosti kojih NEMA u tijelu teksta: {len(osam)}")
    for n in osam[:25]:
        print(f"  {n['vrijednost']:>10}  samo u [{n['sloj']}]  {n['kontekst'][:70]}")
    print()
    print("Ovo je INVENTAR, ne presuda. Za svaku vrijednost koja stoji u dva sloja provjeri")
    print("stoji li ISTO. Za vrijednost koje nema u tijelu provjeri je li prikaz ostao")
    print("nepropagiran nakon popravka teksta.")

    if a.json:
        json.dump({"zajednicke": zaj, "osamljene": osam},
                  open(a.json, "w", encoding="utf-8"), ensure_ascii=False, indent=1)
        print(f"\n[json → {a.json}]")
    return 1 if osam else 0


if __name__ == "__main__":
    sys.exit(main())
