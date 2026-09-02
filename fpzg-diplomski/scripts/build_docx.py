# -*- coding: utf-8 -*-
"""Post-obrada diplomskog rada, usklađena s uzornim radom (Sabatti).

Razlike prema ranijoj verziji: bez uvlake prvog retka, prikazi u 11 pt,
natpisi s točkom, izvori u 10 pt, numeracija stranica od tijela rada,
popisi prikaza kao zasebni naslovi na kraju.
"""
import copy, json, os, re
import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Sve ovisno o pojedinom radu dolazi iz rad.json u radnoj mapi; nista nije zakucano.
import pathlib as _pl
KONF = json.loads(_pl.Path("rad.json").read_text(encoding="utf-8")) \
    if _pl.Path("rad.json").exists() else {}
DAT = KONF.get("docx", "rad.docx")

# Boja je smirena na traženje korisnika: tamna škriljasto-plava umjesto jarke,
# neutralni sivi rubovi i jedva vidljiva traka. Zadržava se raspoznatljivost
# zaglavlja bez dojma prezentacije.
ZAGLAVLJE, TRAKA, RUB = "3A5573", "F1F4F8", "A9B3BF"

PT_TIJELO, PT_NATPIS, PT_IZVOR, PT_TABLICA = 12, 11, 10, 11
PT_H1, PT_H2 = 14, 12
NS_SLIKA = "{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
SIRINA_TWIPS = 9070          # 16,02 cm uz margine 2,54 cm

d = docx.Document(DAT)
SM = {s.name: s for s in d.styles}


def el(tag, **atr):
    e = OxmlElement(tag)
    for k, v in atr.items():
        e.set(qn("w:" + k), v)
    return e


def ima_sliku(p):
    return any(r._element.findall(".//" + NS_SLIKA) for r in p.runs)


def _rpr(bold=False, sz=PT_TIJELO):
    rPr = OxmlElement("w:rPr")
    rPr.append(el("w:rFonts", ascii="Times New Roman", hAnsi="Times New Roman"))
    if bold:
        rPr.append(el("w:b", val="true"))
    rPr.append(el("w:color", val="000000"))
    rPr.append(el("w:sz", val=str(sz * 2)))
    return rPr


# --- 0. margine kao u uzoru -------------------------------------------------
for s in d.sections:
    s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Cm(2.54)

# --- 1. markeri prijeloma stranice -----------------------------------------
pars = d.paragraphs
for i, p in enumerate(pars):
    if p.text.strip() == "[[PB]]":
        for nxt in pars[i+1:]:
            if nxt.text.strip():
                nxt.paragraph_format.page_break_before = True
                break
        p._element.getparent().remove(p._element)

# --- 2. stilovi naslova -----------------------------------------------------
for nm, sz, sb, sa in (("Heading 1", PT_H1, 18, 12), ("Heading 2", PT_H2, 12, 6)):
    st = SM[nm]
    st.font.name = "Times New Roman"; st.font.size = Pt(sz); st.font.bold = True
    st.font.color.rgb = RGBColor(0, 0, 0)
    pf = st.paragraph_format
    pf.space_before = Pt(sb); pf.space_after = Pt(sa)
    pf.line_spacing = 1.5
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.keep_with_next = True

# --- 3. naslovnice i izjava -------------------------------------------------
pars = d.paragraphs
i_izjava = next(i for i, p in enumerate(pars)
                if p.text.strip().startswith("Izjava o akademskoj"))
NASLOV_RADA = KONF.get("naslov", "")
DESNO = ("Mentor:", "Student:", "JMBAG:")

for p in pars[:i_izjava]:
    tx = p.text.strip(); pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(0); pf.space_after = Pt(6)
    pf.first_line_indent = Cm(0); pf.left_indent = Cm(0)
    pf.alignment = (WD_ALIGN_PARAGRAPH.RIGHT if tx.startswith(DESNO)
                    else WD_ALIGN_PARAGRAPH.CENTER)
    for r in p.runs:
        if tx == NASLOV_RADA:
            r.font.size = Pt(16); r.font.bold = True
        elif tx == "Diplomski rad":
            r.font.size = Pt(13); r.font.bold = False
        else:
            r.font.size = Pt(PT_TIJELO); r.font.bold = False

for p in pars[i_izjava:]:
    if p.text.strip().startswith("Izjava o akademskoj"):
        p.paragraph_format.space_after = Pt(12)
        for r in p.runs:
            r.font.bold = True; r.font.size = Pt(PT_TIJELO)
        break

# --- 4. tablice -------------------------------------------------------------
def sjena(cell, hex_):
    cell._tc.get_or_add_tcPr().append(el("w:shd", val="clear", color="auto", fill=hex_))


def rubovi(tbl):
    b = OxmlElement("w:tblBorders")
    for kraj in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b.append(el(f"w:{kraj}", val="single", sz="4", color=RUB))
    tbl._tbl.tblPr.append(b)


# Numerirani prikaz je tablica kojoj NEPOSREDNO PRETHODI natpis "Tablica N.".
# Sve ostale (popis tvrdnji u prilogu, tablica replikacije) su dio teksta, ne
# prikazi: ne dobivaju boju zaglavlja i smiju se prelomiti preko stranica.
# Ranije je pravilo bilo pozicijsko (zadnja tablica), pa je dodavanje jos jednog
# priloga vratilo boju tablici koja je nije smjela imati.
_NATPIS_RE = re.compile(r"^(Tablica|Grafikon|Slika)\s+(\d+)\s*\.")
# Redoslijed je jedini stabilan kljuc: lxml pri svakom pristupu gradi novi proxy,
# pa ni `is` ni `id()` nad elementima nisu pouzdani.
_je_prikaz = []
_zadnji_natpis = False
for _el in d.element.body.iterchildren():
    if _el.tag.endswith("}p"):
        _txt = "".join(_t.text or "" for _t in _el.iter(qn("w:t"))).strip()
        if _txt:
            _zadnji_natpis = bool(_NATPIS_RE.match(_txt))
    elif _el.tag.endswith("}tbl"):
        _je_prikaz.append(_zadnji_natpis)
        _zadnji_natpis = False
for _i, t in enumerate(d.tables):
    _prikaz = _je_prikaz[_i] if _i < len(_je_prikaz) else True
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    rubovi(t)
    if not _prikaz and len(t.columns) == 2:
        # sirine se moraju upisati u w:tblGrid i w:tcW uz fiksni raspored,
        # inace Word racuna stupce sam i "Br." ostaje neopravdano sirok
        t.autofit = False
        lay = OxmlElement("w:tblLayout"); lay.set(qn("w:type"), "fixed")
        t._tbl.tblPr.append(lay)
        grid = t._tbl.find(qn("w:tblGrid"))
        if grid is not None:
            for gc, sirina in zip(grid.findall(qn("w:gridCol")), (907, 8163)):
                gc.set(qn("w:w"), str(sirina))
        for red in t.rows:
            for c, sirina in zip(red.cells, (Cm(1.6), Cm(14.4))):
                c.width = sirina
    for ri, row in enumerate(t.rows):
        trPr = row._tr.get_or_add_trPr()
        trPr.append(el("w:cantSplit", val="true"))
        if ri == 0:
            trPr.append(el("w:tblHeader", val="true"))
        for c in row.cells:
            if _prikaz:
                if ri == 0:
                    sjena(c, ZAGLAVLJE)
                elif ri % 2 == 0:
                    sjena(c, TRAKA)
            for p in c.paragraphs:
                pf = p.paragraph_format
                pf.line_spacing = 1.0
                pf.space_after = Pt(2); pf.space_before = Pt(2)
                pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
                pf.first_line_indent = Cm(0); pf.left_indent = Cm(0)
                pf.keep_with_next = True; pf.keep_together = True
                for r in p.runs:
                    r.font.size = Pt(PT_TABLICA)
                    if ri == 0:
                        r.font.bold = True
                        if _prikaz:
                            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

# --- 5. natpisi, slike, izvori ---------------------------------------------
JE_NATPIS = re.compile(r"^(Tablica|Grafikon|Slika)\s+(\d+)\s*\.")

for p in d.paragraphs[i_izjava:]:
    tx = p.text.strip(); pf = p.paragraph_format
    if ima_sliku(p):
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.space_before = Pt(2); pf.space_after = Pt(2)
        pf.keep_with_next = True; pf.keep_together = True
        pf.first_line_indent = Cm(0)
    elif JE_NATPIS.match(tx):
        pf.space_before = Pt(10); pf.space_after = Pt(4)
        pf.line_spacing = 1.0
        pf.keep_with_next = True; pf.keep_together = True
        pf.first_line_indent = Cm(0)
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in p.runs:
            r.font.size = Pt(PT_NATPIS); r.font.bold = True
    elif tx.startswith(("Izvor:", "Napomena:")):
        pf.space_before = Pt(2); pf.space_after = Pt(12)
        pf.line_spacing = 1.0
        pf.keep_with_next = False
        pf.keep_together = True
        pf.first_line_indent = Cm(0)
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in p.runs:
            # Izvor ide obicnim slogom. Kurziv je ovdje suvisan: redak je vec
            # odvojen manjim stupnjem i polozajem ispod prikaza.
            r.font.size = Pt(PT_IZVOR); r.font.italic = False

# --- 6. geometrija odlomaka: BEZ uvlake, razmak 6 pt (kao u uzoru) ---------
i_lit = next((i for i, p in enumerate(d.paragraphs)
              if p.text.strip() == "Literatura"), len(d.paragraphs))
i_prilog = next((i for i, p in enumerate(d.paragraphs)
                 if p.text.strip().startswith("Prilog 1")), len(d.paragraphs))
i_saz = next((i for i, p in enumerate(d.paragraphs)
              if p.text.strip() == "Sažetak"), len(d.paragraphs))

for i, p in enumerate(d.paragraphs):
    tx = p.text.strip()
    if i < i_izjava or p.style.name.startswith("Heading"):
        continue
    if not tx and not ima_sliku(p):
        continue
    pf = p.paragraph_format

    if i_lit < i < i_prilog:                       # popis literature
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.left_indent = Cm(1.25); pf.first_line_indent = Cm(-1.25)
        pf.line_spacing = 1.5
        pf.space_after = Pt(6); pf.space_before = Pt(0)
        for r in p.runs:
            r.font.size = Pt(PT_TIJELO)
        continue

    if i_prilog < i < i_saz:                       # Prilog: upitnik ide zbijeno
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.first_line_indent = Cm(0)
        if p.style.name == "Compact":              # ponuđeni odgovori
            pf.line_spacing = 1.0
            pf.space_before = Pt(0); pf.space_after = Pt(1)
            pf.left_indent = Cm(1.0)
            pf.keep_together = True
        else:                                      # tekst pitanja i napomene
            pf.line_spacing = 1.15
            pf.space_before = Pt(8); pf.space_after = Pt(3)
            pf.left_indent = Cm(0)
            pf.keep_with_next = True
        for r in p.runs:
            r.font.size = Pt(PT_TIJELO)
        continue

    if i >= i_saz:                                 # sažetak i summary: 11 pt
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.line_spacing = 1.5
        pf.space_after = Pt(6)
        pf.first_line_indent = Cm(0); pf.left_indent = Cm(0)
        for r in p.runs:
            r.font.size = Pt(PT_NATPIS)
        if tx in ("Sažetak", "Summary"):
            for r in p.runs:
                r.font.bold = True; r.font.italic = True
        continue

    if ima_sliku(p) or JE_NATPIS.match(tx) or tx.startswith(("Izvor:", "Napomena:")):
        pf.first_line_indent = Cm(0)
        continue

    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    pf.first_line_indent = Cm(0)                   # uzor nema uvlaku nigdje
    pf.left_indent = Cm(0)

for shp in d.inline_shapes:
    ratio = shp.height / shp.width
    shp.width = Cm(15.5); shp.height = Cm(15.5 * ratio)

# --- 7. unakrsne reference --------------------------------------------------
_id = [1000]


def polje(par, instr, rezultat, poslije=None, bookmark=None, bold=False, sz=PT_TIJELO):
    novi = []
    if bookmark:
        _id[0] += 1
        novi.append(el("w:bookmarkStart", id=str(_id[0]), name=bookmark))
    r1 = OxmlElement("w:r"); r1.append(_rpr(bold, sz)); r1.append(el("w:fldChar", fldCharType="begin"))
    r2 = OxmlElement("w:r"); r2.append(_rpr(bold, sz))
    it = OxmlElement("w:instrText")
    it.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    it.text = instr; r2.append(it)
    r3 = OxmlElement("w:r"); r3.append(_rpr(bold, sz)); r3.append(el("w:fldChar", fldCharType="separate"))
    r4 = OxmlElement("w:r"); r4.append(_rpr(bold, sz))
    t4 = OxmlElement("w:t"); t4.text = rezultat; r4.append(t4)
    r5 = OxmlElement("w:r"); r5.append(_rpr(bold, sz)); r5.append(el("w:fldChar", fldCharType="end"))
    novi += [r1, r2, r3, r4, r5]
    if bookmark:
        novi.append(el("w:bookmarkEnd", id=str(_id[0])))
    ref = poslije if poslije is not None else par._p
    for e in novi:
        ref.addnext(e); ref = e
    return novi


def kljuc(vrsta, broj):
    return f"_Ref_{'tab' if vrsta == 'Tablica' else 'graf'}{broj}"


natpisa = 0
for p in d.paragraphs:
    m = JE_NATPIS.match(p.text.strip())
    if not m:
        continue
    vrsta, broj = m.group(1), m.group(2)
    ostatak = p.text.strip()[m.end():]
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    r0 = p.add_run(vrsta + " "); r0.font.size = Pt(PT_NATPIS); r0.font.bold = True
    polje(p, f" SEQ {vrsta} \\* ARABIC ", broj, poslije=r0._element,
          bookmark=kljuc(vrsta, broj), bold=True, sz=PT_NATPIS)
    rz = p.add_run("." + ostatak); rz.font.size = Pt(PT_NATPIS); rz.font.bold = True
    natpisa += 1

POZIV = re.compile(r"\b([Tt]ablic(?:a|e|i|u|om)|[Gg]rafikon(?:a|u|om|i|e)?)(\s+)(\d+)\b")
refova = 0
for p in d.paragraphs:
    if JE_NATPIS.match(p.text.strip()) or p.style.name.startswith("Heading"):
        continue
    for r in list(p.runs):
        m = POZIV.search(r.text)
        if not m:
            continue
        vrsta = "Tablica" if m.group(1).lower().startswith("tablic") else "Grafikon"
        broj = m.group(3)
        if not 1 <= int(broj) <= 8:
            continue
        rep = r.text[m.end(3):]
        r.text = r.text[:m.end(2)]
        novi = polje(p, f" REF {kljuc(vrsta, broj)} \\h ", broj, poslije=r._element)
        if rep:
            rr = OxmlElement("w:r")
            if r._element.find(qn("w:rPr")) is not None:
                rr.append(copy.deepcopy(r._element.find(qn("w:rPr"))))
            tt = OxmlElement("w:t")
            tt.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            tt.text = rep; rr.append(tt)
            novi[-1].addnext(rr)
        refova += 1

# --- 8. prijelomi stranica --------------------------------------------------
PRIJELOM = ("Sadržaj", "Literatura", "Popis tablica", "Popis grafikona", "Prilog")
for p in d.paragraphs:
    if p.style.name == "Heading 1":
        p.paragraph_format.page_break_before = p.text.strip().startswith(PRIJELOM)

# --- 9. stil popisa prikaza -------------------------------------------------
stils = d.styles.element
if not any(s.get(qn("w:styleId")) == "TableofFigures" for s in stils.findall(qn("w:style"))):
    st = el("w:style", type="paragraph", styleId="TableofFigures")
    st.append(el("w:name", val="table of figures"))
    st.append(el("w:basedOn", val="Normal"))
    pPr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tabs.append(el("w:tab", val="right", leader="dot", pos=str(SIRINA_TWIPS)))
    pPr.append(tabs)
    pPr.append(el("w:spacing", after="0", line="360", lineRule="auto"))
    pPr.append(el("w:jc", val="both"))
    st.append(pPr)
    rPr = OxmlElement("w:rPr")
    rPr.append(el("w:rFonts", ascii="Times New Roman", hAnsi="Times New Roman"))
    rPr.append(el("w:sz", val="24"))
    st.append(rPr)
    stils.append(st)

# --- 9b. stil fusnota -------------------------------------------------------
# pandoc ostavlja FootnoteText prazan, pa fusnote naslijede Normal i ispadnu
# 12 pt s proredom 1,5 — jednake tijelu teksta. Fusnota mora biti manja i zbijena.
_st = d.styles.element
for st in _st.findall(qn("w:style")):
    if st.get(qn("w:styleId")) != "FootnoteText":
        continue
    for tag in ("w:pPr", "w:rPr"):
        for e in st.findall(qn(tag)):
            st.remove(e)
    pPr = OxmlElement("w:pPr")
    pPr.append(el("w:spacing", before="0", after="0", line="240", lineRule="auto"))
    pPr.append(el("w:jc", val="both"))
    pPr.append(el("w:ind", firstLine="0", left="0"))
    st.append(pPr)
    rPr = OxmlElement("w:rPr")
    rPr.append(el("w:rFonts", ascii="Times New Roman", hAnsi="Times New Roman"))
    rPr.append(el("w:sz", val="20"))          # 10 pt
    rPr.append(el("w:szCs", val="20"))
    st.append(rPr)

# --- 10. Sadržaj (prije prijeloma sekcije) ---------------------------------
sek = next(p for p in d.paragraphs if p.text.strip() == "[[SEC]]")
nas = sek.insert_paragraph_before("Sadržaj")
nas.style = SM["Heading 1"]
nas.paragraph_format.page_break_before = True
toc = sek.insert_paragraph_before("")
f = el("w:fldSimple", instr='TOC \\o "1-2" \\h \\z \\u')
r = OxmlElement("w:r"); r.append(_rpr())
t_ = OxmlElement("w:t"); t_.text = "[Sadržaj — u Wordu: desni klik → Update Field]"
r.append(t_); f.append(r); toc._p.append(f)

# --- 11. popisi prikaza (prije Priloga) ------------------------------------
STRANICE = json.load(open("natpisi_stranice.json", encoding="utf-8")) \
    if os.path.exists("natpisi_stranice.json") else {}
NATPISI = [(p.text.strip(), kljuc(*JE_NATPIS.match(p.text.strip()).groups()))
           for p in d.paragraphs if JE_NATPIS.match(p.text.strip())]

def _sidro(dok):
    """Odlomak pred koji idu popisi prikaza.

    Prije je ovdje stajao `next(...)` bez zadane vrijednosti, pa je rad BEZ PRILOGA rušio
    izgradnju s `StopIteration` — a `references/struktura.md` istog skilla priloge navodi
    kao „prilozi (ako postoje)", dakle neobvezne. Nađeno na prihvatnom testu (kolovoz
    2026.) na pravom diplomskom radu.

    Redoslijed zamjena je redoslijed strukture rada s kraja: prilozi, pa sažetak, pa
    summary. Ako ničega od toga nema, popisi idu na kraj dokumenta.
    """
    for uzorak in ("Prilog 1", "Sažetak", "Summary"):
        for p in dok.paragraphs:
            if p.text.strip().startswith(uzorak):
                return p
    return dok.add_paragraph("")


sidro = _sidro(d)


def popis(vrsta, naslov):
    h = sidro.insert_paragraph_before(naslov)
    h.style = SM["Heading 1"]
    h.paragraph_format.page_break_before = True
    stavke = [(t, b, str(STRANICE[b])) for t, b in NATPISI
              if t.startswith(vrsta) and b in STRANICE]

    def red():
        p = sidro.insert_paragraph_before("")
        p.style = SM.get("Table of Figures", SM["Normal"])
        pf = p.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.line_spacing = 1.5; pf.space_after = Pt(0)
        pf.first_line_indent = Cm(0); pf.left_indent = Cm(0)
        pf.tab_stops.add_tab_stop(Cm(16.02), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        return p

    if not stavke:
        p = red()
        fs = el("w:fldSimple", instr=f' TOC \\h \\z \\c "{vrsta}" ')
        rr = OxmlElement("w:r"); rr.append(_rpr())
        tt = OxmlElement("w:t"); tt.text = f"[{naslov} — u Wordu: desni klik → Update Field]"
        rr.append(tt); fs.append(rr); p._p.append(fs)
        return
    for k, (tekst, bm, broj) in enumerate(stavke):
        p = red()
        if k == 0:
            for tip, txt in (("begin", None), (None, f' TOC \\h \\z \\c "{vrsta}" '),
                             ("separate", None)):
                rr = OxmlElement("w:r"); rr.append(_rpr())
                if txt is None:
                    rr.append(el("w:fldChar", fldCharType=tip))
                else:
                    it = OxmlElement("w:instrText")
                    it.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                    it.text = txt; rr.append(it)
                p._p.append(rr)
        rr = OxmlElement("w:r"); rr.append(_rpr())
        tt = OxmlElement("w:t")
        tt.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        tt.text = tekst; rr.append(tt); p._p.append(rr)
        rr = OxmlElement("w:r"); rr.append(_rpr()); rr.append(OxmlElement("w:tab"))
        p._p.append(rr)
        fp = el("w:fldSimple", instr=f" PAGEREF {bm} \\h ")
        rr = OxmlElement("w:r"); rr.append(_rpr())
        tt = OxmlElement("w:t"); tt.text = broj; rr.append(tt); fp.append(rr)
        p._p.append(fp)
        if k == len(stavke) - 1:
            rr = OxmlElement("w:r"); rr.append(_rpr())
            rr.append(el("w:fldChar", fldCharType="end")); p._p.append(rr)


popis("Tablica", "Popis tablica")
popis("Grafikon", "Popis grafikona")

# --- 12. prijelom sekcije i numeracija stranica ----------------------------
# Predtekst ostaje bez broja; numeracija kreće od 1 na „1. Uvod".
sek = next(p for p in d.paragraphs if p.text.strip() == "[[SEC]]")
zadnji_pred = sek._p.getprevious()
kraj_sectPr = d.element.body.find(qn("w:sectPr"))
nova = copy.deepcopy(kraj_sectPr)
for tag in ("w:footerReference", "w:headerReference", "w:pgNumType"):
    for e in nova.findall(qn(tag)):
        nova.remove(e)
pPr = zadnji_pred.find(qn("w:pPr"))
if pPr is None:
    pPr = OxmlElement("w:pPr"); zadnji_pred.insert(0, pPr)
pPr.append(nova)
sek._p.getparent().remove(sek._p)

s2 = d.sections[1]
s2.footer.is_linked_to_previous = False
d.sections[0].footer.is_linked_to_previous = False
for p in list(d.sections[0].footer.paragraphs):
    p.text = ""
fp = s2.footer.paragraphs[0]
fp.text = ""
fp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
fld = el("w:fldSimple", instr=" PAGE ")
rr = OxmlElement("w:r"); rr.append(_rpr())
tt = OxmlElement("w:t"); tt.text = "1"; rr.append(tt); fld.append(rr)
fp._p.append(fld)
pg = el("w:pgNumType", start="1")
s2._sectPr.append(pg)

d.save(DAT)
print(f"build4: {len(d.tables)} tablica, {len(d.inline_shapes)} grafikona, "
      f"{natpisa} natpisa, {refova} unakrsnih referenci, {len(d.sections)} sekcije")
