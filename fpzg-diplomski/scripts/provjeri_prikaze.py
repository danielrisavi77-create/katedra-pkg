#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Provjera predajnog dokumenta: lome li se prikazi, jesu li polja uravnotežena,
je li išta ostalo nedovršeno.

Ne provjerava sadržaj — to rade check_rules.py i audit_all.py. Ovo hvata kvarove
koje se inače primijeti tek kad rad izađe iz pisača.

    python3 provjeri_prikaze.py rad.docx
"""
import argparse
import json
import re
import subprocess
import sys
import zipfile

try:
    import docx
except ImportError:
    sys.exit("nedostaje python-docx")

NATPIS = re.compile(r"^(Tablica|Grafikon|Slika)\s+(\d+)\s*[.:]")


def norm(s):
    return re.sub(r"[^0-9a-zšđčćž]+", " ", s.replace("­", "").lower()).strip()


def stranice_pdfa(pdf):
    n = int(subprocess.run(["pdfinfo", pdf], capture_output=True, text=True)
            .stdout.split("Pages:")[1].split()[0])
    return [norm(subprocess.run(["pdftotext", "-f", str(i), "-l", str(i), pdf, "-"],
                                capture_output=True, text=True).stdout)
            for i in range(1, n + 1)], n


def main(put):
    pdf = put[:-5] + ".pdf"
    d = docx.Document(put)
    nalazi = []

    # ── 1. tablice: ponovljeno zaglavlje na dvije stranice = tablica se lomi ──
    try:
        pg, n = stranice_pdfa(pdf)
    except Exception as e:
        print(f"⚠ PDF nije dostupan ({e}); provjera lomljenja preskočena")
        pg, n = [], 0

    if pg:
        for k, t in enumerate(d.tables, 1):
            zag = norm(" ".join(c.text for c in t.rows[0].cells))
            if not zag:
                continue
            pojave = [i + 1 for i, s in enumerate(pg) if zag in s]
            if len(pojave) > 1:
                nalazi.append(f"tablica {k} se LOMI preko stranica {pojave}")
        print(f"tablica: {len(d.tables)} · lome se: "
              f"{len([x for x in nalazi if 'LOMI' in x]) or 'nijedna ✅'}")

    # ── 2. polja ─────────────────────────────────────────────────────────────
    x = zipfile.ZipFile(put).read("word/document.xml").decode("utf-8")
    b = len(re.findall(r'fldCharType="begin"', x))
    e = len(re.findall(r'fldCharType="end"', x))
    s_ = len(re.findall(r'fldCharType="separate"', x))
    if not (b == e == s_):
        nalazi.append(f"polja NISU uravnotežena: begin={b} separate={s_} end={e}")
    print(f"polja: {b} begin / {s_} separate / {e} end "
          f"{'✅' if b == e == s_ else '❌'}")
    # granica rijeci: bez nje "REF " pogada i unutar "PAGEREF "
    for tip, uzorak in (("SEQ", r"\bSEQ "), ("REF", r"(?<!PAGE)\bREF "),
                        ("PAGEREF", r"\bPAGEREF "), ("TOC", r"\bTOC ")):
        print(f"   {tip:<8} {len(re.findall(uzorak, x))}")

    # ── 3. fusnote ───────────────────────────────────────────────────────────
    z = zipfile.ZipFile(put)
    if "word/footnotes.xml" in z.namelist():
        fx = z.read("word/footnotes.xml").decode("utf-8")
        W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        from xml.etree import ElementTree as ET
        prave = [f for f in ET.fromstring(fx).findall(W + "footnote")
                 if not f.get(W + "type")]
        sidra = re.findall(r'<w:footnoteReference[^>]*w:id="(\d+)"', x)
        if len(prave) != len(sidra):
            nalazi.append(f"fusnote: {len(prave)} tekstova, {len(sidra)} sidara")
        st = z.read("word/styles.xml").decode("utf-8")
        m = re.search(r'<w:style [^>]*w:styleId="FootnoteText".*?</w:style>', st, re.S)
        if m and "w:sz" not in m.group(0):
            nalazi.append("FootnoteText nema postavljenu veličinu — fusnote će "
                          "naslijediti Normal i ispasti jednake tijelu teksta")
        print(f"fusnote: {len(prave)} / sidara {len(sidra)} "
              f"{'✅' if len(prave) == len(sidra) else '❌'}")

    # ── 4. nedovršeno ────────────────────────────────────────────────────────
    tekst = "\n".join(p.text for p in d.paragraphs)
    for t in d.tables:
        for r in t.rows:
            for c in r.cells:
                tekst += "\n" + c.text
    for igla, poruka in (("PROVJERI STR", "citat bez broja stranice"),
                         ("TREBA IZVOR", "tvrdnja bez izvora"),
                         ("[[PB]]", "neobrađen marker prijeloma"),
                         ("[[SEC]]", "neobrađen marker sekcije"),
                         ("IME I PREZIME", "placeholder za ime")):
        k = tekst.count(igla)
        if k:
            nalazi.append(f"{poruka}: {k}×")
    print(f"em crtica (—): {tekst.count('—')}")

    # ── 5. natpisi ───────────────────────────────────────────────────────────
    # Popisi prikaza na kraju rada ponavljaju natpise; oni nisu prikazi i nemaju
    # izvor. Skeniranje staje na prvom naslovu popisa.
    pars = d.paragraphs
    kraj = next((i for i, p in enumerate(pars)
                 if re.match(r"(?i)^popis (tablica|grafikona|ilustracija|slika)$",
                             p.text.strip())), len(pars))
    pars = pars[:kraj]
    natpisi = [p.text.strip() for p in pars if NATPIS.match(p.text.strip())]
    bez_izvora = 0
    for i, p in enumerate(pars):
        if NATPIS.match(p.text.strip()):
            okolina = " ".join(q.text for q in pars[i:i + 6])
            if "Izvor:" not in okolina:
                bez_izvora += 1
    if bez_izvora:
        nalazi.append(f"prikaza bez retka „Izvor:”: {bez_izvora}")
    print(f"natpisa: {len(natpisi)} · bez izvora: {bez_izvora or '0 ✅'}")

    print()
    if nalazi:
        print("NALAZI:")
        for x_ in nalazi:
            print("  ❌", x_)
    else:
        print("✅ nema nalaza")
    return 1 if nalazi else 0


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("rad")
    sys.exit(main(ap.parse_args().rad))
