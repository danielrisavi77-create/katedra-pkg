#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Stilske navike koje se u dužem tekstu nakupe neprimjetno.

Hvata ono što ni provjera pravila ni audit ne gledaju: interpunkcijske tikove
koji se pojedinačno čine bezazlenima, a na pedeset stranica postanu manir.

    python3 provjeri_stil.py rad.docx
    python3 provjeri_stil.py pog*.md
"""
import argparse
import re
import sys
from collections import Counter

PRAG_DVOTOCKA = 4.0        # na 100 rečenica
PRAG_TIRE = 0

NABRAJANJE = re.compile(
    r"(?:dvije|dva|tri|triju|četiri|četiriju|pet|šest|sedam|osam|devet|deset|"
    r"nekoliko|sljedeć\w+|ovako|kako slijedi|sastavnic\w+|kategorij\w+|"
    r"varijabl\w+|cjelin\w+|dimenzij\w+|skupin\w+|razin\w+|korak\w*|faz\w+)"
    r"[^.:]{0,40}$", re.I)


def ucitaj(putovi):
    tekst = []
    for put in putovi:
        if put.endswith(".docx"):
            import docx
            d = docx.Document(put)
            # Mjeri se SAMO tijelo rada. Naslovnice nose "Mentor:" i "Student:",
            # popis literature dvotocke u bibliografskim jedinicama, prilog
            # doslovno prenesena pitanja — nista od toga nije autorska proza.
            poc = next((i for i, p in enumerate(d.paragraphs)
                        if re.match(r"(?i)^1\.\s*uvod$", p.text.strip())), 0)
            kraj = next((i for i, p in enumerate(d.paragraphs)
                         if i > poc and re.match(r"(?i)^(literatura|popis literature)$",
                                                 p.text.strip())), len(d.paragraphs))
            for p in d.paragraphs[poc:kraj]:
                t = p.text.strip()
                if not t or re.match(r"^(Tablica|Grafikon|Slika)\s+\d+\s*[.:]", t):
                    continue
                if t.startswith(("Izvor:", "Napomena:")):
                    continue
                tekst.append(t)
        else:
            t = open(put, encoding="utf-8").read()
            t = re.sub(r"^\|.*$", "", t, flags=re.M)          # tablice
            t = re.sub(r"^\*\*(Tablica|Grafikon|Slika).*$", "", t, flags=re.M)
            t = re.sub(r"^(Izvor|Napomena):.*$", "", t, flags=re.M)
            t = re.sub(r"^#.*$", "", t, flags=re.M)
            t = re.sub(r"^\[\^\w+\]:", "", t, flags=re.M)     # oznake fusnota
            tekst += [r for r in t.split("\n") if r.strip()]
    return tekst


def main(putovi):
    redci = ucitaj(putovi)
    spoj = "\n".join(redci)
    recenica = len(re.findall(r"[.!?](?=\s|$)", spoj))
    nalazi = []

    # ── dvotočka ─────────────────────────────────────────────────────────────
    razrada, nabraja = [], 0
    for red in redci:
        for m in re.finditer(r"(?<!\d):(?!\s*\d)", red):
            prije = red[:m.start()]
            if NABRAJANJE.search(prije):
                nabraja += 1
            else:
                razrada.append((prije[-60:], red[m.end():m.end() + 60].strip()))
    uk = len(razrada) + nabraja
    gust = uk / max(recenica, 1) * 100
    print(f"dvotočaka: {uk} na {recenica} rečenica = {gust:.1f} na 100 "
          f"(prag {PRAG_DVOTOCKA})")
    print(f"   uvodi nabrajanje: {nabraja}   razrađuje prethodnu tvrdnju: {len(razrada)}")
    if gust > PRAG_DVOTOCKA:
        nalazi.append(f"dvotočka {gust:.1f} na 100 rečenica — previše")
    if razrada:
        print("\n   kandidati za točku i novu rečenicu:")
        for a, b in razrada[:12]:
            print(f"     …{a} [:] {b}…")
        if len(razrada) > 12:
            print(f"     … i još {len(razrada) - 12}")

    # ── duga crtica ──────────────────────────────────────────────────────────
    em = spoj.count("—")
    print(f"\ndugih crtica (—): {em}")
    if em > PRAG_TIRE:
        nalazi.append(f"duga crtica: {em}× — u hrvatskom akademskom tekstu ne stoji")

    # ── ponovljeni otvarači rečenica ─────────────────────────────────────────
    otvaraci = Counter()
    for r in re.split(r"(?<=[.!?])\s+", spoj):
        w = r.strip().split()
        if len(w) >= 2:
            otvaraci[" ".join(w[:2]).lower()] += 1
    cesti = [(k, v) for k, v in otvaraci.most_common(8) if v >= 4]
    if cesti:
        print("\nponovljeni otvarači rečenica (≥4×):")
        for k, v in cesti:
            print(f"   {v:>2}×  {k}")

    print()
    if nalazi:
        print("NALAZI:")
        for n in nalazi:
            print("  ❌", n)
    else:
        print("✅ nema nalaza")
    return 1 if nalazi else 0


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("rad", nargs="+")
    sys.exit(main(ap.parse_args().rad))
