#!/usr/bin/env python3
"""
check_argument.py — razina argumenta, ne forma.

Rad koji prođe sve formalne provjere i dalje ne nosi peticu ako nema tezu koja
se provlači kroz poglavlja i zaključak koji zatvara krug. Ovo je jedini alat u
Katedri koji gleda TO, a ne je li font Times New Roman.

Mjeri sedam dimenzija:
  1. TEZA              ima li uvod obranjivu tvrdnju, a ne samo najavu sadržaja
  2. ZAKLJUČAK         vraćaju li se pojmovi iz uvoda u zaključak (preklapanje %)
  3. ISTRAŽIVAČKO PIT. je li pitanje ili cilj rada izrečen izrijekom
  4. PROPORCIJE        raspodjela riječi + samo metodološki relevantni strukturni signali
  5. VLASTITI DOPRINOS prikazi su signal samo gdje ih metodologija očekuje
  6. DESKRIPTIVNOST    soft signal odlomaka bez prepoznate analitičke veze
  7. CITATNA GUSTOĆA   poglavlje bez ijednog citata

SVE JE HEURISTIKA. Alat ne presuđuje nego nudi kandidate i brojke: „kandidat za
tezu", „provjeri", „čita se kao". Konačnu ocjenu donosi čovjek.

    python3 <KATEDRA_SKILL>/scripts/check_argument.py ./rad.docx --metodologija theoretical
    python3 <KATEDRA_SKILL>/scripts/check_argument.py ./rad.docx --profil ./.katedra/resolved_profile.json
    python3 <KATEDRA_SKILL>/scripts/check_argument.py ./rad.docx --citatni-stil ieee
    python3 <KATEDRA_SKILL>/scripts/check_argument.py ./rad.docx --json ./.katedra/arg.json

Izlazni kodovi:
  0  nema nalaza
  1  ima nalaza
  2  greška (nema datoteke / ne može se pročitati)
"""
import argparse
import json
import os
import re
import statistics as st
import sys
from collections import Counter

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import hr_text as H
import argument_methodology as M
import citation_dialects as C

OK, UPOZ, LOSE = "✅", "⚠️", "❌"

# --------------------------------------------------------------- rječnici

STOP = set("""
a ako ali baš bez bi bih bila bili bilo bio biste bismo biti bit ce će ćemo ćete
ću da dakle dalje dio do dok dosta drugi drugo dva i iako ih ili
im ima imaju imati iz iza između izvan ja je jedan jedna jedno jer jest jesu joj
još ju kada kad kako kao koja koje kojeg kojem koji kojih kojima koju koliko
kroz li me među mi mnogo mogu moći moraju mora može mu na nad nakon nam nas naš
naša naše naši ne nego neka neki neko nekoliko nema nešto ni nije nisu niti no
njih njihov njemu o od odnosno oko on ona onda one oni ono ova ovaj ove ovi ovo
ovom ovu pa po pod pored posebno poslije prema pri prije s sa sam samo se sebe
si smo su svaki sve svi svoj svoje svojih ta taj tako takve takvi te tebe
tek ti time to toga tom tome tu u uz uzduž vam vas već vi više za zar zato
zbog što sto the of and in
također upravo osim jedne jednog jednu unutar putem
tijekom uslijed radi glede stoga ipak naime međutim
""".split())

# Vezna sredstva koja nose tvrdnju, ne najavu sadržaja.
TVRDNJSKE = [
    "jer", "stoga", "zbog toga", "zbog čega", "za razliku od", "umjesto",
    "budući da", "s obzirom na to da", "upravo zato", "što znači da",
    "ne zato što", "nego zato što", "posljedica je", "dovodi do", "uzrok",
    "premda", "iako", "no", "međutim", "dok", "nasuprot",
]
# Najavni obrasci — rečenica koja ih sadrži opisuje rad, a ne svijet.
NAJAVNI = [
    "analizirat će se", "analizirati će se", "prikazat će se", "prikazati će se",
    "cilj rada", "cilj ovog rada", "cilj ovoga rada", "svrha rada", "predmet rada",
    "u ovom radu", "u ovome radu", "u radu se", "rad se sastoji", "rad je podijeljen",
    "u prvom poglavlju", "u drugom poglavlju", "u trećem poglavlju",
    "u nastavku rada", "obradit će se", "istražit će se", "bit će prikazan",
    "bit će riječi", "opisat će se", "razmotrit će se", "objasnit će se",
    "definirat će se", "u zaključku", "struktura rada", "posljednje poglavlje",
]
PITANJE = [
    "istraživačko pitanje", "istraživačka pitanja", "cilj rada", "cilj ovog rada",
    "cilj ovoga rada", "svrha rada", "hipotez", "polazna teza", "teza rada",
    "pitanje na koje", "ovaj rad odgovara na",
]
# Uzročno-posljedični veznici — nose analizu, ne opis.
UZROCNI = [
    "jer", "zato što", "budući da", "s obzirom na to da", "stoga", "zbog toga",
    "zbog čega", "uslijed toga", "posljedično", "prema tome", "dakle", "time",
    "čime", "tako da", "dovodi do", "doveo je do", "dovelo je do", "rezultira",
    "rezultat je", "uzrokuje", "uzrokovalo", "utječe na", "utjecalo je",
    "proizlazi", "objašnjava se", "razlog je", "razlog tome", "posljedica",
    "što je omogućilo", "ovisi o", "uvjetovano", "pokazuje da", "upućuje na to da",
    "što znači", "pri čemu", "što upućuje", "što pokazuje", "što potvrđuje",
    "zahvaljujući", "s obzirom na", "unatoč tome", "usprkos tome",
]


# B09: analitičnost nije samo uzročnost. Kontrast, konceptualno razlikovanje,
# inferencija i evaluacija također su analitički signali, osobito u teorijskim,
# pravnim i povijesnim radovima.
ANALITICKI_SIGNALI = tuple(UZROCNI) + (
    "za razliku od", "nasuprot", "u usporedbi s", "s jedne strane",
    "s druge strane", "međutim", "dok", "nego", "razlikuje", "razlikuju",
    "proturječi", "podupire", "pretpostavlja", "implicira", "može se zaključiti",
    "ograničenje", "prednost", "nedostatak", "kritika", "argument",
)

VLASTITI = ["autorov", "autorova", "autorovo", "autora", "vlastit", "izrada autora",
            "izračun autora", "obrada autora", "prema izračunu autora"]
ADAPTIRANI = ["prema "]

# Popis gore je bio doslovan, pa je najkraći i najčešći hrvatski oblik — „Izvor:
# autor" — ispadao iz njega: „autor" nije podniz nijednog zapisa („autora" jest
# duži oblik, ne kraći). Posljedica je bila tvrdnja da rad nema nijedan vlastiti
# prikaz iako su svi bili autorski (nalaz 5 iz zakrpe-katedra.md). Pravilo je
# stoga na KORIJENU riječi uz granicu riječi, ne na popisu oblika.
_VLASTITI_KORIJEN_RE = re.compile(r"\bautor\w*\b", re.IGNORECASE)
# Druga strana: autorstvo koje NIJE studentovo. „Izvor: autorski tim HNB-a" ili
# „Izvor: autori projekta Obzor 2020" pripisuju rad nekome drugome i ne smiju se
# priznati kao vlastiti prikaz. Glavni znak je strukturan — vlastito ime velikim
# slovom iza korijena. Uz njega ide uzak popis zajedničkih imenica jer se
# „autori projekta" od „obrade autorice" ne razlikuje ničim strukturnim: oblik
# `autorice` je istovremeno genitiv jednine (moj rad) i nominativ množine (tuđi).
# Popis je namjerno kratak i radi kao IZNIMKA na inače sigurnom pravilu; ono što
# promakne pogriješi u mekšem smjeru (prikaz se pripiše studentu), a to je samo
# propušten signal — dok je obrnuti smjer lažna tvrdnja „rad nema nijedan
# vlastiti prikaz" na radu u kojem su svi prikazi autorski.
_TUDJE_AUTORSTVO_RE = re.compile(
    r"\bautor\w*\s+(?:tim\w*|skupin\w*|kolektiv\w*|grup\w*|projekt\w*|"
    r"studij\w*|istraživanj\w*|istrazivanj\w*|publikacij\w*)\b"
    r"|\bautor\w*\s+[A-ZČĆŽŠĐ]",
    re.IGNORECASE,
)


def je_vlastiti_prikaz(red: str, norm) -> bool:
    """Je li izvor prikaza studentov vlastiti rad."""
    t = str(red or "")
    if _TUDJE_AUTORSTVO_RE.search(t):
        return False
    if _VLASTITI_KORIJEN_RE.search(t):
        return True
    n = norm(t)
    return any(norm(v) in n for v in VLASTITI)


def je_natpis(t):
    """Natpis prikaza, a ne rečenica koja se na prikaz poziva.

    hr_text.NATPIS hvata oba oblika; „Tablica 1. pokazuje da…" je proza i mora
    ostati u mjerenju, „Tablica 1. Kretanje noćenja" je natpis i ne smije.
    """
    m = H.NATPIS.match(t or "")
    if not m:
        return False
    ostatak = t[m.end():].strip()
    return not ostatak or ostatak[0].isupper() or ostatak[0].isdigit() \
        or ostatak[0] in "„\"'"


def pogl(n):
    """Hrvatski broj poglavlja: 1 poglavlje, 2–4 poglavlja, 5+ poglavlja."""
    return f"{n} poglavlje" if n % 10 == 1 and n % 100 != 11 else f"{n} poglavlja"


def norm(s):
    s = H.bez_dijakritika(str(s or "")).lower()
    s = re.sub(r"^\s*(?:[ivxlc]+|\d+)(?:\.\d+)*\.?\s+", " ", s)
    return re.sub(r"\s+", " ", re.sub(r"[^a-z0-9]+", " ", s)).strip()


# v1.1-fix (Q3d): fiksno kraćenje na 6 znakova ostavljalo je kratke riječi
# netaknutima, pa se „porez"/„poreza", „kriza"/„krize", „zakon"/„zakona" i
# „model"/„modela" nikada nisu poklopili. Prvo se skida čest padežni nastavak,
# tek onda slijedi kraćenje. Namjerno konzervativno: korijen kraći od 4 znaka
# ne nastaje, jer lažno poklapanje ovdje šteti više od promašaja.
PADEZNI_NASTAVCI = ("ovima", "evima", "ima", "ama", "ova", "eva",
                    "om", "em", "a", "e", "i", "u")
MIN_KORIJEN = 4


def korijen(rijec, n=6):
    """Grubi hrvatski „stem": padežni nastavak dolje, pa prvih n znakova."""
    r = H.bez_dijakritika(rijec).lower()
    for nastavak in PADEZNI_NASTAVCI:
        if r.endswith(nastavak) and len(r) - len(nastavak) >= MIN_KORIJEN:
            r = r[:-len(nastavak)]
            break
    return r[:n]


# v1.1-fix (Q3c): STOP se uspoređivao s norm(rijec), a sam je pisan s
# dijakriticima, pa nijedna stop-riječ s dijakritikom nije filtrirana („već",
# „među", „između"). Lista se normalizira istim postupkom kao i tekst, pa
# usporedba vrijedi i prije i poslije popravka hr_text.bez_dijakritika za „đ".
STOP_N = {norm(w) for w in STOP}
STOP_N.discard("")


# v1.1-fix (Q4a): signal se tražio kao prefiks (' ' + signal), pa je trozvučni
# „dok" hvatao „dokument", „dokumentacija", „dokaz" i „doktrina" — posve opisno
# poglavlje ispadalo je 100 % analitično. Sada se svaki signal prevodi u uzorak
# s granicom riječi; prefiksno hvatanje hrvatske fleksije ostaje dopušteno samo
# ako je zadnja riječ signala dovoljno duga da ne hvata tuđe osnove.
MIN_OSNOVA = 5


def _uzorak_signala(izraz):
    """Uzorak za jedan analitički signal; None ako se izraz normalizira u prazno."""
    n = norm(izraz)
    if not n:
        return None
    zadnja = n.split()[-1]
    rep = r"\w*" if len(zadnja) >= MIN_OSNOVA else r"\b"
    return re.compile(r"\b" + re.escape(n) + rep)


ANALITICKI_UZORCI = tuple(u for u in (_uzorak_signala(s) for s in ANALITICKI_SIGNALI)
                          if u is not None)


# ------------------------------------------------------------- čitanje rada

def poglavlja(put, citatni_stil="autor-godina"):
    """Vrati (poglavlja, svi_izvor_redci). Poglavlje = Heading 1 + njegova proza."""
    try:
        from docx import Document
    except ImportError:
        sys.exit("Treba python-docx:  pip install python-docx --break-system-packages")
    try:
        doc = Document(put)
    except Exception as e:
        print(f"❌ {put} se ne može otvoriti kao .docx: {e}", file=sys.stderr)
        sys.exit(2)

    dialect = C.resolve_dialect(citatni_stil)
    footnotes = C.extract_docx_footnotes(put) if dialect == "legal-footnote" else {}
    pog = [{"naslov": "(prije prvog poglavlja)", "odlomci": [], "podnaslovi": [], "_citati": []}]
    izvori = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if not t:
            continue
        try:
            stil = (p.style.name or "").strip()
        except Exception:
            stil = ""
        m = re.match(r"(?i)^(?:heading|naslov)\s*(\d)$", stil)
        razina = int(m.group(1)) if m else 0
        if razina == 1:
            pog.append({"naslov": t, "odlomci": [], "podnaslovi": [], "_citati": []})
            continue
        if razina:
            pog[-1]["podnaslovi"].append(t)
            continue
        if H.IZVOR.match(t):
            izvori.append(t)
            continue
        if je_natpis(t):
            continue
        if len(t) < 40 and not t.endswith((".", "!", "?")):
            continue
        pog[-1]["odlomci"].append(t)
        pog[-1]["_citati"].extend(C.parse_citations(t, citatni_stil))
        if dialect == "legal-footnote":
            pog[-1]["_citati"].extend(C.footnote_citations_for_paragraph(p, footnotes))

    for tab in doc.tables:
        for red in tab.rows:
            for c in red.cells:
                for p in c.paragraphs:
                    t = (p.text or "").strip()
                    if H.IZVOR.match(t):
                        izvori.append(t)

    if len(pog) > 1 and not pog[0]["odlomci"]:
        pog.pop(0)
    for p in pog:
        p["tekst"] = " ".join(p["odlomci"])
        p["rijeci"] = len(H.rijeci(p["tekst"]))
        p["citati"] = len(p["_citati"])
        p["tipovi_izvora"] = dict(Counter(r.source_type for r in p["_citati"] if r.source_type != "unknown"))
    return pog, izvori


def je_literatura(naslov):
    return bool(H.NASLOV_LIT.match(naslov)) or norm(naslov).startswith(
        ("popis izvora", "popis literature", "literatura", "bibliografija",
         "popis tablica", "popis slika", "popis grafikona", "prilog", "prilozi",
         "sazetak", "summary", "abstract", "zivotopis", "sadrzaj"))


# ------------------------------------------------------------------ mjerenja

class Ocjena:
    def __init__(self):
        self.dim = []      # (naziv, stanje, brojka, [redci], [nalazi])

    def dodaj(self, naziv, stanje, brojka, redci=None, nalaz=None):
        self.dim.append({"dimenzija": naziv, "stanje": stanje, "brojka": brojka,
                         "redci": redci or [], "nalaz": nalaz or ""})

    @property
    def nalazi(self):
        return [d for d in self.dim if d["stanje"] != OK]


def dim_teza(oc, uvod):
    if uvod is None:
        oc.dodaj("teza", LOSE, "nema uvoda",
                 ["Uvod nije prepoznat (nema Heading 1 čiji naslov sadrži „uvod\")."],
                 "bez uvoda se teza ne može ni tražiti")
        return []
    kandidati = []
    for r in H.recenice(uvod["tekst"]):
        n = " " + norm(r) + " "
        ima_tvrdnju = any(f" {norm(v)} " in n for v in TVRDNJSKE)
        najava = any(norm(v) in n for v in NAJAVNI)
        if ima_tvrdnju and not najava and len(H.rijeci(r)) >= 8:
            veze = [v for v in TVRDNJSKE if f" {norm(v)} " in n]
            kandidati.append((r, veze))
    if not kandidati:
        oc.dodaj("teza", LOSE, "0 kandidata",
                 ["U uvodu nema nijedne rečenice koja iznosi tvrdnju s obrazloženjem.",
                  "Sve prepoznate rečenice ili najavljuju sadržaj rada "
                  "(„u ovom radu se analizira\") ili ne nose vezno sredstvo tvrdnje.",
                  "Provjeri ručno: teza mora biti rečenica s kojom se netko može ne složiti."],
                 "nema kandidata za tezu u uvodu")
    else:
        redci = ["KANDIDATI za tezu (nisu presuda — potvrdi ili odbaci ručno):"]
        for r, veze in kandidati[:3]:
            redci.append(f"  „{r[:150]}{'…' if len(r) > 150 else ''}\"")
            redci.append(f"     veza: {', '.join(veze[:3])}")
        if len(kandidati) > 3:
            redci.append(f"  … i još {len(kandidati) - 3} kandidata")
        oc.dodaj("teza", OK if len(kandidati) >= 1 else UPOZ,
                 f"{len(kandidati)} kandidata", redci)
    return [k[0] for k in kandidati]


def dim_zakljucak(oc, uvod, zakljucak):
    if uvod is None or zakljucak is None:
        koji = "uvod" if uvod is None else "zaključak"
        oc.dodaj("zaključak zatvara krug", LOSE, f"nema: {koji}",
                 [f"Poglavlje „{koji}\" nije prepoznato po naslovu (Heading 1).",
                  "Bez oba dijela preklapanje pojmova se ne može izračunati."],
                 f"{koji} nije pronađen")
        return {}
    rij = [w for w in re.findall(r"[^\W\d_]{4,}", uvod["tekst"], re.UNICODE)
           if norm(w) not in STOP_N and norm(w)]
    brojac = Counter()
    prikaz = {}
    for w in rij:
        k = korijen(w)
        brojac[k] += 1
        prikaz.setdefault(k, w.lower())
    top = brojac.most_common(10)
    if not top:
        oc.dodaj("zaključak zatvara krug", UPOZ, "uvod prekratak",
                 ["Uvod nema dovoljno sadržajnih riječi za usporedbu."])
        return {}
    z_korijeni = {korijen(w) for w in re.findall(r"[^\W\d_]{4,}", zakljucak["tekst"],
                                                 re.UNICODE)}
    nadjeni = [k for k, _ in top if k in z_korijeni]
    udio = len(nadjeni) / len(top) * 100
    fale = [prikaz[k] for k, _ in top if k not in z_korijeni]
    stanje = OK if udio >= 60 else (UPOZ if udio >= 30 else LOSE)
    redci = [f"top pojmovi uvoda: " + ", ".join(prikaz[k] for k, _ in top),
             f"vraćaju se u zaključak: {len(nadjeni)}/{len(top)} ({udio:.0f} %)"]
    if fale:
        redci.append("ne pojavljuju se u zaključku: " + ", ".join(fale))
        redci.append("Provjeri: pojam koji je uvod najavio, a zaključak ne spominje, "
                     "znak je da rad nije zatvorio krug — ili da uvod obećava previše.")
    oc.dodaj("zaključak zatvara krug", stanje, f"{udio:.0f} % preklapanja", redci,
             "" if stanje == OK else "zaključak se slabo vraća na pojmove iz uvoda")
    return {"udio": round(udio, 1), "pojmovi": [prikaz[k] for k, _ in top],
            "nedostaju": fale}


def dim_pitanje(oc, uvod):
    if uvod is None:
        return
    t = norm(uvod["tekst"])
    nadjeno = [f for f in PITANJE if norm(f) in t]
    upitnik = "?" in uvod["tekst"]
    redci = []
    if nadjeno:
        redci.append("formulacije u uvodu: " + ", ".join(sorted(set(nadjeno))))
    if upitnik:
        pit = [r for r in H.recenice(uvod["tekst"]) if r.rstrip().endswith("?")]
        for r in pit[:2]:
            redci.append(f"izravno pitanje: „{r[:140]}\"")
    if nadjeno or upitnik:
        oc.dodaj("istraživačko pitanje", OK,
                 "izrečeno" if nadjeno else "izravno pitanje", redci)
    else:
        oc.dodaj("istraživačko pitanje", UPOZ, "nije nađeno",
                 ["Uvod ne sadrži ni izravno pitanje ni formulaciju „istraživačko pitanje\", "
                  "„cilj rada\", „svrha rada\" ili „hipoteza\".",
                  "Provjeri: komisija to traži u prve dvije stranice."],
                 "istraživačko pitanje nije izrečeno izrijekom")


def dim_proporcije(oc, sadrzajna, policy):
    if len(sadrzajna) < 2:
        oc.dodaj("proporcije poglavlja", UPOZ, pogl(len(sadrzajna)),
                 ["Premalo poglavlja za usporedbu duljina."])
        return {}
    duljine = [p["rijeci"] for p in sadrzajna]
    ukupno = sum(duljine) or 1
    medijan = st.median(duljine)
    redci = [f"metodologija: {policy.label}",
             f"medijan poglavlja: {medijan:.0f} riječi · ukupno {ukupno} riječi"]
    odstupanja = []
    for p in sadrzajna:
        udio = p["rijeci"] / ukupno * 100
        oznaka = ""
        if medijan and p["rijeci"] > 2 * medijan:
            oznaka = "  ← više od 2× medijana"
            odstupanja.append(p["naslov"])
        elif medijan and p["rijeci"] < 0.4 * medijan:
            oznaka = "  ← manje od 0,4× medijana"
            odstupanja.append(p["naslov"])
        redci.append(f"  {p['naslov'][:46]:<46} {p['rijeci']:>6} rij. "
                     f"{udio:>5.1f} %{oznaka}")

    mjere = {"medijan": medijan, "odstupanja": odstupanja,
             "metodologija": policy.type}
    metodoloski_nalazi = []
    prvi_analiticki = None
    markers = tuple(policy.analytical_title_markers)
    if markers:
        for i, p in enumerate(sadrzajna):
            n = norm(p["naslov"])
            if any(norm(a) in n for a in markers):
                prvi_analiticki = i
                break

    if policy.requires_analytical_section and prvi_analiticki is None:
        redci.append(
            f"Za metodologiju „{policy.label}” nije prepoznato poglavlje čiji naslov "
            "upućuje na analizu/nalaze/rezultate. To je signal za ručnu provjeru, ne presuda."
        )
        metodoloski_nalazi.append("nije prepoznat metodološki očekivan analitički dio")

    # v1.1-fix (Q4b): indeks 0 znači da prije analitičkog poglavlja nema teksta,
    # pa udio teorije NIJE 0 % — nego se ne može izračunati. Prije popravka se
    # zapisivala neistinita brojka 0.0, a provjera praga tiho se preskakala.
    if policy.theory_share_max is not None and prvi_analiticki:
        teorija = sum(p["rijeci"] for p in sadrzajna[:prvi_analiticki])
        analiza = sum(p["rijeci"] for p in sadrzajna[prvi_analiticki:])
        uk = teorija + analiza or 1
        theory_share = teorija / uk
        mjere["teorija_udio"] = round(theory_share * 100, 1)
        redci.append(f"teorija (do „{sadrzajna[prvi_analiticki]['naslov'][:30]}”): "
                     f"{teorija} rij. ({theory_share*100:.0f} %) · "
                     f"analiza: {analiza} rij. ({analiza/uk*100:.0f} %)")
        if theory_share > policy.theory_share_max:
            limit = policy.theory_share_max * 100
            redci.append(
                f"Teorijski dio prelazi metodološki heuristički prag {limit:.0f} % — "
                "provjeri je li empirijska/analitička komponenta dovoljno razvijena."
            )
            metodoloski_nalazi.append(f"teorija >{limit:.0f} %")
    elif policy.theory_share_max is not None:
        redci.append(
            "Udio teorije nije izračunljiv: ispred prvog poglavlja koje se po naslovu "
            "čita kao analitičko nema teksta (analitički dio nije prepoznat ili je "
            f"prvo poglavlje). Prag od {policy.theory_share_max*100:.0f} % zato nije provjeren."
        )
        redci.append(
            "Što napraviti: prebroji ručno koliko rada otpada na teoriju, a koliko na "
            "analizu, ili preimenuj poglavlja tako da se teorijski dio razlikuje od analize."
        )
        metodoloski_nalazi.append("udio teorije nije izračunljiv")

    if odstupanja:
        redci.append("Provjeri poglavlja s oznakom ← : neujednačena duljina obično znači "
                     "da je jedno poglavlje progutalo dva, ili da je jedno ostalo skica.")

    # Metodološki signal sam po sebi je upozorenje; hard failure ostaje rezerviran za
    # kombinaciju ozbiljnog strukturnog odstupanja + relevantnog metodološkog signala.
    if odstupanja and metodoloski_nalazi:
        stanje = LOSE
    elif odstupanja or metodoloski_nalazi:
        stanje = UPOZ
    else:
        stanje = OK
    nalaz_dijelovi = []
    if odstupanja:
        nalaz_dijelovi.append("neujednačena poglavlja")
    nalaz_dijelovi.extend(metodoloski_nalazi)
    oc.dodaj("proporcije poglavlja", stanje,
             f"{pogl(len(sadrzajna))}, medijan {medijan:.0f} rij.", redci,
             " · ".join(nalaz_dijelovi))
    return mjere


def dim_doprinos(oc, izvori, policy):
    expectation = policy.display_contribution
    if not izvori:
        if expectation == "expected":
            oc.dodaj("vlastiti doprinos", UPOZ, "nema redaka „Izvor:”",
                     [f"Metodologija „{policy.label}” često ima podatkovne ili numeričke "
                      "prikaze, ali njihovo odsustvo nije dokaz da vlastiti doprinos ne postoji.",
                      "Provjeri ručno jesu li vlastita obrada, model, izračun ili rezultati "
                      "jasno vidljivi u tekstu ili drugim artefaktima."],
                     f"{policy.label}: vlastiti empirijski doprinos nije vidljiv kroz prikaze")
        else:
            oc.dodaj("vlastiti doprinos", OK, "prikazi nisu obvezan signal",
                     [f"Za metodologiju „{policy.label}” odsustvo tablice/grafikona nije nalaz.",
                      "Vlastiti doprinos može biti konceptualan, interpretativan, pravni, "
                      "povijesni ili tekstualno-analitički i mora se provjeriti u sadržaju."], "")
        return {"ukupno": 0, "vlastiti": 0, "adaptirani": 0,
                "display_expectation": expectation}

    vlastiti, adaptirani, preuzeti = [], [], []
    for red in izvori:
        n = norm(red)
        if je_vlastiti_prikaz(red, norm):
            vlastiti.append(red)
        elif any(norm(v) in n for v in ADAPTIRANI):
            adaptirani.append(red)
        else:
            preuzeti.append(red)
    svoji = len(vlastiti) + len(adaptirani)
    redci = [f"metodologija: {policy.label}",
             f"prikaza s izvorom: {len(izvori)} · vlastitih: {len(vlastiti)} · "
             f"prerađenih („prema…”): {len(adaptirani)} · preuzetih: {len(preuzeti)}"]
    for r in (vlastiti + adaptirani)[:4]:
        redci.append(f"  {r[:90]}")

    if expectation == "expected" and svoji == 0:
        stanje = UPOZ
        nalaz = f"{policy.label}: nema vlastitog/prerađenog prikaza kao empirijskog signala"
        redci.append("Nijedan prikaz nije označen kao autorov/prerađen. To nije dokaz "
                     "odsutnosti doprinosa; provjeri model, izračun, kod, rezultate i tekst.")
    elif expectation == "expected" and svoji == 1:
        stanje = UPOZ
        nalaz = f"{policy.label}: samo jedan vlastiti/prerađeni prikaz"
        redci.append("Samo jedan vlastiti/prerađeni prikaz — provjeri je li empirijski doprinos "
                     "dovoljno vidljiv i izvan samog prikaza.")
    else:
        stanje = OK
        nalaz = ""
        if expectation == "optional" and svoji == 0:
            redci.append("Svi prikazi su preuzeti, ali ova metodologija ne zahtijeva vlastiti "
                         "prikaz kao dokaz doprinosa; doprinos provjeri u argumentu/analizi.")
    oc.dodaj("vlastiti doprinos", stanje, f"{svoji}/{len(izvori)} vlastitih", redci, nalaz)
    return {"ukupno": len(izvori), "vlastiti": len(vlastiti),
            "adaptirani": len(adaptirani), "display_expectation": expectation}


def dim_deskriptivnost(oc, sadrzajna):
    odl = [o for p in sadrzajna for o in p["odlomci"]]
    if not odl:
        oc.dodaj("deskriptivnost", UPOZ, "nema proze", ["Nije pronađen prozni tekst."])
        return {"signal": "unknown"}
    bez = []
    for o in odl:
        n = " " + norm(o) + " "
        # v1.1-fix (Q4a): poklapanje po granici riječi, ne po prefiksu.
        if not any(u.search(n) for u in ANALITICKI_UZORCI):
            bez.append(o)
    udio = len(bez) / len(odl) * 100
    signal = "low" if udio < 35 else ("medium" if udio < 50 else "high")
    # AUD-018: jedan leksički signal nikada sam ne smije biti hard failure.
    stanje = OK if signal == "low" else UPOZ
    redci = [f"odlomaka bez prepoznatog analitičkog signala: "
             f"{len(bez)}/{len(odl)} ({udio:.0f} %)",
             "Ovo je leksički SIGNAL, ne dokaz deskriptivnosti: uzročnost, kontrast, "
             "konceptualno razlikovanje i inferencija mogu biti izraženi na više načina."]
    for o in bez[:3]:
        redci.append(f"  primjer za ručnu provjeru: {o[:110]}…")
    oc.dodaj("deskriptivnost", stanje, f"{udio:.0f} % bez analitičkog signala", redci,
             "" if stanje == OK else "visok udio odlomaka bez prepoznatog analitičkog signala")
    return {"udio_opisnih": round(udio, 1),  # legacy key za user_profile compatibility
            "udio_bez_analitickog_signala": round(udio, 1),
            "odlomaka": len(odl), "signal": signal}


def dim_citati(oc, tijelo, citatni_stil="autor-godina"):
    # v1.1-fix (Q4c): mjeri se cijelo tijelo rada (uvod uključen), a ne podskup
    # bez uvoda. Prije popravka su „citata ukupno" i gustoća na 1000 riječi bile
    # oznake na razini rada nad podskupom poglavlja, a uvod bez ijednog citata
    # nije se mogao ni prijaviti.
    if not tijelo:
        return {}
    prazna = [p for p in tijelo if p["citati"] == 0 and p["rijeci"] >= 60]
    ukupno = sum(p["citati"] for p in tijelo)
    tipovi = Counter()
    for p in tijelo:
        tipovi.update(p.get("tipovi_izvora") or {})
    rij = sum(p["rijeci"] for p in tijelo) or 1
    redci = [f"citata u tijelu rada (uvod uključen, bez popisa literature i priloga): "
             f"{ukupno} · gustoća {ukupno / rij * 1000:.1f} na 1000 riječi"]
    for p in tijelo:
        g = p["citati"] / p["rijeci"] * 1000 if p["rijeci"] else 0
        redci.append(f"  {p['naslov'][:46]:<46} {p['citati']:>3} cit. "
                     f"({g:>5.1f}/1000){'  ← nijedan citat' if p['citati'] == 0 and p['rijeci'] >= 60 else ''}")
    # zaključak bez citata nije nalaz — ondje se ne uvode novi izvori
    tvrdi = [p for p in prazna if "zakljuc" not in norm(p["naslov"])]
    stanje = OK if not tvrdi else LOSE
    if ukupno == 0:
        stanje = LOSE
        redci.append(f"U cijelom radu nije prepoznat nijedan citat za deklarirani stil „{citatni_stil}”. "
                     "Provjeri je li stil točno deklariran i postoje li reference u dokumentu.")
    kratka = [p for p in tijelo if p["citati"] == 0 and 0 < p["rijeci"] < 60]
    if kratka:
        redci.append("Poglavlja kraća od 60 riječi ne ulaze u nalaz o citatima: "
                     + "; ".join(p["naslov"][:40] for p in kratka))
    if tvrdi:
        redci.append("Poglavlje bez ijednog citata: svaka tvrdnja u njemu stoji na "
                     "autorovu autoritetu. Provjeri je li to namjerno (vlastita analiza) "
                     "ili je izostala potpora.")
    oc.dodaj("citatna gustoća", stanje, f"{ukupno} citata, "
             f"{pogl(len(prazna))} bez citata", redci,
             ("bez citata: " + "; ".join(p["naslov"][:40] for p in tvrdi)) if tvrdi else "")
    return {"ukupno": ukupno, "bez_citata": [p["naslov"] for p in tvrdi], "tipovi_izvora": dict(tipovi)}


def dim_pokrivenost_numericka(oc, put, citatni_stil):
    """v1.9 (nalaz 6): u numeričkom dijalektu (vancouver/ieee) popis literature je
    numeriran pa se pokrivenost može izmjeriti bez claim ledgera: svaki citat
    ima stavku, svaka stavka je citirana, prvo pojavljivanje raste. Sva pravila
    žive u citation_dialects.numeric_report (B10). Autor-godina i dalje ide kroz
    verify_sources.py --pokrivenost."""
    try:
        if C.resolve_dialect(citatni_stil) not in C.NUMERIC_DIALECTS:
            return None
        r = C.numeric_report_file(put, citatni_stil)
    except Exception as e:  # pragma: no cover - dijagnostika, ne presuda
        oc.dodaj("pokrivenost popisa", UPOZ, "nije izmjereno", [str(e)], "")
        return None
    redci = [f"popis: {r['popis_stavki']} stavki (1..{r['N_max']}); citata u tekstu: "
             f"{r['citata_u_tekstu']}, različitih brojeva: {r['razlicitih_citiranih']}"]
    losi = []
    if not r["popis_stavki"]:
        losi.append("numerirani popis literature nije prepoznat (očekuje se „1. Autor…\" pod naslovom popisa)")
    if r["sirocad"]:
        losi.append(f"siročad (u popisu, necitirano): {r['sirocad'][:15]}")
    if r["citat_bez_reference"]:
        losi.append(f"citat bez reference: {r['citat_bez_reference'][:15]}")
    if r["popis_prazni_brojevi"] or r["popis_dupli"]:
        losi.append(f"popis: prazni brojevi {r['popis_prazni_brojevi'][:10]}, dupli {r['popis_dupli'][:10]}")
    savjeti = []
    if r["skokovi_redoslijeda"]:
        s = r["skokovi_redoslijeda"][0]
        savjeti.append(f"redoslijed prvog pojavljivanja nije uzlazan: prvi skok na {s[0]} "
                       f"(očekivano {s[1]}), ukupno {len(r['skokovi_redoslijeda'])} skokova")
    if r["raspon_sa_spojnicom"]:
        savjeti.append(f"raspon sa spojnicom umjesto en-crtice: {r['raspon_sa_spojnicom'][:5]}")
    if r["citat_bez_razmaka"]:
        savjeti.append(f"bez razmaka iza zareza (Vancouver traži „(67, 68)\"): {r['citat_bez_razmaka'][:5]}")
    if r["nabrajanje_umjesto_raspona"]:
        savjeti.append(f"nabrajanje umjesto raspona: {r['nabrajanje_umjesto_raspona'][:5]}")
    stanje = LOSE if losi else (UPOZ if savjeti else OK)
    oc.dodaj("pokrivenost popisa", stanje,
             f"{len(r['sirocad'])} siročadi, {len(r['citat_bez_reference'])} bez reference",
             redci + losi + savjeti, "; ".join(losi))
    return {"stil": citatni_stil, "popis_stavki": r["popis_stavki"],
            "sirocad": r["sirocad"], "citat_bez_reference": r["citat_bez_reference"],
            "skokovi_redoslijeda": len(r["skokovi_redoslijeda"]),
            "stilski": {"raspon_sa_spojnicom": len(r["raspon_sa_spojnicom"]),
                        "citat_bez_razmaka": len(r["citat_bez_razmaka"]),
                        "nabrajanje_umjesto_raspona": len(r["nabrajanje_umjesto_raspona"])}}


# --------------------------------------------------------------------- ispis

def ispis(oc, naziv, pog, sadrzajna):
    print("=" * 78)
    print(f"ARGUMENT — {naziv}")
    print("=" * 78)
    print(f"poglavlja (Heading 1): {len(pog)} · sadržajnih (bez uvoda i popisa): "
          f"{len(sadrzajna)} · riječi: {sum(p['rijeci'] for p in pog)}")
    print("Sve niže je HEURISTIKA: brojke i kandidati, ne presuda. Potvrdi ručno.")
    print()
    for d in oc.dim:
        print(f"{d['stanje']}  {d['dimenzija'].upper()}  —  {d['brojka']}")
        for r in d["redci"]:
            print(f"      {r}")
        print()
    nalazi = oc.nalazi
    print("-" * 78)
    print("SAŽETAK")
    if not nalazi:
        print("  Nema nalaza po argumentacijskim dimenzijama. Forma i argument su druga "
              "stvar — ovo ne govori ništa o pravilima fakulteta.")
    for d in nalazi:
        print(f"  {d['stanje']} {d['dimenzija']}: {d['nalaz'] or d['brojka']}")
    print("  Ništa od ovoga nije dokaz. Alat čita obrasce u tekstu, ne smisao.")
    print()


def main():
    ap = argparse.ArgumentParser(
        description="Teza, zaključak i metodološki svjesne heuristike argumenta.")
    ap.add_argument("rad", help="putanja do .docx-a")
    ap.add_argument("--metodologija", choices=M.METHODOLOGY_TYPES,
                    help="metodološki policy; nadjačava profil.metodologija.type")
    ap.add_argument("--profil", help="resolved profile JSON s opcionalnim metodologija.type")
    ap.add_argument("--citatni-stil", choices=C.CITATION_STYLES,
                    help="citatni stil; nadjačava profil.citiranje.stil")
    ap.add_argument("--json", dest="json_out", metavar="PUT",
                    help="zapiši mjere za user_profile.py learn --argument")
    a = ap.parse_args()

    if not os.path.isfile(a.rad):
        print(f"❌ rad ne postoji: {a.rad}", file=sys.stderr)
        return 2
    if not a.rad.lower().endswith(".docx"):
        print("❌ očekuje se .docx", file=sys.stderr)
        return 2

    try:
        citation_style, citation_dialect, citation_source = C.resolve_style(a.citatni_stil, a.profil)
    except C.CitationDialectError as exc:
        print(f"❌ {exc}", file=sys.stderr)
        return 2

    try:
        method_policy, method_source = M.resolve_methodology(a.metodologija, a.profil)
    except M.MethodologyError as exc:
        print(f"❌ {exc}", file=sys.stderr)
        return 2

    pog, izvori = poglavlja(a.rad, citation_style)
    if not pog or all(not p["rijeci"] for p in pog):
        print(f"❌ {os.path.basename(a.rad)}: nije pronađen prozni tekst.", file=sys.stderr)
        print("   Provjeri koristi li dokument stilove Heading za naslove i ima li tijela.",
              file=sys.stderr)
        return 2

    uvod = next((p for p in pog if re.search(r"\buvod", norm(p["naslov"]))), None)
    zakljucak = next((p for p in pog if "zakljuc" in norm(p["naslov"])), None)
    sadrzajna = [p for p in pog
                 if not je_literatura(p["naslov"])
                 and p is not uvod and p["rijeci"] >= 1
                 and p["naslov"] != "(prije prvog poglavlja)"]
    if not sadrzajna:
        sadrzajna = [p for p in pog if p["rijeci"] >= 1]
    # v1.1-fix (Q4c): citatna gustoća mjeri se nad tijelom rada, dakle i nad uvodom.
    tijelo = [p for p in pog
              if not je_literatura(p["naslov"]) and p["rijeci"] >= 1
              and p["naslov"] != "(prije prvog poglavlja)"]
    if not tijelo:
        tijelo = sadrzajna

    oc = Ocjena()
    if len(pog) == 1 and pog[0]["naslov"] == "(prije prvog poglavlja)":
        oc.dodaj("struktura", UPOZ, "nema Heading 1",
                 ["Dokument nema nijedan naslov u stilu Heading 1.",
                  "Poglavlja se ne mogu razdvojiti pa proporcije i citatna gustoća "
                  "po poglavlju nisu mjerodavne."],
                 "nema strukture naslova")

    teze = dim_teza(oc, uvod)
    krug = dim_zakljucak(oc, uvod, zakljucak)
    dim_pitanje(oc, uvod)
    prop = dim_proporcije(oc, [p for p in sadrzajna if p is not zakljucak] or sadrzajna, method_policy)
    dopr = dim_doprinos(oc, izvori, method_policy)
    desk = dim_deskriptivnost(oc, sadrzajna)
    cit = dim_citati(oc, tijelo, citation_style)
    pokr = dim_pokrivenost_numericka(oc, a.rad, citation_style)
    if pokr is not None:
        cit["pokrivenost"] = pokr

    ispis(oc, os.path.basename(a.rad), pog, sadrzajna)

    if a.json_out:
        os.makedirs(os.path.dirname(os.path.abspath(a.json_out)) or ".", exist_ok=True)
        with open(a.json_out, "w", encoding="utf-8") as f:
            json.dump({
                "alat": "check_argument",
                "napomena": "sve su vrijednosti heuristika, ne presuda",
                "rad": os.path.abspath(a.rad),
                "metodologija": {"type": method_policy.type, "source": method_source},
                "citiranje": {"stil": citation_style, "dialect": citation_dialect, "source": citation_source},
                "rijeci": sum(p["rijeci"] for p in pog),
                "poglavlja": [{"naslov": p["naslov"], "rijeci": p["rijeci"],
                               "citati": p["citati"]} for p in pog],
                "dimenzije": [{"dimenzija": d["dimenzija"], "stanje": d["stanje"],
                               "brojka": d["brojka"], "nalaz": d["nalaz"]}
                              for d in oc.dim],
                "kandidati_za_tezu": teze[:5],
                "zakljucak": krug,
                "proporcije": prop,
                "doprinos": dopr,
                "deskriptivnost": desk,
                "citati": cit,
                "broj_nalaza": len(oc.nalazi),
            }, f, ensure_ascii=False, indent=1)
        print(f"[mjere → {a.json_out}]")

    return 1 if oc.nalazi else 0


if __name__ == "__main__":
    sys.exit(main())
