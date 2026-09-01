#!/usr/bin/env python3
"""Snapshoti i usporedba verzija rada — što se tiho izgubilo između dvije verzije.

`.docx` je u gitu binarni blob bez upotrebljivog diffa: bez snapshota rollback ne
postoji. Zato se prije BILO KAKVE izmjene dokumenta radi snapshot, a poslije se
verzije uspoređuju po odlomcima.

Najvažniji nalaz nije broj izmijenjenih odlomaka nego IZGUBLJEN CITAT: tekst
poslije prepisivanja izgleda bolje, a tvrdnja je ostala bez potpore. Odmah iza
njega idu izgubljene brojke (postotci, iznosi, godine iz podataka).

Uporaba:
  python3 <KATEDRA_SKILL>/scripts/diff_versions.py --snapshot ./rad.docx --biljeska "prije faze G"
  python3 <KATEDRA_SKILL>/scripts/diff_versions.py ./rad_v2.docx ./rad_v3.docx
  python3 <KATEDRA_SKILL>/scripts/diff_versions.py ./rad_v2.docx ./rad_v3.docx --za-mentora > ./izmjene.md
  python3 <KATEDRA_SKILL>/scripts/diff_versions.py --popis
  python3 <KATEDRA_SKILL>/scripts/diff_versions.py --vrati v3 --u ./rad_vraceno.docx

Izlazni kodovi:
  0  gotovo; usporedba nije našla izgubljene citate
  1  usporedba je našla izgubljene citate ili izgubljene brojke
  2  datoteka se ne može pročitati, ili bi se nešto prepisalo bez --force
"""
import argparse
import datetime
import difflib
import hashlib
import json
import os
import re
import shutil
import sys
import unicodedata
import zipfile
from collections import Counter

HERE = os.path.dirname(os.path.abspath(__file__))
if HERE not in sys.path:
    sys.path.insert(0, HERE)
import hr_text as H  # noqa: E402
from context import NesigurnaPutanja, atomic_write_json, atomic_write_text, resolve_state_dir  # noqa: E402
from artifact_state import record_artifact  # noqa: E402

VERZIJA = 1
# v1.1-fix Q19: separator unutar broja mora biti ODMAH ispred znamenke —
# „15, 20 i 25" su tri broja, a „19,6" je jedan (hrvatski decimalni zarez).
BROJKA_RE = re.compile(r"\d+(?:[.,  ]\d+)*")
# Praćenje nabrajanja („45, 67 i 89" kao cjelina, bez obzira na prag) POSTOJALO je
# ovdje i povučeno je: pravilo je ovisilo o interpunkciji OKO broja, pa je bilo
# asimetrično između dviju uspoređivanih verzija. Obrazloženje i podjela posla s
# `verify_rewrite` stoje u docstringu `brojke()` niže. Komentar i implementacija
# preživjeli su povlačenje jedan krug predugo i tvrdili suprotno od koda — to je
# točno onaj kvar zbog kojeg sljedeći čitatelj vjeruje komentaru, pa su obrisani.


# ------------------------------------------------------------ čitanje teksta

def _nfc(s):
    """v1.1-fix Q19: „Čavlek" u NFC i „Čavlek" u NFD su ISTI tekst za čovjeka,
    a različite niske za Python.

    Usporedba se dosad radila nad sirovim niskama, pa je verzija spremljena na
    macOS-u (NFD) protiv verzije iz Worda (NFC) davala „izmijenjeno" za svaki
    odlomak, „IZGUBLJENI CITATI" za citat koji doslovno stoji u novom tekstu i
    izlaz 1 — uz ispis dviju vizualno identičnih niski, pa se nalaz nije dao ni
    razriješiti gledanjem. Normalizacija je jedini način da usporedba bude o
    sadržaju, a ne o načinu zapisa dijakritika.
    """
    return unicodedata.normalize("NFC", s or "")


def _je_wordov_paket(put):
    """Je li datoteka OOXML paket s word/document.xml, bez obzira na nastavak?

    2. krug: format se određivao ISKLJUČIVO po nastavku imena. Nakon što je
    --snapshot počeo čuvati izvorni nastavak, isti Wordov dokument spremljen pod
    imenom „rad.doc" (Word to ime nudi kad se rad radi iz starijeg predloška)
    prestao je ići u .docx granu, iako se prije, pod imenom „rad.docx",
    uspoređivao ispravno. Provjerava se zato SADRŽAJ — zip s word/document.xml je
    Wordov dokument ma kako se zvao. Pravilo je strukturno, ne popis nastavaka.
    """
    try:
        with zipfile.ZipFile(put) as z:
            return "word/document.xml" in z.namelist()
    except (zipfile.BadZipFile, OSError, EOFError):
        return False


def odlomci_svi(put):
    """Vrati (svi, tijelo): svi odlomci redom, i oni bez popisa literature.

    Diff ide po SVIM odlomcima (izmjena u popisu izvora je isto izmjena), a citati
    i brojke se broje samo u tijelu — inače bi zapis „Hall, C. M. … (2018.)" iz
    popisa maskirao citat koji je iz teksta nestao.
    """
    if not os.path.isfile(put):
        print(f"❌ nema datoteke: {put}\n"
              f"   Što napraviti: provjeri putanju; skripta se pokreće iz scripts/.",
              file=sys.stderr)
        return None
    if _je_wordov_paket(put) or put.lower().endswith(".docx"):
        try:
            from docx import Document
        except ImportError:
            print("❌ treba python-docx:  pip install python-docx --break-system-packages",
                  file=sys.stderr)
            return None
        try:
            d = Document(put)
        except Exception as e:
            print(f"❌ {os.path.basename(put)} se ne može otvoriti kao .docx: {e}\n"
                  f"   Što napraviti: otvori u Wordu i spremi ponovno; ako je datoteka "
                  f"prazna ili prekinuta u prijenosu, uzmi je iz snapshota "
                  f"(--popis pa --vrati).", file=sys.stderr)
            return None
        redci = [_nfc(p.text).strip() for p in d.paragraphs]
        celije = []
        try:
            for t in d.tables:
                for red in t.rows:
                    for c in red.cells:
                        celije.append(_nfc(c.text).strip())
        except Exception:
            pass
    else:
        try:
            redci = [_nfc(r).strip() for r in open(put, encoding="utf-8").read().split("\n")]
        except OSError as e:
            print(f"❌ {put} se ne može pročitati: {e}", file=sys.stderr)
            return None
        except UnicodeDecodeError:
            # 2. krug (nedovršen popravak): tekstualna grana je hvatala samo
            # OSError, pa je BINARAN rukopis (.doc, .odt, .rtf — na hrvatskim
            # fakultetima svakodnevni) izlazio golim UnicodeDecodeError
            # traceback-om i izlazom 1. Alat taj lanac SAM predloži: --snapshot
            # rad.doc ispiše „--vrati v1 --u ./rad_vraceno.doc", pa naredba koju
            # je alat napisao završi rušenjem. Uredna poruka i izlaz 2 su isti
            # ugovor koji vrijedi za neispravan .docx.
            print(f"❌ {os.path.basename(put)} nije tekstualna datoteka ni Word paket "
                  f"(.docx) — usporedba po odlomcima ne može je pročitati.\n"
                  f"   Što napraviti: otvori je u Wordu (ili LibreOfficeu) i spremi "
                  f"kao .docx, pa ponovi usporedbu s tom datotekom.", file=sys.stderr)
            return None
        celije = []

    tijelo, u_popisu = [], False
    for t in redci:
        if not t:
            continue
        if H.NASLOV_LIT.match(t.lstrip("# ").strip()):
            u_popisu = True
            continue
        if not u_popisu:
            tijelo.append(t)
    svi = [t for t in redci + celije if t]
    tijelo += [t for t in celije if t]      # tablice su uvijek tijelo rada
    return svi, tijelo


def bez_citata(tekst):
    """Ukloni citate iz teksta — godina iz citata nije podatak nego referenca."""
    t = H.CITAT_ZAGRADNI.sub(" ", tekst)
    return H.CITAT_NARATIVNI.sub(r"\1", t)


def brojke(tekst):
    """Multiskup brojki s više od 2 znamenke (postotci, iznosi, godine iz podataka).

    v1.1-fix Q19-brojke, drugi krug: praćenje nabrajanja je POVUČENO. Prvi popravak
    je — uz ispravan tokenizator, koji više ne lijepi „15, 20” u jedan token „15,20” —
    dodao i pravilo „broj u nabrajanju prati se bez obzira na prag". To je pravilo
    ovisilo o interpunkciji OKO broja, pa je bilo asimetrično između dviju verzija
    koje se uspoređuju: obično prestrukturiranje rečenice („45, 67 i 89" → „45, 67 te
    čak 89", ili razdvajanje u tri rečenice) mijenja članstvo samo na jednoj strani,
    pa je gate prijavljivao „IZGUBLJENE BROJKE" za brojke koje doslovno stoje u novom
    tekstu — i to u `stil` zahvatu, čija je cijela svrha prestrukturiranje.

    Prag >2 znamenke je token-lokalan i zato simetričan: preživljava svako
    preslagivanje. Ostaje dokumentirano ograničenje (jedno- i dvoznamenkasti brojevi
    se ne prate), ali je ono pošteno i predvidljivo, dok je lažna blokada skuplja od
    propuštenog malog broja. Popravak tokenizatora se ZADRŽAVA — fantomski token
    „15,20” je bio stvarna greška.
    """
    c = Counter()
    for m in BROJKA_RE.finditer(tekst):
        s = m.group().strip()
        if len(re.sub(r"\D", "", s)) > 2:
            c[re.sub(r"[   ]", "", s)] += 1
    return c

def citljiva_putanja(put):
    """Relativna putanja ako je kratka, inače apsolutna — bez „../../../.." lanaca."""
    try:
        rel = os.path.relpath(put, os.getcwd())
    except ValueError:
        return os.path.abspath(put)
    return rel if not rel.startswith(os.path.join("..", "..")) else os.path.abspath(put)


def _putanja_za_izvjestaj(put, korijen):
    """Putanja kakva SMIJE stajati u JSON izvještaju koji putuje s radom.

    2. krug (Q13): polja „stari" i „novi" pisala su se `os.path.abspath`, pa je
    izvještaj koji student pošalje mentoru ili commita u repozitorij nosio punu
    putanju s korisničkim imenom (/home/ime/…, /Users/ime/…). Prijašnja odluka
    bila je NE DIRATI, uz obrazloženje da bi popravak razbio zajednički oblik
    koji dijele svi JSON izvještaji; to obrazloženje više ne stoji jer je isti
    prijelaz u međuvremenu napravljen u `artifact_state._norm_path` /
    `_prikaz_putanje` i pokazao se ispravnim.

    Pravilo je strukturno (leži li datoteka unutar korijena projekta), ne popis
    sumnjivih imena mapa: rad unutar projekta zapisuje se cijelom relativnom
    putanjom, a rad izvan njega svede se na zadnje dvije komponente uz oznaku
    „…/", pa se datoteka i dalje prepoznaje, a struktura tuđeg diska ne izlazi.
    """
    try:
        p = os.path.realpath(put)
        r = os.path.realpath(korijen)
        rel = os.path.relpath(p, r)
    except (OSError, ValueError):
        return os.path.basename(put)
    dijelovi = rel.split(os.sep)
    if ".." not in dijelovi:
        return "/".join(dijelovi)
    stvarne = [d for d in dijelovi if d not in ("", "..", ".")]
    return "…/" + "/".join(stvarne[-2:])


def sha256(put):
    h = hashlib.sha256()
    with open(put, "rb") as f:
        for blok in iter(lambda: f.read(1 << 16), b""):
            h.update(blok)
    return h.hexdigest()


# --------------------------------------------------------------- verzije.json

def put_verzija(kat):
    return os.path.join(kat, "verzije.json")


def ucitaj_verzije(kat):
    put = put_verzija(kat)
    if not os.path.isfile(put):
        return {"verzija": VERZIJA, "snapshoti": []}
    try:
        with open(put, encoding="utf-8") as f:
            d = json.load(f)
    except (OSError, json.JSONDecodeError) as e:
        print(f"❌ {put} nije čitljiv JSON: {e}\n"
              f"   Što napraviti: ne krpaj ručno — preimenuj ga; snapshoti u "
              f"{os.path.join(kat, 'verzije')} ostaju netaknuti.", file=sys.stderr)
        return None
    if not isinstance(d.get("snapshoti"), list):
        print(f"❌ {put} nema ključ „snapshoti\".", file=sys.stderr)
        return None
    return d


def spremi_verzije(kat, d):
    """v1.1-fix Q14: zapis kroz zajednički atomarni helper.

    Prije se pisalo u fiksni „verzije.json.tmp" pa os.replace. Unaprijed
    podmetnuta simbolička poveznica na toj putanji se slijedila, pa je popis
    snapshota — jedini zapis o tome gdje je rollback — završio izvan projekta, a
    sam verzije.json ostao poveznica do kraja rada. Druga rupa istog obrasca:
    dvije istovremene naredbe dijele ime tmp datoteke i gubitnik pukne.
    """
    os.makedirs(kat, exist_ok=True)
    return atomic_write_json(put_verzija(kat), d)


def _num(s):
    m = re.search(r"\d+", str(s or ""))
    return int(m.group()) if m else 0


def snapshot(put, biljeska, kat):
    if not os.path.isfile(put):
        print(f"❌ nema datoteke: {put}", file=sys.stderr)
        return 2
    d = ucitaj_verzije(kat)
    if d is None:
        return 2
    sha = sha256(put)
    for s in d["snapshoti"]:
        if s.get("sha256") == sha:
            root = os.path.dirname(os.path.abspath(kat))
            try:
                record_artifact(root, put, kind="document", version_id=s.get("id"),
                                snapshot_id=s.get("id"), snapshot_path=s.get("datoteka"),
                                note=s.get("biljeska") or "")
            except (OSError, ValueError) as e:
                print(f"❌ artifact manifest nije ažuriran: {e}", file=sys.stderr)
                return 2
            print(f"= sadržaj je identičan snapshotu {s.get('id')} "
                  f"({s.get('datum', '')[:16]}, „{s.get('biljeska', '')}“).")
            print("  Novi snapshot nije stvoren — nema što spremiti.")
            return 0
    sada = datetime.datetime.now()
    # v1.1-fix Q19: zadrži izvorni nastavak — inače --vrati vrati .md/.txt pod
    # imenom .docx, a sljedeći alat takvu datoteku odbije
    osnova, nastavak = os.path.splitext(os.path.basename(put))
    nastavak = nastavak or ".docx"
    vrijeme = sada.strftime("%Y%m%d_%H%M")
    mapa = os.path.join(kat, "verzije")
    os.makedirs(mapa, exist_ok=True)
    cilj = os.path.join(mapa, f"{osnova}_{vrijeme}{nastavak}")
    n = 1
    while os.path.exists(cilj):
        n += 1
        cilj = os.path.join(mapa, f"{osnova}_{vrijeme}_{n}{nastavak}")
    shutil.copy2(put, cilj)
    vid = f"v{max([_num(s.get('id')) for s in d['snapshoti']] or [0]) + 1}"
    d["verzija"] = VERZIJA
    d["snapshoti"].append({
        "id": vid,
        "datoteka": os.path.join("verzije", os.path.basename(cilj)),
        "sha256": sha,
        "biljeska": biljeska or "",
        "datum": sada.replace(microsecond=0).isoformat(),
    })
    try:
        spremi_verzije(kat, d)
    except NesigurnaPutanja as e:
        print(f"❌ popis snapshota nije zapisan: {e}", file=sys.stderr)
        print(f"   Kopija dokumenta ipak postoji: {citljiva_putanja(cilj)}", file=sys.stderr)
        return 2
    except OSError as e:
        print(f"❌ popis snapshota nije zapisan ({e}).", file=sys.stderr)
        print("   Što napraviti: provjeri prava pisanja nad .katedra/ pa ponovi naredbu.",
              file=sys.stderr)
        return 2
    root = os.path.dirname(os.path.abspath(kat))
    try:
        record_artifact(root, put, kind="document", version_id=vid, snapshot_id=vid,
                        snapshot_path=os.path.join("verzije", os.path.basename(cilj)),
                        note=biljeska or "")
    except (OSError, ValueError) as e:
        print(f"❌ snapshot je spremljen, ali artifact manifest nije ažuriran: {e}", file=sys.stderr)
        return 2
    print(f"✅ snapshot {vid} → {citljiva_putanja(cilj)}")
    print(f"   sha256 {sha[:16]}…  ·  bilješka: {biljeska or '(bez bilješke)'}")
    # v1.1-fix Q19: predložena naredba je hardkodirala „.docx". Za snapshot .md/.txt
    # rukopisa student je tako dobio ASCII datoteku pod imenom .docx, koju sljedeći
    # alat u lancu (pa i sam diff_versions) odbija kao neispravan Word dokument.
    print(f"   povratak: python3 <KATEDRA_SKILL>/scripts/diff_versions.py --vrati {vid} "
          f"--u ./rad_vraceno{nastavak}")
    return 0


def popis(kat):
    d = ucitaj_verzije(kat)
    if d is None:
        return 2
    if not d["snapshoti"]:
        print(f"Nema nijednog snapshota u {put_verzija(kat)}.")
        print("Prije prve izmjene dokumenta:")
        print("  python3 <KATEDRA_SKILL>/scripts/diff_versions.py --snapshot ./rad.docx --biljeska \"prije faze G\"")
        return 0
    print("=" * 78)
    print(f"SNAPSHOTI — {os.path.abspath(put_verzija(kat))}")
    print("=" * 78)
    for s in d["snapshoti"]:
        put = os.path.join(kat, s.get("datoteka", ""))
        ima = "" if os.path.isfile(put) else "  ⚠️ datoteka nedostaje"
        print(f"{s.get('id', '?'):<5} {s.get('datum', '')[:16]:<17} "
              f"{s.get('sha256', '')[:12]}  {s.get('biljeska', '')}{ima}")
        print(f"      {s.get('datoteka', '')}")
    return 0


def vrati(vid, u, kat, force):
    d = ucitaj_verzije(kat)
    if d is None:
        return 2
    nadeni = next((s for s in d["snapshoti"] if str(s.get("id")) == vid), None)
    if not nadeni:
        print(f"❌ nema snapshota „{vid}\".\n"
              f"   Postojeći: {', '.join(str(s.get('id')) for s in d['snapshoti']) or '—'}\n"
              f"   Popis s bilješkama: python3 <KATEDRA_SKILL>/scripts/diff_versions.py --popis", file=sys.stderr)
        return 2
    izvor = os.path.join(kat, nadeni["datoteka"])
    if not os.path.isfile(izvor):
        print(f"❌ snapshot je zapisan, ali datoteka nedostaje: {izvor}", file=sys.stderr)
        return 2
    # v1.1-fix Q19: zadano ime nasljeđuje nastavak snimljene datoteke.
    # 2. krug: isti nastavak mora stajati i u SAVJETIMA ispod — hardkodirani
    # „.docx" u njima vodio je studenta natrag na krivo ime datoteke, tj. točno
    # na kvar koji je popravak u snapshot() već zatvorio.
    nastavak = os.path.splitext(nadeni.get("datoteka", ""))[1] or ".docx"
    if not u:
        u = os.path.join(os.path.dirname(os.path.abspath(kat)),
                         f"rad_vraceno_{vid}{nastavak}")
        print(f"[--u nije zadan → {u}]")
    if os.path.exists(u) and not force:
        print(f"❌ {u} već postoji. Snapshot se NE vraća preko postojeće datoteke.\n"
              f"   Što napraviti: zadaj drugo ime (--u ./rad_vraceno{nastavak}) ili, ako "
              f"stvarno želiš prepisati, dodaj --force.", file=sys.stderr)
        return 2
    os.makedirs(os.path.dirname(os.path.abspath(u)) or ".", exist_ok=True)
    shutil.copy2(izvor, u)
    print(f"✅ {vid} ({nadeni.get('biljeska', '')}) → {u}")
    print("   Original nije diran. Usporedi prije nego što ga uzmeš u rad:")
    print(f"   python3 <KATEDRA_SKILL>/scripts/diff_versions.py {u} ./rad{nastavak}")
    return 0


# ------------------------------------------------------------------ usporedba

def usporedi(stari, novi):
    sm = difflib.SequenceMatcher(None, stari, novi, autojunk=False)
    dodani, uklonjeni, izmijenjeni = [], [], []
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            continue
        if tag == "insert":
            dodani.extend(novi[j1:j2])
        elif tag == "delete":
            uklonjeni.extend(stari[i1:i2])
        elif tag == "replace":
            n = min(i2 - i1, j2 - j1)
            for k in range(n):
                izmijenjeni.append((stari[i1 + k], novi[j1 + k]))
            if (i2 - i1) > n:
                uklonjeni.extend(stari[i1 + n:i2])
            if (j2 - j1) > n:
                dodani.extend(novi[j1 + n:j2])
    return dodani, uklonjeni, izmijenjeni


def analiza(put_a, put_b):
    ra = odlomci_svi(put_a)
    rb = odlomci_svi(put_b)
    if ra is None or rb is None:
        return None
    a, tijelo_a = ra
    b, tijelo_b = rb
    if not a and not b:
        print("⚠️  obje verzije su prazne (nema nijednog odlomka) — nema što usporediti.",
              file=sys.stderr)
    ta = H.bez_lokatora("\n".join(tijelo_a))
    tb = H.bez_lokatora("\n".join(tijelo_b))
    dodani, uklonjeni, izmijenjeni = usporedi(a, b)

    ca, cb = H.kljucevi_citata(ta), H.kljucevi_citata(tb)
    izgubljeni_citati = sorted([k for k in ca if cb.get(k, 0) == 0])
    rjedi_citati = sorted([(k, ca[k], cb[k]) for k in ca if 0 < cb.get(k, 0) < ca[k]])
    novi_citati = sorted([k for k in cb if ca.get(k, 0) == 0])

    ba, bb = brojke(bez_citata(ta)), brojke(bez_citata(tb))
    izgubljene_brojke = sorted([n for n in ba if bb.get(n, 0) == 0])
    nove_brojke = sorted([n for n in bb if ba.get(n, 0) == 0])

    rij_a = sum(len(H.rijeci(p)) for p in tijelo_a)
    rij_b = sum(len(H.rijeci(p)) for p in tijelo_b)
    return {
        "a": put_a, "b": put_b,
        "odlomaka_a": len(a), "odlomaka_b": len(b),
        "rijeci_a": rij_a, "rijeci_b": rij_b,
        "dodani": dodani, "uklonjeni": uklonjeni, "izmijenjeni": izmijenjeni,
        "izgubljeni_citati": izgubljeni_citati,
        "rjedi_citati": rjedi_citati,
        "novi_citati": novi_citati,
        "izgubljene_brojke": izgubljene_brojke,
        "nove_brojke": nove_brojke,
    }


def mn(n, jednina, dvojina, mnozina):
    """Hrvatska sročnost uz broj: 1 odlomak · 2 odlomka · 5 odlomaka."""
    z, d = n % 10, n % 100
    if 11 <= d <= 14:
        return mnozina
    return jednina if z == 1 else (dvojina if z in (2, 3, 4) else mnozina)


def skrati(s, n=100):
    s = re.sub(r"\s+", " ", s or "").strip()
    return s if len(s) <= n else s[:n - 1] + "…"


def ispis(r):
    print("=" * 78)
    print(f"USPOREDBA  {os.path.basename(r['a'])}  →  {os.path.basename(r['b'])}")
    print("=" * 78)
    d = r["rijeci_b"] - r["rijeci_a"]
    print(f"odlomci  {r['odlomaka_a']} → {r['odlomaka_b']}    "
          f"riječi  {r['rijeci_a']} → {r['rijeci_b']} ({d:+d})   "
          f"[riječi i citati: tijelo rada, bez popisa izvora]")
    print(f"dodano {len(r['dodani'])} · uklonjeno {len(r['uklonjeni'])} "
          f"· izmijenjeno {len(r['izmijenjeni'])} "
          f"{mn(len(r['izmijenjeni']), 'odlomak', 'odlomka', 'odlomaka')}")
    print()

    if r["izgubljeni_citati"]:
        print("❌ IZGUBLJENI CITATI — najvažniji nalaz")
        print("   Tvrdnja je vjerojatno ostala u tekstu, a potpora joj je nestala.")
        for prez, god in r["izgubljeni_citati"]:
            print(f"     ({prez.capitalize()}, {god}.)")
        print("   Provjeri svaki: je li i tvrdnja maknuta, ili je citat samo ispao?")
        print()
    else:
        print("✅ nijedan citat nije nestao između verzija")

    if r["rjedi_citati"]:
        print(f"⚠️  citati koji se pojavljuju rjeđe nego prije ({len(r['rjedi_citati'])}):")
        for (prez, god), sa, sb in r["rjedi_citati"][:10]:
            print(f"     ({prez.capitalize()}, {god}.)  {sa}× → {sb}×")
    if r["novi_citati"]:
        print(f"   novi citati ({len(r['novi_citati'])}): "
              + ", ".join(f"{p.capitalize()} {g}" for p, g in r["novi_citati"][:10]))

    if r["izgubljene_brojke"]:
        print()
        print(f"⚠️  IZGUBLJENE BROJKE ({len(r['izgubljene_brojke'])}) — postotci, iznosi, "
              f"podaci koji više nisu nigdje u tekstu:")
        print("     " + ", ".join(r["izgubljene_brojke"][:25])
              + (" …" if len(r["izgubljene_brojke"]) > 25 else ""))
    if r["nove_brojke"]:
        print(f"   nove brojke ({len(r['nove_brojke'])}): "
              + ", ".join(r["nove_brojke"][:15])
              + (" …" if len(r["nove_brojke"]) > 15 else ""))

    print()
    if r["uklonjeni"]:
        print(f"UKLONJENI ODLOMCI ({len(r['uklonjeni'])}):")
        for p in r["uklonjeni"][:8]:
            print(f"  − {skrati(p)}")
        if len(r["uklonjeni"]) > 8:
            print(f"    … još {len(r['uklonjeni']) - 8}")
    if r["dodani"]:
        print(f"DODANI ODLOMCI ({len(r['dodani'])}):")
        for p in r["dodani"][:8]:
            print(f"  + {skrati(p)}")
        if len(r["dodani"]) > 8:
            print(f"    … još {len(r['dodani']) - 8}")
    if r["izmijenjeni"]:
        print(f"IZMIJENJENI ODLOMCI ({len(r['izmijenjeni'])}):")
        for stari, novi in r["izmijenjeni"][:5]:
            print(f"  ~ prije: {skrati(stari, 90)}")
            print(f"    sada:  {skrati(novi, 90)}")
        if len(r["izmijenjeni"]) > 5:
            print(f"    … još {len(r['izmijenjeni']) - 5}")
    return 1 if (r["izgubljeni_citati"] or r["izgubljene_brojke"]) else 0


def za_mentora(r):
    d = r["rijeci_b"] - r["rijeci_a"]
    smjer = ("proširen za" if d > 0 else "skraćen za") if d else "isti opseg,"
    print("# Što je izmijenjeno")
    print()
    print(f"Usporedba: **{os.path.basename(r['a'])}** → **{os.path.basename(r['b'])}**  ")
    print(f"Datum: {datetime.date.today().isoformat()}")
    print()
    if d:
        print(f"Tekst je {smjer} {abs(d)} riječi (s {r['rijeci_a']} na {r['rijeci_b']}).")
    else:
        print(f"Opseg teksta je nepromijenjen ({r['rijeci_b']} riječi).")
    print()

    print("## Sadržajno")
    print()
    if r["dodani"]:
        print(f"Dodano je {len(r['dodani'])} "
              f"{mn(len(r['dodani']), 'novi odlomak', 'nova odlomka', 'novih odlomaka')}. "
              f"Najvažniji:")
        print()
        for p in r["dodani"][:5]:
            print(f"- {skrati(p, 160)}")
        print()
    if r["izmijenjeni"]:
        print(f"Prerađeno je {len(r['izmijenjeni'])} "
              f"{mn(len(r['izmijenjeni']), 'odlomak', 'odlomka', 'odlomaka')}. Primjer:")
        print()
        for stari, novi in r["izmijenjeni"][:3]:
            print(f"- *prije:* {skrati(stari, 120)}")
            print(f"  *sada:* {skrati(novi, 120)}")
        print()
    if r["uklonjeni"]:
        print(f"Uklonjeno je {len(r['uklonjeni'])} "
              f"{mn(len(r['uklonjeni']), 'odlomak', 'odlomka', 'odlomaka')}:")
        print()
        for p in r["uklonjeni"][:5]:
            print(f"- {skrati(p, 160)}")
        print()
    if not (r["dodani"] or r["izmijenjeni"] or r["uklonjeni"]):
        print("Nijedan odlomak nije dodan, uklonjen ni prerađen.")
        print()

    print("## Izvori i podaci")
    print()
    if r["izgubljeni_citati"]:
        print("**Pozor — citati kojih više nema u tekstu:**")
        print()
        for prez, god in r["izgubljeni_citati"]:
            print(f"- ({prez.capitalize()}, {god}.)")
        print()
        print("Ako je i tvrdnja uklonjena, sve je u redu; ako je tvrdnja ostala, "
              "treba joj vratiti izvor.")
        print()
    else:
        print("Svi izvori citirani u prethodnoj verziji citirani su i dalje.")
        print()
    if r["novi_citati"]:
        print("Novouvedeni izvori: "
              + ", ".join(f"{p.capitalize()} ({g}.)" for p, g in r["novi_citati"]) + ".")
        print()
    if r["izgubljene_brojke"]:
        print("Brojčani podaci kojih više nema u tekstu: "
              + ", ".join(r["izgubljene_brojke"][:20]) + ".")
        print()
    return 1 if (r["izgubljeni_citati"] or r["izgubljene_brojke"]) else 0


# ---------------------------------------------------------------------- main

def main():
    ap = argparse.ArgumentParser(
        description="Snapshoti i usporedba verzija rada (.docx).")
    ap.add_argument("stari", nargs="?", help="starija verzija")
    ap.add_argument("novi", nargs="?", help="novija verzija")
    ap.add_argument("--snapshot", metavar="RAD.DOCX", help="spremi snimku prije izmjene")
    ap.add_argument("--biljeska", metavar="TEKST", default="",
                    help="zašto se snimka radi (npr. \"prije faze G\")")
    ap.add_argument("--popis", action="store_true", help="ispiši postojeće snapshote")
    ap.add_argument("--vrati", metavar="ID", help="vrati snapshot pod novim imenom")
    ap.add_argument("--u", metavar="PUT", help="kamo vratiti snapshot")
    ap.add_argument("--force", action="store_true", help="dopusti prepisivanje pri --vrati")
    ap.add_argument("--za-mentora", action="store_true",
                    help="markdown sažetak izmjena za mentora (bez tehničkog šuma)")
    ap.add_argument("--kat", metavar="PUT", default=None,
                    help="eksplicitna mapa .katedra (ima prednost nad project rootom)")
    ap.add_argument("--project-root", default=None,
                    help="korijen rada; zadano: KATEDRA_PROJECT_ROOT ili trenutni direktorij")
    ap.add_argument("--json", dest="json_out", metavar="PUT", help="zapiši usporedbu u JSON")
    a = ap.parse_args()
    kat = resolve_state_dir(a.kat, a.project_root)

    if a.snapshot:
        return snapshot(a.snapshot, a.biljeska, kat)
    if a.popis:
        return popis(kat)
    if a.vrati:
        return vrati(a.vrati, a.u, kat, a.force)

    if not (a.stari and a.novi):
        ap.print_usage()
        print("\nTrebaju dvije verzije za usporedbu, ili jedna od naredbi "
              "--snapshot / --popis / --vrati.\n"
              "  python3 <KATEDRA_SKILL>/scripts/diff_versions.py ./rad_v2.docx ./rad_v3.docx", file=sys.stderr)
        return 2

    r = analiza(a.stari, a.novi)
    if r is None:
        return 2

    if a.json_out:
        korijen = os.path.dirname(os.path.abspath(kat))
        payload = {
            "alat": "diff_versions",
            "stari": _putanja_za_izvjestaj(a.stari, korijen),
            "novi": _putanja_za_izvjestaj(a.novi, korijen),
            "rijeci": {"prije": r["rijeci_a"], "poslije": r["rijeci_b"]},
            "odlomci": {"dodani": len(r["dodani"]), "uklonjeni": len(r["uklonjeni"]),
                        "izmijenjeni": len(r["izmijenjeni"])},
            "izgubljeni_citati": [f"{p} {g}" for p, g in r["izgubljeni_citati"]],
            "novi_citati": [f"{p} {g}" for p, g in r["novi_citati"]],
            "izgubljene_brojke": r["izgubljene_brojke"],
        }
        # v1.1-fix Q14: i izvještaj ide kroz atomarni zapis — fiksni „<put>.tmp"
        # se slijedio kroz simboličku poveznicu jednako kao i kod stanja.
        try:
            atomic_write_text(a.json_out, json.dumps(payload, ensure_ascii=False, indent=1) + "\n")
        except NesigurnaPutanja as e:
            print(f"❌ izvještaj nije zapisan: {e}", file=sys.stderr)
            return 2
        except OSError as e:
            print(f"❌ izvještaj nije zapisan ({e}).", file=sys.stderr)
            return 2

    return za_mentora(r) if a.za_mentora else ispis(r)


if __name__ == "__main__":
    try:
        sys.exit(main())
    except BrokenPipeError:
        # ispis presječen (npr. `| head`) — to nije greška, samo tiho izađi
        os._exit(0)
    except KeyboardInterrupt:
        sys.exit(130)
