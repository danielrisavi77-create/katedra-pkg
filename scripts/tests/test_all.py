#!/usr/bin/env python3
"""Regresijski testovi za rad-audit skripte (nova/izmijenjena logika).

Uporaba:  python3 test_all.py

Gradi fixture (make_fixtures.py) u privremeni folder, poziva funkcije skripti
IZRAVNO (bez subprocessa) hvatajući stdout, i provjerava očekivane ishode.
Ne zahtijeva pytest — čisti stdlib + python-docx (isto ograničenje kao i
ostatak skilla). Exit kod 0 = sve prošlo, 1 = barem jedan test pao.
"""
import sys
import os
import io
import contextlib
import tempfile
import shutil

HERE = os.path.dirname(os.path.abspath(__file__))
SCRIPTS = os.path.dirname(HERE)
sys.path.insert(0, SCRIPTS)

import make_fixtures  # noqa: E402

RESULTS = []


def check(name, condition, detail=""):
    RESULTS.append((name, bool(condition), detail))


def capture(fn, *args):
    buf = io.StringIO()
    code = 0
    try:
        with contextlib.redirect_stdout(buf):
            code = fn(*args) or 0
    except SystemExit as e:
        code = e.code if isinstance(e.code, int) else 0
    return buf.getvalue(), code


def main():
    tmp = tempfile.mkdtemp(prefix="rad_audit_fixtures_")
    fx = os.path.join(tmp, "fixtures")
    make_fixtures.build(fx)
    src = os.path.join(fx, "izvori")

    # --- 1) apply_safe_fixes: navodnici po odlomku, inč-oznaka preskočena, warning za nesparen ---
    import apply_safe_fixes
    out = os.path.join(tmp, "quotes_fixed.docx")
    txt, code = capture(apply_safe_fixes.main,
                         [os.path.join(fx, "quotes.docx"), out, "--no-mult", "--no-breaks", "--no-autofit"])
    from docx import Document
    d = Document(out)
    texts = [p.text for p in d.paragraphs]
    check("quotes: par ispravno pretvoren", texts[0] == 'On je rekao „pozdrav” i otišao.', texts[0])
    check("quotes: inč-oznaka NETAKNUTA", '12"' in texts[1], texts[1])
    check("quotes: nesparen odlomak i dalje ima jedan „", texts[2].count("„") == 1, texts[2])
    check("quotes: razdvojen kroz 3 runa ispravno spojen/pretvoren",
          texts[3] == 'Prvi „razdvojen navodnik” u tri runa.', texts[3])
    check("quotes: upozorenje za neparan odlomak ispisano", "neparan broj navodnika" in txt)

    # --- 2) check_citations: numerirani naslov LITERATURE prepoznat, bez lažnih nalaza ---
    import check_citations
    txt, code = capture(check_citations.main, os.path.join(fx, "ieee_numbered_heading.docx"))
    check("check_citations: numerirani naslov prepoznat (bez rupa/siročadi)",
          "rupe u numeraciji: nema" in txt and "SIROČAD" in txt and code == 0, txt)

    # --- 3) check_citations_authoryear: nalazi siroče + citat bez reference ---
    import check_citations_authoryear
    txt, code = capture(check_citations_authoryear.main, os.path.join(fx, "author_year.docx"))
    check("authoryear: nalazi siroče horvat/2018", "'horvat', '2018'" in txt, txt)
    check("authoryear: nalazi citat bez reference kovač/2021", "'kovač', '2021'" in txt, txt)
    check("authoryear: exit code != 0 (ima nalaza)", code != 0)

    # --- 4) common.detect_citation_style ---
    from common import detect_citation_style, load_docx_text
    body, cells, _ = load_docx_text(os.path.join(fx, "author_year.docx"), include_tables=True)
    style, counts = detect_citation_style(body + "\n" + "\n".join(cells))
    check("detect_citation_style: prepoznaje authoryear", style == "authoryear", (style, counts))

    # --- 5) common.load_supplementary_text: fusnota se čita ---
    from common import load_supplementary_text
    sup = load_supplementary_text(os.path.join(fx, "footnote.docx"))
    check("load_supplementary_text: fusnota pročitana", "Ivić, 2020" in sup["footnotes"], sup)

    # --- 6) check_fields: neprihvaćene izmjene detektirane ---
    import check_fields
    txt, code = capture(check_fields.main, os.path.join(fx, "tracked_changes.docx"))
    check("check_fields: w:ins detektiran", "NEPRIHVAĆENE IZMJENE" in txt and code != 0, txt[:200])
    txt2, code2 = capture(check_fields.main, os.path.join(fx, "quotes.docx"))
    check("check_fields: bez tracked changes -> nema upozorenja", "NEPRIHVAĆENE IZMJENE" not in txt2 and code2 == 0)

    # --- 7) domains: auto-detekcija elektro ---
    from domains import detect_domain
    body, cells, _ = load_docx_text(os.path.join(fx, "elektro.docx"), include_tables=True)
    dom, scores = detect_domain(body + "\n" + "\n".join(cells))
    check("detect_domain: elektro prepoznat", dom == "elektro", (dom, scores))

    # --- 8) numbers_inventory: koristi elektro domenu ---
    import numbers_inventory
    txt, code = capture(numbers_inventory.main, os.path.join(fx, "elektro.docx"))
    check("numbers_inventory: domena elektro u ispisu", "elektro" in txt, txt[:200])

    # --- 9) cross_check: kontekst prikazan, lažni pozitivac vidljiv u kontekstu ---
    import cross_check
    txt, code = capture(cross_check.main, [os.path.join(fx, "cross_fp.docx"), src])
    check("cross_check: prikazuje kontekst (└─)", "└─" in txt, txt[:300])
    check("cross_check: lažni pozitivac '40 t' vidljiv u kontekstu s 'tvrtke'",
          any("tvrtke" in line for line in txt.split("\n") if "40" in line or "└─" in line), txt)

    # --- 10) check_overlap: neoznačeno vs označeno vs parafraza ---
    import check_overlap
    txt, code = capture(check_overlap.main, [os.path.join(fx, "overlap.docx"), src])
    check("check_overlap: odlomak #1 (neoznačeno) flagiran",
          "#1" in txt and "NEMA vidljive oznake" in txt, txt[:400])
    check("check_overlap: odlomak #2 (označeno) prepoznat kao izgleda označeno",
          "izgleda označeno" in txt, txt[:600])
    check("check_overlap: odlomak #3 (parafraza) NIJE flagiran",
          "#3" not in txt.split("SAŽETAK")[0])

    # --- 11) common.sentences: kratice ne razbijaju rečenicu ---
    from common import sentences
    s = sentences("Firma d.o.o. posluje dobro. Vidi npr. sljedeći primjer.")
    check("sentences: 'd.o.o.' ne razbija rečenicu", s[0] == "Firma d.o.o. posluje dobro.", s)
    check("sentences: 'npr.' ne razbija rečenicu", s[1] == "Vidi npr. sljedeći primjer.", s)

    # =====================================================================
    # Runda 2 — regresijski testovi za nalaze drugog audita (17 fixeva)
    # =====================================================================
    import zipfile as _zip

    # --- R1) citat samo u ćeliji tablice NIJE siroče; [2020] nije citat ---
    txt, code = capture(check_citations.main, os.path.join(fx, "table_cite.docx"))
    check("R1 tablica: [2] iz ćelije NIJE siroče", "SIROČAD (u popisu, ne citirano): nema" in txt, txt)
    check("R1 godina: [2020] prijavljen kao godina, ne citat",
          "vjerojatno godina" in txt and "[2020]" in txt, txt)
    check("R1: interno konzistentno (exit 0)", code == 0)

    # --- R2) autor-godina citat SAMO u fusnoti se vidi (end-to-end) ---
    txt, code = capture(check_citations_authoryear.main, os.path.join(fx, "fn_ay_cite.docx"))
    check("R2 fusnota: citat iz fusnote prebrojan", "Citirano u tekstu (uklj. fusnote/endnote): 1" in txt, txt)
    check("R2 fusnota: nema lažnog siročeta", "SIROČAD (u popisu, ne citirano): nema" in txt, txt)

    # --- R3) --no-indent ne kontaminira rPr; XML valjan raspored ---
    out_r3 = os.path.join(tmp, "rpr_fixed.docx")
    txt, code = capture(apply_safe_fixes.main,
                         [os.path.join(fx, "rpr_spacing.docx"), out_r3, "--no-indent", "--no-quotes", "--no-mult"])
    import re as _re
    x = _zip.ZipFile(out_r3).read("word/document.xml").decode()
    check("R3: NEMA w:after u rPr spacingu", not _re.search(r"<w:rPr>[^<]*<w:spacing[^>]*w:after", x), None)
    check("R3: w:after JEST u pPr", bool(_re.search(r'<w:pPr>.*?<w:spacing[^>]*w:after="120"', x, _re.S)), None)

    # --- R4) engleski “…” par ostaje; njemački „…“ se popravi ---
    out_r4 = os.path.join(tmp, "eng_fixed.docx")
    capture(apply_safe_fixes.main,
            [os.path.join(fx, "eng_quotes.docx"), out_r4, "--no-mult", "--no-breaks", "--no-autofit"])
    from docx import Document as _D
    pts = [p.text for p in _D(out_r4).paragraphs]
    check("R4: engleski par netaknut", "“smart control”" in pts[0], pts[0])
    check("R4: njemački par popravljen u „…”", "„citat”" in pts[1], pts[1])

    # --- R5) hex literal preskočen, prava multiplikacija pretvorena ---
    out_r5 = os.path.join(tmp, "hex_fixed.docx")
    capture(apply_safe_fixes.main,
            [os.path.join(fx, "hex.docx"), out_r5, "--no-quotes", "--no-breaks", "--no-autofit"])
    t5 = _D(out_r5).paragraphs[0].text
    check("R5: 0x41/0xFF00 netaknuti", "0x41" in t5 and "0xFF00" in t5, t5)
    check("R5: 80x80 → 80 × 80", "80 × 80" in t5, t5)

    # --- R6) _normal_firstline0 NE dira tuđi stil ---
    out_r6 = os.path.join(tmp, "styles_fixed.docx")
    capture(apply_safe_fixes.main,
            [os.path.join(fx, "styles_extra.docx"), out_r6, "--no-indent", "--no-quotes", "--no-mult"])
    s6 = _zip.ZipFile(out_r6).read("word/styles.xml").decode()
    m6 = _re.search(r'styleId="Citat9".*?</w:style>', s6, _re.S)
    check("R6: Citat9 firstLine=709 netaknut", m6 and 'firstLine="709"' in m6.group(0), None)

    # --- R7) tekst POSLIJE inline textboxa se obrađuje (ne preskače tiho) ---
    out_r7 = os.path.join(tmp, "txbx_fixed.docx")
    capture(apply_safe_fixes.main,
            [os.path.join(fx, "txbx.docx"), out_r7, "--no-mult", "--no-breaks", "--no-autofit"])
    x7 = _zip.ZipFile(out_r7).read("word/document.xml").decode()
    t7 = _re.findall(r"<w:t[^>]*>([^<]*)</w:t>", x7)
    check("R7: navodnici PRIJE textboxa pretvoreni", any("„prvi citat”" in t for t in t7), t7[:3])
    check("R7: navodnici POSLIJE textboxa pretvoreni", any("„drugi citat”" in t for t in t7), t7[:3])

    # --- R8) %/°/V/Hz jedinice + detekcija sukoba vrijednosti ---
    txt, code = capture(numbers_inventory.main, os.path.join(fx, "elektro_konflikt.docx"))
    check("R8: V i Hz u inventaru", "V    :" in txt.replace("V   :", "V    :") or " V " in txt, txt)
    check("R8: % u inventaru", _re.search(r"%\s*:", txt) is not None, txt)
    check("R8: ° u inventaru", _re.search(r"°\s*:", txt) is not None, txt)
    check("R8: sukob 'napon' 230 vs 400 flagiran",
          "'napon'" in txt and "230" in txt and "400" in txt and "⚠" in txt, txt)
    check("R8: sekcija rečenica NIJE prazna (skriveni filter maknut)",
          "Sustav radi na naponu" in txt, txt)

    # --- R9) mixed stil: generate_report pokreće OBA B checkera + napomenu ---
    import generate_report as _gr
    out_j9 = os.path.join(tmp, "mixed.json")
    capture(_gr.main, [os.path.join(fx, "mixed_style.docx"), "--out",
                        os.path.join(tmp, "mixed.md"), "--json", out_j9])
    import json as _json
    p9 = _json.load(open(out_j9, encoding="utf-8"))
    b_phases = [k for k in p9["phase_exit_codes"] if k.startswith("B")]
    check("R9: mixed → IEEE i autor-godina faze + napomena",
          len(b_phases) == 3 and any("Napomena" in k for k in b_phases), b_phases)

    # --- R10) cross_check exit kod 1 kad ima nepotvrđenih tvrdnji ---
    # (elektro_konflikt ima 230 V koji NE postoji u izvorima → mora dati exit 1;
    #  cross_fp fixture namjerno ima sve "pronađeno" pa nije prikladan ovdje)
    txt, code = capture(cross_check.main, [os.path.join(fx, "elektro_konflikt.docx"), src])
    check("R10: cross_check exit != 0 uz promašaje",
          code != 0 and "nije nađeno u izvorima" in txt, (code, txt[-200:]))

    # --- R11) '-ost.' kraj rečenice se NE guta ('st.' kratica) ---
    s11 = sentences("Provjerena je nosivost. Rezultat je dobar.")
    check("R11: 'nosivost.' ispravno završava rečenicu", len(s11) == 2, s11)

    # --- R12) golo 'S' nije claim; oznaka čelika s kvalitetom jest ---
    from cross_check import auto_claims as _ac
    cl12, _dom = _ac("Čelik S 355 i stup tipa S te kvaliteta S355J2 u konstrukciji od čelika s vijkom i pločom uz profil.")
    check("R12: golo 'S' NIJE claim", "S" not in cl12, cl12)
    check("R12: 'S 355' i 'S355J2' jesu claimovi", "S 355" in cl12 and "S355J2" in cl12, cl12)

    # --- 12) generate_report: end-to-end, kritični nalazi bucketirani ---
    import generate_report
    out_md = os.path.join(tmp, "report.md")
    out_json = os.path.join(tmp, "report.json")
    txt, code = capture(generate_report.main,
                         [os.path.join(fx, "author_year.docx"), "--sources", src,
                          "--out", out_md, "--json", out_json])
    check("generate_report: .md spremljen", os.path.exists(out_md))
    check("generate_report: .json spremljen", os.path.exists(out_json))
    import json
    with open(out_json, encoding="utf-8") as f:
        payload = json.load(f)
    check("generate_report: siroče/citat-bez-reference u KRITIČNO bucketu",
          payload["counts"]["kritično"] >= 2, payload["counts"])

    # --- R16 (v1.9): Vancouver (N) dijalekt ---
    from common import parse_vancouver_citations, vancouver_is_decimal
    body16, cells16, _ = load_docx_text(os.path.join(fx, "vancouver.docx"), include_tables=True)
    style16, counts16 = detect_citation_style(body16 + "\n" + "\n".join(cells16))
    check("R16: detect_citation_style prepoznaje vancouver", style16 == "vancouver", (style16, counts16))
    check("R16: autor-godina NE vidi Rec(2003)24 kao citat", counts16["authoryear"] == 0, counts16)
    check("R16: tablična ćelija „158 (77,8)\" nije citat",
          vancouver_is_decimal("77,8", "158 ") and parse_vancouver_citations("158 (77,8)") == [])
    check("R16: „(67,68)\" jest citat, „(12,35)\" iza brojke nije",
          parse_vancouver_citations("drugdje (67,68)")[0][2] == {67, 68}
          and parse_vancouver_citations("158 (12,35)") == [])
    check("R16: svezak(broj) „53(3-4)\" nije citat", parse_vancouver_citations("2013;53(3-4):367") == [])
    check("R16: raspon „(3–7)\" se širi", parse_vancouver_citations("x (3–7)")[0][2] == {3, 4, 5, 6, 7})
    txt16, code16 = capture(check_citations.main, os.path.join(fx, "vancouver.docx"))
    check("R16: check_citations bira Vancouver", "[Vancouver (N)]" in txt16, txt16)
    check("R16: popis 1..68 prepoznat, prilog NIJE stavka", "Definirano u LITERATURI: 7" in txt16, txt16)
    check("R16: siroče 5", "SIROČAD (u popisu, ne citirano): [5]" in txt16, txt16)
    check("R16: citat bez reference 6", "CITAT BEZ REFERENCE: [6]" in txt16, txt16)
    check("R16: redoslijed prekršen (3 prije 2)", "krši rastući redoslijed" in txt16, txt16)
    check("R16: bez razmaka iza zareza „67,68\"", "bez razmaka iza zareza" in txt16 and "67,68" in txt16, txt16)
    check("R16: sedam autora bez „i sur.\" → stavka 4", "bez „i sur.\"/„et al.\": stavke [4]" in txt16, txt16)
    check("R16: exit code != 0 (ima nalaza)", code16 != 0)
    # IEEE fixture i dalje prolazi kroz istu skriptu s eksplicitnim stilom
    txt16b, code16b = capture(check_citations.main, os.path.join(fx, "ieee_numbered_heading.docx"), "ieee")
    check("R16: IEEE fixture s eksplicitnim stilom bez nalaza", code16b == 0 and "[IEEE [N]]" in txt16b, txt16b)

    shutil.rmtree(tmp, ignore_errors=True)

    # --- R13: zakrpe nađene na obranjenom FPZG radu (kolovoz 2026.) ---
    # Svih pet ima isti oblik: alat je bio kalibriran na jedan citatni dijalekt, pa
    # rad koji radi nešto drukčije — ali ispravno — prijavljuje kao pogrešan. To je
    # najskuplja vrsta kvara ovdje: student dobije popis nepostojećih grešaka i,
    # ako mu vjeruje, pokvari rad koji je bio dobar.
    import common as C13
    import check_citations_authoryear as B13

    def _zagradni(t):
        nadjeno = C13.CITE_AY_RE.findall(t)
        return C13.parse_ay_citation_group(nadjeno[0]) if nadjeno else set()

    # (1) lokator stranice iza godine — FPZG Upute propisuju baš taj oblik
    check("R13: (Becker, 2007: 45) je citat",
          ("becker", "2007") in _zagradni("(Becker, 2007: 45)"))
    check("R13: (Streeck, 2014: xiv) je citat",
          ("streeck", "2014") in _zagradni("(Streeck, 2014: xiv)"))
    check("R13: narativni lokator (Krippner (2005: 174))",
          ("krippner", "2005") in C13.parse_ay_narrative("Krippner (2005: 174) tvrdi"))

    # (2) sufiks godine je dio identiteta, inače se dva rada slijevaju u jedan ključ
    check("R13: sufiks 2013a/2013b preživi zagradni oblik",
          C13.parse_ay_segment("Becker, 2013a") == ("becker", "2013a")
          and C13.parse_ay_segment("Becker, 2013b") == ("becker", "2013b"))

    # (3) popis literature u hrvatskim/FPZG oblicima
    for redak, kljuc in [
        ("Becker, Gary (2007) Ekonomski pristup. Zagreb: Naklada.", ("becker", "2007")),
        ("Becker, G. (2007) Ekonomski pristup. Zagreb: Naklada.", ("becker", "2007")),
        ("Van der Zwan, Natascha (2014) Making sense. SER 12(1).", ("zwan", "2014")),
        ("HNB (Hrvatska narodna banka) (2023) Izvješće. Zagreb: HNB.", ("hnb", "2023")),
        ("easyJet plc (2025.), Full year results, Luton: easyJet plc", ("easyjet", "2025")),
        ("Podravka d.d. (2024.) Izvješće. Koprivnica: Podravka.", ("podravka", "2024")),
    ]:
        kljucevi, _ = B13.extract_biblio_keys(redak)
        check(f"R13: popis literature — {redak[:34]}", kljuc in kljucevi, kljucevi)

    check("R13: obična rečenica u popisu NIJE referenca",
          B13.extract_biblio_keys(
              "ovo je obična rečenica koja spominje 2024. godinu") == (set(), 1))

    # (4) tekst i popis moraju dati ISTI ključ za istu referencu
    iz_teksta = _zagradni("(Van der Zwan, 2014)")
    iz_popisa, _ = B13.extract_biblio_keys("Van der Zwan, Natascha (2014) Making sense.")
    check("R13: Van der Zwan daje isti ključ iz teksta i iz popisa",
          bool(iz_teksta & iz_popisa), (iz_teksta, iz_popisa))

    # (5) institucija s malom riječi u imenu, uz granicu da proza ne postane citat
    check("R13: Europska komisija (2021.) je citat",
          ("europska", "2021") in C13.parse_ay_narrative("Europska komisija (2021.) je odobrila"))
    check("R13: proza sa zagradnom godinom NIJE citat",
          C13.parse_ay_narrative(
              "Analiza je provedena u promatranom razdoblju (2021.)") == set())

    # --- report ---
    passed = sum(1 for _, ok, _ in RESULTS if ok)
    print("=" * 70)
    print(f"REZULTATI TESTOVA: {passed}/{len(RESULTS)} prošlo")
    print("=" * 70)
    for name, ok, detail in RESULTS:
        mark = "✓" if ok else "✗ FAIL"
        print(f"  {mark:8} {name}")
        if not ok and detail:
            print(f"           detalj: {detail}")
    return 0 if passed == len(RESULTS) else 1


if __name__ == "__main__":
    sys.exit(main())
