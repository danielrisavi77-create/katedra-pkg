#!/usr/bin/env python3
"""Centralni citatni parser za Katedru.

B10 contract: consumeri ne smiju imati vlastite regexe za citation dialecte.
Ovaj modul prepoznaje citatne *reference* i njihove tipove; ne provjerava postoji
li izvor niti podržava li tvrdnju. To pripada source/evidence slojevima.
"""
from __future__ import annotations

from collections import Counter
import json
from pathlib import Path
import re
from typing import NamedTuple
import unicodedata
import zipfile
import xml.etree.ElementTree as ET


CITATION_STYLES = (
    'autor-godina', 'apa', 'apa-hr', 'harvard', 'ieee', 'vancouver', 'legal-footnote',
)
STYLE_TO_DIALECT = {
    'autor-godina': 'author-year',
    'apa': 'author-year',
    'apa-hr': 'author-year',
    'harvard': 'author-year',
    'ieee': 'ieee',
    'vancouver': 'vancouver',
    'legal-footnote': 'legal-footnote',
}
# v1.9 (nalaz 6): dijalekti čiji je ključ REDNI BROJ stavke u popisu literature.
# Potrošači koji granaju „numerički vs. autor-godina" pitaju ovo, ne nabrajaju
# dijalekte — inače svaki novi numerički dijalekt traži obilazak svih consumera.
NUMERIC_DIALECTS = ('ieee', 'vancouver')
SOURCE_TYPES = (
    'journal_article', 'book', 'chapter', 'law', 'regulation', 'court_decision',
    'eu_act', 'official_report', 'dataset', 'web_page', 'unknown',
)


class CitationDialectError(ValueError):
    pass


class CitationRef(NamedTuple):
    dialect: str
    key: str
    raw: str
    source_type: str = 'unknown'
    year: str = ''
    author: str = ''
    year_has_dot: bool | None = None


# Author-year compatibility patterns are re-exported by hr_text.
CITAT_ZAGRADNI = re.compile(
    r"\(([^()]{2,160}?,\s*\d{4}\.?[a-z]?"
    r"(?:\s*;\s*[^()]{2,160}?,\s*\d{4}\.?[a-z]?)*)\)"
)
# Q13-fix: rastavljač koautora je isti u zagradnom i narativnom obliku, pa se
# piše na JEDNOM mjestu. Prije su „&" i „and" nedostajali, a bili su prisutni u
# `verify_sources._AUTOR_SEP_RE` — ista pojava, dva različita popisa.
_KOAUTOR_SEP = r"(?:\s*,\s*|\s+(?:i|te|and)\s+|\s*&\s*)"
# „et al." se veže bez rastavljača („Čavlek et al. (2020.)"), za razliku od
# hrvatskih „i sur."/"i dr." koji dolaze iza veznika.
_ET_AL = r"(?:\s+et\s+al\.)"
# Q13, drugi krug: uzrok Q13 je imenovan kao „uzorak traži token s velikim
# početnim slovom prije zagrade", ali je zatvorena samo grana „et al.". Prezime
# s malim plemićkim/patronimskim prefiksom („de Vries", „van Dijk", „van der
# Borg", „von Neumann") i dalje je iz narativnog citata davalo ključ „vries" /
# „dijk", dok je popis literature preko `verify_sources._PREZIME_RE` davao „de
# vries" / „van dijk" — isti izvor prijavljen ISTOVREMENO kao „nigdje citiran" i
# kao „citiran bez izvora", tj. točno ona dvostruka prijava zbog koje je Q13 i
# otvoren. Čestica je dio prezimena, ne riječ rečenice.
# Popis je zatvoren razred (nije semantička odluka nego pravopisna činjenica) i
# `verify_sources` ga preuzima odavde, da ne postoje opet dva popisa — dva
# popisa i jesu bila uzrok Q13.
# Granica riječi s obje strane je nužna: bez nje bi „…tvrde Vries (2018.)"
# uhvatilo „de" iz „tvrde" i dalo drukčiji opseg na ispravnom ulazu.
# Popis SMIJE sadržavati i riječi koje su ujedno hrvatske („do", „da"), jer se
# čestica ionako skida iz ključa (v. `kljuc_prvog_autora`): u „…rast je trajao
# do Marić (2019.)" uzorak pokupi „do Marić", ali ključ ostaje „marić", isti kao
# iz popisa literature. Ovdje uzorak služi samo tome da se u višeautorskom
# narativnom citatu („de Vries i van Dijk (2018.)") uhvati PRVI autor, a ne
# drugi — bez čestice regex se usidri na „Dijk" i ključ je bio „dijk".
PARTIKULA = (r"(?:v[ao]n|der|den|dell[aeio]|degli|dei|del|de|di|du|dalla|dal|"
             r"da|la|le|lo|do|dos|das|ter|ten|bin|ibn)")
_PREZIME = r"(?:\b" + PARTIKULA + r"\b\s+)*[A-ZČĆŽŠĐ][\wÀ-ɏ'’\-]+"
_KOAUTORI = (r"(?:" + _KOAUTOR_SEP + r"(?:sur\.|dr\.|" + _PREZIME + r")|" + _ET_AL + r")*")
CITAT_NARATIVNI = re.compile(
    r"(" + _PREZIME + _KOAUTORI + r")"
    r"\s*\((\d{4})\.?([a-z]?)\)"
)
AUTHOR_YEAR_PAREN_DETAIL = re.compile(
    r"\(([^()]{2,200}?,\s*\d{4}\.?[a-z]?"
    r"(?:\s*;\s*[^()]{2,200}?,\s*\d{4}\.?[a-z]?)*)"
    r"(?:\s*[,:]\s*(?:str|s|p|pp|ss)\.?\s*[\dIVXivx][^()]{0,24})?\s*\)"
)
AUTHOR_YEAR_NARRATIVE_DETAIL = re.compile(
    r"(" + _PREZIME + _KOAUTORI + r")"
    r"\s*\((\d{4})(\.?)([a-z]?)\)"
)
# D4-fix (v1.0.1): lokator sa završnom točkom („, str. 41.") prije je ostavljao
# tu točku u tekstu, pa je nastajalo „(Čavlek, 1998..)" i citat se više NIJE
# prepoznavao — verify_rewrite je uspoređivao Counter() s Counter() i javljao
# „0 referenci, identično" iako je citat u međuvremenu zamijenjen.
_LOK_BROJ = r"\d+(?:\s*[–—-]\s*\d+)?"
_LOK_REP = r"(?:\s*\.?\s*(?:i|te)\s*" + _LOK_BROJ + r")?\s*\.?"
LOKATOR = re.compile(
    r"((?:1[89]|20)\d{2}\.?[a-z]?)"
    r"\s*(?:,\s*(?:str|ss|pp|s|p)\.?\s*" + _LOK_BROJ + _LOK_REP +
    r"|:\s*" + _LOK_BROJ + _LOK_REP + r")"
)
# Ako je uz godinu ipak ostala dvostruka točka, skupi je prije prepoznavanja.
_DVOSTRUKA_TOCKA = re.compile(r"((?:1[89]|20)\d{2}\.[a-z]?)\.+")

IEEE_RANGE_SEPARATE = re.compile(r"\[(\d+)\]\s*[–—-]\s*\[(\d+)\]")
IEEE_GROUP = re.compile(r"\[\s*(\d+(?:\s*[,;–—-]\s*\d+)*)\s*\]")

# v1.9 (nalaz 6): Vancouver — brojevi u OVALNIM zagradama: „(1)", „(2, 5)",
# „(3–7)", „(12,13)". Zdravstveni fakulteti (HKS-FZS, MEF, ZVU…) traže baš ovaj
# oblik, a paket ga dosad nije poznavao pa je takav rad prolazio kao rad BEZ
# citata. Najviše tri znamenke: „(2003)" je godina, ne citat.
# Isti se uzorak koristi za tijelo, ćelije tablica i fusnote; iz njega se
# izvode i sve stilske provjere (razmak iza zareza, spojnica u rasponu,
# nabrajanje umjesto raspona) — B10: consumeri nemaju vlastite regexe.
VANCOUVER_GROUP = re.compile(
    r"\(\s*(\d{1,3}(?:\s*[,;]\s*\d{1,3}|\s*[–—-]\s*\d{1,3})*)\s*\)"
)
# Decimale u ovalnoj zagradi koje NISU citat (pravila prenesena iz
# katedra-lite-dodaci/scripts/provjeri_vancouver.py, gdje su dokazana na radu
# sa 75 referenci i tablicom „n (%)"):
#   • jedna znamenka iza zareza — „(77,8)", „(50,0)" — uvijek decimala;
#   • dvije znamenke iza zareza — „(12,35)" — decimala samo ako joj u istom
#     odlomku/ćeliji NEPOSREDNO prethodi brojka: tablična ćelija „158 (12,35)".
#     Bez te brojke „(67,68)" je citat bez razmaka iza zareza (stilski nalaz).
#   • zagrada zalijepljena za brojku — „2013;53(3-4):367-76" — je svezak(broj)
#     u popisu literature, ne citat.
_VANC_DECIMAL_1 = re.compile(r"\d{1,3},\d")
_VANC_DECIMAL_2 = re.compile(r"\d{1,3},\d{2}")
_VANC_BROJKA_ISPRED = re.compile(r"\d\s*$")
_VANC_ZALIJEPLJENA = re.compile(r"\d$")


class VancouverGroup(NamedTuple):
    """Jedna ovalna zagrada s citatima i sve što stilske provjere trebaju."""
    start: int
    raw: str                    # sadržaj zagrade, bez zagrada
    nums: tuple                 # ekspandirani brojevi (raspon → svi članovi)
    bez_razmaka: bool           # „(67,68)" — Vancouver traži „(67, 68)"
    spojnica: bool              # „(5-7)" — raspon spojnicom umjesto en-crtice
    nabrajanje: bool            # „(5, 6, 7)" — tri i više uzastopnih bez raspona


def vancouver_je_decimala(sadrzaj: str, prefiks: str) -> bool:
    """Je li „(sadrzaj)" decimala/postotak ili svezak(broj), a ne citat."""
    if _VANC_DECIMAL_1.fullmatch(sadrzaj):
        return True
    if _VANC_DECIMAL_2.fullmatch(sadrzaj) and _VANC_BROJKA_ISPRED.search(prefiks):
        return True
    if _VANC_ZALIJEPLJENA.search(prefiks):
        return True
    return False


def vancouver_groups(text: str) -> list[VancouverGroup]:
    """Sve Vancouver zagrade u tekstu, bez decimala i bez svezak(broj) oblika."""
    text = text or ''
    out: list[VancouverGroup] = []
    for m in VANCOUVER_GROUP.finditer(text):
        g = m.group(1)
        if vancouver_je_decimala(g, text[:m.start()]):
            continue
        nums: list[int] = []
        for part in re.split(r'\s*[,;]\s*', g):
            nums.extend(_expand_numeric_part(part))
        if not nums:
            continue
        ns = [int(x) for x in re.findall(r'\d+', g)]
        nabrajanje = (len(ns) >= 3 and not re.search(r'[–—-]', g)
                      and ns == list(range(ns[0], ns[0] + len(ns))))
        out.append(VancouverGroup(
            m.start(), g, tuple(nums),
            bool(re.search(r'\d,\d', g)),
            bool(re.search(r'\d\s*-\s*\d', g)),
            nabrajanje,
        ))
    return out


def parse_vancouver(text: str) -> list[CitationRef]:
    refs: list[CitationRef] = []
    for grp in vancouver_groups(text):
        raw = f'({grp.raw})'
        for n in sorted(grp.nums):
            refs.append(CitationRef('vancouver', str(n), raw))
    return refs


# Numerirani popis literature: „1. Autor A…", „1) Autor A…", „[1] Autor A…".
# Vrijedi za oba numerička dijalekta; ključ stavke je isti kao ključ citata.
NUMERIC_LIST_ITEM = re.compile(r"^\s*(?:(\d{1,3})\s*[.)]|\[\s*(\d{1,3})\s*\])\s+\S")
# Sljedeći naslov prve razine iza popisa („9. PRILOZI", „ŽIVOTOPIS") — kraj popisa
# kad struktura dokumenta ne nosi stilove naslova.
_KRAJ_POPISA = re.compile(r"^\s*(?:\d{1,2}\.\s+)?[A-ZČĆŽŠĐ][A-ZČĆŽŠĐ ]{3,}$")


def numeric_list_items(lines) -> tuple[dict[int, str], list[int]]:
    """(stavke {broj: tekst}, dupli brojevi) iz redaka numeriranog popisa."""
    stavke: dict[int, str] = {}
    dupli: list[int] = []
    for t in lines:
        m = NUMERIC_LIST_ITEM.match(t or '')
        if not m:
            continue
        n = int(m.group(1) or m.group(2))
        if n in stavke:
            dupli.append(n)
        stavke[n] = t.strip()
    return stavke, sorted(set(dupli))


def split_reference_list(paragraphs, od_naslova: str | None = None):
    """(tijelo, popis) po naslovu popisa literature (hr_text.NASLOV_LIT).

    Popis završava na sljedećem verzalnom naslovu prve razine ili na kraju.
    Ako naslova nema, sve je tijelo, a popis je prazan.
    """
    try:
        import hr_text as H
        naslov_lit = H.NASLOV_LIT
    except Exception:  # pragma: no cover - samostalno pokretanje izvan scripts/
        naslov_lit = re.compile(r"(?i)^\s*(?:\d+\.?\s*)?(?:popis\s+)?(?:citirane\s+)?"
                                r"(?:literatura|literature|reference|references|bibliografija)\s*$")
    paras = [(p or '') for p in paragraphs]
    start = None
    for i, t in enumerate(paras):
        s = t.strip()
        if od_naslova:
            if s.upper().endswith(od_naslova.strip().upper()):
                start = i
                break
        elif naslov_lit.match(s):
            start = i
            break
    if start is None:
        return paras, []
    kraj = len(paras)
    for j in range(start + 1, len(paras)):
        s = paras[j].strip()
        # Verzalni naslov („9. PRILOZI") jest oblika „N. tekst", ali referenca
        # nikad nije cijela verzalna i bez interpunkcije — zato se ovdje NE
        # isključuje ono što NUMERIC_LIST_ITEM prepoznaje.
        if s and _KRAJ_POPISA.match(s):
            kraj = j
            break
    return paras[:start] + paras[kraj:], paras[start + 1:kraj]


def _ieee_groups(text: str) -> list[VancouverGroup]:
    """IEEE „[n]" grupe u istom obliku kao Vancouver, da `numeric_report` radi
    za oba numerička dijalekta. Stilske zastavice su specifične za Vancouver
    (razmak iza zareza, en-crtica) pa ostaju False."""
    out: list[VancouverGroup] = []
    po_grupi: dict[tuple[int, str], list[int]] = {}
    for r in parse_ieee(text):
        po_grupi.setdefault((0, r.raw), []).append(int(r.key))
    for (_, raw), nums in po_grupi.items():
        out.append(VancouverGroup(0, raw.strip('[]'), tuple(nums), False, False, False))
    return out


def numeric_report(paragraphs, cells=(), footnotes=(), od_naslova: str | None = None,
                   style: str = 'vancouver') -> dict:
    """Numerički citati protiv numeriranog popisa, nad već pročitanim tekstom.

    Ista pravila kao katedra-lite-dodaci/scripts/provjeri_vancouver.py:
    siročad, citat bez reference, redoslijed prvog pojavljivanja, rasponi,
    razmak iza zareza, prazni/dupli brojevi u popisu. Ne provjerava sadržaj
    reference (verify_sources) ni „i sur." pravilo (kućni stil).
    Bez ovisnosti o python-docx — pozivatelj daje odlomke, ćelije i fusnote.
    """
    dialect = resolve_dialect(style)
    if dialect not in NUMERIC_DIALECTS:
        raise CitationDialectError(f'numeric_report traži numerički dijalekt, ne {style}')
    grupe_iz = vancouver_groups if dialect == 'vancouver' else _ieee_groups
    tijelo, popis = split_reference_list(paragraphs, od_naslova)
    popis_set = set(popis)
    tekst_za_citate = [t for t in list(tijelo) + list(cells) + list(footnotes)
                       if t and t not in popis_set]
    redoslijed: list[int] = []
    grupe: list[VancouverGroup] = []
    for t in tekst_za_citate:
        for grp in grupe_iz(t):
            grupe.append(grp)
            redoslijed.extend(grp.nums)
    citirani = set(redoslijed)
    stavke, dupli = numeric_list_items(popis)
    n_max = max(stavke) if stavke else 0
    prazni = [n for n in range(1, n_max + 1) if n not in stavke]
    prvi_put: list[int] = []
    videno: set[int] = set()
    for n in redoslijed:
        if n not in videno:
            videno.add(n)
            prvi_put.append(n)
    skokovi: list[tuple[int, int]] = []
    ocekivan = 1
    for n in prvi_put:
        if n > ocekivan:
            skokovi.append((n, ocekivan))
        ocekivan = max(ocekivan, n + 1) if n >= ocekivan else ocekivan
    return {
        'popis_stavki': len(stavke), 'N_max': n_max,
        'popis_prazni_brojevi': prazni, 'popis_dupli': dupli,
        'citata_u_tekstu': len(grupe), 'razlicitih_citiranih': len(citirani),
        'sirocad': sorted(n for n in stavke if n not in citirani),
        'citat_bez_reference': sorted(n for n in citirani if n not in stavke),
        'prvo_pojavljivanje': prvi_put[:200],
        'skokovi_redoslijeda': skokovi[:20],
        'raspon_sa_spojnicom': [g.raw for g in grupe if g.spojnica],
        'nabrajanje_umjesto_raspona': [g.raw for g in grupe if g.nabrajanje],
        'citat_bez_razmaka': [g.raw for g in grupe if g.bez_razmaka],
    }


def vancouver_report(paragraphs, cells=(), footnotes=(), od_naslova: str | None = None) -> dict:
    return numeric_report(paragraphs, cells, footnotes, od_naslova, 'vancouver')


def numeric_report_file(path: str | Path, style: str = 'vancouver',
                        od_naslova: str | None = None) -> dict:
    """`numeric_report` nad .docx-om: odlomci + ćelije tablica + fusnote."""
    p = Path(path)
    if p.suffix.lower() != '.docx':
        return numeric_report(p.read_text(encoding='utf-8').splitlines(), (), (), od_naslova, style)
    from docx import Document
    doc = Document(p)
    paragraphs = [(x.text or '') for x in doc.paragraphs]
    cells: list[str] = []

    def _celije(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    cells.append(cell.text or '')
                    _celije(cell.tables)
    _celije(doc.tables)
    footnotes = list(extract_docx_footnotes(p).values())
    return numeric_report(paragraphs, cells, footnotes, od_naslova, style)

_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def _norm(s: str) -> str:
    s = unicodedata.normalize('NFC', str(s or '')).strip().lower()
    return re.sub(r'\s+', ' ', s)


def resolve_dialect(style: str | None) -> str:
    value = (style or 'autor-godina').strip().lower()
    if value not in STYLE_TO_DIALECT:
        raise CitationDialectError(
            f'nepoznat citatni stil: {style}; dopušteno: ' + ', '.join(CITATION_STYLES)
        )
    return STYLE_TO_DIALECT[value]


def load_style_from_profile(profile_path: str | Path | None) -> str | None:
    if not profile_path:
        return None
    p = Path(profile_path)
    if not p.is_file():
        raise CitationDialectError(f'profil ne postoji: {p}')
    try:
        data = json.loads(p.read_text(encoding='utf-8'))
    except Exception as exc:
        raise CitationDialectError(f'profil nije valjan JSON: {p}: {exc}') from exc
    citation = data.get('citiranje')
    if citation is None:
        return None
    if not isinstance(citation, dict):
        raise CitationDialectError('citiranje u profilu mora biti objekt')
    value = citation.get('stil')
    return str(value) if value else None


def resolve_style(explicit: str | None, profile_path: str | Path | None = None):
    if explicit:
        style, source = explicit, 'cli'
    else:
        from_profile = load_style_from_profile(profile_path)
        style, source = (from_profile, 'profile') if from_profile else ('autor-godina', 'default')
    dialect = resolve_dialect(style)
    return style.lower(), dialect, source


def bez_lokatora(text: str) -> str:
    return LOKATOR.sub(r'\1', text or '')


# Q13-fix: ključ citata prije je nastajao iz cijelog niza autora podijeljenog
# samo na „,", „ i " i „ te ". Zato je „(Müller & Schmidt, 2020)" davao ključ
# „müller & schmidt", a „(Čavlek et al., 2011)" ključ „čavlek et al." — dva
# različita ključa za istog prvog autora, ovisno o tome kako je citat napisan.
# Isti izvor tako u pokrivenosti i u otisku citata figurira kao dva različita.
# U popisu literature autori se razdvajaju i točkom-zarezom; u tekućem tekstu se
# točka-zarez NE smije brojati kao rastavljač autora (inače „…zaključak; Perić
# (2020.)" daje ključ „zaključak"), pa je ona samo ovdje, a ne i u citatnim
# uzorcima gore.
_KOAUTOR_SEP_RE = re.compile(_KOAUTOR_SEP.replace(r"\s*,\s*", r"\s*[,;]\s*"), re.IGNORECASE)
_ITD_RE = re.compile(r"\s*(?:et\s+al\.?|i\s+sur\.?|i\s+dr\.?)\s*$", re.IGNORECASE)
# Q13, drugi krug: ključ se svodi na jezgru prezimena, bez vodećih čestica.
# Popis literature i tekući tekst NIKAD ne mogu imati identičan popis čestica —
# `verify_sources._PREZIME_RE` u popisu prihvaća bilo koju malu riječ ispred
# prezimena, a u tekućem tekstu ista ta riječ može biti obična riječ rečenice
# („…rast je trajao do Marić (2019.)"). Zato se pravilo ne oslanja na to da oba
# uzorka pogode isti opseg: čestica se skida iz KLJUČA, pa „de Vries" i „Vries",
# „van der Borg" i „Borg" daju isti ključ bez obzira na to koja ga je strana
# proizvela i je li pisac uopće napisao česticu. Tolerancija ide u sigurnom
# smjeru — najgore što se može dogoditi je da se dva zapisa poklope (nema
# prijave), a ne da se isti izvor prijavi dvaput u dvije suprotne rubrike.
_VODECE_CESTICE_RE = re.compile(r"^(?:\b" + PARTIKULA + r"\b\s+)+", re.IGNORECASE)


def kljuc_prvog_autora(autori: str) -> str:
    """Prezime prvog autora iz prikaznog niza („Müller & Schmidt" → „müller").

    Jedno mjesto za pravilo: koriste ga i citatni parser (ključ citata) i
    `verify_sources.pokrivenost` (ključ jedinice u popisu literature). Dok su
    to bila dva popisa rastavljača, ista je jedinica dobivala dva ključa i
    prijavljivala se istovremeno kao „nigdje citirana" i kao „citirana bez
    izvora". Iz istog razloga se skidaju i vodeće čestice („de Vries" → „vries").
    Prikazno ime izvora ostaje netaknuto — mijenja se samo ključ za usporedbu.
    """
    s = _ITD_RE.sub("", str(autori or "").strip())
    prvi = _KOAUTOR_SEP_RE.split(s)[0]
    prvi = _ITD_RE.sub("", prvi).strip().lower()
    jezgra = _VODECE_CESTICE_RE.sub("", prvi).strip()
    # Ako je od prezimena ostala samo čestica („De", „Van"), ključ je izvorni
    # niz — bolje malo preširok ključ nego prazan.
    return jezgra or prvi


def _first_author(authors: str) -> str:
    return kljuc_prvog_autora(authors)


def parse_author_year(text: str) -> list[CitationRef]:
    refs: list[tuple[int, CitationRef]] = []
    cleaned = _DVOSTRUKA_TOCKA.sub(r'\1', bez_lokatora(text or ''))
    for m in AUTHOR_YEAR_PAREN_DETAIL.finditer(cleaned):
        for seg in m.group(1).split(';'):
            part = seg.strip()
            mm = re.match(r"\s*(.+?),\s*(\d{4})(\.?)([a-z]?)\s*$", part)
            if not mm:
                continue
            author, year, dot, suffix = mm.groups()
            first = _first_author(author)
            y = year + suffix
            refs.append((m.start(), CitationRef('author-year', f'{first}:{y}', m.group(0),
                                                'unknown', y, first, bool(dot))))
    for m in AUTHOR_YEAR_NARRATIVE_DETAIL.finditer(cleaned):
        authors, year, dot, suffix = m.groups()
        first = _first_author(authors)
        y = year + suffix
        refs.append((m.start(), CitationRef('author-year', f'{first}:{y}', m.group(0),
                                            'unknown', y, first, bool(dot))))
    refs.sort(key=lambda item: item[0])
    return [r for _, r in refs]


def _expand_numeric_part(part: str) -> list[int]:
    part = part.strip()
    m = re.fullmatch(r'(\d+)\s*[–—-]\s*(\d+)', part)
    if m:
        a, b = map(int, m.groups())
        if b >= a and b - a <= 100:
            return list(range(a, b + 1))
        return [a, b]
    return [int(part)] if part.isdigit() else []


def parse_ieee(text: str) -> list[CitationRef]:
    text = text or ''
    refs: list[tuple[int, CitationRef]] = []
    blocked: list[tuple[int, int]] = []
    for m in IEEE_RANGE_SEPARATE.finditer(text):
        a, b = map(int, m.groups())
        nums = range(a, b + 1) if b >= a and b - a <= 100 else (a, b)
        for n in nums:
            refs.append((m.start(), CitationRef('ieee', str(n), m.group(0))))
        blocked.append(m.span())
    for m in IEEE_GROUP.finditer(text):
        if any(m.start() >= a and m.end() <= b for a, b in blocked):
            continue
        content = m.group(1)
        nums: list[int] = []
        for part in re.split(r'\s*[,;]\s*', content):
            nums.extend(_expand_numeric_part(part))
        for n in nums:
            refs.append((m.start(), CitationRef('ieee', str(n), m.group(0))))
    refs.sort(key=lambda item: (item[0], int(item[1].key)))
    return [r for _, r in refs]


COURT_RE = re.compile(
    r"\b(?:U-(?:I|II|III|IV)-\d+/\d{2,4}|(?:Rev|Gž|Pž|Usž)-?\d+/\d{2,4}|[CT]-\d+/\d{2,4})\b",
    re.I,
)
EU_RE = re.compile(
    r"\b(?:Uredba|Direktiva|Odluka|Regulation|Directive)\s*\(?(?:EU|EZ|EEZ)\)?"
    r"(?:\s*(?:br\.)?)?\s*\d{2,4}/\d+",
    re.I,
)
LAW_RE = re.compile(r"\bZakon\b|\bNarodne\s+novine\b|\bNN\s*(?:br\.?\s*)?\d+/\d+", re.I)
REGULATION_RE = re.compile(r"\b(?:Pravilnik|Uredba|Poslovnik)\b", re.I)


def classify_source_type(text: str) -> str:
    """Vrsta izvora iz teksta reference; uvijek član `SOURCE_TYPES`.

    `SOURCE_TYPES` je bio deklariran i nigdje korišten, pa je ključ pravne
    reference (`f'{source_type}:{ident}'`) mogao dobiti vrijednost izvan
    objavljenog skupa a da to nigdje ne pukne. Skup je sada ugovor, provjeren na
    jedinom mjestu koje ga proizvodi — to je jeftinije nego obrisati ga i s njim
    izgubiti namjeru.
    """
    t = text or ''
    if COURT_RE.search(t):
        vrsta = 'court_decision'
    elif EU_RE.search(t):
        vrsta = 'eu_act'
    elif LAW_RE.search(t):
        vrsta = 'law'
    elif REGULATION_RE.search(t):
        vrsta = 'regulation'
    else:
        vrsta = 'unknown'
    if vrsta not in SOURCE_TYPES:      # pragma: no cover - obrana ugovora
        raise CitationDialectError(
            f"nepoznata vrsta izvora: {vrsta!r}; dopušteno: {', '.join(SOURCE_TYPES)}")
    return vrsta


# Q18-fix (v1.0.1): ključ pravne reference je IDENTIFIKATOR propisa (broj NN,
# broj predmeta, broj akta), a ne cijeli tekst odlomka. Prije je preformulacija
# „i" u „te se" u bilo kojoj rečenici koja spominje zakon mijenjala ključ, pa je
# verify_rewrite javljao „citati odstupaju" upravo u modu čija je svrha
# preformulacija.
NN_BROJ_RE = re.compile(
    r"(?:Narodne\s+novine|\bNN\b)\s*,?\s*(?:br\.?\s*)?(\d+/\d+)", re.I)
AKT_BROJ_RE = re.compile(r"(\d{2,4}/\d+)")
# Kad broja nema, ključ je naziv akta („Zakon o javnoj nabavi") — najviše tri
# riječi iza „o", što je stabilno na preformulaciju ostatka rečenice.
_AKT_VRSTA = (r"\b(?:Zakon|Pravilnik|Uredba|Poslovnik|Odluka|Direktiva|Regulation|Directive)\b")
_AKT_PREDMET = r"(?:\s+o\b(?:\s+[^\W\d_]+){1,3})"
AKT_NAZIV_RE = re.compile(_AKT_VRSTA + _AKT_PREDMET + r"?", re.I | re.UNICODE)
# Q18-fix, treći krug: gola vrsta akta („zakon", „uredba") NIJE identifikator.
# `parse_legal_text(strict=True)` je dosad na svaku rečenicu tijela koja usput
# spominje riječ „zakon" emitirala referencu s ključem „law:zakon", pa je otisak
# nosio onoliko takvih referenci koliko odlomaka spominje propis. Brisanje ili
# preformulacija te riječi u jednom odlomku mijenja broj, a verify_rewrite to
# javlja kao „citati odstupaju" — točno onaj lažni blok u modu 3 koji je Q18
# trebao zatvoriti, samo sužen na golu vrstu akta. Vlastito ime propisa ima
# predmet („Zakon o javnoj nabavi"); bez njega nema što citirati.
AKT_NAZIV_STROGI_RE = re.compile(_AKT_VRSTA + _AKT_PREDMET, re.I | re.UNICODE)


def legal_identifier_strict(text: str) -> str | None:
    """Konkretan identifikator pravnog izvora, ili None ako ga nema.

    Q18-fix, drugi krug: `legal_identifier()` na kraju pada na `_norm(cijeli
    tekst)`, što je upravo izvor Q18 defekta — svaka preformulacija odlomka
    mijenja ključ. Za TIJELO teksta taj fallback se ne smije koristiti: ondje
    referenca ulazi u otisak samo ako ima nešto stabilno (broj predmeta, broj
    EU akta, NN broj ili naziv propisa). Fusnota je citat po definiciji pa
    ondje fallback ostaje.
    """
    t = text or ''
    m = COURT_RE.search(t)
    if m:
        return _norm(m.group(0))
    m = EU_RE.search(t)
    if m:
        broj = AKT_BROJ_RE.search(m.group(0))
        return _norm(broj.group(1) if broj else m.group(0))
    m = NN_BROJ_RE.search(t)
    if m:
        return _norm(m.group(1))
    m = AKT_NAZIV_STROGI_RE.search(t)
    if m:
        return _norm(m.group(0))
    return None


def legal_identifier(text: str) -> str:
    """Identifikator pravnog izvora: broj predmeta, broj akta ili naziv akta.

    Fusnotni put — ondje je i gola vrsta akta bolji ključ od cijelog teksta,
    pa se stroža provjera dopunjuje starim, širim uzorkom naziva.
    """
    konkretan = legal_identifier_strict(text)
    if konkretan is not None:
        return konkretan
    m = AKT_NAZIV_RE.search(text or '')
    if m:
        return _norm(m.group(0))
    return _norm(text or '')


def parse_legal_text(text: str, strict: bool = False) -> list[CitationRef]:
    """Pravne reference iz teksta.

    `strict=True` (tijelo teksta) emitira referencu samo ako postoji konkretan
    identifikator — spominjanje propisa bez ijednog broja ili naziva nije citat
    i ne smije ući u otisak, jer bi ga svaka preformulacija promijenila.
    """
    source_type = classify_source_type(text)
    if source_type == 'unknown':
        return []
    ident = legal_identifier_strict(text) if strict else legal_identifier(text)
    if ident is None:
        return []
    key = f'{source_type}:{ident}'
    return [CitationRef('legal-footnote', key, text.strip(), source_type)]


def parse_citations(text: str, style: str = 'autor-godina') -> list[CitationRef]:
    """Citatne reference iz TIJELA teksta.

    Q18, treći krug: pravni put je ovdje `strict`, jer su svi pozivatelji ove
    funkcije tijelo rada, a ne fusnote — `verify_rewrite` (.md/.txt rukopisi i
    ćelije tablica u .docx-u), `check_argument.poglavlja`, `check_rules` i
    `hr_text`. Prije je gola vrsta akta („…Ovršni zakon propisuje…") i ovdje
    emitirala referencu s ključem `law:zakon`, pa je preformulacija rečenice
    („Zakon je donesen…" → „Propis je donesen…") mijenjala otisak i
    `verify_rewrite.usporedi()` je vraćao „citati odstupaju (1 → 0)" — lažan
    blok u modu čija je svrha upravo preformulacija. Popravak prvog kruga
    zatvorio je samo `parse_legal_text(strict=True)`, do kojega se dolazi jedino
    iz `citation_fingerprint_file` za .docx, pa su test i kvar gledali dva
    različita puta.
    Fusnotni put NAMJERNO ostaje širi (`legal_identifier` u
    `footnote_citations_for_paragraph`): fusnota je citat po definiciji, ondje
    je i gola vrsta akta bolji ključ od ničega.
    """
    dialect = resolve_dialect(style)
    if dialect == 'author-year':
        return parse_author_year(text)
    if dialect == 'ieee':
        return parse_ieee(text)
    if dialect == 'vancouver':
        return parse_vancouver(text)
    return parse_legal_text(text, strict=True)


def extract_docx_footnotes(path: str | Path) -> dict[str, str]:
    p = Path(path)
    try:
        with zipfile.ZipFile(p) as z:
            if 'word/footnotes.xml' not in z.namelist():
                return {}
            root = ET.fromstring(z.read('word/footnotes.xml'))
    except (OSError, zipfile.BadZipFile, ET.ParseError):
        return {}
    result: dict[str, str] = {}
    for footnote in root.findall(f'.//{{{_W}}}footnote'):
        fid = footnote.attrib.get(f'{{{_W}}}id')
        if fid is None or fid.startswith('-'):
            continue
        text = ''.join(node.text or '' for node in footnote.iter(f'{{{_W}}}t')).strip()
        if text:
            result[fid] = re.sub(r'\s+', ' ', text)
    return result


def paragraph_footnote_ids(paragraph) -> list[str]:
    ids: list[str] = []
    p = getattr(paragraph, '_p', None)
    if p is None:
        return ids
    for node in p.iter(f'{{{_W}}}footnoteReference'):
        fid = node.get(f'{{{_W}}}id')
        if fid is not None and not fid.startswith('-'):
            ids.append(fid)
    return ids


def footnote_citations_for_paragraph(paragraph, footnotes: dict[str, str]) -> list[CitationRef]:
    refs: list[CitationRef] = []
    for fid in paragraph_footnote_ids(paragraph):
        text = footnotes.get(fid, '')
        source_type = classify_source_type(text)
        key_text = legal_identifier(text) if text else f'id:{fid}'
        refs.append(CitationRef('legal-footnote', f'{source_type}:{key_text}', text or f'footnote {fid}', source_type))
    return refs


def citation_fingerprint_text(text: str, style: str) -> Counter:
    return Counter(r.key for r in parse_citations(text or '', style))


def _paragraphs_in_tables(tables) -> list:
    """Odlomci iz ćelija tablica, uključujući ugniježđene tablice.

    D2-fix (v1.0.1): obrisan citat u ćeliji tablice prije je bio nevidljiv, a
    verify_rewrite je svejedno potvrđivao „sadržaj očuvan". Ćelija je tijelo
    rada. Traversal je isti kao u diff_versions.odlomci_svi.
    """
    out: list = []
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                out.extend(cell.paragraphs)
                out.extend(_paragraphs_in_tables(cell.tables))
    return out


def all_paragraphs(doc) -> list:
    """Svi odlomci dokumenta: tijelo + ćelije tablica."""
    out = list(doc.paragraphs)
    try:
        out.extend(_paragraphs_in_tables(doc.tables))
    except Exception:
        pass
    return out


def citation_fingerprint_file(path: str | Path, style: str) -> Counter:
    p = Path(path)
    dialect = resolve_dialect(style)
    if p.suffix.lower() != '.docx':
        return citation_fingerprint_text(p.read_text(encoding='utf-8'), style)
    try:
        from docx import Document
        doc = Document(p)
    except Exception:
        return Counter()
    paragraphs = all_paragraphs(doc)
    if dialect != 'legal-footnote':
        # D2-fix, ostatak: fusnote su se čitale samo u legal-footnote dijalektu,
        # pa je u apa-hr/autor-godina radu citat koji živi ISKLJUČIVO u fusnoti
        # bio nevidljiv otisku — obrisan ili zamijenjen, otisak je ostajao
        # identičan. Fusnota je citat u svakom dijalektu, ne samo u pravnom.
        dijelovi = [(x.text or '') for x in paragraphs]
        dijelovi.extend(extract_docx_footnotes(p).values())
        return citation_fingerprint_text(' '.join(dijelovi), style)
    # Q18-fix (v1.0.1), drugi krug nakon nezavisne revizije: prva verzija je
    # pravne reference uzimala ISKLJUČIVO iz fusnota, čime je zamjena propisa u
    # tijelu teksta („Zakon o javnoj nabavi (NN 120/16)" → „Zakon o zaštiti
    # okoliša (NN 80/13)") postala nevidljiva — otisak {} vs {}, pa
    # verify_rewrite javlja „0 referenci, identično" na potpuno zamijenjen
    # propis. To je bilo šire od onoga što je Q18 tražio: Q18 je tražio STABILAN
    # ključ, a ne uklanjanje detekcije. Tijelo se sada opet čita, ali kroz
    # `strict` identifikator, pa preformulacija odlomka ne mijenja ključ dok god
    # broj/naziv propisa ostaje isti.
    footnotes = extract_docx_footnotes(p)
    refs: list[CitationRef] = []
    for paragraph in paragraphs:
        refs.extend(footnote_citations_for_paragraph(paragraph, footnotes))
        refs.extend(parse_legal_text(paragraph.text or '', strict=True))
    return Counter(r.key for r in refs)
