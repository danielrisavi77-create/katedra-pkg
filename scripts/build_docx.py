#!/usr/bin/env python3
"""Generiraj .docx kostur rada usklađen s profilom fakulteta.

Do audita je Katedra rad ISKLJUČIVO ocjenjivala: nijedna od 43 skripte nije
imala nijedan `Document()` ni `.save()`. Student je pisao prozu u razgovoru,
dokument slagao ručno u Wordu, a Katedra mu je poslije govorila što nije u redu.
Šest od četrnaest blokirajućih stavki iz `references/predaja.md` §2 su čisto
mehaničke — naslovnica, sadržaj kao POLJE, rimska pa arapska numeracija,
prijelom pred poglavljem, natpis+prikaz+„Izvor:" koji se ne smiju razdvojiti —
i sve ih je jeftinije proizvesti ispravno nego naknadno prijavljivati.

Sve vrijednosti dolaze iz razriješenog profila fakulteta (`profile_rules`), ne
iz koda: font, veličina, prored, poravnanje, margine, veličina naslova
poglavlja, prijelom pred poglavljem, obavezni dijelovi i shema numeracije. Novi
fakultet time dobiva generiranje besplatno, čim mu profil prođe readiness gate.

Uporaba:
  python3 <KATEDRA_SKILL>/scripts/build_docx.py --fakultet efzg --tip zavrsni \\
      --tema "Utjecaj pandemije na turoperatore" --autor "Ime Prezime" \\
      --mentor "prof. dr. sc. Ime Prezime" --out ./rad.docx
  python3 <KATEDRA_SKILL>/scripts/build_docx.py --profil ./.katedra/resolved_profile.json \\
      --plan ./.katedra/plan.json --out ./rad.docx
  python3 <KATEDRA_SKILL>/scripts/build_docx.py --fakultet efzg --tip zavrsni \\
      --out ./rad.docx --provjeri

`--provjeri` nakon generiranja pokreće `check_rules.py` nad vlastitim izlazom i
pada ako dokument ne prolazi profil po kojem je napravljen. To je jedini oblik
jamstva koji nešto vrijedi: generator i provjera dijele isti izvor istine.

Izlazni kodovi:
  0  dokument napisan (i, uz --provjeri, prošao provjeru)
  1  dokument napisan, ali --provjeri je našao kršenje
  2  ulaz se ne može pročitati (profil, plan, nepoznat tip rada)
"""
from __future__ import annotations

import argparse
import json
import os
import re
import subprocess
import sys

HERE = os.path.dirname(os.path.abspath(__file__))
if HERE not in sys.path:
    sys.path.insert(0, HERE)

import hr_text as H  # noqa: E402
from profile_rules import ProfileRuleError, resolve_profile, resolve_work_type  # noqa: E402

FAKULTETI = os.path.join(HERE, "..", "references", "fakulteti")

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


class GreskaUlaza(Exception):
    """Ulaz se ne može pročitati — izlazni kod 2."""


# ----------------------------------------------------------------- XML pomoć

def _el(tag: str, **atributi):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    e = OxmlElement(tag)
    for k, v in atributi.items():
        e.set(qn(f"w:{k}"), str(v))
    return e


def _polje(odlomak, kod: str, rezerva: str = "") -> None:
    """Ubaci pravo Wordovo POLJE (fldChar/instrText) u odlomak.

    Sadržaj koji je samo natipkan kao tekst nije polje: ne osvježava se, a
    `references/predaja.md` izričito traži „sadržaj je TOC POLJE, ne ručno
    tipkan popis". Isto vrijedi za broj stranice.
    """
    from docx.oxml.ns import qn

    r1 = odlomak.add_run()._r
    r1.append(_el("w:fldChar", fldCharType="begin"))

    r2 = odlomak.add_run()._r
    instr = _el("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = kod
    r2.append(instr)

    r3 = odlomak.add_run()._r
    r3.append(_el("w:fldChar", fldCharType="separate"))

    if rezerva:
        odlomak.add_run(rezerva)

    r4 = odlomak.add_run()._r
    r4.append(_el("w:fldChar", fldCharType="end"))


def _numeracija_stranica(sekcija, format_: str, pocetak: int | None = None) -> None:
    """Postavi format numeracije stranica na sekciju (upperRoman / decimal)."""
    sectPr = sekcija._sectPr
    from docx.oxml.ns import qn

    for stari in sectPr.findall(qn("w:pgNumType")):
        sectPr.remove(stari)
    atributi = {"fmt": format_}
    if pocetak is not None:
        atributi["start"] = pocetak
    sectPr.append(_el("w:pgNumType", **atributi))


def _postavi_font_svugdje(run_ili_stil, ime: str) -> None:
    """Postavi font za ascii, hAnsi I cs.

    Audit nalaz: `check_rules.font_iz_rpr` je čitao samo `w:ascii`, a upravo
    `w:hAnsi` je slot koji u Wordu pokriva hrvatske dijakritike. Generator zato
    postavlja sva tri, inače bi vlastiti dokument mogao ispasti neusklađen na
    slovima č/ć/ž/š/đ.
    """
    from docx.oxml.ns import qn

    rPr = run_ili_stil.element.rPr if hasattr(run_ili_stil, "element") else None
    if rPr is None:
        return
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = _el("w:rFonts")
        rPr.append(rFonts)
    for slot in ("ascii", "hAnsi", "cs"):
        rFonts.set(qn(f"w:{slot}"), ime)


def _drzi_uz_sljedeci(odlomak, cijepanje: bool = False) -> None:
    """keepNext (+ opcionalno keepLines) — natpis, prikaz i „Izvor:" ostaju skupa."""
    pf = odlomak.paragraph_format
    pf.keep_with_next = True
    if cijepanje:
        pf.keep_together = True


def _tablica_se_ne_lomi(tablica) -> None:
    """cantSplit na svakom retku — prikaz se ne smije lomiti preko stranica."""
    from docx.oxml.ns import qn

    for red in tablica.rows:
        trPr = red._tr.get_or_add_trPr()
        if trPr.find(qn("w:cantSplit")) is None:
            trPr.append(_el("w:cantSplit"))




def _je_popis_literature(naslov: str) -> bool:
    n = str(naslov or "").strip().lower().lstrip("0123456789. ")
    return n.startswith("popis") and any(k in n for k in ("literatur", "izvor", "bibliograf", "referenc"))


def _renderiraj(d, blokovi, _poglavlje, prikazi_pravila, font, velicina,
                meta=None, uvlaka_popisa=False):
    """Blokovi rukopisa → odlomci dokumenta.

    Naslov prve razine ide kroz `_poglavlje`, pa dobiva prijelom i stil koje
    profil traži — inače bi rukopis zaobišao upravo ona pravila zbog kojih
    generator postoji. Natpis, prikaz i „Izvor:" drže se zajedno (`keepNext`,
    `cantSplit`): to je sklop koji `check_rules.provjeri_prikaze` provjerava.

    `uvlaka_popisa`: odlomci iza naslova „POPIS LITERATURE" dobivaju viseću
    uvlaku (profil `citiranje.uvlaka_u_popisu`) — inače provjeri_literaturu
    javlja 0/N jedinica s uvlakom, a rukopis to ne može izraziti.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    meta = meta or {}
    u_popisu = False

    def _dodaj(odlomak, dijelovi):
        for tekst, podebljano, kurziv in dijelovi:
            if not tekst:
                continue
            r = odlomak.add_run(tekst)
            r.bold = bool(podebljano)
            r.italic = bool(kurziv)
        return odlomak

    for blok in blokovi:
        vrsta = blok[0]
        if vrsta == "naslov":
            razina, tekst = blok[1], blok[2]
            if razina == 1:
                _poglavlje(tekst)
                u_popisu = bool(uvlaka_popisa) and _je_popis_literature(tekst)
            else:
                d.add_heading(tekst, level=min(razina, 4))
        elif vrsta == "odlomak":
            p = _dodaj(d.add_paragraph(), blok[1])
            if u_popisu:
                p.paragraph_format.left_indent = Cm(1.25)
                p.paragraph_format.first_line_indent = Cm(-1.25)
                p.paragraph_format.space_after = Pt(6)
        elif vrsta == "citat":
            # Stil „Quote" je STRUKTURNI signal da ovo nije prozni odlomak; sama
            # uvlaka nije, pa bi blok-citat inače padao na pravilu „najmanje dvije
            # rečenice po odlomku". Uvlaka ostaje kao izgled.
            try:
                p = _dodaj(d.add_paragraph(style="Quote"), blok[1])
            except KeyError:
                p = _dodaj(d.add_paragraph(), blok[1])
            p.paragraph_format.left_indent = Cm(1.25)
        elif vrsta == "popis":
            stil = "List Number" if blok[2] else "List Bullet"
            for stavka in blok[1]:
                try:
                    _dodaj(d.add_paragraph(style=stil), stavka)
                except KeyError:
                    _dodaj(d.add_paragraph(), stavka)
        elif vrsta == "prijelom":
            d.add_page_break()
        elif vrsta == "sekcija":
            # Sekcije su vlasništvo generatora (prednji dio vs. tijelo, numeracija).
            # Marker iz rukopisa se namjerno preskače umjesto da se poštuje: dvije
            # strane koje obje dijele sekcije daju numeraciju koju nitko ne kontrolira.
            continue
        elif vrsta == "tablica":
            redci, natpis, izvor = blok[1], blok[2], blok[3]
            if not redci:
                continue
            if natpis:
                p = d.add_paragraph(natpis)
                _drzi_uz_sljedeci(p)
            t = d.add_table(rows=len(redci), cols=max(len(r) for r in redci))
            try:
                t.style = "Table Grid"
            except KeyError:
                pass
            for i, red in enumerate(redci):
                for j, celija in enumerate(red):
                    t.cell(i, j).text = celija
            oboji_tablicu(t, razrijesi_paletu(meta.get("tablice_boja")))
            _tablica_se_ne_lomi(t)
            if izvor:
                p = d.add_paragraph(izvor)
                p.runs[0].font.size = Pt(max(8, velicina - 2))


# --------------------------------------------------------------------------- #
# Sjenčanje tablica — polje `tablice_boja` iz intakea (§ 0.4)
#
# Do v1.4 je to polje bilo MRTVO: intake ga je pitao, `stanje_init.py` ga je
# zapisivao, i nijedna skripta ga nikad nije pročitala. Generator je hardkodirao
# „Table Grid". Pitati korisnika za nešto što nigdje ne djeluje gore je nego ne
# pitati — odgovor stvara očekivanje koje dokument ne ispuni.
#
# Tonovi su namjerno blijedi: tablica u akademskom radu nosi podatke, ne dizajn,
# a jaka boja u crno-bijelom ispisu postaje siva mrlja koja guta tekst.
PALETE_TABLICA = {
    "bez boje": (None, None),
    "sivo": ("D9D9D9", "F2F2F2"),
    "rozo-sivo": ("EDE3E3", "F7F2F2"),
    "plavo": ("DCE6F1", "F2F6FB"),
    "zeleno": ("E2EFDA", "F3F8F0"),
}


def razrijesi_paletu(vrijednost):
    """(boja_zaglavlja, boja_zebre) iz slobodnog teksta koji je korisnik zadao.

    Zebra (naizmjenično sjenčani redci) uključuje se samo ako je izrijekom
    zatražena — na tablici od tri retka ona je šum, a `check_paragraphs` ionako
    mjeri čitljivost teksta, ne tablice.
    """
    v = H.bez_dijakritika(str(vrijednost or "")).lower().strip()
    if not v:
        return (None, None)
    zebra = "zebra" in v
    v = v.replace("zebra", "").replace("+", " ").strip(" -,") or "bez boje"
    # Točno podudaranje prije djelomičnog, a djelomično od NAJDULJEG ključa:
    # „sivo" je podniz od „rozo-sivo", pa bi kraći ključ inače pojeo dulji i
    # korisnik koji je tražio rozo-sivo dobio bi sivo, bez ijedne poruke.
    kljucevi = sorted(PALETE_TABLICA, key=len, reverse=True)
    for kljuc in kljucevi:
        if H.bez_dijakritika(kljuc).lower() == v:
            zaglavlje, tijelo = PALETE_TABLICA[kljuc]
            return (zaglavlje, tijelo if zebra else None)
    for kljuc in kljucevi:
        if H.bez_dijakritika(kljuc).lower() in v:
            zaglavlje, tijelo = PALETE_TABLICA[kljuc]
            return (zaglavlje, tijelo if zebra else None)
    # Nepoznat naziv boje NIJE tiha greška: vraća se siva i zapisuje se poruka,
    # jer bi tiho ignoriranje opet dalo dokument koji ne odgovara odgovoru.
    return ("D9D9D9", "F2F2F2" if zebra else None)


def _sjencaj(celija, boja_hex):
    if not boja_hex:
        return
    from docx.oxml.ns import qn
    tcPr = celija._tc.get_or_add_tcPr()
    shd = _el("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), boja_hex)
    tcPr.append(shd)


def oboji_tablicu(t, paleta, ima_zaglavlje=True):
    zaglavlje, zebra = paleta
    for i, red in enumerate(t.rows):
        if i == 0 and ima_zaglavlje:
            for c in red.cells:
                _sjencaj(c, zaglavlje)
        elif zebra and i % 2 == 0:
            for c in red.cells:
                _sjencaj(c, zebra)


def ocekuje_prikaz(metodologija):
    """Očekuje li rad ove metodologije autorov prikaz (tablicu/sliku/grafikon)?

    Oslanja se na B09 policy iz `argument_methodology`, koji je već izvor istine
    za isto pitanje u `check_argument`. Nepoznata/nenavedena metodologija →
    True, jer je neutralni `generic` slučaj i prazan primjer je lakše obrisati
    nego se sjetiti da ga treba dodati.
    """
    if not metodologija:
        return True
    try:
        import argument_methodology as AM
        policy = AM.POLICIES.get(str(metodologija).strip().lower())
    except Exception:
        return True
    if policy is None:
        return True
    return bool(getattr(policy, "requires_analytical_section", True))


# ---------------------------------------------------- numeracija iz profila

WORD_NUMERACIJA = {
    "rimski": "upperRoman",
    "rimski_mala": "lowerRoman",
    "arapski": "decimal",
    "bez": None,
}


def shema_numeracije(profil):
    """Vrati (prednji_format, tijelo_format, tijelo_pocinje_od) ili None.

    Do ovog popravka generator je BEZUVJETNO radio rimsko-pa-arapsku podjelu.
    To je EFZG pravilo, ali se primjenjivalo i na FPZG, čiji profil to pravilo
    UOPĆE NEMA — konvencija jednog fakulteta nametala se drugome bez osnove.
    Ako profil ne propisuje shemu, vraća se None i dokument ostaje na
    jedinstvenoj arapskoj numeraciji: bolje ne raditi ništa nego izmisliti
    pravilo.

    `format.numeracija` je od ovog popravka STRUKTURIRAN objekt
    (`prednji_dio`/`tijelo`/`tijelo_pocinje_od`). Slobodan tekst ostaje podržan
    kao naslijeđeni oblik jer ga profili trećih fakulteta još mogu imati, ali se
    na njega ne oslanjamo: obrazac je tražio doslovnu riječ „rimski", pa je
    drukčije formuliran ali jednako obvezujući propis („paginacija prednjeg
    dijela velikim latinskim brojkama") tiho davao None. Kad je oblik naslijeđen
    i neprepoznat, to je nalaz koji se vidi (⚠️ pri gradnji), ne tišina.
    """
    sirovo = (profil.get("format") or {}).get("numeracija")
    if isinstance(sirovo, dict):
        prednji = WORD_NUMERACIJA.get(str(sirovo.get("prednji_dio") or ""), "?")
        tijelo = WORD_NUMERACIJA.get(str(sirovo.get("tijelo") or ""), "?")
        if prednji == "?" or tijelo == "?":
            return None
        try:
            od = int(sirovo.get("tijelo_pocinje_od") or 1)
        except (TypeError, ValueError):
            od = 1
        if prednji is None and tijelo is None:
            return None
        return (prednji, tijelo, max(1, od))

    opis = str(sirovo or "").strip()
    if not opis:
        return None
    n = H.bez_dijakritika(opis.lower())
    if "rimski" not in n:
        return None
    m = re.search(r"(?:po[cč]ev[sš]i\s+od|od)\s+(\d+)", n)
    return ("upperRoman", "decimal", int(m.group(1)) if m else 1)


def upozorenje_numeracije(profil):
    """Poruka kad profil PROPISUJE numeraciju koju generator nije razumio.

    Tišina je ovdje najgori ishod: rad bez propisane paginacije izgleda uredno
    dok ga mentor ne otvori. Ako je pravilo zapisano, a shema ispala None, to se
    mora vidjeti pri gradnji.
    """
    sirovo = (profil.get("format") or {}).get("numeracija")
    if not sirovo or shema_numeracije(profil):
        return None
    if isinstance(sirovo, dict):
        return ("⚠️  format.numeracija ima nepoznate vrijednosti "
                f"({sirovo.get('prednji_dio')!r} → {sirovo.get('tijelo')!r}); "
                "dopušteno je rimski / rimski_mala / arapski / bez. "
                "Dokument je ostao bez podjele numeracije.")
    return ("⚠️  format.numeracija je zapisan slobodnim tekstom koji generator nije "
            f"prepoznao: {str(sirovo)[:80]!r}. Što napraviti: zapiši ga strukturirano "
            "(\"prednji_dio\": \"rimski\", \"tijelo\": \"arapski\", \"tijelo_pocinje_od\": 1). "
            "Dokument je zasad ostao bez podjele numeracije.")


# ------------------------------------------------------------------ profil

def ucitaj_profil(args) -> dict:
    if args.profil:
        if not os.path.isfile(args.profil):
            raise GreskaUlaza(f"nema profila: {args.profil}")
        try:
            with open(args.profil, encoding="utf-8") as f:
                return json.load(f)
        except (OSError, json.JSONDecodeError) as exc:
            raise GreskaUlaza(f"{args.profil} nije čitljiv JSON: {exc}") from exc
    if not args.fakultet:
        raise GreskaUlaza(
            "navedi --fakultet ili --profil.\n"
            "   Što napraviti: npr. --fakultet efzg --tip zavrsni"
        )
    try:
        return resolve_profile(
            args.fakultet, faculty_dir=FAKULTETI, work_type=args.tip
        ).profile
    except ProfileRuleError as exc:
        raise GreskaUlaza(str(exc)) from exc


def ucitaj_plan(put: str | None) -> list[dict]:
    if not put:
        return []
    if not os.path.isfile(put):
        raise GreskaUlaza(f"nema plana: {put}")
    try:
        with open(put, encoding="utf-8") as f:
            plan = json.load(f)
    except (OSError, json.JSONDecodeError) as exc:
        raise GreskaUlaza(f"{put} nije čitljiv JSON: {exc}") from exc
    return [p for p in (plan.get("poglavlja") or []) if isinstance(p, dict)]


# --------------------------------------------------------------- generiranje

_FRONT_MATTER_PRIJE_SADRZAJA = (
    "naslovnica", "izjava", "sažetak", "sazetak", "summary", "abstract",
)


def _je_sadrzaj(naziv: str) -> bool:
    n = naziv.lower()
    return n.startswith("sadržaj") or n.startswith("sadrzaj") or n.startswith("kazalo")


def _je_uvod(naziv: str) -> bool:
    return naziv.lower().lstrip("0123456789. ").startswith("uvod")


def _je_razrada(naziv: str) -> bool:
    return naziv.lower().startswith("razrada")


def gradi(profil: dict, meta: dict, poglavlja: list[dict], izlaz: str,
          rukopis: list[dict] | None = None) -> dict:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    fmt = profil.get("format") or {}
    struktura = profil.get("struktura") or {}
    prikazi_pravila = struktura.get("prikazi") or {}

    fontovi = fmt.get("font") or ["Times New Roman"]
    font = fontovi[0] if isinstance(fontovi, list) else str(fontovi)
    velicina = float(fmt.get("velicina_pt") or 12)
    prored = float(fmt.get("prored") or 1.5)
    naslov_pt = float(fmt.get("naslov_poglavlja_pt") or velicina + 2)
    prijelom = bool(fmt.get("prijelom_pred_poglavljem"))
    poravnanje = (fmt.get("poravnanje") or "obostrano").lower()
    margine = fmt.get("margine_cm") or {}

    d = Document()

    # --- osnovni stil ---
    normal = d.styles["Normal"]
    normal.font.name = font
    normal.font.size = Pt(velicina)
    _postavi_font_svugdje(normal, font)
    pf = normal.paragraph_format
    pf.line_spacing = prored
    pf.space_after = Pt(0)
    if poravnanje.startswith("obostran"):
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for razina, uvecanje in ((1, 0), (2, -1), (3, -2)):
        try:
            st = d.styles[f"Heading {razina}"]
        except KeyError:
            continue
        st.font.name = font
        st.font.size = Pt(max(velicina, naslov_pt + uvecanje))
        st.font.bold = True
        st.font.color.rgb = None
        _postavi_font_svugdje(st, font)
        st.paragraph_format.line_spacing = prored
        st.paragraph_format.keep_with_next = True
        if razina == 1:
            st.paragraph_format.page_break_before = prijelom

    # --- margine na svim sekcijama ---
    def _margine(sek):
        if margine.get("gore") is not None:
            sek.top_margin = Cm(float(margine["gore"]))
        if margine.get("dolje") is not None:
            sek.bottom_margin = Cm(float(margine["dolje"]))
        if margine.get("lijevo") is not None:
            sek.left_margin = Cm(float(margine["lijevo"]))
        if margine.get("desno") is not None:
            sek.right_margin = Cm(float(margine["desno"]))

    prva = d.sections[0]
    _margine(prva)

    obavezni = list(struktura.get("obavezni_dijelovi") or [])
    napravljeno: list[str] = []

    # ---------------------------------------------------------- front matter
    for dio in obavezni:
        if _je_sadrzaj(dio) or _je_uvod(dio) or _je_razrada(dio):
            continue
        n = dio.lower()
        if not any(k in n for k in _FRONT_MATTER_PRIJE_SADRZAJA):
            continue
        if "naslovnica" in n:
            engleska = "englesk" in n
            p = d.add_paragraph(meta["fakultet_pun"] if not engleska
                                else meta["fakultet_pun_en"])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for _ in range(6):
                d.add_paragraph()
            t = d.add_paragraph()
            t.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = t.add_run(meta["tema"] if not engleska else meta["tema_en"])
            run.bold = True
            run.font.size = Pt(naslov_pt + 2)
            pod = d.add_paragraph(
                ("ZAVRŠNI RAD" if meta["tip"] == "zavrsni" else
                 "DIPLOMSKI RAD" if meta["tip"] == "diplomski" else
                 "SEMINARSKI RAD") if not engleska else "FINAL THESIS")
            pod.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for _ in range(6):
                d.add_paragraph()
            for redak in ((f"Student: {meta['autor']}", f"Mentor: {meta['mentor']}")
                          if not engleska else
                          (f"Student: {meta['autor']}", f"Supervisor: {meta['mentor']}")):
                r = d.add_paragraph(redak)
                r.alignment = WD_ALIGN_PARAGRAPH.CENTER
            g = d.add_paragraph(meta["mjesto_godina"])
            g.alignment = WD_ALIGN_PARAGRAPH.CENTER
            d.add_page_break()
            napravljeno.append(dio)
            continue

        d.add_heading(dio.upper(), level=1)
        if "izjava" in n:
            d.add_paragraph(
                f"Izjavljujem da sam rad pod naslovom „{meta['tema']}” izradio/la "
                f"samostalno, koristeći se navedenim izvorima, te da rad ne sadrži "
                f"dijelove tuđih radova bez navođenja izvora."
            )
            d.add_paragraph()
            potpis = d.add_paragraph(f"{meta['mjesto_godina']}\t\t\t{meta['autor']}")
            potpis.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif ("summary" not in n and "abstract" not in n
              and meta.get("sazetak_blokovi")):
            # Gotov sažetak iz `.katedra/sazetak.md` (mod 2) — ugrađuje se
            # umjesto predloška; naslov prve razine iz datoteke se preskače jer
            # ga generator upravo ispisao.
            _renderiraj(d, [b for b in meta["sazetak_blokovi"]
                            if not (b[0] == "naslov" and b[1] == 1)],
                        lambda t: d.add_heading(t.upper(), level=1),
                        prikazi_pravila, font, velicina, meta=meta)
        else:
            d.add_paragraph(
                "[Sažetak: 150–250 riječi. Predmet, cilj, metoda, glavni nalaz, "
                "zaključak — svaka stavka jednom rečenicom.]"
                if "summary" not in n and "abstract" not in n else
                "[Abstract: 150–250 words. Subject, aim, method, main finding, conclusion.]"
            )
            d.add_paragraph()
            d.add_paragraph("Ključne riječi: [pet pojmova]" if "summary" not in n
                            and "abstract" not in n else "Keywords: [five terms]")
        napravljeno.append(dio)

    # -------------------------------------------------------------- sadržaj
    sadrzaj_dio = next((x for x in obavezni if _je_sadrzaj(x)), "Sadržaj")
    d.add_heading(sadrzaj_dio.upper(), level=1)
    toc = d.add_paragraph()
    _polje(toc, 'TOC \\o "1-3" \\h \\z \\u',
           "[Sadržaj se generira: desni klik → Update Field / Ažuriraj polje]")
    napravljeno.append(sadrzaj_dio)

    shema = shema_numeracije(profil)
    if shema and shema[0]:
        _numeracija_stranica(prva, shema[0], 1)

    # -------------------------------------- prijelom sekcije: arapski od Uvoda
    tijelo = d.add_section(WD_SECTION.NEW_PAGE)
    _margine(tijelo)
    if shema and shema[1]:
        _numeracija_stranica(tijelo, shema[1], shema[2])
    # Obje sekcije moraju eksplicitno odvezati podnožje PRIJE upisa: dok je
    # `is_linked_to_previous` True, python-docx vraća naslijeđeno podnožje i
    # upis se ne materijalizira u vlastiti footer part — prednji dio rada tada
    # ostane bez broja stranice, pa se rimska numeracija nigdje ne vidi iako je
    # `pgNumType` ispravno postavljen. Uhvaćeno tek renderiranjem u PDF.
    # `add_section()` prestrukturira tijelo dokumenta, pa referenca na prvu
    # sekciju uzeta PRIJE toga postaje stara: upis u njezino podnožje tada tiho
    # nestane pri spremanju (nikad se ne stvori `footerReference`), a rimska se
    # numeracija nigdje ne vidi iako je `pgNumType` ispravan. Zato se prva
    # sekcija ovdje dohvaća ponovno. Uhvaćeno tek renderiranjem u PDF —
    # nijedna provjera nad .docx-om to ne bi pokazala.
    prva_sada = d.sections[0]
    for sek in (prva_sada, tijelo):
        sek.footer.is_linked_to_previous = False
        f = sek.footer.paragraphs[0]
        f.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if not f.runs:
            _polje(f, "PAGE")

    # ---------------------------------------------------------------- tijelo
    broj = 0

    def _poglavlje(naslov: str) -> None:
        nonlocal broj
        # Izjave (o autorstvu, o korištenju AI alata) nisu sadržajna poglavlja:
        # generator ih u prednjem dijelu ispisuje bez broja, pa ih ni iz rukopisa
        # ne smije numerirati — inače rad dobije „7. IZJAVA …" iza literature.
        if naslov.strip().lower().startswith("izjava"):
            d.add_heading(naslov.upper(), level=1)
            return
        broj += 1
        d.add_heading(f"{broj}. {naslov.upper()}", level=1)

    uvod_dio = next((x for x in obavezni if _je_uvod(x)), "Uvod")

    # Naslovi prve razine koje rukopis već donosi — kostur ih ne smije duplirati
    # (dupli ZAKLJUČAK i POPIS LITERATURE s uputama u uglatim zagradama bili su
    # najčešći zaostali placeholder u sastavljenom radu).
    naslovi_rukopisa: list[str] = []
    for pog in (rukopis or []):
        for b in (pog.get("blokovi") or []):
            if b[0] == "naslov" and b[1] == 1:
                naslovi_rukopisa.append(str(b[2]).strip().lower())

    def _rukopis_ima(dio: str) -> bool:
        n = dio.lower()
        if n.startswith(("zaključ", "zakljuc")):
            return any(t.startswith(("zaključ", "zakljuc")) for t in naslovi_rukopisa)
        if n.startswith("popis") or n.startswith("literatura"):
            return any(_je_popis_literature(t) for t in naslovi_rukopisa)
        if "izjava" in n:
            return any(t.startswith("izjava") for t in naslovi_rukopisa)
        return False

    uvlaka_popisa = bool((profil.get("citiranje") or {}).get("uvlaka_u_popisu"))

    if rukopis:
        # Rukopis je izvor istine: poglavlja dolaze iz `.katedra/poglavlja/*.md`,
        # a ne iz praznih rezerviranih mjesta. Prazan kostur ostaje za slučaj kad
        # rukopisa još nema — student tada dobije okvir, ne izmišljen sadržaj.
        for pog in rukopis:
            blokovi = pog.get("blokovi") or []
            prvi_naslov = next((b for b in blokovi if b[0] == "naslov"), None)
            if prvi_naslov is None:
                _poglavlje(str(pog.get("naslov") or pog.get("kljuc") or "Poglavlje"))
            _renderiraj(d, blokovi, _poglavlje, prikazi_pravila, font, velicina,
                        meta=meta, uvlaka_popisa=uvlaka_popisa)
        napravljeno.append(uvod_dio)
        # Tijelo teksta JEST rukopis; uvjetni dijelovi („ako postoje") ne dobivaju
        # prazan predložak kad ih rukopis nema.
        for dio in obavezni:
            n = dio.lower()
            if n.startswith("tijelo") or "ako postoj" in n or _rukopis_ima(dio):
                napravljeno.append(dio)
    else:
        # Bez rukopisa ostaje prazan kostur — student dobiva okvir s Uvodom, ne
        # izmišljen sadržaj. (Uvod se ovdje ispisuje izrijekom; kad je pisan u
        # rukopisu, dolazi iz njega.)
        _poglavlje(uvod_dio.split(".", 1)[-1].strip() or "Uvod")
        d.add_paragraph(
            "[Uvod: predmet i motivacija, problem i teza, cilj, metoda, struktura "
            "rada. Teza mora biti tvrdnja s kojom se netko može ne složiti.]"
        )
        napravljeno.append(uvod_dio)

    if not rukopis and poglavlja:
        for pog in poglavlja:
            naslov = str(pog.get("naslov") or pog.get("naziv") or "").strip()
            if not naslov or _je_uvod(naslov):
                continue
            if naslov.lower().startswith("zaključ") or naslov.lower().startswith("zakljuc"):
                continue
            _poglavlje(naslov)
            sadrzaj_pog = str(pog.get("sadrzaj") or "").strip()
            d.add_paragraph(sadrzaj_pog or "[Sadržaj poglavlja prema planu.]")
            for pot in (pog.get("potpoglavlja") or []):
                if isinstance(pot, dict):
                    pn = str(pot.get("naslov") or pot.get("naziv") or "").strip()
                    if pn:
                        d.add_heading(pn, level=2)
                        d.add_paragraph(str(pot.get("sadrzaj") or "").strip()
                                        or "[Sadržaj potpoglavlja prema planu.]")
    elif not rukopis:
        for naziv in ("Teorijski okvir", "Metodologija", "Analiza i rasprava"):
            _poglavlje(naziv)
            d.add_paragraph("[Sadržaj poglavlja prema planu.]")

    # Primjer prikaza NIJE bezuvjetan. Profil propisuje samo OBLIK prikaza
    # (natpis iznad, „Izvor:" ispod, mora biti spomenut, ne smije se lomiti) —
    # nigdje ne stoji da rad MORA imati tablicu. Treba li je uopće, ovisi o
    # metodologiji, a to paket već modelira: `argument_methodology` (B09) za
    # teorijski, doktrinarno-pravni, povijesni i pregledni rad izričito NE traži
    # empirijsko poglavlje ni autorov prikaz, i takav rad se zbog izostanka
    # prikaza ne smije kažnjavati. Generator je do ovog popravka svakom radu
    # ubacivao tablicu, pa i doktrinarno-pravnom — čime je studenta nagovarao na
    # prikaz koji njegovoj vrsti rada ne treba.
    # Uz rukopis primjer prikaza ne ide: rukopis je izvor istine, a predložak
    # „Tablica 1. Naziv tablice" iza gotovog teksta ostaje kao zaostali placeholder.
    if prikazi_pravila and meta.get("s_prikazom") and not rukopis:
        iznad = str(prikazi_pravila.get("natpis") or "iznad").lower().startswith("iznad")
        # Poziv u tekstu je obavezan po profilu (`mora_biti_spomenut_u_tekstu`),
        # pa ga kostur mora imati — inače generator proizvodi dokument koji pada
        # na vlastitoj provjeri.
        if prikazi_pravila.get("mora_biti_spomenut_u_tekstu"):
            d.add_paragraph(
                "U Tablici 1. prikazani su podaci o kojima je riječ u ovom "
                "potpoglavlju. Njihova se struktura komentira u nastavku, uz "
                "obrazloženje odstupanja koja se u promatranom razdoblju javljaju."
            )
        natpis = d.add_paragraph("Tablica 1. Naziv tablice")
        if iznad:
            _drzi_uz_sljedeci(natpis, cijepanje=True)
        t = d.add_table(rows=2, cols=2)
        t.style = "Table Grid"
        oboji_tablicu(t, razrijesi_paletu(meta.get("tablice_boja")))
        t.cell(0, 0).text = "[Stupac]"
        t.cell(0, 1).text = "[Stupac]"
        t.cell(1, 0).text = "[Vrijednost]"
        t.cell(1, 1).text = "[Vrijednost]"
        if prikazi_pravila.get("ne_smije_se_lomiti"):
            _tablica_se_ne_lomi(t)
        if prikazi_pravila.get("izvor_ispod", True):
            izv = d.add_paragraph("Izvor: [autor, godina, str.] / izrada autora")
            izv.paragraph_format.keep_together = True
        if not iznad:
            _drzi_uz_sljedeci(natpis, cijepanje=True)

    zakljucak_dio = next((x for x in obavezni
                          if x.lower().startswith(("zaključ", "zakljuc"))), "Zaključak")
    if zakljucak_dio not in napravljeno:
        _poglavlje(zakljucak_dio)
        d.add_paragraph(
            "[Zaključak: odgovor na tezu iz uvoda, glavni nalazi, ograničenja, "
            "vlastiti doprinos. Bez novih izvora i bez novih tvrdnji.]"
        )
        napravljeno.append(zakljucak_dio)

    # ------------------------------------------------------------ back matter
    for dio in obavezni:
        n = dio.lower()
        if dio in napravljeno or _je_razrada(dio):
            continue
        if any(k in n for k in _FRONT_MATTER_PRIJE_SADRZAJA) or _je_sadrzaj(dio):
            continue
        d.add_heading(dio.upper(), level=1)
        if n.startswith("popis izvora") or n.startswith("literatura"):
            primjer = (profil.get("citiranje") or {}).get("popis_primjer")
            d.add_paragraph(str(primjer) if primjer else "[Izvori abecedno.]")
        elif n.startswith("popis"):
            p = d.add_paragraph()
            _polje(p, 'TOC \\h \\z \\c "Tablica"', "[Popis se generira iz natpisa.]")
        else:
            d.add_paragraph("[Sadržaj.]")
        napravljeno.append(dio)

    os.makedirs(os.path.dirname(os.path.abspath(izlaz)) or ".", exist_ok=True)
    d.save(izlaz)
    return {"dijelovi": napravljeno, "poglavlja": broj}


# ------------------------------------------------------------------- provjera

def provjeri(izlaz: str, args) -> int:
    naredba = [sys.executable, os.path.join(HERE, "check_rules.py"), izlaz]
    if args.profil:
        naredba += ["--profil", args.profil]
    else:
        naredba += ["--fakultet", args.fakultet]
    naredba += ["--tip", args.tip]
    proc = subprocess.run(naredba, text=True, capture_output=True)
    sys.stdout.write(proc.stdout)
    if proc.stderr.strip():
        sys.stderr.write(proc.stderr)
    return proc.returncode


def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(
        description="Generiraj .docx kostur rada usklađen s profilom fakulteta."
    )
    ap.add_argument("--fakultet", help="slug fakulteta (npr. efzg)")
    ap.add_argument("--profil", help="razriješeni profil (JSON) umjesto --fakultet")
    ap.add_argument("--tip", default="zavrsni",
                    help="vrsta rada (zavrsni, diplomski, seminarski…)")
    ap.add_argument("--tema", default="[Naslov rada]")
    ap.add_argument("--tema-en", dest="tema_en", default=None)
    ap.add_argument("--autor", default="[Ime Prezime]")
    ap.add_argument("--mentor", default="[Mentor]")
    ap.add_argument("--mjesto", default="Zagreb")
    ap.add_argument("--godina", default=None)
    ap.add_argument("--plan", help=".katedra/plan.json — poglavlja iz plana")
    ap.add_argument("--rukopis", nargs="?", const="", default=None,
                    help="mapa s poglavljima u markdownu (.katedra/poglavlja). "
                         "Markdown je izvor istine — .docx se iz njega SASTAVLJA, "
                         "pa se ručno dotjerivanje u Wordu gubi pri sljedećem "
                         "sastavljanju.")
    ap.add_argument("--project-root", dest="project_root")
    ap.add_argument("--out", required=True, help="izlazna .docx datoteka")
    ap.add_argument("--metodologija", default=None,
                    help="theoretical|quantitative|qualitative|mixed_methods|case_study|"
                         "doctrinal_legal|historical|review — određuje treba li kostur "
                         "primjer prikaza; inače se čita iz profila")
    ap.add_argument("--tablice-boja", dest="tablice_boja", default=None,
                    help="sjenčanje tablica: bez boje | sivo | rozo-sivo | plavo | zeleno "
                         "(dodaj „zebra\" za naizmjenične retke). Bez ove zastavice "
                         "vrijednost se čita iz .katedra/stanje.json (intake § 0.4).")
    ap.add_argument("--bez-prikaza", dest="bez_prikaza", action="store_true",
                    help="ne ubacuj primjer prikaza ni kad ga metodologija očekuje")
    ap.add_argument("--provjeri", action="store_true",
                    help="nakon generiranja pokreni check_rules.py nad izlazom")
    args = ap.parse_args(argv)

    try:
        profil = ucitaj_profil(args)
        poglavlja = ucitaj_plan(args.plan)
    except GreskaUlaza as exc:
        print(f"❌ {exc}", file=sys.stderr)
        return 2

    try:
        resolve_work_type(profil, args.tip)
    except ProfileRuleError as exc:
        print(f"❌ {exc}", file=sys.stderr)
        return 2
    except Exception:
        pass

    godina = args.godina or "[godina]"
    naziv = str(profil.get("naziv") or profil.get("fakultet") or args.fakultet or "").strip()
    metodologija = args.metodologija or (profil.get("metodologija") or {}).get("type")
    s_prikazom = ocekuje_prikaz(metodologija) and not args.bez_prikaza

    tablice_boja = args.tablice_boja
    if tablice_boja is None:
        # Odgovor iz intakea živi u stanju projekta; zastavica ga samo nadjačava.
        try:
            import context as _ctx
            with open(_ctx.resolve_state_file("stanje.json",
                                              project_root=args.project_root),
                      encoding="utf-8") as _f:
                tablice_boja = (json.load(_f) or {}).get("tablice_boja")
        except (OSError, json.JSONDecodeError, ValueError):
            tablice_boja = None

    meta = {
        "tablice_boja": tablice_boja,
        "s_prikazom": s_prikazom,
        "metodologija": metodologija,
        "tip": args.tip,
        "tema": args.tema,
        "tema_en": args.tema_en or "[Thesis title in English]",
        "autor": args.autor,
        "mentor": args.mentor,
        "fakultet_pun": naziv.upper() or "[FAKULTET]",
        "fakultet_pun_en": (naziv.upper() or "[FACULTY]"),
        "mjesto_godina": f"{args.mjesto}, {godina}.",
    }

    rukopis = None
    if args.rukopis is not None:
        import rukopis as R
        mapa = args.rukopis or R.mapa_rukopisa(args.project_root)
        try:
            pogl = R.poglavlja(mapa)
        except R.GreskaRukopisa as exc:
            print(f"❌ {exc}", file=sys.stderr)
            return 2
        if not pogl:
            print(f"❌ rukopis je prazan: {mapa}", file=sys.stderr)
            print("   Što napraviti: `rukopis.py init --plan .katedra/plan.json`, "
                  "pa napiši poglavlja.", file=sys.stderr)
            return 2
        rukopis = []
        for pg in pogl:
            with open(pg["put"], encoding="utf-8") as f:
                rukopis.append({**pg, "blokovi": R.parsiraj(f.read())})
        # Gotov sažetak (.katedra/sazetak.md, uz mapu poglavlja) ide u prednji
        # dio umjesto predloška „[Sažetak: 150–250 riječi…]".
        sazetak_put = os.path.join(os.path.dirname(os.path.abspath(mapa)), "sazetak.md")
        if os.path.isfile(sazetak_put):
            with open(sazetak_put, encoding="utf-8") as f:
                meta["sazetak_blokovi"] = R.parsiraj(f.read())

    try:
        rez = gradi(profil, meta, poglavlja, args.out, rukopis=rukopis)
    except ImportError:
        print("❌ treba python-docx:  pip install python-docx --break-system-packages",
              file=sys.stderr)
        return 2

    print(f"[rad → {args.out}] {len(rez['dijelovi'])} obaveznih dijelova, "
          f"{rez['poglavlja']} poglavlja")
    if rukopis:
        rijeci = sum(p.get("rijeci", 0) for p in rukopis)
        print(f"   sastavljeno iz rukopisa: {len(rukopis)} datoteka, ~{rijeci} riječi.")
        print("   Markdown je izvor istine — ručne izmjene u Wordu gube se pri "
              "sljedećem sastavljanju.")
    print("   sadržaj i popisi su POLJA — otvori u Wordu i osvježi (Ctrl+A, F9).")
    upozorenje = upozorenje_numeracije(profil)
    if upozorenje:
        print(upozorenje)

    if args.provjeri:
        print()
        return 1 if provjeri(args.out, args) else 0
    return 0


if __name__ == "__main__":
    try:
        sys.exit(main())
    except BrokenPipeError:
        os._exit(0)
    except KeyboardInterrupt:
        sys.exit(130)
