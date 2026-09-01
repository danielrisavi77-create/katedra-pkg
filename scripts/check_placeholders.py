#!/usr/bin/env python3
"""Provjera je li u radu ostao ijedan placeholder prije predaje.

`SKILL.md` (željezno pravilo 2) i `references/pisanje.md` nalažu da se tvrdnja bez
potpore u priloženoj građi označi s `[TREBA IZVOR]`, a stranica koja se ne može
potvrditi s `[PROVJERI STR.]` — umjesto da se izmišlja izvor. To je ispravno
ponašanje tijekom pisanja. `references/predaja.md` §2 zatim navodi „nijedan
[TREBA IZVOR] ni [PROVJERI STR.] nije ostao u tekstu" kao **blokirajuću** stavku
pred predaju.

Do audita tu stavku nije provjeravao nijedan alat: token se u cijelom paketu
spominjao samo u jednoj savjetodavnoj poruci `plan_state.py`. Rad se dakle mogao
uvezati i predati s uglatom zagradom u tekstu, što je najsramotniji način da
formalno uredan rad padne.

Provjeravaju se i tablične ćelije i fusnote, ne samo odlomci — placeholder uz
brojku u tablici („19,6 % [PROVJERI STR.]") je tipičan slučaj, a upravo je
propuštanje tablica bio zaseban nalaz audita nad `verify_rewrite.py`.

Uporaba:
  python3 <KATEDRA_SKILL>/scripts/check_placeholders.py ./rad.docx
  python3 <KATEDRA_SKILL>/scripts/check_placeholders.py ./rad.docx --json ./.katedra/placeholders.json
  python3 <KATEDRA_SKILL>/scripts/check_placeholders.py ./poglavlje.md --dodatni "[MOJ TAG]"

Izlazni kodovi:
  0  nijedan placeholder nije pronađen
  1  pronađen barem jedan — rad nije spreman za predaju
  2  datoteka se ne može pročitati
"""
from __future__ import annotations

import argparse
import json
import os
import re
import sys
import zipfile

HERE = os.path.dirname(os.path.abspath(__file__))
if HERE not in sys.path:
    sys.path.insert(0, HERE)

# Kanonski tokeni koje skill sam emitira (SKILL.md pravilo 2, pisanje.md).
KANONSKI = ["[TREBA IZVOR]", "[PROVJERI STR.]"]

# Šire mreže: generički TODO markeri i uglata zagrada s VELIKIM slovima, koja je
# u hrvatskom akademskom tekstu gotovo uvijek ostatak radnog procesa, a ne
# legitiman sadržaj. Legitimne iznimke (npr. „[sic]", „[ur.]", „[prev. aut.]")
# pišu se malim slovima pa ne upadaju u ovaj obrazac.
_GENERICKI_RE = re.compile(
    r"\[(?:[A-ZČĆŽŠĐ][A-ZČĆŽŠĐ0-9 .,_/–-]{2,40})\]"
)
_TODO_RE = re.compile(r"\b(?:TODO|FIXME|XXX|HACK|TBD)\b")

NASLOV_STIL_RE = re.compile(r"^(Heading|Naslov)\s*\d*", re.IGNORECASE)


class GreskaUlaza(Exception):
    """Ulazna datoteka se ne može pročitati — izlazni kod 2."""


def _tekst_fusnota(put: str) -> list[str]:
    """Izvuci tekst fusnota iz .docx zipa.

    python-docx ne izlaže fusnote kroz javni API, pa se `word/footnotes.xml`
    čita izravno. Ako dijela nema (dokument bez fusnota), vraća se prazan popis.
    """
    try:
        with zipfile.ZipFile(put) as z:
            if "word/footnotes.xml" not in z.namelist():
                return []
            xml = z.read("word/footnotes.xml").decode("utf-8", "replace")
    except (OSError, zipfile.BadZipFile, KeyError):
        return []
    # Sirovo izvlačenje <w:t>…</w:t> je dovoljno: tražimo doslovne tokene.
    return [
        re.sub(r"<[^>]+>", "", dio)
        for dio in re.findall(r"<w:t[^>]*>.*?</w:t>", xml, re.DOTALL)
    ]


def _iz_docxa(put: str) -> list[tuple[str, str, str]]:
    """Vrati (mjesto, kontekst, tekst) za svaki tekstualni blok .docx-a.

    `mjesto` je zadnji viđeni naslov (ili „<bez naslova>"), `kontekst` je vrsta
    bloka: odlomak, tablica, fusnota.
    """
    try:
        import docx  # noqa: PLC0415 — namjerno lijeno, skripta radi i za .md/.txt
    except ImportError as exc:
        raise GreskaUlaza(
            "za .docx treba python-docx.\n"
            "   Što napraviti: pip install python-docx, ili predaj .md/.txt izvoz rada."
        ) from exc
    try:
        d = docx.Document(put)
    except Exception as exc:  # noqa: BLE001 — vanjska biblioteka, poruka ide korisniku
        raise GreskaUlaza(
            f"{put} se ne može otvoriti kao .docx: {exc}\n"
            "   Što napraviti: provjeri je li datoteka stvarno .docx (ne .doc preimenovan)."
        ) from exc

    blokovi: list[tuple[str, str, str]] = []
    trenutni_naslov = "<bez naslova>"
    for p in d.paragraphs:
        stil = (p.style.name if p.style is not None else "") or ""
        tekst = p.text or ""
        if NASLOV_STIL_RE.match(stil) and tekst.strip():
            trenutni_naslov = tekst.strip()
        if tekst.strip():
            blokovi.append((trenutni_naslov, "odlomak", tekst))

    def _prodji_tablice(tablice, naslov: str) -> None:
        for t in tablice:
            for red in t.rows:
                for celija in red.cells:
                    for p in celija.paragraphs:
                        if (p.text or "").strip():
                            blokovi.append((naslov, "tablica", p.text))
                    if celija.tables:
                        _prodji_tablice(celija.tables, naslov)

    _prodji_tablice(d.tables, trenutni_naslov)

    for tekst in _tekst_fusnota(put):
        if tekst.strip():
            blokovi.append(("<fusnota>", "fusnota", tekst))

    return blokovi


def _iz_teksta(put: str) -> list[tuple[str, str, str]]:
    try:
        sadrzaj = open(put, encoding="utf-8").read()
    except (OSError, UnicodeDecodeError) as exc:
        raise GreskaUlaza(
            f"{put} se ne može pročitati: {exc}\n"
            "   Što napraviti: datoteka mora biti .docx, ili .md/.txt u UTF-8."
        ) from exc
    blokovi = []
    naslov = "<bez naslova>"
    for red in sadrzaj.split("\n"):
        if red.startswith("#"):
            naslov = red.lstrip("#").strip() or naslov
        if red.strip():
            blokovi.append((naslov, "odlomak", red))
    return blokovi


def ucitaj_blokove(put: str) -> list[tuple[str, str, str]]:
    if not os.path.isfile(put):
        raise GreskaUlaza(f"nema datoteke: {put}")
    if put.lower().endswith(".docx"):
        return _iz_docxa(put)
    return _iz_teksta(put)


def pronadji(blokovi, dodatni: list[str], samo_kanonski: bool) -> list[dict]:
    """Vrati popis nalaza; svaki nosi mjesto, vrstu bloka, token i kontekst."""
    tokeni = list(KANONSKI) + list(dodatni or [])
    nalazi = []
    for redni, (mjesto, vrsta, tekst) in enumerate(blokovi):
        pogodci: list[tuple[str, int]] = []
        for tok in tokeni:
            start = 0
            while True:
                i = tekst.find(tok, start)
                if i < 0:
                    break
                pogodci.append((tok, i))
                start = i + len(tok)
        if not samo_kanonski:
            for m in _GENERICKI_RE.finditer(tekst):
                if m.group(0) not in tokeni:
                    pogodci.append((m.group(0), m.start()))
            for m in _TODO_RE.finditer(tekst):
                pogodci.append((m.group(0), m.start()))
        for tok, i in sorted(pogodci, key=lambda x: x[1]):
            lijevo = max(0, i - 45)
            nalazi.append({
                "blok": redni,
                "mjesto": mjesto,
                "vrsta": vrsta,
                "token": tok,
                "kanonski": tok in KANONSKI,
                "kontekst": ("…" if lijevo else "") + tekst[lijevo:i + len(tok) + 45].strip(),
            })
    return nalazi


def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(
        description="Provjeri je li u radu ostao ijedan placeholder ([TREBA IZVOR], "
                    "[PROVJERI STR.], TODO…) prije predaje."
    )
    ap.add_argument("rad", help="rad (.docx) ili poglavlje (.md/.txt)")
    ap.add_argument("--dodatni", action="append", default=[],
                    help="dodatni doslovan token (može više puta)")
    ap.add_argument("--samo-kanonski", action="store_true",
                    help="traži samo [TREBA IZVOR] i [PROVJERI STR.], bez TODO/generičkih")
    ap.add_argument("--json", dest="json_out", metavar="PUT", help="zapiši nalaze u JSON")
    args = ap.parse_args(argv)

    try:
        blokovi = ucitaj_blokove(args.rad)
    except GreskaUlaza as exc:
        print(f"❌ {exc}", file=sys.stderr)
        return 2

    nalazi = pronadji(blokovi, args.dodatni, args.samo_kanonski)

    print("=" * 78)
    print(f"PLACEHOLDERI — {os.path.basename(args.rad)}")
    print("=" * 78)
    print(f"pregledano blokova: {len(blokovi)} (odlomci + tablične ćelije + fusnote)")

    if nalazi:
        kanonskih = sum(1 for n in nalazi if n["kanonski"])
        print(f"\n❌ pronađeno {len(nalazi)} placeholder(a), od toga {kanonskih} "
              f"kanonskih — rad NIJE spreman za predaju:\n")
        for n in nalazi[:40]:
            print(f"  [{n['vrsta']}] {n['token']}  ·  {n['mjesto']}")
            print(f"      „{n['kontekst']}”")
        if len(nalazi) > 40:
            print(f"  … i još {len(nalazi) - 40} (svi su u --json izlazu)")
        print("\n   Što napraviti: svaki [TREBA IZVOR] razriješi stvarnim izvorom iz "
              "građe\n   (ili ukloni tvrdnju), a svaki [PROVJERI STR.] potvrdi u izvorniku.")
    else:
        print("\n✅ nijedan placeholder nije ostao u tekstu.")

    if args.json_out:
        os.makedirs(os.path.dirname(os.path.abspath(args.json_out)) or ".", exist_ok=True)
        tmp = args.json_out + ".tmp"
        with open(tmp, "w", encoding="utf-8") as f:
            json.dump({
                "schema_version": 1,
                "alat": "check_placeholders",
                "rad": os.path.abspath(args.rad),
                "blokova": len(blokovi),
                "nalazi": nalazi,
            }, f, ensure_ascii=False, indent=1)
            f.write("\n")
        os.replace(tmp, args.json_out)
        print(f"\n[placeholderi → {args.json_out}]")

    return 1 if nalazi else 0


if __name__ == "__main__":
    try:
        sys.exit(main())
    except BrokenPipeError:
        os._exit(0)
    except KeyboardInterrupt:
        sys.exit(130)
