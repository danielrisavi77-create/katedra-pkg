"""Zajedničke funkcije za audit skripte. Ovisi samo o python-docx (+ standardna lib)."""
import re
import sys
import zipfile
import html


def die(msg):
    print(msg, file=sys.stderr)
    sys.exit(1)


def load_docx_text(path, include_tables=True):
    """Vrati čist tekst rada (odlomci + opcionalno ćelije tablica) preko python-docx.
    NE koristi regex po XML-u (povlači markup)."""
    try:
        from docx import Document
    except ImportError:
        die("Treba python-docx:  pip install python-docx --break-system-packages")
    d = Document(path)
    paras = [p.text for p in d.paragraphs]
    cells = []
    if include_tables:
        for t in d.tables:
            for row in t.rows:
                for c in row.cells:
                    cells.append(c.text)
    return "\n".join(paras), cells, d


def read_document_xml(path):
    """Sirovi word/document.xml (za provjeru polja, prijeloma, stilova)."""
    with zipfile.ZipFile(path) as z:
        return z.read("word/document.xml").decode("utf-8")


def read_part(path, part):
    with zipfile.ZipFile(path) as z:
        try:
            return z.read(part).decode("utf-8")
        except KeyError:
            return ""


_W_T_RE = re.compile(r"<w:t\b[^>]*>(.*?)</w:t>", re.S)


def _text_from_part_xml(xml):
    if not xml:
        return ""
    return "".join(html.unescape(m) for m in _W_T_RE.findall(xml))


def load_supplementary_text(path):
    """Tekst iz fusnota/endnota/headera/footera — python-docx .paragraphs ih NE
    pokriva, pa citiranje u fusnotama (Chicago stil) i brojke u zaglavljima/podnožjima
    ostaju nevidljive ostatku audita ako se ovo ne doda posebno.

    Vraća dict: {'footnotes': str, 'endnotes': str, 'headers': str, 'footers': str}
    (svaki str je spljošteni tekst iz odgovarajućih XML dijelova, prazan ako dio ne postoji)."""
    out = {"footnotes": "", "endnotes": "", "headers": "", "footers": ""}
    try:
        with zipfile.ZipFile(path) as z:
            names = z.namelist()

            def read(part):
                try:
                    return z.read(part).decode("utf-8")
                except KeyError:
                    return ""

            out["footnotes"] = _text_from_part_xml(read("word/footnotes.xml"))
            out["endnotes"] = _text_from_part_xml(read("word/endnotes.xml"))
            headers = sorted(n for n in names if re.match(r"word/header\d+\.xml$", n))
            footers = sorted(n for n in names if re.match(r"word/footer\d+\.xml$", n))
            out["headers"] = "\n".join(_text_from_part_xml(read(n)) for n in headers)
            out["footers"] = "\n".join(_text_from_part_xml(read(n)) for n in footers)
    except Exception:
        pass  # read-only pomoćna funkcija — ne rušimo audit ako dio nedostaje/je oštećen
    return out


ABBREVIATIONS = [
    "npr.", "tzv.", "d.o.o.", "j.d.o.o.", "god.", "str.", "sl.", "itd.", "itsl.",
    "tj.", "odn.", "br.", "prof.", "doc.", "dr.", "mr.", "sc.", "ing.", "vol.",
    "gl.", "\u010dl.", "st.", "sur.", "ur.", "izd.", "nakl.", "i.e.", "e.g.",
]
_ABBR_PLACEHOLDER = "\x01"  # sentinel koji se ne pojavljuje u normalnom tekstu

# Regex koji matcha kraticu SAMO kao samostalnu rije\u010d (ispred nje ne smije biti
# slovo) \u2014 obi\u010dan substring replace bi progutao to\u010dku na kraju svake rije\u010di
# koja ZAVR\u0160AVA kraticom: "nosivost." sadr\u017ei "st.". U hrvatskom su -ost imenice
# na kraju re\u010denice posvuda, pa bi statistika re\u010denica bila sustavno iskrivljena.
_ABBR_RE = re.compile(
    r"(?<![\w\u010d\u0107\u017e\u0161\u0111\u010c\u0106\u017d\u0160\u0110])("
    + "|".join(re.escape(ab) for ab in sorted(ABBREVIATIONS, key=len, reverse=True))
    + r")",
    re.IGNORECASE,
)


def sentences(text):
    """Grubo dijeljenje na re\u010denice (hr).

    Prije dijeljenja za\u0161titi uobi\u010dajene kratice (npr., tzv., d.o.o., god., str.\u2026)
    tako da njihova to\u010dka ne bude pogre\u0161no protuma\u010dena kao kraj re\u010denice.
    Kratica se \u0161titi samo kao samostalna rije\u010d (v. _ABBR_RE), ne kao sufiks
    ("nosivost." NIJE "st.")."""
    text = re.sub(r"\s+", " ", text)
    protected = _ABBR_RE.sub(lambda m: m.group(0).replace(".", _ABBR_PLACEHOLDER), text)
    parts = re.split(r"(?<=[\.\!\?])\s+", protected)
    parts = [p.replace(_ABBR_PLACEHOLDER, ".") for p in parts]
    return [s.strip() for s in parts if len(s.strip()) > 3]


def parse_citation_group(inner):
    """'[19, 21]' / '[19–22]' -> set brojeva."""
    nums = set()
    for part in re.split(r",", inner):
        part = part.strip()
        rng = re.split(r"[–\-]", part)
        if len(rng) == 2 and rng[0].strip().isdigit() and rng[1].strip().isdigit():
            nums.update(range(int(rng[0]), int(rng[1]) + 1))
        elif part.isdigit():
            nums.add(int(part))
    return nums


# ---------------------------------------------------------------------------
# Autor-godina (APA/Harvard) citiranje — pored numeričkog IEEE [N] stila.
# ---------------------------------------------------------------------------

IEEE_CITE_RE = re.compile(r"\[\d{1,3}(?:[\s,–\-]+\d{1,3})*\]")

# Lokator stranice iza godine: "(Becker, 2007: 45)", "(Streeck, 2014: xiv)".
# FPZG Upute propisuju BAŠ taj oblik s dvotočkom. Bez njega ispada svaki citat sa
# stranicom — a to je citat koji je BOLJE napisan od golog — pa referenca postane
# „siroče" i izvještaj prijavi nepostojeće greške. Najgora vrsta kvara: alat
# kažnjava rad zato što je precizniji.
LOKATOR = r"(?:\s*:\s*[\dxivlcdmXIVLCDM]+(?:\s*[-–]\s*[\dxivlcdmXIVLCDM]+)?)?"

# Parentetički citat koji završava s ", GODINA" (dopušta ; za više grupa u istoj zagradi)
# HR APA: točka iza godine je dopuštena — "(Čavlek i sur., 2011.)"
CITE_AY_RE = re.compile(
    r"\(([^()]{2,160}?,\s*\d{4}\.?[a-z]?" + LOKATOR +
    r"(?:\s*;\s*[^()]{2,160}?,\s*\d{4}\.?[a-z]?" + LOKATOR + r")*)\)"
)

# Narativni citat: "Faulkner (2001.)", "Hall, Prayag i Amore (2018.)".
# U hrvatskim radovima čini VEĆINU citata; bez njega je brojanje besmisleno.
# Mora obuhvatiti i višečlana imena ("TUI AG", "UN Tourism",
# "Načinović Braje") te čestice ("van Heiningen") — inače se sidri na
# zadnju riječ i ključ ispadne "ag" umjesto "tui".
# Institucionalni autor nosi malu riječ u imenu („Europska komisija", „Hrvatska
# narodna banka", „Državni zavod za statistiku"). Uzorak je dopuštao samo velike
# riječi i čestice, pa je takav narativni citat ispadao, a njegov redak u popisu
# literature postajao lažno „siroče". Broj malih riječi je OGRANIČEN na dvije:
# bez granice bi „Analiza je provedena u razdoblju (2021.)" prošlo kao citat.
_MALA_RIJEC = r"[a-zčćžšđ][\wÀ-ɏ'’\-]+"
CITE_AY_NARRATIVE_RE = re.compile(
    r"\b([A-ZČĆŽŠĐ][\wÀ-ɏ'’\-]+"
    r"(?:[\s,]+(?:i|te|sur\.|suradnici|dr\.|van|von|de|del|di|da"
    r"|[A-ZČĆŽŠĐ][\wÀ-ɏ'’\-]+))*"
    r"(?:\s+" + _MALA_RIJEC + r"){0,2})"
    r"\s*\((\d{4})\.?([a-z]?)" + LOKATOR + r"\)"
)

# Čestice u prezimenu ("van der Zwan", "de Vries", "von Hayek"). Ključ se gradi od
# PRVE riječi koja nije čestica: "TUI AG" → "tui" (institucija, prva riječ nosi
# identitet), "Van der Zwan" → "zwan". Prije je pravilo bilo „uzmi prvu riječ", pa
# je isto prezime davalo „van" iz teksta i „zwan" iz popisa literature — dvije
# funkcije istog alata, dva ključa, i uredna referenca ispadne siroče.
CESTICE = {"van", "von", "de", "del", "della", "di", "da", "dos", "der", "den",
           "la", "le", "el", "al", "ten", "ter", "af", "av", "bin", "ibn", "mac"}


def kljuc_prezimena(ime):
    """Ključ autora iz imena: prva riječ koja nije čestica, mala slova."""
    rijeci = [r for r in re.split(r"[\s,]+", (ime or "").strip()) if r]
    for r in rijeci:
        cista = r.strip(".").lower()
        if cista and cista not in CESTICE:
            return cista
    return rijeci[0].strip(".").lower() if rijeci else ""


def parse_ay_narrative(text):
    """Skup (prvo_prezime, godina) iz narativnih citata."""
    out = set()
    for imena, god, sufiks in CITE_AY_NARRATIVE_RE.findall(text):
        out.add((kljuc_prezimena(imena), (god + sufiks).lower()))
    return out


def parse_ay_segment(seg):
    """'Ivić i Perić, 2020a' -> ('ivić', '2020a'). None ako ne prepozna oblik."""
    # Sufiks (2013a / 2013b) je DIO identiteta jedinice: dva rada istog autora iz
    # iste godine inače se slijevaju u jedan ključ i jedan od njih uvijek ispadne
    # siroče. Narativni parser sufiks je zadržavao, zagradni ga je odbacivao — pa
    # su dvije funkcije istog alata davale različite ključeve za isti citat.
    # Lokator stranice iza godine ("Becker, 2007: 45") ovdje se prepoznaje i
    # odbacuje: on je oznaka mjesta u izvoru, ne dio identiteta.
    m = re.match(r"\s*(.+?),\s*(\d{4})\.?([a-z]?)" + LOKATOR + r"\s*$", seg.strip())
    if not m:
        return None
    author_part, year = m.group(1), m.group(2) + m.group(3)
    kljuc = kljuc_prezimena(author_part)
    if not kljuc or not kljuc[:1].isalpha():
        return None
    return (kljuc, year.lower())


def parse_ay_citation_group(inner):
    """Sadržaj cijele zagrade (može imati više '; '-odvojenih grupa) -> set (prezime, godina)."""
    keys = set()
    for seg in re.split(r"\s*;\s*", inner):
        k = parse_ay_segment(seg)
        if k:
            keys.add(k)
    return keys


def detect_citation_style(text):
    """Heuristička detekcija: 'ieee' (numerički [N]), 'authoryear' (Prezime, GODINA),
    'mixed' (oboje u sličnoj mjeri) ili 'unknown' (nema dovoljno signala).
    Vraća (stil, {'ieee': n, 'authoryear': n})."""
    ieee_n = len(IEEE_CITE_RE.findall(text))
    ay_n = sum(len(parse_ay_citation_group(m)) for m in CITE_AY_RE.findall(text))
    counts = {"ieee": ieee_n, "authoryear": ay_n}
    if ieee_n == 0 and ay_n == 0:
        return "unknown", counts
    if ieee_n >= ay_n * 1.5:
        return "ieee", counts
    if ay_n >= ieee_n * 1.5:
        return "authoryear", counts
    return "mixed", counts
